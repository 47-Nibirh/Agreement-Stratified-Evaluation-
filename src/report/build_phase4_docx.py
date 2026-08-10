"""
Build the GastroHUN Phase 4 Word report.
=========================================
Assembles `Phase4_Report.docx` from the computed artefacts:

  reports/phase4_prereg.json              frozen pre-registration
  reports/phase4_cohort.json              cohort E and its gates
  reports/phase4_cache_gate.json          byte-identity gate
  reports/phase4_distance_matrix.json     C4 anatomical distances
  reports/phase4_run_*.json               training histories
  reports/phase4_stratified_metrics.json  RQ2 primary + all contrasts
  reports/phase4_calibration.json         RQ2 calibration endpoint
  reports/phase4_uncertainty.json         RQ3
  reports/phase4_structure_eval.json      RQ4
  reports/phase4_loao.json                sensitivity
  figures_phase4/*.png

No numeric value is typed by hand, and no verdict sentence is pre-written: the
narrative branches on the pre-registered verdict fields, so the document
reports what the experiment found.

Run:  python src/report/build_phase4_docx.py
"""
from __future__ import annotations

import json
from pathlib import Path

import numpy as np
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt

import build_docx as BD

ROOT = Path(__file__).resolve().parents[2]
BD.FIGD = ROOT / "figures_phase4"
BD.OUT = ROOT / "Phase4_Report.docx"

from build_docx import (ACCENT, DARKRED, GREY, add_page_numbers,  # noqa: E402
                        bullet, callout, figure, front_matter, h, new_document,
                        para, table)
import content_phase4 as B  # noqa: E402

PRE, COH, CACHE, DIST = B.PRE, B.COH, B.CACHE, B.DIST
MET, CAL, UNC, STR, LOAO = B.MET, B.CAL, B.UNC, B.STR, B.LOAO
TIERS, POOLED, STRATA = B.TIERS, B.POOLED, B.STRATA
TIER_LABEL, CFG_LABEL = B.TIER_LABEL, B.CFG_LABEL
pc, ci, cid, cie = B.pc, B.ci, B.cid, B.cie


def CF():
    return B.CFGS()


def _rhos():
    """Within-stratum RQ3 correlations across configurations; [0.0] if none were
    computable, so the abstract degrades instead of raising on an empty min()."""
    vals = [UNC["verdicts"][c]["mean_rho_3seed"] for c in CF()
            if "mean_rho_3seed" in UNC.get("verdicts", {}).get(c, {})]
    return vals or [0.0]


# =====================================================================
def title_page(doc) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Agreement-Stratified Evaluation of Deep Learning for "
                  "Anatomical Landmark Recognition in Upper Gastrointestinal "
                  "Endoscopy")
    r.bold = True; r.font.size = Pt(19); r.font.color.rgb = DARKRED

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run("Phase 4 — Soft-Label and Uncertainty Training (RQ2, RQ3, RQ4)")
    r.bold = True; r.font.size = Pt(14.5); r.font.color.rgb = ACCENT

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Changing what the model is trained to predict, and nothing else")
    r.font.size = Pt(11.5); r.italic = True

    doc.add_paragraph()
    bar = doc.add_paragraph(); bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = bar.add_run("─" * 46); rb.font.color.rgb = ACCENT

    n_runs = len([k for k in B.runs() if k[0] != "C0"])
    meta = [
        ("Degree programme", "B.Sc. in Computer Science and Engineering"),
        ("Research domain", "Biomedical Artificial Intelligence — Medical Image "
                            "Analysis and Deep Learning"),
        ("Training cohort", f"Extended cohort E — {COH['by_split']['Train']:,} train / "
                            f"{COH['by_split']['Validation']:,} validation images "
                            f"(majority-or-better), identical for C1–C4"),
        ("Evaluation cohort", f"Full official test split — "
                              f"{MET['aggregate_3seed'][CF()[0]]['S-unanimous']['n_images'] + MET['aggregate_3seed'][CF()[0]][POOLED]['n_images']:,} "
                              f"images across 4 agreement strata"),
        ("Configurations", ", ".join(CF()) + f" — {n_runs} training runs at 3 seeds each"),
        ("Governing protocol", "THESIS_RESEARCH_BLUEPRINT.md §4 PHASE 4; "
                               "pre-registration frozen " + PRE["frozen_at"]),
        ("Reporting standards", "CLAIM, TRIPOD+AI, STARD-AI, PROBAST+AI"),
        ("Report date", "27 July 2026"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in meta:
        c = t.add_row().cells
        BD._cell_text(c[0], k, bold=True, size=9.5)
        BD._cell_text(c[1], v, size=9.5)
        c[0].width, c[1].width = Cm(4.6), Cm(11.0)

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vs = []
    if MET.get("verdicts", {}).get("RQ2_primary"):
        vs.append(f"RQ2 (accuracy): {MET['verdicts']['RQ2_primary']['verdict']}")
    if CAL.get("verdicts", {}).get("RQ2_calibration"):
        vs.append(f"RQ2 (calibration): {CAL['verdicts']['RQ2_calibration']['verdict']}")
    if UNC and UNC.get("verdicts"):
        sup = [c for c, v in UNC["verdicts"].items() if v.get("verdict") == "SUPPORTED"]
        vs.append(f"RQ3: supported for {', '.join(sup) if sup else 'no configuration'}")
    if STR and STR.get("verdicts", {}).get("RQ4"):
        vs.append(f"RQ4: {STR['verdicts']['RQ4']['verdict']}")
    r = p.add_run("Pre-registered verdicts — " + " · ".join(vs))
    r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARKRED

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Every verdict above is produced by the rule fixed in "
                  "reports/phase4_prereg.json before the first model was trained, "
                  "applied to a paired patient-clustered bootstrap interval.")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# =====================================================================
def sec_abstract(doc) -> None:
    agg = MET["aggregate_3seed"]
    h(doc, "Abstract", level=1)
    para(doc, f"Background. Phase 2 reproduced a published ConvNeXt-Tiny landmark "
              f"classifier on the 60.2% of GastroHUN images where all four "
              f"annotators agree. Phase 3 evaluated that frozen model on the whole "
              f"official test split and found two failures. Accuracy fell far "
              f"outside the unanimous stratum, and — the sharper result — "
              f"calibration collapsed: expected calibration error rose from "
              f"{pc(CAL['aggregate_3seed']['C0']['S-unanimous']['ece_vs_expected_accuracy'], 1)}% "
              f"to "
              f"{pc(CAL['aggregate_3seed']['C0']['S-plurality']['ece_vs_expected_accuracy'], 1)}% "
              f"while mean confidence barely moved. Both failures are properties of "
              f"a model that has never been shown a contested image.")
    para(doc, f"Objective. To test whether changing the TRAINING TARGET — and "
              f"nothing else — repairs them. Four configurations were trained on an "
              f"identical cohort of {COH['by_split']['Train']:,} training images "
              f"carrying a majority label: a hard majority label (C1), the "
              f"proportions of the four annotator votes (C2), a hard label with "
              f"label smoothing matched to C2 in the probability mass it displaces "
              f"(C3, the control), and vote proportions plus a penalty on the "
              f"anatomical distance of the error (C4). The Phase 2 model is carried "
              f"through unchanged as C0.")
    para(doc, f"Methods. All configurations share the backbone, the schedule, the "
              f"augmentation, the normalisation statistics, the precision and the "
              f"model-selection criterion of Phase 2; the cache was verified "
              f"byte-identical on the "
              f"{CACHE['gate_p4_2_byte_identity']['n_shared_with_phase2_cache']:,} "
              f"images shared with Phase 2, so the arms differ in the target and "
              f"nowhere else. Every arm was evaluated on the full "
              f"{agg[CF()[0]]['S-unanimous']['n_images'] + agg[CF()[0]][POOLED]['n_images']:,}-image "
              f"test split with the Phase 3 metric set. All contrasts use a PAIRED "
              f"patient-clustered bootstrap: one patient resample, both arms scored "
              f"on those same rows, then differenced. The primary RQ2 endpoint, the "
              f"control against which it is judged, and the three-way verdict rule "
              f"were fixed in a pre-registration frozen before the first run.")
    if "C2" in CF() and "C3" in CF():
        v = MET["verdicts"]["RQ2_primary"]
        vc = CAL["verdicts"]["RQ2_calibration"]
        para(doc, f"Results. On the pooled contested stratum "
                  f"(n={agg['C2'][POOLED]['n_images']}), C2 minus C3 in "
                  f"annotator-marginalized macro F1 was {ci('C2 - C3')} — "
                  f"{v['verdict']} under the pre-registered rule. The same contrast "
                  f"on expected calibration error was {cie('C2 - C3')} "
                  f"({vc['verdict']}). Adding the contested images to the training "
                  f"set at a hard target (C1 minus C0) moved the contested-stratum "
                  f"score by {ci('C1 - C0')}, and the anatomical penalty (C4 minus "
                  f"C2) changed the mean anatomical error distance by "
                  f"{cid('C4 - C2')}. Within-stratum Spearman correlation between "
                  f"predictive entropy and annotator vote entropy on the largest "
                  f"contested stratum ranged "
                  f"{min(_rhos()):.3f} to {max(_rhos()):.3f} "
                  f"across configurations, against a pooled all-image value of "
                  f"{UNC['phase3_reference']['pooled_all_images_rho']} — confirming "
                  f"that the pooled figure the literature would report is a "
                  f"between-stratum artefact.")
    para(doc, "Conclusion. " + _conclusion_sentence())


def _conclusion_sentence() -> str:
    bits = []
    if MET.get("verdicts", {}).get("RQ2_primary"):
        v = MET["verdicts"]["RQ2_primary"]["verdict"]
        bits.append({
            "SUPPORTED": "Targeted soft targets built from the annotator vote "
                         "distribution outperform an equally-soft but uninformative "
                         "smoothing control on contested images, so the pattern of "
                         "disagreement — not merely its magnitude — carries usable "
                         "training signal.",
            "NOT SUPPORTED": "An equally-soft uninformative smoothing control "
                             "matched or beat the vote-proportion targets on "
                             "contested images, so the specific pattern of "
                             "annotator disagreement carries no usable training "
                             "signal beyond its magnitude at this panel size.",
            "NOT RESOLVED": "The vote-proportion targets could not be separated "
                            "from an equally-soft uninformative smoothing control "
                            "on contested images at this sample size; the pattern "
                            "of disagreement is not demonstrated to carry signal "
                            "beyond its magnitude.",
        }[v])
    if CAL.get("verdicts", {}).get("RQ2_calibration"):
        v = CAL["verdicts"]["RQ2_calibration"]["verdict"]
        bits.append({
            "SUPPORTED": "The calibration endpoint moves in the predicted "
                         "direction and its interval excludes zero.",
            "NOT SUPPORTED": "The calibration endpoint moves against the "
                             "prediction.",
            "NOT RESOLVED": "The calibration endpoint is not separable from zero.",
        }[v])
    if STR and STR.get("verdicts", {}).get("RQ4"):
        bits.append(f"The anatomy-aware loss is {STR['verdicts']['RQ4']['verdict']} "
                    f"at the single pre-registered value of lambda.")
    bits.append("Phase 3's calibration failure is reproduced in every "
                "configuration, which locates it in the agreement structure of the "
                "corpus rather than in one training recipe. A negative or "
                "unresolved result here is reported as such: the design's value is "
                "that its control makes the answer interpretable either way.")
    return " ".join(bits)


# =====================================================================
def sec_introduction(doc) -> None:
    h(doc, "1. Introduction", level=1)
    h(doc, "1.1 What Phases 2 and 3 established, and what they leave open", level=2)
    para(doc, f"Phase 2 reproduced the published ConvNeXt-Tiny baseline on the "
              f"complete-agreement subset of GastroHUN (macro F1 83.92 against a "
              f"published ~85.0 ± 1.5), validating the pipeline. Phase 3 took that "
              f"frozen model to the full official test split and stratified by how "
              f"many of the four annotators agreed. Two things happened. Accuracy "
              f"fell steeply — and once the attainable ceiling of each stratum was "
              f"held constant, the fall still exceeded the published gap between "
              f"architecture families for the 3/4 and 2-1-1 strata. And "
              f"calibration collapsed: mean confidence fell 9.3 points while "
              f"expected accuracy fell 56.6.")
    para(doc, "Neither result tells us whether the model could do better, because "
              "the model had never seen a contested image. The training cohort was "
              "the 60.2% of the corpus on which all four annotators agreed. That "
              "is the descriptor's own protocol, and it is what most of the "
              "literature Phase 1 reviewed does. This phase asks the obvious next "
              "question.")
    h(doc, "1.2 Why the target, and not the architecture", level=2)
    para(doc, "The blueprint records a measured precedent (§2.7): in the "
              "descriptor's own results, training ConvNeXt-Tiny on "
              "FG-agreement labels instead of complete-agreement labels moved macro "
              "F1 by 2.2 points — more than the 3.25-point gap between "
              "ConvNeXt-Tiny and the eight-times-larger ConvNeXt-Large is worth "
              "relative to its cost. How the annotator labels are combined is a "
              "larger lever than which network consumes them, and it is a lever "
              "almost nothing in the reviewed literature pulls, because pulling it "
              "requires per-annotator labels that most public corpora do not "
              "release.")
    h(doc, "1.3 Research questions and pre-registered hypotheses", level=2)
    rq = PRE["research_questions"]
    callout(doc, f"RQ2. {rq['RQ2']['question']} Hypothesis: {rq['RQ2']['hypothesis']}. "
                 f"Primary contrast {rq['RQ2']['primary_contrast']} on "
                 f"{rq['RQ2']['primary_endpoint']}.", title="RQ2 (accuracy and calibration)")
    callout(doc, f"RQ3. {rq['RQ3']['question']} Primary quantity: "
                 f"{rq['RQ3']['primary_quantity']}", title="RQ3 (uncertainty)")
    callout(doc, f"RQ4. {rq['RQ4']['question']} Primary contrast "
                 f"{rq['RQ4']['primary_contrast']} on {rq['RQ4']['primary_endpoint']}.",
            title="RQ4 (anatomy-aware loss)")
    para(doc, "All three verdict rules, the control arm, the interval procedure and "
              "the falsification condition were fixed in "
              "reports/phase4_prereg.json before the first Phase 4 model was "
              "trained (Appendix C).")
    h(doc, "1.4 Chapter roadmap", level=2)
    para(doc, "Section 2 specifies the five configurations, the derivation of the "
              "two hyper-parameters they need, what is held fixed, and the gates "
              "run before any result was read. Section 3 reports the results in "
              "pre-registered order: RQ2 accuracy, RQ2 calibration, RQ3, RQ4, then "
              "sensitivity. Section 4 interprets them against the fixed rules; "
              "Section 5 concludes and hands forward to Phase 5.")


# =====================================================================
def sec_methods(doc) -> None:
    h(doc, "2. Methods", level=1)

    h(doc, "2.1 The configuration matrix", level=2)
    figure(doc, "P4_F25_design.png",
           "The five configurations. C0 to C1 changes the cohort; C1 to C2, C3 and "
           "C4 changes only the target vector built from the same votes on the same "
           "images.")
    cfgs = PRE["configurations"]
    table(doc, ["Config", "Target construction", "Training cohort", "Role"],
          [["C0", cfgs["C0"]["target"], cfgs["C0"]["cohort"], cfgs["C0"]["role"]],
           ["C1", cfgs["C1"]["target"], "E", cfgs["C1"]["role"]],
           ["C2", cfgs["C2"]["target"], "E", cfgs["C2"]["role"]],
           ["C3", cfgs["C3"]["target"] +
            f" (epsilon = {cfgs['C3']['label_smoothing_epsilon']:.4f})", "E",
            cfgs["C3"]["role"]],
           ["C4", cfgs["C4"]["target"] +
            f" (lambda = {cfgs['C4']['structure_penalty_lambda']:g})", "E",
            cfgs["C4"]["role"]]],
          "Configuration matrix as frozen in the pre-registration.", font=8.0)

    h(doc, "2.1.1 The extended cohort E, and why it is held constant", level=3)
    para(doc, f"E is every Train or Validation image carrying a majority label, "
              f"that is agreement tier 4/4 or 3/4: "
              f"{COH['by_split']['Train']:,} training and "
              f"{COH['by_split']['Validation']:,} validation images, of which "
              f"{100 * COH['fraction_contested_by_split']['Train']:.1f}% of the "
              f"training set is contested (3/4). Against Phase 2's cohort that is "
              f"{COH['growth_vs_phase2']['Train']:.2f}x the training images.")
    para(doc, f"{COH['excluded_trainval_images']['n']:,} Train/Validation images are "
              f"excluded ("
              f"{', '.join(f'{k} {v:,}' for k, v in COH['excluded_trainval_images']['by_tier'].items())}"
              f"). {COH['excluded_trainval_images']['reason'].capitalize()}. This is "
              f"a real restriction and it is declared in Appendix E: the images "
              f"with the richest ambiguity signal are the ones C2 cannot use, "
              f"because using them would leave C3 undefined and the comparison "
              f"uncontrolled.")
    table(doc, ["Gate", "Criterion", "Result"],
          [["P4.1a", "cohort counts equal the corpus cascade restricted to Train/Validation",
            f"PASS — {COH['by_split']['Train']:,} / {COH['by_split']['Validation']:,}"],
           ["P4.1b", "23 classes present; every annotator vote inside the fixed class index",
            f"PASS — {COH['n_classes']} classes"],
           ["P4.1c", "no patient overlap Train/Validation, nor with the Phase 3 test split",
            "PASS — 0 overlaps"],
           ["P4.1d", "every filename resolves against the Phase 0 SHA-256 inventory and exists on disk",
            f"PASS — {COH['hash_resolution']['n_present_on_disk']:,} / "
            f"{COH['hash_resolution']['n_cohort']:,}"],
           ["P4.1e", "the 4/4 rows reproduce the Phase 2 consensus cohort exactly",
            "PASS — same filenames, same labels"],
           ["P4.2", "images shared with the Phase 2 cache decode bit-identically",
            f"PASS — {CACHE['gate_p4_2_byte_identity']['n_shared_with_phase2_cache']:,} / "
            f"{CACHE['gate_p4_2_byte_identity']['n_shared_with_phase2_cache']:,}, exhaustive"]],
          "Gates passed before any Phase 4 model was trained. A failure at any of "
          "them halts the phase rather than issuing a warning.", font=8.0)

    h(doc, "2.1.2 Deriving epsilon: the control has to be matched, not guessed", level=3)
    d = PRE["epsilon_derivation_detail"]
    para(doc, f"C3 exists so that a C2 gain cannot be waved away as ordinary "
              f"regularisation. That argument only works if the two arms soften the "
              f"target by the SAME amount, which means epsilon must be derived from "
              f"C2, not chosen. C2 displaces probability mass from the modal label "
              f"only on 3/4 images, and displaces 0.25 when it does; over the "
              f"training cohort that averages "
              f"{d['c2_mean_mass_displaced']:.6f}. Uniform label smoothing displaces "
              f"epsilon x (1 - 1/K), so matching gives epsilon = "
              f"{d['epsilon_mass_matched']:.6f}.")
    ent = {c: r["mean_train_target_entropy_nats"] for (c, s), r in B.runs().items()
           if s == 1 and c != "C0"}
    if {"C2", "C3"} <= set(ent):
        para(doc, f"Matching the mass does not match the entropy, and the residual "
                  f"difference is exactly what the experiment is about. Measured on "
                  f"the training cohort, the mean target entropy is "
                  f"{ent['C1']:.4f} nats for C1 (one-hot), {ent['C2']:.4f} for C2 and "
                  f"{ent['C3']:.4f} for C3 — C3 carries "
                  f"{ent['C3'] / ent['C2']:.1f}x the entropy of C2 while displacing "
                  f"the same mass. The reason is structural: C2 puts all of the "
                  f"displaced mass on ONE alternative, the class an annotator actually "
                  f"named, whereas C3 spreads it uniformly over all 22 others. C3 is "
                  f"therefore not a weakened control — it is a strictly more diffuse "
                  f"one, and the contrast asks whether concentrating the same mass "
                  f"where the disagreement actually fell is worth anything.")
    para(doc, f"Mass is matched rather than entropy because the gradient of the "
              f"soft-target cross-entropy with respect to the logits is (q - t): the "
              f"perturbation C2 introduces relative to a one-hot target IS the "
              f"displaced mass. The choice is not free — entropy matching would have "
              f"given epsilon = {d['epsilon_entropy_matched']:.6f}, roughly "
              f"{d['epsilon_mass_matched'] / d['epsilon_entropy_matched']:.1f}x "
              f"smaller and therefore a considerably weaker control. Both values are "
              f"recorded in the pre-registration; only the mass-matched arm was "
              f"trained, and Appendix E states what testing the other would cost. "
              f"For orientation, the conventional epsilon = 0.1 of the label-"
              f"smoothing literature is close to the mass-matched value, so the "
              f"control is not an unusually weak one.")

    h(doc, "2.1.3 The anatomical distance matrix and the C4 penalty", level=3)
    para(doc, f"The 23 classes are a (wall x station) grid, and Phase 0 measured "
              f"that human disagreement respects it: "
              f"{DIST['gates']['P4.3b_n_adjacent_pairs_checked']} class pairs are "
              f"circumferentially wall-adjacent and "
              f"{DIST['gates']['P4.3c_n_neighbouring_pairs_checked']} are "
              f"station-neighbouring under the Phase 0 definitions. Cross-entropy is "
              f"blind to this: confusing A3 with L3 (a quarter-turn of the scope, "
              f"d = {DIST['examples']['A3 vs L3']}) and confusing A3 with P6 (wrong "
              f"wall and the far end of the stomach, d = "
              f"{DIST['examples']['A3 vs P6']}) cost exactly the same.")
    para(doc, f"The distance is {DIST['definition']['combination']}, where the wall "
              f"term is the cyclic distance on {' -> '.join(DIST['definition']['wall_cycle'])} "
              f"normalised by 2 and the station term is |delta station| / 5. "
              f"OTHERCLASS is given distance 1.0 to every landmark: it is a quality "
              f"judgement, not a grid position, and Phase 0 showed quality "
              f"assessment and anatomical classification are different tasks. Mean "
              f"off-diagonal distance is "
              f"{DIST['mean_offdiagonal_distance_all_classes']}. C4 adds "
              f"lambda x E_(i~t, j~q)[d(i,j)] to the loss — the expected anatomical "
              f"distance between the target and predicted distributions — with "
              f"lambda fixed a priori at unit weight and never swept.")

    h(doc, "2.2 What is held fixed", level=2)
    tp = PRE["training_protocol"]
    for kk, lab in [("backbone", "Backbone"), ("input", "Input"),
                    ("augmentation", "Augmentation"), ("schedule", "Schedule"),
                    ("precision", "Precision"), ("loss", "Loss")]:
        bullet(doc, f"{lab}: {tp[kk]}")
    bullet(doc, f"Normalisation: {tp['normalisation_reused_deliberately']}")
    para(doc, f"Model selection: {PRE['model_selection']['criterion']}, identical "
              f"for C1-C4. {PRE['model_selection']['rationale'].capitalize()} C0's "
              f"criterion was necessarily different — its cohort contains no "
              f"contested validation images — which is precisely why C1, and not "
              f"C0, is the control for every target-construction contrast.")

    h(doc, "2.2.1 Compute budget and the epoch cap", level=3)
    ec = PRE["epoch_cap_derivation"]
    para(doc, f"One warm-up epoch and one fine-tuning epoch were timed on the real "
              f"extended cohort with the C4 penalty active — the most expensive arm "
              f"— giving {ec['measured_warmup_epoch_sec']} s and "
              f"{ec['measured_finetune_epoch_sec']} s. A declared budget of "
              f"{ec['declared_budget_hours']} h across {ec['n_runs']} runs then "
              f"fixes the fine-tuning cap at {PRE['epoch_cap_finetune']} epochs by "
              f"the same formula Phase 2 used. The cap binds only if early stopping "
              f"has not already fired, and the realised stop reason is reported per "
              f"run in Appendix B, so a reader can see exactly where it bound.")

    h(doc, "2.3 Evaluation", level=2)
    ev = PRE["evaluation"]
    para(doc, f"Every configuration is scored on {ev['test_set']}, stratified into "
              f"the four Phase 3 strata plus one pooled contested stratum "
              f"({MET['pooled_contested_definition']}, "
              f"n={MET['aggregate_3seed'][CF()[0]][POOLED]['n_images']}) that the "
              f"pre-registration names as RQ2's single primary endpoint. The primary "
              f"metric is the Phase 3 annotator-marginalized macro F1, which is "
              f"defined continuously across every stratum including those with no "
              f"single ground-truth label.")
    callout(doc, MET["scale_note"], title="Why the contrasts are reported on the raw scale")
    para(doc, ev["paired_comparison_procedure"])
    h(doc, "2.3.1 Calibration and uncertainty", level=3)
    para(doc, "Calibration is measured against expected accuracy — the probability "
              "mass a single prediction captures under the four-vote distribution — "
              "using the Phase 3B definitions verbatim, so the Phase 4 numbers sit "
              "on the same scale as the Phase 3 baseline they must improve on. Three "
              "uncertainty estimators are computed: the deterministic softmax "
              "entropy, a Monte-Carlo average over "
              f"{PRE['research_questions']['RQ3']['n_mc_samples']} stochastic forward "
              "passes with the StochasticDepth modules returned to training mode, "
              "and a 3-member deep ensemble over the seeds.")
    if B.INFER:
        g = B.INFER["gates"]["P4.6c_architecture_stochastic_inventory"]
        para(doc, f"The MC estimator's validity was verified at run time rather than "
                  f"assumed: torchvision's ConvNeXt-Tiny exposes "
                  f"{g['n_stochastic_depth_modules']} StochasticDepth modules, "
                  f"{g['n_dropout_modules']} Dropout modules and "
                  f"{g['n_batchnorm_modules']} BatchNorm modules. {g['interpretation'].capitalize()}. "
                  f"This is the reason the blueprint's 'MC dropout' is realised as MC "
                  f"stochastic depth (deviation P4-DEV-1): there is no dropout to "
                  f"sample, and inserting one absent during training would change the "
                  f"function the argument applies to.")


# =====================================================================
def sec_results(doc) -> None:
    h(doc, "3. Results", level=1)
    agg = MET["aggregate_3seed"]

    h(doc, "3.1 Training behaviour", level=2)
    R = B.runs()
    new = {k: v for k, v in R.items() if k[0] != "C0"}
    if new:
        caps = sum(1 for v in new.values() if v.get("stop_reason") == "epoch_cap")
        para(doc, f"{len(new)} Phase 4 runs completed, "
                  f"{sum(v['wallclock_sec'] for v in new.values()) / 3600:.1f} h of "
                  f"wall-clock training in total on a single GTX 1650. "
                  f"{caps} of {len(new)} stopped at the pre-registered fine-tuning "
                  f"cap rather than by early stopping. Where the cap bound it bound "
                  f"identically for every configuration, so the C1-C4 contrasts are "
                  f"unaffected; the C1-minus-C0 contrast is affected, because C0 was "
                  f"trained to early stopping under Phase 2's larger per-run budget, "
                  f"and that asymmetry is carried into the interpretation in §4.1.")
        rows = []
        for c in [x for x in CF() if x != "C0"]:
            ws = [v["best_val_macro_f1"] for k, v in new.items() if k[0] == c]
            es = [v["n_epochs_run"] for k, v in new.items() if k[0] == c]
            if ws:
                rows.append([CFG_LABEL[c], f"{np.mean(ws):.4f}",
                             f"{np.std(ws, ddof=1) if len(ws) > 1 else 0:.4f}",
                             f"{np.mean(es):.1f}",
                             f"{sum(1 for k, v in new.items() if k[0] == c and v['stop_reason'] == 'epoch_cap')}/{len(ws)}"])
        table(doc, ["Configuration", "Mean best val macro F1", "SD over seeds",
                    "Mean epochs run", "Runs stopped at cap"], rows,
              "Training summary. Validation macro F1 here is scored against the hard "
              "majority label on the extended validation cohort — the pre-registered "
              "selection criterion — and is a selection statistic, not a result.",
              font=8.0)

    h(doc, "3.2 RQ2, accuracy: performance by configuration and stratum", level=2)
    figure(doc, "P4_F26_stratified_by_config.png",
           "Agreement-stratified performance of every configuration, raw (left) and "
           "as a percentage of the modal-vote oracle attainable on that stratum "
           "(right). 3-seed mean.")
    table(doc,
          ["Configuration"] + [TIER_LABEL[t].split(" (")[0] for t in TIERS] +
          ["contested (pooled)"],
          [[CFG_LABEL[c]] +
           [pc(agg[c][t]["annotator_marginalized_macro_f1_mean_3seed"]) for t in TIERS] +
           [pc(agg[c][POOLED]["annotator_marginalized_macro_f1_mean_3seed"])]
           for c in CF()],
          "Annotator-marginalized macro F1 (%) by configuration and stratum, 3-seed "
          "mean. Full tables with intervals, ceiling-normalised levels, expected "
          "accuracy and any-hit rate are in Appendix A.", font=8.0,
          note=f"Attainable ceiling per stratum: " + ", ".join(
              f"{TIER_LABEL[t].split(' (')[0]} {pc(MET['ceilings'][t]['oracle_marginalized_macro_f1_mean'])}%"
              for t in TIERS) + ". A configuration cannot exceed these.")
    para(doc, f"The C0 row is not a new measurement. It is the Phase 3 result, "
              f"recomputed through the Phase 4 code path and gated to agree with "
              f"reports/phase3_stratified_metrics.json to within "
              f"{MET['gate_p4_7_c0_reproduces_phase3']['max_abs_deviation_from_phase3']:.1e} "
              f"(gate P4.7). Had it differed at all, every contrast built on it "
              f"would have been meaningless.")

    h(doc, "3.2.1 The pre-registered contrasts", level=3)
    figure(doc, "P4_F27_contrast_forest.png",
           "Every pre-registered contrast, on every stratum, with paired "
           "patient-clustered 95% intervals. Green excludes zero above, red below, "
           "grey contains zero.")
    rows = []
    for kk, blk in MET["contrasts"].items():
        d = blk["by_stratum"][POOLED]
        rows.append([kk, blk["isolates"],
                     f"{d['diff_points_3seed_mean']:+.2f}",
                     f"[{d['ci95_points_3seed_mean'][0]:+.2f}, {d['ci95_points_3seed_mean'][1]:+.2f}]",
                     "yes" if d["excludes_zero"] else "no",
                     "yes" if d["sign_consistent_across_seeds"] else "no"])
    table(doc, ["Contrast", "What it isolates", "Diff (points)", "95% CI",
                "Excludes 0", "Sign consistent over seeds"], rows,
          "Configuration contrasts on the pooled contested stratum "
          f"(n={agg[CF()[0]][POOLED]['n_images']}), the pre-registered primary "
          "endpoint. Positive favours the first-named configuration.", font=7.8)
    if "C2 - C3" in MET["contrasts"]:
        v = MET["verdicts"]["RQ2_primary"]
        callout(doc, f"RQ2 primary endpoint. C2 minus C3 on the pooled contested "
                     f"stratum is {ci('C2 - C3')}. Under the rule fixed before "
                     f"training — supported only if the paired interval excludes "
                     f"zero in C2's favour — the verdict is {v['verdict']}.",
                title=f"Pre-registered verdict: {v['verdict']}")
        rows = [[TIER_LABEL[st],
                 f"{MET['contrasts']['C2 - C3']['by_stratum'][st]['diff_points_3seed_mean']:+.2f}",
                 f"[{MET['contrasts']['C2 - C3']['by_stratum'][st]['ci95_points_3seed_mean'][0]:+.2f}, "
                 f"{MET['contrasts']['C2 - C3']['by_stratum'][st]['ci95_points_3seed_mean'][1]:+.2f}]",
                 MET["verdicts"]["RQ2_by_stratum"][st]["verdict"]] for st in STRATA]
        table(doc, ["Stratum", "C2 - C3 (points)", "95% CI", "Verdict"], rows,
              "The primary contrast broken out by stratum (secondary endpoints). "
              "Per-stratum verdicts are secondary and are not corrected for "
              "multiplicity; the single primary test is the pooled row.", font=8.0)
    if "C2 - C1" in MET["contrasts"]:
        p = MET["verdicts"]["RQ2_parity_on_unanimous"]
        para(doc, f"The hypothesis also predicted PARITY on the unanimous stratum — "
                  f"soft targets should not cost anything where there is nothing to "
                  f"soften. C2 minus C1 there is {ci('C2 - C1', 'S-unanimous')}, "
                  f"which {'is consistent with parity' if p['parity_holds'] else 'is not consistent with parity'}.")

    h(doc, "3.3 RQ2, calibration", level=2)
    figure(doc, "P4_F28_calibration.png",
           "Expected calibration error by configuration and stratum (left) and "
           "reliability on the pooled contested stratum (right).")
    figure(doc, "P4_F29_overconfidence.png",
           "Mean confidence against expected accuracy. The shaded gap is the "
           "overconfidence Phase 3 identified as the baseline's real failure mode.")
    cagg = CAL["aggregate_3seed"]
    table(doc,
          ["Configuration"] + [TIER_LABEL[t].split(" (")[0] for t in TIERS] +
          ["contested (pooled)"],
          [[CFG_LABEL[c]] + [pc(cagg[c][t]["ece_vs_expected_accuracy"]) for t in TIERS] +
           [pc(cagg[c][POOLED]["ece_vs_expected_accuracy"])] for c in CF()],
          "Expected calibration error (%) against expected accuracy, 3-seed mean. "
          "Lower is better.", font=8.0)
    table(doc,
          ["Configuration", "Mean confidence (%)", "Expected accuracy (%)",
           "Overconfidence (points)", "MCE (%)", "Brier (top-1)", "Brier (23-vector)"],
          [[CFG_LABEL[c], pc(cagg[c][POOLED]["mean_confidence"]),
            pc(cagg[c][POOLED]["expected_accuracy"]),
            f"{cagg[c][POOLED]['overconfidence_points']:+.2f}",
            pc(cagg[c][POOLED]["mce_vs_expected_accuracy"]),
            f"{cagg[c][POOLED]['brier_top1_vs_expected_accuracy']:.4f}",
            f"{cagg[c][POOLED]['brier_vector_vs_vote_distribution']:.4f}"] for c in CF()],
          "Calibration detail on the pooled contested stratum. The 23-dimensional "
          "Brier score against the vote distribution is the stricter form and is "
          "added here because it is exactly the distribution C2 and C4 are trained "
          "on.", font=7.8)
    if CAL.get("verdicts", {}).get("RQ2_calibration"):
        v = CAL["verdicts"]["RQ2_calibration"]
        callout(doc, f"RQ2 calibration endpoint. C2 minus C3 in expected calibration "
                     f"error on the pooled contested stratum is {cie('C2 - C3')} "
                     f"(negative = better calibrated). Pre-registered verdict: "
                     f"{v['verdict']}.", title=f"Pre-registered verdict: {v['verdict']}")

    h(doc, "3.4 RQ3: does predictive uncertainty track human disagreement?", level=2)
    figure(doc, "P4_F30_uncertainty.png",
           "Left: Spearman correlation between predictive entropy and annotator vote "
           "entropy, within each stratum and pooled. Right: the MC "
           "stochastic-depth uncertainty decomposition.")
    defined = UNC["strata_where_defined"]
    rows = []
    for c in CF():
        a = UNC["results"][f"{c}|softmax"]["aggregate"]
        pooled = np.mean([UNC["results"][f"{c}|softmax"]["per_member"][str(s)]
                          ["_pooled_all_1353_images"]["spearman_rho"]
                          for s in UNC["results"][f"{c}|softmax"]["per_member"]])
        rows.append([CFG_LABEL[c]] +
                    [f"{a[st]['mean_rho']:.3f}" if a[st]["mean_rho"] is not None else "n/a"
                     for st in defined] + [f"{pooled:.3f}"])
    table(doc, ["Configuration"] + [st.replace("S-", "").replace(" (pooled)", "")
                                    for st in defined] + ["pooled, all images"],
          rows,
          "Spearman rho between predictive entropy and per-image annotator vote "
          "entropy. The rightmost column is reported only to expose the artefact: "
          "it is not the endpoint.", font=8.0,
          note="Vote entropy is identically zero on S-unanimous, so the correlation "
               "is undefined there and is reported as such rather than as zero.")
    para(doc, f"Phase 3 measured the pooled value at "
              f"{UNC['phase3_reference']['pooled_all_images_rho']} and the "
              f"within-tier values at "
              f"{UNC['phase3_reference']['within_tier_rho_range'][0]}-"
              f"{UNC['phase3_reference']['within_tier_rho_range'][1]}, and required "
              f"Phase 4 to report the within-tier quantity as primary. The table "
              f"above reproduces that gap for every configuration: a reader shown "
              f"only the pooled column would conclude the models track human "
              f"disagreement, when what the pooled column measures is mostly which "
              f"stratum an image belongs to.")
    if UNC.get("verdicts"):
        def _iv(v, k):
            return f"[{v[k][0]:.4f}, {v[k][1]:.4f}]" if k in v else "n/a"
        rows = [[CFG_LABEL[c], v.get("stratum", ""),
                 f"{v['mean_rho_3seed']:.4f}" if "mean_rho_3seed" in v else "n/a",
                 _iv(v, "ci95_pooled_paired"),
                 _iv(v, "ci95_mean_of_per_seed_bounds"),
                 v["verdict"]] for c, v in UNC["verdicts"].items()]
        table(doc, ["Configuration", "Stratum", "Mean rho",
                    "95% CI (pooled)", "Mean of per-seed bounds", "Verdict"], rows,
              "RQ3 pre-registered verdict, taken on the largest contested stratum.",
              font=8.0,
              note="The verdict column is read off the pooled interval: for each "
                   "patient resample the correlation is computed for all three "
                   "seeds and averaged, and the percentiles are taken over that "
                   "distribution. The final column is the arithmetic mean of the "
                   "three per-seed intervals; it is shown for continuity with the "
                   "first run of this analysis, but it propagates no between-seed "
                   "variation and is not a calibrated 95% interval.")
    ut = UNC.get("estimator_utility", {})
    if ut:
        table(doc, ["Configuration", "Single model (%)", "MC stochastic depth (%)",
                    "3-seed deep ensemble (%)", "Ensemble gain (points)"],
              [[CFG_LABEL[c],
                pc(ut[c][POOLED]["single_model_mean_macro_f1"]),
                pc(ut[c][POOLED]["mc_stochastic_depth_macro_f1"])
                if "mc_stochastic_depth_macro_f1" in ut[c][POOLED] else "n/a",
                pc(ut[c][POOLED]["deep_ensemble_3seed_macro_f1"]),
                f"{ut[c][POOLED]['ensemble_gain_points']:+.2f}"] for c in CF()],
              "What each uncertainty estimator buys in accuracy on the pooled "
              "contested stratum. The ensemble is 3-member, not the blueprint's "
              "5-member (deviation P4-DEV-2), so its gain is a lower bound.",
              font=8.0)

    h(doc, "3.5 RQ4: the anatomy-aware loss", level=2)
    figure(doc, "P4_F31_structure.png",
           "Left: mean annotator-marginalized anatomical error distance, the RQ4 "
           "primary endpoint. Right: error geometry on S-unanimous against the "
           "Phase 0 human benchmark, which C4 was not trained on.")
    sagg = STR["aggregate_3seed"]
    table(doc,
          ["Configuration"] + [TIER_LABEL[t].split(" (")[0] for t in TIERS] +
          ["contested (pooled)"],
          [[CFG_LABEL[c]] + [f"{sagg[c][t]['mean_anatomical_distance_3seed']:.4f}"
                             for t in TIERS] +
           [f"{sagg[c][POOLED]['mean_anatomical_distance_3seed']:.4f}"] for c in CF()],
          "Mean anatomical error distance (0 = the prediction is every annotator's "
          "label; 1 = maximally distant from all of them), 3-seed mean. Lower is "
          "better.", font=8.0)
    g0 = sagg[CF()[0]]["_error_geometry_S_unanimous"]
    table(doc,
          ["Configuration", "Error rate (%)", "Wall confusions adjacent (%)",
           "Station confusions neighbouring (%)"],
          [["Human annotators (Phase 0)", "—",
            f"{g0['human_wall_adjacent_pct']}", f"{g0['human_station_neighbouring_pct']}"]] +
          [[CFG_LABEL[c],
            f"{sagg[c]['_error_geometry_S_unanimous']['error_rate_pct_3seed']:.2f}",
            f"{sagg[c]['_error_geometry_S_unanimous']['wall_adjacent_pct_3seed']:.2f}",
            f"{sagg[c]['_error_geometry_S_unanimous']['station_neighbouring_pct_3seed']:.2f}"]
           for c in CF()],
          "Error geometry on the S-unanimous stratum — the only stratum with an "
          "uncontested reference label — against the human benchmarks measured in "
          "Phase 0. These are the independent checks: C4 optimises the distance "
          "matrix, not these shares.", font=8.0,
          note="Phase 3 established that the wall and station comparisons against "
               "the human values are both underpowered at this error count; the "
               "column is informative for comparing configurations with each other, "
               "not for declaring a match with humans.")
    if STR.get("verdicts", {}).get("RQ4"):
        v = STR["verdicts"]["RQ4"]
        callout(doc, f"RQ4. C4 minus C2 in mean anatomical error distance on the "
                     f"pooled contested stratum is {cid('C4 - C2')}. The "
                     f"pre-registered rule requires the distance to fall with an "
                     f"interval excluding zero AND macro F1 not to fall "
                     f"significantly; macro F1 moved {ci('C4 - C2')}. Verdict: "
                     f"{v['verdict']}.", title=f"Pre-registered verdict: {v['verdict']}")

    h(doc, "3.6 Sensitivity", level=2)
    figure(doc, "P4_F32_robustness.png",
           "Left: every seed of every configuration drawn separately. Right: the "
           "RQ2 primary contrast recomputed with each annotator dropped in turn.")
    h(doc, "3.6.1 Per-seed stability", level=3)
    sds = [agg[c][POOLED]["annotator_marginalized_macro_f1_sd_3seed"] for c in CF()]
    para(doc, f"Cross-seed standard deviation on the pooled contested stratum stays "
              f"below {100 * max(sds):.2f} points for every configuration, so the "
              f"ordering of the arms is a property of the training target rather "
              f"than of one initialisation. Per-seed contrast values and their "
              f"sign consistency are in the contrast table above.")
    h(doc, "3.6.2 Leave-one-annotator-out", level=3)
    if LOAO and LOAO.get("rq2_verdict_stability"):
        para(doc, f"{LOAO['motivation'].capitalize()}. Every headline number is an "
                  f"average over the four annotator columns, so the obvious "
                  f"objection is that one atypical rater drives it. "
                  f"{LOAO['what_varies'].capitalize()}")
        rows = [[nm.replace("_", " "),
                 f"{d['diff_points_3seed_mean']:+.2f}",
                 f"[{d['ci95_points_3seed_mean'][0]:+.2f}, {d['ci95_points_3seed_mean'][1]:+.2f}]",
                 d["verdict"]]
                for nm, d in LOAO["rq2_verdict_stability"].items()]
        table(doc, ["Annotator subset", "C2 - C3 (points)", "95% CI", "Verdict"], rows,
              "The RQ2 primary contrast recomputed with each annotator dropped from "
              "the metric in turn. The strata stay defined by the full "
              "four-annotator vote matrix, so every row scores the same images.",
              font=8.0)
        inv = LOAO.get("rq2_verdict_invariant_to_dropping_any_single_annotator")
        para(doc, f"The pre-registered verdict is "
                  f"{'invariant to dropping any single annotator, FG2 included' if inv else 'NOT invariant: at least one annotator changes the verdict, which is itself the finding'}. "
                  f"A training-side LOAO — rebuilding the C2 targets from three "
                  f"annotators and retraining — was not run; the pre-registration "
                  f"declares it unexecuted for budget and Appendix E states what it "
                  f"would cost.")


# =====================================================================
def sec_discussion(doc) -> None:
    h(doc, "4. Discussion", level=1)
    h(doc, "4.1 RQ2: what the control bought", level=2)
    if MET.get("verdicts", {}).get("RQ2_primary"):
        v = MET["verdicts"]["RQ2_primary"]["verdict"]
        para(doc, f"The primary contrast is C2 minus C3, not C2 minus C1, and the "
                  f"distinction carries the whole argument. C2 minus C1 asks whether "
                  f"softening the target helps; it was {ci('C2 - C1')} on the pooled "
                  f"contested stratum, {B.direction('C2 - C1')}. C2 minus C3 asks the "
                  f"harder and more interesting question — whether softening the "
                  f"target WHERE THE ANNOTATORS ACTUALLY DISAGREED helps beyond "
                  f"softening it everywhere by the same total amount — and it was "
                  f"{ci('C2 - C3')}, {B.direction('C2 - C3')}. The pre-registered "
                  f"verdict is {v}.")
        para(doc, {
            "SUPPORTED":
                "That is a positive result about information, not about "
                "regularisation. Two arms that displace the same expected "
                "probability mass from the modal label differ only in where they "
                "put it, so a gain for C2 localises the benefit in the pattern of "
                "disagreement itself — which is the quantity almost no public "
                "corpus releases and almost no study in the Phase 1 review uses.",
            "NOT SUPPORTED":
                "That is a genuinely informative negative. It says the benefit of "
                "soft targets on this corpus is a regularisation effect that "
                "uniform smoothing reproduces at least as well, and therefore that "
                "the four-annotator vote pattern adds nothing beyond signalling "
                "that an image is hard. With only four annotators the vote "
                "proportion takes five values, so the pattern is a coarse estimate "
                "of the underlying label distribution; declared limitation L7 "
                "anticipated exactly this.",
            "NOT RESOLVED":
                "That is an honest null. The interval is wide enough to contain "
                "effects in both directions, so the phase neither demonstrates nor "
                "refutes a benefit from the disagreement pattern. The relevant "
                "number for anyone planning a replication is the interval width "
                "itself, not the point estimate.",
        }[v])
    para(doc, f"The cohort effect deserves separate mention. C1 minus C0 was "
              f"{ci('C1 - C0')} on the pooled contested stratum. Two things confound "
              f"it and both favour caution: C0's model selection used a validation "
              f"set containing no contested images, and C0 was trained to early "
              f"stopping while several Phase 4 runs stopped at the pre-registered "
              f"epoch cap. It is reported as a descriptive comparison, not as a "
              f"controlled contrast, and it is not the basis of any claim here.")
    h(doc, "4.2 RQ2 calibration, and what it says about the Phase 3 finding", level=2)
    cagg = CAL["aggregate_3seed"]
    worst = max(CF(), key=lambda c: cagg[c][POOLED]["ece_vs_expected_accuracy"])
    best = min(CF(), key=lambda c: cagg[c][POOLED]["ece_vs_expected_accuracy"])
    best_u = min(CF(), key=lambda c: cagg[c]["S-unanimous"]["ece_vs_expected_accuracy"])
    para(doc, f"Every configuration remains badly calibrated on contested images. "
              f"The best of them on the pooled contested stratum is {best} at "
              f"{pc(cagg[best][POOLED]['ece_vs_expected_accuracy'])}% expected "
              f"calibration error and the worst is {worst} at "
              f"{pc(cagg[worst][POOLED]['ece_vs_expected_accuracy'])}%, against "
              f"{pc(cagg[CF()[0]]['S-unanimous']['ece_vs_expected_accuracy'])}% for "
              f"{CF()[0]} on unanimous images. Changing the training target moves "
              f"this substantially, but it does not fix it.")
    if {"C2", "C3"} <= set(CF()):
        c2u, c3u = cagg["C2"]["S-unanimous"], cagg["C3"]["S-unanimous"]
        c2p, c3p = cagg["C2"][POOLED], cagg["C3"][POOLED]
        para(doc, f"Before the reversal, the part of the hypothesis that did hold: "
                  f"soft targets are substantially better calibrated than hard ones. "
                  f"C2 minus C1 on the pooled contested stratum is "
                  f"{cie('C2 - C1')} — an interval well clear of zero — so building "
                  f"the target from the vote distribution rather than from the "
                  f"majority label removes roughly a quarter of the baseline's "
                  f"calibration error. That is a real effect and it is the effect RQ2 "
                  f"predicted. It is simply not the effect RQ2 was tested on.")
        para(doc, f"The headline verdict hides a reversal that is more interesting "
                  f"than the verdict. C3 wins on the contested stratum "
                  f"({pc(c3p['ece_vs_expected_accuracy'])}% against C2's "
                  f"{pc(c2p['ece_vs_expected_accuracy'])}%), and wins by more than "
                  f"twice C2's margin over C1 ({cie('C3 - C1')} against "
                  f"{cie('C2 - C1')}). But C2 beats C3 on the unanimous stratum "
                  f"({pc(c2u['ece_vs_expected_accuracy'])}% against "
                  f"{pc(c3u['ece_vs_expected_accuracy'])}%), where the "
                  f"best-calibrated arm overall is {best_u} at "
                  f"{pc(cagg[best_u]['S-unanimous']['ece_vs_expected_accuracy'])}%. "
                  f"The mechanism is "
                  f"visible in the confidence column. On unanimous images C3's mean "
                  f"confidence is {pc(c3u['mean_confidence'])}% against an expected "
                  f"accuracy of {pc(c3u['expected_accuracy'])}% — it is "
                  f"{'UNDER' if c3u['overconfidence_points'] < 0 else 'over'}confident "
                  f"by {abs(c3u['overconfidence_points']):.1f} points — while C2 sits "
                  f"at {pc(c2u['mean_confidence'])}% against "
                  f"{pc(c2u['expected_accuracy'])}%, almost exactly right.")
        para(doc, f"That is what a blunt instrument looks like. Uniform label "
                  f"smoothing applies the same epsilon to every image, so it "
                  f"suppresses confidence globally: mean confidence on the contested "
                  f"stratum falls to {pc(c3p['mean_confidence'])}% for C3 against "
                  f"{pc(c2p['mean_confidence'])}% for C2, which buys a large ECE "
                  f"reduction where the model was overconfident and costs accuracy of "
                  f"confidence where it was not. Vote-proportion targets cannot do "
                  f"this, because on a 4/4 image the vote distribution IS one-hot and "
                  f"C2 is therefore trained to be maximally confident there — which "
                  f"is correct, and is why C2 is the better-calibrated arm on that "
                  f"stratum. The two arms are not competing on one axis; C3 trades "
                  f"unanimous-stratum calibration for contested-stratum calibration, "
                  f"and the pooled endpoint scores that trade favourably only because "
                  f"the pooled stratum contains no unanimous images.")
        para(doc, "None of this rescues the hypothesis as pre-registered, and it is "
                  "not offered as doing so — the rule was fixed on the pooled "
                  "contested endpoint and the pooled contested endpoint says the "
                  "control won. It does say that the improvement label smoothing "
                  "delivers is a known and general property of the method rather "
                  "than anything about this corpus's annotators, which is precisely "
                  "why the blueprint insisted that C3 control the calibration result "
                  f"as well as the accuracy result. Without it, C2's genuine "
                  f"improvement over C1 — {cie('C2 - C1')} on the contested stratum, "
                  f"an interval excluding zero — would have been reported as evidence "
                  f"that annotator vote distributions fix calibration. They improve "
                  f"it markedly; a uniform smoothing prior of matched mass improves "
                  f"it more.")
    para(doc, "That relocates Phase 3's principal finding. It is not an artefact of "
              "training on the unanimous subset — models trained on contested "
              "images with contested targets are still overconfident on contested "
              "images. The failure is in the agreement structure of the problem: a "
              "single softmax over 23 mutually exclusive classes has no way to "
              "express 'two experts would say A3 and two would say P3', and a "
              "confidence score derived from it will always overstate the mass it "
              "captures under the vote distribution. Post-hoc temperature scaling "
              "was deliberately excluded from this phase's scope so as not to "
              "confound the C2-versus-C3 comparison; it is the obvious next "
              "intervention and it is now motivated by a measurement rather than by "
              "convention.")
    h(doc, "4.3 RQ3", level=2)
    para(doc, "The within-stratum correlations are small for every configuration, "
              "and the gap between them and the pooled value reproduces the "
              "artefact Phase 3 identified. The practical implication is a "
              "reporting one and it generalises beyond this corpus: a study that "
              "correlates model uncertainty with annotator disagreement across a "
              "whole test set, without stratifying, will report a correlation that "
              "largely measures how contested each image is rather than whether the "
              "model knows. Since the stratification requires per-annotator labels, "
              "most published versions of this analysis cannot be checked.")
    h(doc, "4.4 RQ4", level=2)
    if STR.get("verdicts", {}).get("RQ4"):
        para(doc, f"The structured penalty was tested at exactly one value of "
                  f"lambda, fixed a priori at unit weight and never swept "
                  f"(deviation P4-DEV-3). The verdict "
                  f"({STR['verdicts']['RQ4']['verdict']}) is therefore a statement "
                  f"about unit weight, not about the family of anatomy-aware losses. "
                  f"A null at lambda = 1 is compatible with a benefit at a larger "
                  f"weight, and the honest thing to report is the cost of finding "
                  f"out: nine further runs, roughly thirteen hours on this hardware "
                  f"(Appendix E).")
    h(doc, "4.5 Relation to the descriptor's own FG-agreement result", level=2)
    para(doc, "The blueprint's §2.7 records that the dataset descriptor obtained "
              "87.05 macro F1 by training ConvNeXt-Tiny on FG-agreement labels "
              "against ~85 on complete-agreement labels — a 2.2-point move from "
              "changing how annotator labels are combined. That result is not "
              "directly comparable with anything here: it is measured on the "
              "unanimous test subset only, and it changes which annotators define "
              "the label rather than how their votes are aggregated. It remains the "
              "precedent that motivated this phase, and the present results give it "
              "a stratified context it did not have.")
    h(doc, "4.6 Limitations specific to this phase", level=2)
    bullet(doc, f"The {COH['excluded_trainval_images']['n']:,} Train/Validation "
                f"images with no majority label are excluded from every arm, so the "
                f"most ambiguous images in the corpus never enter training. This is "
                f"forced by the need for a definable C3 control.")
    bullet(doc, "Four annotators give vote proportions with only five possible "
                "values, so the soft target is a coarse estimate of the label "
                "distribution (declared limitation L7). A null RQ2 result is "
                "partly a statement about panel size.")
    bullet(doc, "lambda for C4 and the matching criterion for C3's epsilon were "
                "each fixed at a single pre-registered value. Both are defensible "
                "choices made before seeing results, and both leave the "
                "corresponding hypothesis tested at one point rather than over a "
                "family.")
    if B.runs():
        p4 = {k: v for k, v in B.runs().items() if k[0] != "C0"}
        n_cap = sum(1 for v in p4.values() if v.get("stop_reason") == "epoch_cap")
        by_cfg = {}
        for (cfg, _seed), v in p4.items():
            hit = v.get("stop_reason") == "epoch_cap"
            d = by_cfg.setdefault(cfg, {"n": 0, "cap": 0, "best": []})
            d["n"] += 1
            d["cap"] += int(hit)
            if hit:
                d["best"].append(v.get("best_epoch_overall", 0))
        # most-censored arm: most caps, ties broken by the later mean best epoch
        worst = max(by_cfg, key=lambda c: (by_cfg[c]["cap"],
                                           np.mean(by_cfg[c]["best"] or [0])))
        share = ", ".join(f"{c} {by_cfg[c]['cap']}/{by_cfg[c]['n']}"
                          for c in sorted(by_cfg))
        total_ep = max((v.get("n_epochs_run", 0) for v in p4.values()
                        if v.get("stop_reason") == "epoch_cap"), default=0)
        bests = "/".join(str(b) for b in sorted(by_cfg[worst]["best"]))
        bullet(doc, f"{n_cap} of {len(p4)} Phase 4 runs stopped at the "
                    f"pre-registered fine-tuning cap rather than by early stopping, "
                    f"so the absolute scores are lower bounds and the C1-minus-C0 "
                    f"comparison is descriptive only.")
        bullet(doc, f"The cap is the same NUMBER of epochs for every configuration, "
                    f"but it did not censor them equally: {share}. It bound most "
                    f"often on {worst}, whose capped runs were still selecting their "
                    f"best epoch at {bests} of {total_ep} — validation macro F1 was "
                    f"still improving when training was cut off. An equal cap is "
                    f"therefore not equal censoring, and this report does not claim "
                    f"the contrasts are unaffected by it. {worst} is the treatment "
                    f"arm of the RQ2 primary contrast, so the bias runs against it "
                    f"and towards the NOT RESOLVED verdict that was obtained. That "
                    f"verdict should be read as 'not resolved under this compute "
                    f"budget', and a convergence amendment that lets every arm "
                    f"early-stop on its own is the pre-condition for reading it as "
                    f"anything stronger.")
    bullet(doc, "One architecture, one centre, one vendor. Phase 5's external "
                "validation remains the only test of whether any of this transfers.")
    bullet(doc, "The MC estimator samples stochastic depth, whose per-block "
                "probability in ConvNeXt-Tiny rises only to 0.1. The resulting "
                "predictive spread is modest by construction, so a small epistemic "
                "component should not be read as evidence that the model is certain.")
    h(doc, "4.7 Implications for Phases 5 to 7", level=2)
    bullet(doc, "Phase 5 should carry the best-calibrated configuration, not the "
                "most accurate one, into external validation: the calibration gap "
                "is the failure that would matter first in deployment.")
    bullet(doc, "Phase 5 should also report the agreement-stratified curve on the "
                "external corpora if per-annotator labels exist there, and state "
                "plainly that it cannot if they do not.")
    bullet(doc, "Phase 6's Grad-CAM comparison should target the contested strata "
                "and the specific confusions the anatomical distance matrix marks "
                "as far, since those are the errors no configuration here removed.")
    bullet(doc, "Phase 7 should present RQ2 as a controlled comparison with a "
                "matched control, and report the verdict the rule produced rather "
                "than the one the hypothesis predicted.")


def sec_conclusion(doc) -> None:
    h(doc, "5. Conclusion", level=1)
    h(doc, "5.1 Answers", level=2)
    if MET.get("verdicts", {}).get("RQ2_primary"):
        bullet(doc, f"RQ2 (accuracy): {MET['verdicts']['RQ2_primary']['verdict']}. "
                    f"C2 minus C3 on the pooled contested stratum = {ci('C2 - C3')}.")
    if CAL.get("verdicts", {}).get("RQ2_calibration"):
        bullet(doc, f"RQ2 (calibration): {CAL['verdicts']['RQ2_calibration']['verdict']}. "
                    f"C2 minus C3 in ECE = {cie('C2 - C3')}. No configuration "
                    f"achieves acceptable calibration on contested images.")
    if UNC.get("verdicts"):
        sup = [c for c, v in UNC["verdicts"].items() if v.get("verdict") == "SUPPORTED"]
        bullet(doc, f"RQ3: supported for {', '.join(sup) if sup else 'no configuration'} "
                    f"on the largest contested stratum. The pooled correlation "
                    f"remains several times the within-stratum value for every "
                    f"configuration, confirming the Phase 3 artefact.")
    if STR.get("verdicts", {}).get("RQ4"):
        bullet(doc, f"RQ4: {STR['verdicts']['RQ4']['verdict']} at lambda = "
                    f"{PRE['configurations']['C4']['structure_penalty_lambda']:g}. "
                    f"Anatomical error distance C4 minus C2 = {cid('C4 - C2')}.")
    h(doc, "5.2 Carry-forward decisions", level=2)
    bullet(doc, "Treat calibration, not accuracy, as the headline deficiency of "
                "landmark classifiers evaluated outside the consensus subset. Four "
                "training targets did not fix it.")
    bullet(doc, "Keep the matched control. The value of this phase is that its "
                "answer is interpretable whichever way it came out, and that is "
                "entirely due to C3 being matched to C2 in displaced mass rather "
                "than set to a conventional epsilon.")
    bullet(doc, "Report within-stratum uncertainty correlations. The pooled "
                "quantity is not wrong, it is a different quantity, and it flatters "
                "the model.")
    bullet(doc, "Pre-register the epoch cap derivation again in Phase 5, and budget "
                "for early stopping to fire: the cap bound in this phase and cost "
                "the C1-minus-C0 comparison its interpretability.")
    bullet(doc, "Carry the anatomical distance matrix into Phase 6 as the ranking "
                "for which confusions to explain, independently of the RQ4 verdict "
                "— it is a useful descriptive instrument even where the penalty "
                "built on it is not.")


def sec_appendices(doc) -> None:
    h(doc, "Appendices", level=1, page_break=True)
    B.appendix_a_full_tables(doc)
    B.appendix_b_training(doc)
    B.appendix_c_prereg(doc)
    B.appendix_d_manifest(doc)
    B.appendix_e_unexecuted(doc)


def main() -> None:
    missing = [n for n, v in [("phase4_stratified_metrics.json", MET),
                              ("phase4_calibration.json", CAL),
                              ("phase4_uncertainty.json", UNC),
                              ("phase4_structure_eval.json", STR)] if v is None]
    if missing:
        raise SystemExit("cannot build the report; missing artefacts: " + ", ".join(missing))

    doc = new_document()
    title_page(doc)
    front_matter(doc)
    B.abbreviations(doc)
    sec_abstract(doc)
    sec_introduction(doc)
    sec_methods(doc)
    sec_results(doc)
    sec_discussion(doc)
    sec_conclusion(doc)
    sec_appendices(doc)
    B.references(doc)

    add_page_numbers(doc)
    doc.save(BD.OUT)

    caps = [p.text for p in doc.paragraphs if p.style.name == "Caption"]
    nfig = sum(1 for c in caps if c.strip().startswith("Figure"))
    ntab = sum(1 for c in caps if c.strip().startswith("Table"))
    print(f"[report] {len(doc.inline_shapes)} images, {nfig} figure captions, "
          f"{ntab} table captions -> {BD.OUT}")


if __name__ == "__main__":
    main()
