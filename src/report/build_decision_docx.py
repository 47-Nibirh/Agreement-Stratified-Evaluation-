"""
Render DATASET_DECISION_REPORT.md to Word.

A small GitHub-flavoured-Markdown subset renderer: ATX headings, pipe tables,
bullet/ordered lists, blockquotes, horizontal rules, and inline **bold**,
*italic*, `code` and [text](url) links.  Deliberately narrow -- it handles the
constructs the decision report actually uses and nothing else, so that the
output stays predictable.

Run:  python src/report/build_decision_docx.py
Then: python src/report/finalise_decision.py   (Word COM -> PDF)
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
SRC_MD = ROOT / "DATASET_DECISION_REPORT.md"
OUT_DOCX = ROOT / "Dataset_Decision_Report.docx"

ACCENT = RGBColor(0x1F, 0x38, 0x64)
MUTED = RGBColor(0x44, 0x44, 0x44)

INLINE = re.compile(
    r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))"
)
LINK = re.compile(r"\[([^\]]+?)\]\(([^)]+?)\)")


def shade(cell, hex_fill: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def add_runs(par, text: str, bold=False, italic=False, size=None) -> None:
    """Write `text` into `par`, honouring inline markdown."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        b, i, mono = bold, italic, False
        body = tok
        if tok.startswith("**") and tok.endswith("**"):
            body, b = tok[2:-2], True
        elif tok.startswith("*") and tok.endswith("*"):
            body, i = tok[1:-1], True
        elif tok.startswith("`") and tok.endswith("`"):
            body, mono = tok[1:-1], True
        else:
            m = LINK.fullmatch(tok)
            if m:
                body = m.group(1)
                if m.group(2) not in body:
                    body = f"{m.group(1)} <{m.group(2)}>"
        run = par.add_run(body)
        run.bold = b
        run.italic = i
        if mono:
            run.font.name = "Consolas"
            run.font.size = Pt((size or 10) - 1)
        elif size:
            run.font.size = Pt(size)


def split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_sep(line: str) -> bool:
    return bool(re.fullmatch(r"\|?[\s:\-|]+\|?", line.strip())) and "-" in line


def add_table(doc, rows: list[list[str]]) -> None:
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci in range(ncols):
            text = row[ci] if ci < len(row) else ""
            cell = cells[ci]
            cell.text = ""
            par = cell.paragraphs[0]
            par.paragraph_format.space_before = Pt(2)
            par.paragraph_format.space_after = Pt(2)
            add_runs(par, text, bold=(ri == 0), size=8.5)
            if ri == 0:
                shade(cell, "1F3864")
                for run in par.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif ri % 2 == 0:
                shade(cell, "F2F4F8")
    doc.add_paragraph()


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, colour in (
        ("Heading 1", 18, ACCENT),
        ("Heading 2", 14, ACCENT),
        ("Heading 3", 12, ACCENT),
        ("Heading 4", 11, MUTED),
    ):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = colour
        st.font.bold = True


def add_footer_pagenum(doc: Document) -> None:
    par = doc.sections[0].footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    for instr in ("begin", "PAGE", "end"):
        el = OxmlElement("w:fldChar" if instr != "PAGE" else "w:instrText")
        if instr == "PAGE":
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        else:
            el.set(qn("w:fldCharType"), instr)
        run._r.append(el)


def add_toc(doc: Document) -> None:
    par = doc.add_paragraph()
    run = par.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose Update Field to build the contents."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, sep, placeholder, end):
        run._r.append(el)


def build() -> int:
    if not SRC_MD.exists():
        print(f"missing {SRC_MD}")
        return 1

    lines = SRC_MD.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style_doc(doc)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.8)
        s.top_margin = s.bottom_margin = Inches(0.8)
    add_footer_pagenum(doc)

    # --- title page -------------------------------------------------------
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Dataset Viability Assessment\nand Replacement Recommendation")
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("Phase 1.5 — Decision Report")
    rs.font.size = Pt(14)
    rs.font.color.rgb = MUTED
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = d.add_run(
        "Fourth-year B.Sc. thesis · AI for upper gastrointestinal endoscopy\n26 July 2026"
    )
    rd.font.size = Pt(11)
    rd.font.color.rgb = MUTED

    doc.add_page_break()
    h = doc.add_paragraph("Contents", style="Heading 1")
    h.paragraph_format.space_after = Pt(10)
    add_toc(doc)
    doc.add_page_break()

    i, n = 0, len(lines)
    skipped_front_matter = False
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # pipe table
        if stripped.startswith("|") and i + 1 < n and is_sep(lines[i + 1]):
            rows = [split_row(stripped)]
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i].strip()))
                i += 1
            add_table(doc, rows)
            continue

        if not stripped:
            i += 1
            continue

        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", stripped):
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level, text = len(m.group(1)), m.group(2)
            # the markdown title block is replaced by the title page above
            if not skipped_front_matter:
                if text.startswith(("Dataset Viability", "Phase 1.5")):
                    i += 1
                    continue
                skipped_front_matter = True
            if level == 1:
                doc.add_page_break()
            par = doc.add_paragraph(style=f"Heading {min(level, 4)}")
            add_runs(par, text)
            i += 1
            continue

        if stripped.startswith(">"):
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Inches(0.35)
            par.paragraph_format.space_before = Pt(6)
            par.paragraph_format.space_after = Pt(6)
            add_runs(par, stripped.lstrip("> ").strip(), italic=True)
            for run in par.runs:
                run.font.color.rgb = ACCENT
            i += 1
            continue

        m = re.match(r"^[-*]\s+(.*)$", stripped)
        if m:
            par = doc.add_paragraph(style="List Bullet")
            add_runs(par, m.group(1))
            i += 1
            continue

        m = re.match(r"^\d+\.\s+(.*)$", stripped)
        if m:
            par = doc.add_paragraph(style="List Number")
            add_runs(par, m.group(1))
            i += 1
            continue

        par = doc.add_paragraph()
        add_runs(par, stripped)
        i += 1

    doc.save(OUT_DOCX)
    print(f"[build] wrote {OUT_DOCX}")
    return 0


if __name__ == "__main__":
    sys.exit(build())
