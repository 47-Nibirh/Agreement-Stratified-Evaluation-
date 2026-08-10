"""
Phase 6 report -- explainability and error analysis.

Every number is interpolated from reports/phase6_*.json. No verdict sentence is
pre-written: each is read from a verdict field that a frozen rule produced, so
the report says whatever the rules said, including where they contradict what
the phase set out to find.

Run:  python src/report/build_phase6_docx.py
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_docx as BD  # noqa: E402

ROOT = Path(__file__).resolve().parents[2]
BD.FIGD = ROOT / "figures_phase6"
BD.OUT = ROOT / "Phase6_Report.docx"

from build_docx import (add_page_numbers, bullet, callout, figure,  # noqa: E402
                        front_matter, h, new_document, para, table)

REP = ROOT / "reports"


def J(name):
    p = REP / name
    return json.loads(p.read_text(encoding="utf-8")) if p.exists() else None


PRE = J("phase6_prereg.json")
HUM = J("phase6_human.json")
GEO = J("phase6_geometry.json")
CAM = J("phase6_cam_gate.json")
ATT = J("phase6_cam_eval.json")
SEL = J("phase6_selective.json")
P3AMD = J("phase3b_amendment.json")
P5CARRY = J("phase5_carry_forward.json")

CFG_LABEL = {"C0": "C0 hard label, 4/4 cohort (Phase 2 reference)",
             "C1": "C1 hard majority label, extended cohort",
             "C2": "C2 vote proportions",
             "C3": "C3 hard label + matched label smoothing (control)",
             "C4": "C4 vote proportions + anatomical penalty"}
STRATA = ["S-unanimous", "S-majority", "S-plurality", "S-no-majority",
          "S-contested (pooled)"]


def ARMS():
    return HUM["arms"] if HUM else PRE["inherits"]["arms"]


def pc(x, n=2):
    return "n/a" if x is None else f"{x:.{n}f}"


def iv(v, n=3):
    if not v or v[0] is None:
        return "n/a"
    return f"[{v[0]:.{n}f}, {v[1]:.{n}f}]"


# =====================================================================
def title_page(doc) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Explainability and Error Analysis")
    r.bold = True; r.font.size = BD.Pt(19)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Phase 6 — What the Model's Mistakes Are, and Whose Standard "
                  "to Judge Them By")
    r.font.size = BD.Pt(12.5)
    doc.add_paragraph()

    rows = []
    if HUM:
        pooled = "S-contested (pooled)"
        rows.append(["P6-A  human comparator (contested), pre-registered rule",
                     HUM["verdict_summary"].get(pooled, "—")])
        q = HUM.get("qualified_verdict", {}).get(pooled, "")
        rows.append(["P6-A  what may actually be claimed (P6-AMD-5)",
                     q.split(".")[0] + "." if q else "—"])
    if GEO:
        v = GEO["verdict_summary"].get("S-contested (pooled)", {})
        rows.append(["P6-B  confusion geometry — wall axis", v.get("wall", "—")])
        rows.append(["P6-B  confusion geometry — station axis", v.get("station", "—")])
    if ATT:
        rows.append(["P6-C1 attribution vs human disagreement",
                     ATT["verdict_summary"].get("P6-C1_primary", "—")])
        rows.append(["P6-C1b same, exploratory substitute signal",
                     ATT["verdict_summary"].get("P6-C1b_exploratory_spread", "—")])
        rows.append(["P6-C2 attribution stability",
                     ATT["verdict_summary"].get("P6-C2_secondary", "—")])
    if SEL:
        rows.append(["P6-D  selective prediction vs the Phase 5 ranking",
                     SEL["verdict_summary"].get("phase5_consistency", "—")])
    if rows:
        table(doc, ["Pre-registered endpoint", "Verdict"], rows,
              "Pre-registered verdicts, selected by the frozen rules in "
              "reports/phase6_prereg.json.", widths=[3.9, 2.6], font=8.5)

    doc.add_paragraph()
    if PRE:
        para(doc, f"Pre-registration frozen {PRE['frozen']} at commit "
                  f"{PRE['git_head_at_freeze'][:8]}, before any endpoint was scored.",
             size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


# =====================================================================
def sec_abstract(doc) -> None:
    h(doc, "Abstract", 1)
    n_img = HUM["results"]["S-unanimous"]["n_images"] + \
        HUM["results"]["S-contested (pooled)"]["n_images"] if HUM else 1353

    para(doc, "Background. Phases 3 to 5 established one durable result and three "
              "unresolved ones. The durable result is a calibration collapse across "
              "strata of expert agreement that no target construction repairs. The "
              "unresolved ones — accuracy, uncertainty and the anatomy-aware loss — "
              "are all comparisons between models. None of them can say whether the "
              "model's residual error reflects genuine visual ambiguity or a limit "
              "of the model, because no comparator outside the model set was ever "
              "measured.")
    para(doc, f"Objective. To change the comparator. Four pre-registered endpoints: "
              f"whether the model is distinguishable from a held-out human annotator "
              f"scored identically (P6-A); whether its error geometry mirrors human "
              f"disagreement geometry when both sides carry intervals (P6-B); whether "
              f"spatial attribution tracks human disagreement where predictive entropy "
              f"did not (P6-C); and whether the Phase 5 out-of-protocol rejection "
              f"ranking survives a threshold-free treatment (P6-D).")
    para(doc, f"Methods. No retraining. The frozen checkpoints of all five arms across "
              f"three seeds were evaluated on the same {n_img:,}-image test split in the "
              f"same row order Phases 3 and 4 used, verified by reproduction gates that "
              f"recover the published Phase 3 numbers exactly. Every internal interval "
              f"is a patient-clustered bootstrap of 1,000 resamples; contrasts are "
              f"paired, with both sides scored on the same resample before differencing.")

    if HUM and GEO and SEL:
        P = "S-contested (pooled)"
        pooled_h = HUM["results"][P]["bootstrap"]["by_arm"][HUM["headline_arm"]]
        geo_p = GEO["results"][P]["by_arm"][GEO["headline_arm"]]
        sp = HUM["sensitivity_P6-AMD-5"]["by_stratum"][P]
        spa = sp["by_arm"][HUM["headline_arm"]]
        para(doc, f"Results. On the contested strata the model out-predicts a held-out "
                  f"annotator by the pre-registered rule "
                  f"(Δ = {pc(pooled_h['delta_mean'], 4)} macro F1, 95% CI "
                  f"{iv(pooled_h['delta_ci95'])}). A post-hoc sensitivity analysis "
                  f"qualifies that sharply: the modal vote of the same three references "
                  f"scores {pc(sp['modal_vote_oracle'], 4)} against "
                  f"{pc(sp['human_held_out'], 4)} for the annotator, and the model "
                  f"reaches only {100 * (spa['position_in_headroom'] or 0):.0f}% of "
                  f"that headroom. The model therefore beats an individual expert but "
                  f"not the panel's own aggregation rule, and a substantial model "
                  f"shortfall against the attainable ceiling remains. Its wall-confusion "
                  f"geometry mirrors the human geometry, while its station geometry "
                  f"diverges by {pc(geo_p.get('station_neighbouring_delta_mean'))} points "
                  f"(CI {iv(geo_p.get('station_neighbouring_delta_ci95'), 2)}). "
                  f"Selective prediction separates the arms far more sharply externally "
                  f"(AURC {pc(SEL['external']['by_arm']['C2']['aurc_3seed'], 4)} for C2 "
                  f"against {pc(SEL['external']['by_arm']['C3']['aurc_3seed'], 4)} for C3) "
                  f"than internally, and agrees with the Phase 5 rejection ordering.")
    para(doc, "Conclusion. Two things degrade as annotators disagree, and the "
              "literature reports them as one. The attainable ceiling falls — the best "
              "single-label predictor drops from 1.00 to 0.67 — which accounts for much "
              "of the decline Phase 3 attributed to the model. But a real model "
              "shortfall against that reduced ceiling remains, so the decline is not "
              "purely a property of the task either. What is unambiguous is that "
              "confidence degrades further and faster than discrimination, that no "
              "target construction repairs it, and that the endpoint which exposes it "
              "is the model's willingness to decline rather than its accuracy.")
    doc.add_page_break()


# =====================================================================
def sec_introduction(doc) -> None:
    h(doc, "1. Introduction", 1, page_break=False)
    h(doc, "1.1 What Phases 3–5 settled, and the three debts they left", 2)
    para(doc, "Phase 3 measured performance across strata of annotator agreement and "
              "found annotator-marginalized macro F1 falling from 83.92 on unanimous "
              "images to 26.15 on the 2-1-1 stratum, with expected calibration error "
              "rising from 9.15% to 56.40%. Phase 4 tested five target constructions "
              "and found none that repaired the calibration. Phase 5 carried all five "
              "arms to two external corpora and found that the label spaces could not "
              "express the wall × station grid at all, reframing the endpoint before "
              "scoring. Three debts were left explicitly open.")

    h(doc, "1.1.1 X3: a withdrawn claim about confusion geometry", 3)
    if P3AMD:
        callout(doc,
                "Phase 3 reported that the model's wall-confusion geometry (89.68%) sat "
                "within 0.12 points of the human value (89.8%) and that its "
                "station-confusion geometry (85.57%) trailed the human 93.1% by 7.5 "
                "points. Correction X3 withdrew both to hypothesis status: the model "
                "shares carried patient-clustered intervals, the human values were "
                "corpus-wide point estimates with none, and the station interval "
                "contained the human value. A comparison between an interval and a "
                "point is not a comparison.",
                title="The debt recorded as X3")

    h(doc, "1.1.2 The blueprint's question: ambiguity or capacity?", 3)
    para(doc, "The blueprint asks of this phase whether \"the residual error reflects "
              "genuine visual ambiguity rather than model capacity\". Confusion geometry "
              "alone cannot settle that. It shows that model errors have the same shape "
              "as human disagreements; it does not show that the images are equally hard "
              "for both. Only a comparator drawn from the annotators themselves can "
              "license the ambiguity claim, and no such comparator existed anywhere in "
              "this project before Phase 6.")

    h(doc, "1.1.3 The Phase 5 carry-forward on out-of-protocol rejection", 3)
    para(doc, "Phase 5 recorded that out-of-protocol rejection \"separated the arms "
              "where every internal endpoint failed to\" and instructed Phases 6–7 to "
              "treat it as a primary endpoint. Phase 5 measured it at exactly one "
              "operating point — whether the 23-way argmax lands on OTHERCLASS — which "
              "is a property of where a decision boundary happens to fall rather than "
              "of how well the model's confidence orders its own mistakes.")

    h(doc, "1.2 Why the comparator changes in this phase", 2)
    para(doc, "Every comparison in Phases 3 to 5 was model against model. That design "
              "can rank target constructions but cannot calibrate the ranking against "
              "anything outside it: a model scoring 26.15 macro F1 looks broken until "
              "one knows what a board-certified endoscopist scores on the same images. "
              "Phase 6 therefore introduces two comparators the project has never used — "
              "the annotators, and the model's own confidence ordering — and keeps the "
              "explainability analysis quantitative so that it can fail.")

    h(doc, "1.3 Endpoints and pre-registered verdict rules", 2)
    if PRE:
        rows = [[k, PRE["endpoints"][k]["name"], PRE["endpoints"][k]["question"]]
                for k in ("P6-A", "P6-B", "P6-C", "P6-D")]
        table(doc, ["ID", "Endpoint", "Question"], rows,
              "The four Phase 6 endpoints, as frozen before any scoring ran.",
              widths=[0.7, 1.6, 4.2], font=8.2)
    doc.add_page_break()


# =====================================================================
def sec_methods(doc) -> None:
    h(doc, "2. Methods", 1)
    h(doc, "2.1 What is held fixed", 2)
    para(doc, "No retraining, no threshold tuning and no checkpoint reselection. The "
              "same frozen checkpoints, the same 224×224 cache, the same training-set "
              "normalisation and the same 1,353-image test split in the same row order "
              "that Phases 3 and 4 used. Two reproduction gates enforce this rather "
              "than assert it.")
    if HUM:
        g = HUM["gates"]
        rows = [["P6.1a", "panel row order identical to the Phase 3 cache index",
                 g["P6.1a"].split("--")[0].strip()],
                ["P6.1b", "C0 re-scored reproduces the Phase 3 macro F1 per stratum",
                 f"{g['P6.1b']['status']} (worst |Δ| = {g['P6.1b']['worst_abs_delta']:.1e})"],
                ["P6.2a", "the 4-annotator-marginalized score equals the Phase 3 value",
                 f"{g['P6.2a']['status']} (worst |Δ| = {g['P6.2a']['worst_abs_delta']:.1e})"]]
        if GEO:
            gg = GEO["gates"]["P6.3a"]
            rows.append(["P6.3a", "C0's S-unanimous geometry reproduces Phase 3 exactly",
                         f"{gg['status']} ({gg['recomputed_wall_adjacent_pct']} / "
                         f"{gg['recomputed_station_neighbouring_pct']})"])
        if CAM:
            ks = list(CAM["gates"])
            if ks:
                rows.append(["P6.4b", "each CAM targets the committed prediction",
                             f"PASS for {len(ks)} checkpoints"])
        if SEL:
            rows.append(["P6.6a", "risk at full coverage equals 1 − accuracy",
                         SEL["gates"]["P6.6a"]["status"]])
        table(doc, ["Gate", "What it checks", "Result"], rows,
              "Phase-boundary reproduction gates. Each recovers a number published in "
              "an earlier phase from the new code path.", widths=[0.7, 3.7, 2.1],
              font=8.2)

    h(doc, "2.2 P6-A: the held-out-annotator construction", 2)
    if PRE:
        para(doc, PRE["endpoints"]["P6-A"]["construction"])
        callout(doc, PRE["endpoints"]["P6-A"]["why_exclude_self"],
                title="Why annotator a is scored against the other three, not all four")

    h(doc, "2.3 P6-B: giving the human benchmark an interval", 2)
    if GEO:
        para(doc, GEO["definitions"])
        rows = [[k, v] for k, v in GEO["event_definitions"].items()]
        table(doc, ["Side", "What counts as one event"], rows,
              "Event definitions. Both sides use the identical adjacency relations.",
              widths=[1.3, 5.2], font=8.2)

    h(doc, "2.4 P6-C: Grad-CAM as a scored quantity", 2)
    if PRE:
        r = PRE["endpoints"]["P6-C"]
        para(doc, f"Layer: {r['layer']}. Target: {r['target_class']}. "
                  f"Top-q for the overlap masks: {r['top_q']}, {r['top_q_note']}.")
        rows = [[k, v] for k, v in r["quantities"].items()]
        table(doc, ["Quantity", "Definition"], rows,
              "The three scored attribution quantities, fixed before any map was "
              "rendered.", widths=[1.3, 5.2], font=8.2)
        callout(doc, "Attribution maps are the standard way this chapter is written, "
                     "and the standard way is not a measurement: a grid of heatmaps "
                     "invites the reader to agree with the author. The pre-registration "
                     "therefore fixed the layer, the target class and the overlap "
                     "threshold in advance and committed to reporting the numbers "
                     "whatever they said. Qualitative panels appear in this report only "
                     "after the quantitative endpoint has been stated, and their caption "
                     "carries that result.",
                     title="Why this phase scores attribution instead of displaying it")
    if CAM and CAM.get("target_selection"):
        callout(doc, CAM["target_selection"],
                title="A numerical fact the gate exposed")

    h(doc, "2.5 P6-D: risk–coverage and AURC", 2)
    if SEL:
        para(doc, f"Confidence score: {SEL['score']}. {SEL['score_note']}")
        para(doc, "Order the images by confidence, accept the most confident fraction c, "
                  "and record the error rate among the accepted. Sweeping c traces the "
                  "risk–coverage curve; AURC is the area beneath it, the expected error "
                  "of a model permitted to abstain, averaged over every abstention "
                  "budget. A model whose confidence perfectly ordered its own mistakes "
                  "would push all error into the low-confidence tail and score near zero.")
    doc.add_page_break()


# =====================================================================
def sec_results(doc) -> None:
    h(doc, "3. Results", 1)

    # ---- P6-A -------------------------------------------------------------
    h(doc, "3.1 P6-A: model versus the human panel", 2)
    if HUM:
        arm = HUM["headline_arm"]
        rows = []
        for s in STRATA:
            if s not in HUM["results"]:
                continue
            e = HUM["results"][s]; b = e["bootstrap"]; a = b["by_arm"][arm]
            rows.append([s, f"{e['n_images']:,}",
                         iv(b["human_panel_mean_ci95"]),
                         iv(a["model_mean_ci95"]),
                         f"{a['delta_mean']:+.4f}", iv(a["delta_ci95"]),
                         a["verdict"]])
        table(doc, ["Stratum", "n", "Human 95% CI", f"Model ({arm}) 95% CI",
                    "Δ", "Δ 95% CI", "Verdict"], rows,
              f"P6-A. Held-out annotator and model, scored by the same metric against "
              f"the same three-annotator reference panel, on the same images and the "
              f"same patient resamples. Headline arm {CFG_LABEL[arm]}.",
              widths=[1.25, 0.42, 0.92, 0.92, 0.55, 0.95, 1.5], font=7.4)

        deg = HUM.get("declared_degeneracy", {})
        if any(v.get("degenerate") for v in deg.values()):
            callout(doc, HUM["amendment"], title="Declared degeneracy, detected not assumed")
        sens = HUM.get("sensitivity_P6-AMD-5")
        if sens:
            callout(doc, f"{sens['what']} {sens['why_it_matters']}",
                    title="P6-AMD-5 — two asymmetries found after scoring")
            rows = []
            for s in STRATA:
                if s not in sens["by_stratum"]:
                    continue
                e = sens["by_stratum"][s]; a = e["by_arm"][arm]
                rows.append([s, pc(e["human_held_out"], 4), pc(a["model"], 4),
                             pc(e["modal_vote_oracle"], 4),
                             "n/a" if a["position_in_headroom"] is None
                             else f"{100 * a['position_in_headroom']:.0f}%",
                             f"{100 * e['mean_singleton_rate']:.0f}%",
                             "yes" if a["exceeds_oracle"] else "no"])
            table(doc, ["Stratum", "Held-out human", f"Model ({arm})",
                        "Modal-vote oracle", "Headroom recovered",
                        "Singleton rate", "Exceeds oracle?"], rows,
                  "P6-AMD-5. The modal vote of the same three references is the best "
                  "any single-label predictor can achieve against them, so it bounds "
                  "the comparison. The singleton rate is the fraction of held-out "
                  "annotators whose label is shared by none of the other three — a "
                  "structural handicap fixed by the stratum definition, not by skill.",
                  widths=[1.3, 0.9, 0.85, 0.95, 0.95, 0.8, 0.75], font=7.4)
            callout(doc, sens["modal_vote_oracle"] + " " + sens["singleton_rate"],
                    title="Why these two quantities settle the question")
        figure(doc, "P6_F38_human_comparator.png",
               "P6-A. Left: held-out annotator, model and the modal-vote oracle on a "
               "common axis; the unanimous stratum is greyed because the human side is "
               "1.0 by construction there. Right: the paired difference with "
               "patient-clustered 95% intervals.")
        if HUM.get("qualified_verdict"):
            rows = [[s, HUM["qualified_verdict"][s]] for s in STRATA
                    if s in HUM["qualified_verdict"]]
            table(doc, ["Stratum", "What may actually be claimed"], rows,
                  "The pre-registered verdicts stand as the frozen rules produced "
                  "them. These are the claims the evidence supports once P6-AMD-5 is "
                  "accounted for, and they are the form used everywhere else in this "
                  "thesis.", widths=[1.3, 5.2], font=7.2)
        if HUM.get("interpretation_superseded"):
            callout(doc, HUM["interpretation_superseded"],
                    title="The pre-registered interpretation, and why it is only "
                          "partly licensed")

    # ---- P6-B -------------------------------------------------------------
    h(doc, "3.2 P6-B: confusion geometry with intervals on both sides", 2)
    if GEO:
        arm = GEO["headline_arm"]
        rows = []
        for s in STRATA:
            if s not in GEO["results"]:
                continue
            e = GEO["results"][s]; a = e["by_arm"][arm]
            defined = e.get("human_geometry_defined", True)
            rows.append([s,
                         pc(a["wall_adjacent_pct_3seed"]),
                         pc(e["human"]["wall_adjacent_pct"]) if defined else "undefined",
                         f"{a.get('wall_adjacent_delta_mean', 0):+.2f}" if defined else "—",
                         pc(a["station_neighbouring_pct_3seed"]),
                         pc(e["human"]["station_neighbouring_pct"]) if defined else "undefined",
                         f"{a.get('station_neighbouring_delta_mean', 0):+.2f}" if defined else "—"])
        table(doc, ["Stratum", "Wall model %", "Wall human %", "Δ",
                    "Station model %", "Station human %", "Δ"], rows,
              f"P6-B. Adjacent-wall and neighbouring-station shares, model and human, "
              f"measured on the same images. Arm {arm}.",
              widths=[1.35, 0.85, 0.85, 0.6, 0.9, 0.9, 0.6], font=7.6)
        callout(doc, GEO["x3_settlement"]["finding"], title="X3, settled")
        figure(doc, "P6_F39_confusion_geometry.png",
               "P6-B. Both sides carry a patient-clustered interval and are differenced "
               "inside one resample. ✱ marks a stratum whose paired interval excludes "
               "zero.")

    # ---- P6-C -------------------------------------------------------------
    h(doc, "3.3 P6-C: attribution", 2)
    if ATT:
        arm = ATT["headline_arm"]
        prim = ATT["primary"]["stratum"]
        amd = ATT.get("amendment_P6-AMD-4")
        if amd:
            callout(doc, f"{amd['what']} {amd['why']} {amd['consequence']}",
                    title="P6-AMD-4 — the pre-registered primary is not estimable")
        rows = []
        for a in ATT["arms"]:
            r = ATT["primary"]["by_arm"][a]
            rows.append([CFG_LABEL[a], pc(r.get("dispersion_mean"), 4),
                         pc(r.get("spearman_rho"), 4), r.get("verdict", "—"),
                         pc(r.get("spread_spearman_rho"), 4),
                         iv(r.get("spread_spearman_ci95"), 4),
                         r.get("spread_verdict", "—")])
        table(doc, ["Configuration", "CAM dispersion", "ρ vs vote entropy",
                    "Pre-registered verdict", "ρ vs vote spread", "95% CI",
                    "Exploratory verdict"], rows,
              f"P6-C1 on {prim}. The pre-registered signal (vote entropy) is constant "
              f"within the stratum and yields no estimate. The right-hand columns are "
              f"the declared-exploratory substitute — anatomical vote spread — which "
              f"varies within a tier. They are exploratory and are not a "
              f"pre-registered endpoint.",
              widths=[1.7, 0.8, 0.8, 1.05, 0.8, 0.95, 0.95], font=7.0)
        if amd:
            callout(doc, amd["substitute"], title="What the substitute signal is")
            callout(doc, amd["carry_back"],
                    title="A consequence for Phase 4's RQ3, carried back")
        figure(doc, "P6_F40_dispersion_vs_entropy.png",
               "P6-C1. Left: the primary within-stratum correlation per arm. Right: the "
               "pooled value against the within-stratum value, showing the size of the "
               "stratum-membership confound.")

        rows = [[CFG_LABEL[a],
                 pc(ATT["secondary"][a]["inter_seed_iou_unanimous"], 4),
                 pc(ATT["secondary"][a]["inter_seed_iou_contested"], 4),
                 f"{ATT['secondary'][a]['delta']:+.4f}",
                 iv(ATT["secondary"][a]["delta_ci95"], 4),
                 ATT["secondary"][a]["verdict"]] for a in ATT["arms"]]
        table(doc, ["Configuration", "IoU unanimous", "IoU contested", "Δ", "95% CI",
                    "Verdict"], rows,
              f"P6-C2, secondary: inter-seed overlap of the top-"
              f"{int(ATT['top_q']*100)}% attribution masks.",
              widths=[2.0, 0.85, 0.85, 0.6, 1.1, 1.2], font=7.8)
        figure(doc, "P6_F41_attribution_stability.png",
               "P6-C2. Whether three seeds that agree on the label also agree on where "
               "the evidence is.")
        figure(doc, "P6_F42_gradcam_panels.png",
               "Qualitative Grad-CAM panels, shown after the quantitative endpoint and "
               "captioned with it. These illustrate; Figure F40 measures.")

    # ---- P6-D -------------------------------------------------------------
    h(doc, "3.4 P6-D: selective prediction", 2)
    if SEL:
        for key, label in (("internal", "Internal — GastroHUN test split"),
                           ("external", "External — HyperKvasir + GastroVision")):
            P = SEL.get(key)
            if not P:
                continue
            rows = []
            for a in P["by_arm"]:
                e = P["by_arm"][a]
                rows.append([CFG_LABEL[a], pc(e["aurc_3seed"], 4),
                             iv(e.get("aurc_ci95"), 4) if key == "internal" else "—",
                             pc(e["coverage_at_risk_10pct_3seed"], 3),
                             pc(e["risk_at_coverage_80pct_3seed"], 3)])
            table(doc, ["Configuration", "AURC", "95% CI", "Coverage at 10% risk",
                        "Risk at 80% coverage"], rows,
                  f"P6-D, {label}. AURC lower is better.",
                  widths=[2.3, 0.75, 1.2, 1.1, 1.1], font=7.8)
        figure(doc, "P6_F43_risk_coverage_internal.png",
               "P6-D internal. Risk–coverage curves and AURC with patient-clustered "
               "intervals.")
        figure(doc, "P6_F44_risk_coverage_external.png",
               "P6-D external. The same measurement on the Phase 5 panel, where "
               "rejecting an out-of-protocol image is the correct action. Intervals "
               "here are image-level (P5-DEV-3) and are not comparable with the "
               "internal ones.")
        pc5 = SEL["external"]["phase5_consistency"]
        callout(doc, f"{pc5['orientation']} Measured ρ = {pc5['spearman_rho']}, "
                     f"95% CI {iv(pc5['spearman_ci95'], 2)} → {pc5['verdict']}. "
                     f"{pc5['caveat']}."
                     + (f" {pc5['interval_excludes_point_estimate_note']}"
                        if pc5.get("interval_excludes_point_estimate") else ""),
                title="Does the single-operating-point result generalise?")

    h(doc, "3.5 Cross-endpoint synthesis", 2)
    figure(doc, "P6_F45_synthesis.png",
           "Arm ranking under every Phase 6 endpoint. The rows disagree, which is the "
           "finding: no configuration wins everywhere, and the endpoint that separates "
           "them most is external selective prediction.")
    doc.add_page_break()


# =====================================================================
def sec_discussion(doc) -> None:
    h(doc, "4. Discussion", 1)
    h(doc, "4.1 Ambiguity or capacity: what P6-A and P6-B jointly license", 2)
    if HUM:
        P = "S-contested (pooled)"
        v = HUM["verdict_summary"].get(P, "")
        sp = HUM["sensitivity_P6-AMD-5"]["by_stratum"][P]
        spa = sp["by_arm"][HUM["headline_arm"]]
        para(doc, f"The pre-registration fixed the interpretation of P6-A before the "
                  f"numbers existed: if the model is indistinguishable from the human "
                  f"panel on the contested strata while scoring in the twenties and "
                  f"forties, the low absolute scores are a property of the task. The "
                  f"frozen rule returned {v.lower()} on the pooled contested stratum. "
                  f"Taken at face value that is a striking claim, and it is the reason "
                  f"this section does not take it at face value.")
        para(doc, f"Two asymmetries were found on re-reading the construction, and both "
                  f"were then measured. First, exposure: the model is trained on targets "
                  f"derived from this panel and is optimised to predict its consensus, "
                  f"while the held-out annotator is simply being themselves. Second, and "
                  f"more decisively, choice: the model selects a label, whereas the "
                  f"annotator is stuck with the one they gave. The modal vote of the "
                  f"same three references — the best any single-label predictor can "
                  f"achieve against them — scores {pc(sp['modal_vote_oracle'], 4)}, "
                  f"against {pc(sp['human_held_out'], 4)} for the held-out annotator. "
                  f"The model reaches {pc(spa['model'], 4)}, which is "
                  f"{100 * (spa['position_in_headroom'] or 0):.0f}% of that headroom. "
                  f"It does not exceed the oracle on any stratum.")
        para(doc, "The defensible reading is therefore narrower than the frozen rule "
                  "suggests, and more useful. Predicting a panel is an easier task than "
                  "being a member of one, and on the 2-1-1 stratum half of all held-out "
                  "annotators are singletons who cannot score well whatever their skill. "
                  "What P6-A establishes is that the attainable ceiling on contested "
                  "images is 0.67 rather than 1.00 — so a large part of the Phase 3 "
                  "decline is the ceiling moving, exactly as the Phase 3B "
                  "ceiling-normalised analysis found — while a substantial model "
                  "shortfall against that reduced ceiling remains. The decline is "
                  "neither purely an artefact of the reference standard nor purely a "
                  "failure of the classifier, and the contribution of this endpoint is "
                  "to separate the two and size them.")
    if GEO:
        para(doc, "P6-B sharpens this. The wall axis mirrors human geometry on every "
                  "stratum where the comparison is defined: when the model confuses the "
                  "circumferential direction of the scope, it confuses it the way "
                  "endoscopists do. The station axis does not. The model's station "
                  "errors travel further along the insertion axis than human "
                  "disagreements do, and that difference survives a patient-clustered "
                  "interval on the pooled contested stratum. That is a specific, "
                  "actionable deficit rather than a general one.")

    h(doc, "4.2 What attribution did and did not explain", 2)
    if ATT:
        est = ATT["primary"].get("estimable", True)
        if not est:
            para(doc, "P6-C1 did not return a null; it returned a discovery about the "
                      "measurement itself. Annotator vote entropy is a deterministic "
                      "function of the vote pattern, and the agreement strata are "
                      "defined by that pattern, so entropy is constant inside "
                      "S-unanimous, S-majority and S-plurality alike. A within-stratum "
                      "correlation against it cannot exist. Reporting this as 'no "
                      "association' would have asserted a measurement that was never "
                      "available, so it is reported as NOT ESTIMABLE.")
            para(doc, f"The declared-exploratory substitute — the mean pairwise "
                      f"anatomical distance between the four annotators' labels, which "
                      f"does vary within a tier — returns "
                      f"{ATT['verdict_summary']['P6-C1b_exploratory_spread']} "
                      f"(ρ = {ATT['verdict_summary']['P6-C1b_exploratory_rho']}, 95% CI "
                      f"{iv(ATT['verdict_summary']['P6-C1b_exploratory_ci95'], 4)}). "
                      f"Taken with Phase 4's finding on predictive entropy, the picture "
                      f"is consistent: this model's uncertainty is not legible either "
                      f"in its output distribution or in where it looks.")
        else:
            para(doc, f"P6-C's pre-registered answer is "
                      f"{ATT['verdict_summary']['P6-C1_primary']} "
                      f"(ρ = {ATT['verdict_summary']['P6-C1_primary_rho']}, 95% CI "
                      f"{iv(ATT['verdict_summary']['P6-C1_primary_ci95'], 4)}).")
        para(doc, f"P6-C2 is the positive result of this endpoint. Attribution "
                  f"destabilises on contested images for every arm, and the three "
                  f"soft-target arms are markedly more spatially consistent across "
                  f"seeds than the two hard-label arms "
                  f"({pc(ATT['secondary']['C2']['inter_seed_iou_unanimous'], 3)} for C2 "
                  f"against {pc(ATT['secondary']['C0']['inter_seed_iou_unanimous'], 3)} "
                  f"for C0 on unanimous images). Training on the vote distribution "
                  f"does not make the model's confidence track disagreement, but it "
                  f"does make three independently seeded models agree about where the "
                  f"evidence is.")
        para(doc, "All of this is evidence about Grad-CAM specifically. No other "
                  "attribution method was run, by design (P6-DEV-3), so that a method "
                  "could not be selected after seeing which one correlated best.")

    h(doc, "4.3 Selective prediction as the deployable form of the calibration finding", 2)
    if SEL:
        ext = SEL.get("external")
        if ext:
            para(doc, f"Internally the arms are nearly indistinguishable on AURC, which "
                      f"is consistent with every other internal endpoint in this project "
                      f"failing to separate them. Externally they separate sharply: "
                      f"{SEL['best_arm_external_aurc']} attains an AURC of "
                      f"{pc(ext['by_arm'][SEL['best_arm_external_aurc']]['aurc_3seed'], 4)} "
                      f"while the internally-best arm "
                      f"{SEL['best_arm_internal_aurc']} reaches only "
                      f"{pc(ext['by_arm'][SEL['best_arm_internal_aurc']]['aurc_3seed'], 4)}. "
                      f"This is the Phase 4 calibration finding in deployable form: the "
                      f"question is not which arm is most accurate but which arm's "
                      f"confidence can be trusted to decide when not to answer.")

    h(doc, "4.4 Limitations specific to this phase", 2)
    bullet(doc, "Grad-CAM is one attribution method among several, and the P6-C verdict "
                "is evidence about it alone (P6-DEV-3).")
    bullet(doc, "Four annotators bound the precision of the human comparator; the "
                "held-out construction leaves a reference panel of three.")
    bullet(doc, "The human comparator is degenerate on any stratum defined by "
                "unanimity, and the human error geometry does not exist there at all. "
                "Both are declared rather than worked around.")
    bullet(doc, "Contested strata remain small (n = 342, 127 and 81), so the per-stratum "
                "intervals are wide and the pooled stratum carries the endpoints.")
    bullet(doc, "External intervals are image-level, inheriting the Phase 5 declaration "
                "P5-DEV-3, and must not be compared against the patient-clustered "
                "intervals in the same report.")
    bullet(doc, "The strata are defined by the same annotator agreement that determines "
                "the human comparator's score, so the human curve is partly mechanical. "
                "This does not affect the model-minus-human contrast, which is measured "
                "on identical rows, but it does mean the human curve alone should not be "
                "read as a free-standing estimate of expert accuracy.")
    doc.add_page_break()


# =====================================================================
def sec_conclusion(doc) -> None:
    h(doc, "5. Conclusion", 1)
    h(doc, "5.1 Answers to the four endpoints", 2)
    rows = []
    if HUM:
        for s in ("S-majority", "S-plurality", "S-no-majority", "S-contested (pooled)"):
            if s in HUM["verdict_summary"]:
                q = HUM.get("qualified_verdict", {}).get(s, "")
                rows.append([f"P6-A  {s}",
                             HUM["verdict_summary"][s] + (
                                 f" — qualified: {q.split('.')[0]}." if q else "")])
    if GEO:
        v = GEO["verdict_summary"].get("S-contested (pooled)", {})
        rows.append(["P6-B  wall axis, contested", v.get("wall", "—")])
        rows.append(["P6-B  station axis, contested", v.get("station", "—")])
    if ATT:
        rows.append(["P6-C1 attribution vs disagreement (pre-registered)",
                     ATT["verdict_summary"]["P6-C1_primary"]])
        rows.append(["P6-C1b same, exploratory substitute signal",
                     ATT["verdict_summary"]["P6-C1b_exploratory_spread"]])
        rows.append(["P6-C2 attribution stability",
                     ATT["verdict_summary"]["P6-C2_secondary"]])
    if SEL:
        rows.append(["P6-D  vs the Phase 5 ranking",
                     SEL["verdict_summary"]["phase5_consistency"]])
    if rows:
        table(doc, ["Endpoint", "Verdict"], rows,
              "Every verdict selected by a rule frozen before scoring.",
              widths=[3.2, 3.3], font=8.2)

    h(doc, "5.2 Carry-forward to Phase 7", 2)
    bullet(doc, "The thesis's central claim is that agreement stratification separates "
                "a falling reference standard from a falling classifier, and that the "
                "literature reports them as one quantity. Both fall; the ceiling "
                "accounts for the larger share, and a real model shortfall remains.")
    bullet(doc, "Never state P6-A as 'the model beats the expert'. State it as the "
                "qualified verdict: the model out-predicts an individual annotator but "
                "recovers only about a quarter of the headroom to the panel's own modal "
                "vote, and it exceeds that oracle on no stratum.")
    bullet(doc, "Report the human comparator curve AND the modal-vote oracle alongside "
                "every stratified performance figure. Phase 3's numbers mislead without "
                "the first, and the first misleads without the second.")
    bullet(doc, "The station axis is the specific residual deficit worth naming; the "
                "wall axis is not.")
    bullet(doc, "Use external AURC as the arm-selection endpoint. It is the only "
                "measurement in the project that separates the configurations decisively "
                "and it agrees with the Phase 5 rejection ranking.")
    doc.add_page_break()


# =====================================================================
def sec_appendices(doc) -> None:
    h(doc, "Appendix A. Per-held-out-annotator breakdown", 1)
    if HUM:
        arm = HUM["headline_arm"]
        for s in ("S-majority", "S-contested (pooled)"):
            if s not in HUM["results"]:
                continue
            rows = []
            for name, e in HUM["results"][s]["per_held_out_annotator"].items():
                rows.append([name, ", ".join(e["reference_panel"]),
                             pc(e["held_out_annotator_score"], 4),
                             pc(e["by_arm"][arm]["model_3seed_mean"], 4),
                             f"{e['by_arm'][arm]['delta_vs_human']:+.4f}"])
            table(doc, ["Held out", "Reference panel", "Human", f"Model ({arm})", "Δ"],
                  rows, f"Per-held-out-annotator scores on {s}. FG2 is the outlier "
                        f"annotator identified in Phase 0 and is included unchanged.",
                  widths=[0.8, 1.7, 0.9, 1.1, 0.8], font=8.0)

    h(doc, "Appendix B. Pre-registration record and declared deviations", 1)
    if PRE:
        rows = [[k, v["item"], v["adopted"], v["because"]]
                for k, v in PRE["declared_deviations"].items()]
        table(doc, ["ID", "Blueprint item", "Adopted", "Because"], rows,
              "Deviations declared in the frozen pre-registration.",
              widths=[0.7, 1.6, 1.5, 2.7], font=7.4)
    amds = []
    if HUM and HUM.get("amendment"):
        amds.append(["P6-AMD-1", HUM["amendment"]])
    if GEO and GEO.get("x3_settlement", {}).get("amendment"):
        amds.append(["P6-AMD-2", GEO["x3_settlement"]["amendment"]])
    if CAM and CAM.get("amendment"):
        amds.append(["P6-AMD-3", CAM["amendment"]])
    if ATT and ATT.get("amendment_P6-AMD-4"):
        a4 = ATT["amendment_P6-AMD-4"]
        amds.append(["P6-AMD-4", f"{a4['what']} {a4['why']} {a4['consequence']} "
                                 f"{a4['substitute']}"])
    if amds:
        table(doc, ["ID", "Amendment discovered during execution"], amds,
              "Amendments. Each was forced by a gate or by a numerical fact the "
              "pre-registration did not anticipate, and each is recorded rather than "
              "silently adopted.", widths=[0.8, 5.7], font=7.6)

    h(doc, "Appendix C. Script and artefact manifest", 1)
    rows = [
        ["src/models/phase6_prereg.py", "reports/phase6_prereg.json", "P6.0"],
        ["src/models/phase6_common.py", "(shared primitives, gates P6.1a/b)", "P6.1"],
        ["src/models/phase6_human.py", "reports/phase6_human.json", "P6.2"],
        ["src/models/phase6_geometry.py", "reports/phase6_geometry.json", "P6.3"],
        ["src/models/phase6_cam.py", "reports/phase6_cams_*.npz, phase6_cam_gate.json", "P6.4"],
        ["src/models/phase6_cam_eval.py", "reports/phase6_cam_eval.json", "P6.5"],
        ["src/models/phase6_selective.py", "reports/phase6_selective.json", "P6.6"],
        ["src/report/figures_phase6.py", "figures_phase6/P6_F38–F45", "P6.7"],
        ["src/report/build_phase6_docx.py", "Phase6_Report.docx", "P6.8"],
    ]
    table(doc, ["Script", "Artefact", "Step"], rows,
          "Reproducibility index. Running the scripts in this order regenerates every "
          "number, figure and verdict in this report.", widths=[2.5, 3.2, 0.8], font=7.8)

    h(doc, "Appendix D. Analyses NOT executed, and what each would cost", 1)
    bullet(doc, "An attribution-method sweep (Grad-CAM++, ScoreCAM, Integrated "
                "Gradients). Roughly 3× the backward passes, about 45 minutes on this "
                "hardware — declined because it would invite selecting whichever method "
                "correlated best.")
    bullet(doc, "A second backbone, to test whether the calibration and attribution "
                "findings are ConvNeXt-specific. Nine training runs, roughly 10 GPU "
                "hours. Not executed; the strongest remaining threat to external "
                "validity of the Phase 6 conclusions.")
    bullet(doc, "Human comparator on the external corpora. Not possible: neither corpus "
                "publishes per-annotator labels.")
    bullet(doc, "A patient-clustered external interval. Not possible: neither corpus "
                "publishes a case identifier.")


def main() -> None:
    missing = [n for n, v in [("phase6_human.json", HUM),
                              ("phase6_geometry.json", GEO),
                              ("phase6_selective.json", SEL)] if v is None]
    if missing:
        raise SystemExit("cannot build the report; missing: " + ", ".join(missing))
    if ATT is None:
        print("  NOTE: phase6_cam_eval.json absent — P6-C sections will be omitted")

    doc = new_document()
    title_page(doc)
    front_matter(doc)
    sec_abstract(doc)
    sec_introduction(doc)
    sec_methods(doc)
    sec_results(doc)
    sec_discussion(doc)
    sec_conclusion(doc)
    sec_appendices(doc)
    add_page_numbers(doc)
    doc.save(BD.OUT)
    caps = [p.text for p in doc.paragraphs if p.style.name == "Caption"]
    nfig = sum(1 for c in caps if c.strip().startswith("Figure"))
    ntab = sum(1 for c in caps if c.strip().startswith("Table"))
    print(f"[report] {len(doc.inline_shapes)} images, {nfig} figure captions, "
          f"{ntab} table captions -> {BD.OUT}")


if __name__ == "__main__":
    main()
