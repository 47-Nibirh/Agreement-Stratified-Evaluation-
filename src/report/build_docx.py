"""
Build the Phase 0 / Phase 1 Word report.
========================================
Assembles `Phase0_Phase1_Report.docx` from the computed artefacts:
  reports/phase0_results.json      - every Phase 0 measurement
  literature/extraction_table.csv  - the 50 included studies with APA strings
  literature/prisma_counts.json    - PRISMA stage counts
  figures/*.png                    - 19 generated figures

No numeric value is typed by hand; all are interpolated from those artefacts.

Run:  python src/report/build_docx.py
"""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
FIGD = ROOT / "figures"
OUT = ROOT / "Phase0_Phase1_Report.docx"

R = json.loads((ROOT / "reports" / "phase0_results.json").read_text(encoding="utf-8"))
P = json.loads((ROOT / "literature" / "prisma_counts.json").read_text(encoding="utf-8"))
LITDF = pd.read_csv(ROOT / "literature" / "extraction_table.csv")

B = R["battery"]
L = R["leakage"]
LA = R["label_audit"]
PT = R["permutation_test"]
PROV = R["provenance"]
POW = R["power"]
ETH = R["ethics"]

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
DARKRED = RGBColor(0x8B, 0x1A, 0x1A)
GREY = RGBColor(0x59, 0x59, 0x59)

FIGN = {"n": 0}
TABN = {"n": 0}


# ==========================================================================
# Low-level helpers
# ==========================================================================
def _field(paragraph, instr: str) -> None:
    """Insert a Word field code (used for TOC, SEQ, PAGE)."""
    r = paragraph.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    sp = OxmlElement("w:fldChar"); sp.set(qn("w:fldCharType"), "separate")
    tx = OxmlElement("w:t"); tx.text = " "
    en = OxmlElement("w:fldChar"); en.set(qn("w:fldCharType"), "end")
    for el in (fc, it, sp, tx, en):
        r._r.append(el)


def _shade(cell, hexcolor: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear"); el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def _cell_text(cell, text, *, bold=False, size=8.0, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1.5)
    p.paragraph_format.space_after = Pt(1.5)
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def _repeat_header(row) -> None:
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader"); el.set(qn("w:val"), "true")
    trPr.append(el)


def h(doc, text, level=1, *, page_break=False):
    if page_break:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT if level > 1 else DARKRED
    return p


def para(doc, text, *, size=10.5, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         space_after=7, bold=False, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.18
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def rich(doc, chunks, *, size=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=7):
    """chunks: list of (text, {'b':bool,'i':bool}) tuples."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.18
    for text, fmt in chunks:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = fmt.get("b", False)
        r.italic = fmt.get("i", False)
        if fmt.get("c"):
            r.font.color.rgb = fmt["c"]
    return p


def bullet(doc, text, *, level=0, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.28 + 0.25 * level)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def callout(doc, text, *, title=None):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.cell(0, 0)
    _shade(c, "F2F4F8")
    c.text = ""
    if title:
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title); r.bold = True; r.font.size = Pt(9.5)
        r.font.color.rgb = DARKRED
        p2 = c.add_paragraph()
    else:
        p2 = c.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run(text); r.font.size = Pt(9.5); r.italic = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def figure(doc, filename, caption, *, width=6.1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    # Bind the image to its caption. Without this a tall figure can be pushed to
    # the next page while its caption stays behind, which orphans both.
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(FIGD / filename), width=Inches(width))

    FIGN["n"] += 1
    cp = doc.add_paragraph(style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(11)
    r = cp.add_run("Figure "); r.bold = True; r.font.size = Pt(9)
    _field(cp, r" SEQ Figure \* ARABIC ")
    r2 = cp.add_run(f". {caption}"); r2.font.size = Pt(9)
    r2.font.color.rgb = GREY
    return cp


def table(doc, headers, rows, caption, *, widths=None, font=8.0,
          head_font=8.0, align_right=None, note=None):
    TABN["n"] += 1
    cp = doc.add_paragraph(style="Caption")
    cp.paragraph_format.space_before = Pt(8)
    cp.paragraph_format.space_after = Pt(3)
    r = cp.add_run("Table "); r.bold = True; r.font.size = Pt(9)
    _field(cp, r" SEQ Table \* ARABIC ")
    r2 = cp.add_run(f". {caption}"); r2.font.size = Pt(9); r2.bold = False
    r2.font.color.rgb = GREY

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    hdr = t.rows[0]
    _repeat_header(hdr)
    for i, htxt in enumerate(headers):
        _shade(hdr.cells[i], "1F3A5F")
        _cell_text(hdr.cells[i], htxt, bold=True, size=head_font,
                   color=RGBColor(0xFF, 0xFF, 0xFF))

    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            al = (WD_ALIGN_PARAGRAPH.RIGHT
                  if align_right and ci in align_right else None)
            _cell_text(cells[ci], val, size=font, align=al)
        if ri % 2 == 1:
            for c in cells:
                _shade(c, "F5F7FA")

    if widths:
        for ri in range(len(t.rows)):
            for ci, w in enumerate(widths):
                t.rows[ri].cells[ci].width = Cm(w)

    if note:
        np_ = doc.add_paragraph()
        np_.paragraph_format.space_before = Pt(2)
        np_.paragraph_format.space_after = Pt(10)
        nr = np_.add_run(f"Note. {note}")
        nr.font.size = Pt(8); nr.italic = True; nr.font.color.rgb = GREY
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


# ==========================================================================
# Document setup
# ==========================================================================
def new_document() -> Document:
    doc = Document()

    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    for s in doc.sections:
        s.page_width, s.page_height = Inches(8.27), Inches(11.69)   # A4
        s.top_margin = s.bottom_margin = Inches(0.9)
        s.left_margin = s.right_margin = Inches(0.95)

    # update all fields (TOC / LoF / LoT) when the document is opened
    el = OxmlElement("w:updateFields"); el.set(qn("w:val"), "true")
    doc.settings.element.append(el)
    return doc


def add_page_numbers(doc) -> None:
    for s in doc.sections:
        p = s.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Page "); r.font.size = Pt(8.5); r.font.color.rgb = GREY
        _field(p, " PAGE ")
        r2 = p.add_run(" of "); r2.font.size = Pt(8.5); r2.font.color.rgb = GREY
        _field(p, " NUMPAGES ")


def title_page(doc) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Automatic Gastrointestinal Disease Classification from Upper "
                  "Gastrointestinal Endoscopy Reports Using Natural Language "
                  "Processing and Machine Learning")
    r.bold = True; r.font.size = Pt(19); r.font.color.rgb = DARKRED

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run("Phase 0 — Data Provenance and Integrity Gate\n"
                  "Phase 1 — Literature Review and Problem Framing")
    r.bold = True; r.font.size = Pt(13.5); r.font.color.rgb = ACCENT

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Interim Technical Report")
    r.font.size = Pt(11.5); r.italic = True

    doc.add_paragraph()
    bar = doc.add_paragraph(); bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = bar.add_run("─" * 46); rb.font.color.rgb = ACCENT

    meta = [
        ("Degree programme", "B.Sc. in Computer Science and Engineering"),
        ("Research domain", "Biomedical Artificial Intelligence — Clinical NLP "
                            "and Machine Learning"),
        ("Dataset under audit", f"Peptic Ulcer_Dataset.xlsx "
                                f"({PROV['n_rows']:,} records × {PROV['n_cols']} fields)"),
        ("Dataset SHA-256", PROV["sha256"]),
        ("Governing protocol", "THESIS_RESEARCH_BLUEPRINT.md (v2.0)"),
        ("Reporting standards", "TRIPOD+AI, PROBAST+AI, PRISMA 2020, CRISP-DM"),
        ("Report date", "26 July 2026"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in meta:
        c = t.add_row().cells
        _cell_text(c[0], k, bold=True, size=9.5)
        _cell_text(c[1], v, size=9.5)
        c[0].width, c[1].width = Cm(4.6), Cm(11.0)

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("STATUS: CONDITIONAL — the integrity gate did not clear. "
                  "Route A (reframe) is in force.")
    r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARKRED

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def front_matter(doc) -> None:
    for title, instr in [
        ("Table of Contents", r'TOC \o "1-3" \h \z \u'),
        ("List of Figures", r'TOC \h \z \c "Figure"'),
        ("List of Tables", r'TOC \h \z \c "Table"'),
    ]:
        # Styled to match Heading 1 but deliberately NOT a heading style, so the
        # front-matter titles do not list themselves inside the tables.
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = DARKRED
        tp = doc.add_paragraph()
        _field(tp, instr)
        note = doc.add_paragraph()
        nr = note.add_run("If this list appears blank, select all (Ctrl+A) and press F9 to "
                          "update the field.")
        nr.font.size = Pt(8); nr.italic = True; nr.font.color.rgb = GREY
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def main() -> None:
    doc = new_document()
    title_page(doc)
    front_matter(doc)

    import content_phase0 as C0
    import content_phase1 as C1

    C0.sec_executive_summary(doc)
    C0.sec_phase0(doc)
    C1.sec_phase1(doc)
    C1.sec_methodology(doc)
    C1.sec_results(doc)
    C1.sec_discussion(doc)
    C1.sec_conclusion(doc)
    C1.sec_references(doc)
    C1.sec_appendix(doc)

    add_page_numbers(doc)
    doc.save(OUT)

    caps = [p.text for p in doc.paragraphs if p.style.name == "Caption"]
    nfig = sum(1 for c in caps if c.strip().startswith("Figure"))
    ntab = sum(1 for c in caps if c.strip().startswith("Table"))
    print(f"[report] {len(doc.inline_shapes)} images, {nfig} figure captions, "
          f"{ntab} table captions -> {OUT}")


if __name__ == "__main__":
    main()
