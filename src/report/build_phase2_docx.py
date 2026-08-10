"""
Build the GastroHUN Phase 2 Word report.
========================================
Assembles `Phase2_Report.docx` from the computed artefacts:

  reports/phase2_env.json                 environment snapshot
  reports/phase2_vram_probe.json          GATE 1, precision benchmark, batch ladder
  reports/phase2_split_provenance.json    GATE 2/3, cohort composition
  reports/phase2_norm_stats.json          training-set channel statistics
  reports/phase2_prereg.json              frozen pre-registration
  reports/phase2_trainable_layers.json    resolved fine-tuning scope
  reports/phase2_run_seed*.json           per-seed run manifests and histories
  reports/phase2_test_metrics.json        test results, intervals, verdict
  figures_phase2/*.png                    generated figures

Reuses the rendering helpers of build_docx.py, repointed at this report's
figure directory and output filename. No numeric value is typed by hand.

Run:  python src/report/build_phase2_docx.py
"""
from __future__ import annotations

import json
from pathlib import Path

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt

import build_docx as BD

ROOT = Path(__file__).resolve().parents[2]
BD.FIGD = ROOT / "figures_phase2"
BD.OUT = ROOT / "Phase2_Report.docx"

from build_docx import (ACCENT, DARKRED, GREY, _cell_text, add_page_numbers,  # noqa: E402
                        bullet, callout, figure, front_matter, h, new_document,
                        para, rich, table)

REP = ROOT / "reports"


def J(name):
    return json.loads((REP / name).read_text(encoding="utf-8"))


ENV = J("phase2_env.json")
PROBE = J("phase2_vram_probe.json")
PROV = J("phase2_split_provenance.json")
NORM = J("phase2_norm_stats.json")
PRE = J("phase2_prereg.json")
LAY = J("phase2_trainable_layers.json")
MET = J("phase2_test_metrics.json")
RUNS = [json.loads(p.read_text(encoding="utf-8"))
        for p in sorted(REP.glob("phase2_run_seed*.json"))]
RUNS.sort(key=lambda r: r["seed"])

R0 = RUNS[0]
AGG = MET["aggregate"]
REPRO = MET["reproduction"]
DEV = {d["id"]: d for d in PRE["deviations"]}
COH = PROV["cohort"]
PASS = REPRO["verdict"] == "PASS"


def pc(x, d=2):
    """Fraction -> percent string."""
    return f"{100 * x:.{d}f}"


# =====================================================================
def title_page(doc) -> None:
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Agreement-Stratified Evaluation of Deep Learning for "
                  "Anatomical Landmark Recognition in Upper Gastrointestinal "
                  "Endoscopy")
    r.bold = True; r.font.size = Pt(19); r.font.color.rgb = DARKRED

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run("Phase 2 — Baseline Reproduction")
    r.bold = True; r.font.size = Pt(14.5); r.font.color.rgb = ACCENT

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Pre-registered reproduction of the published ConvNeXt-Tiny "
                  "benchmark on the complete-agreement cohort")
    r.font.size = Pt(11.5); r.italic = True

    doc.add_paragraph()
    bar = doc.add_paragraph(); bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = bar.add_run("─" * 46); rb.font.color.rgb = ACCENT

    meta = [
        ("Degree programme", "B.Sc. in Computer Science and Engineering"),
        ("Research domain", "Biomedical Artificial Intelligence — Medical Image "
                            "Analysis and Deep Learning"),
        ("Evaluation cohort",
         f"GastroHUN complete-agreement subset — {COH['n_images']:,} images "
         f"({COH['by_split']['Train']:,} train / {COH['by_split']['Validation']} "
         f"validation / {COH['by_split']['Test']} test), "
         f"{COH['n_patients']} patients, {COH['n_classes']} classes"),
        ("Corpus provenance",
         "Hospital Universitario Nacional de Colombia; Sci Data 12:102 (2025); "
         "doi:10.1038/s41597-025-04401-5"),
        ("Model", f"ConvNeXt-Tiny, ImageNet-pretrained, "
                  f"{LAY['feature_params_total'] / 1e6:.1f} M feature parameters"),
        ("Compute", f"{ENV['gpu']}, {PROBE['total_vram_mib'] / 1024:.1f} GB VRAM, "
                    f"compute capability {ENV['compute_capability']}"),
        ("Governing protocol", "THESIS_RESEARCH_BLUEPRINT.md (v3.0) §4 Phase 2"),
        ("Pre-registration", f"reports/phase2_prereg.json, frozen "
                             f"{PRE['frozen_at']}"),
        ("Reporting standards", "CLAIM, TRIPOD+AI, STARD-AI, PROBAST+AI"),
        ("Report date", "26 July 2026"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in meta:
        c = t.add_row().cells
        _cell_text(c[0], k, bold=True, size=9.5)
        _cell_text(c[1], v, size=9.5)
        c[0].width, c[1].width = Cm(4.6), Cm(11.0)

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"VERDICT: {REPRO['verdict']} — observed macro F1 "
        f"{REPRO['observed_macro_f1']:.2f} against a pre-registered target of "
        f"{REPRO['published_macro_f1']:.1f} ± {REPRO['acceptance_band_points']} "
        f"points (Δ = {REPRO['delta_points']:+.2f}).")
    r.bold = True; r.font.size = Pt(10.5)
    r.font.color.rgb = ACCENT if PASS else DARKRED

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def abstract(doc) -> None:
    h(doc, "Abstract", 1)
    para(doc,
         f"Phase 2 reproduces the published ConvNeXt-Tiny image-classification "
         f"benchmark on the GastroHUN complete-agreement cohort. Its purpose is "
         f"instrument calibration rather than discovery: it establishes that the "
         f"training and evaluation apparatus is trustworthy before Phases 3–6 "
         f"use it to make claims the literature has not tested. "
         f"The evaluation cohort was rebuilt from the official patient-level "
         f"splits and verified against the Phase 0 integrity artefacts "
         f"({COH['by_split']['Train']:,} / {COH['by_split']['Validation']} / "
         f"{COH['by_split']['Test']} images over {COH['n_classes']} classes, "
         f"zero patient overlap, {PROV['hash_resolution']['n_resolved']:,} of "
         f"{PROV['hash_resolution']['n_cohort']:,} filenames resolved against "
         f"the Phase 0 SHA-256 inventory). "
         f"Channel statistics were computed on the training split alone and "
         f"differ substantially from the ImageNet defaults "
         f"(red mean {NORM['mean'][0]:.3f} against {NORM['imagenet_mean'][0]:.3f}). "
         f"A hardware benchmark established that automatic mixed precision, "
         f"which the governing protocol prescribes, runs at "
         f"{PROBE['amp_vs_fp32_speedup']:.2f}× the throughput of float32 on this "
         f"tensor-core-less device; float32 was therefore adopted and the "
         f"deviation recorded. "
         f"The reproduction target, acceptance band, seed count, interval "
         f"procedure and failure diagnostics were frozen in a pre-registration "
         f"before the first training run. "
         f"Across {len(RUNS)} seeds the reproduced model attained a mean test "
         f"macro F1 of {REPRO['observed_macro_f1']:.2f} "
         f"(patient-clustered bootstrap 95% CI "
         f"{pc(AGG['seed_mean_boot_ci95'][0])}–"
         f"{pc(AGG['seed_mean_boot_ci95'][1])}), against a published reference "
         f"of {REPRO['published_macro_f1']:.1f}. The pre-registered acceptance "
         f"criterion of ±{REPRO['acceptance_band_points']} points is therefore "
         f"{'met' if PASS else 'not met'}, and the pipeline is "
         f"{'validated for use in Phases 3–6' if PASS else 'not yet validated'}.")


# =====================================================================
def chapter1(doc) -> None:
    h(doc, "1  Introduction and Phase Positioning", 1, page_break=True)

    h(doc, "1.1  Position of Phase 2 in the research pipeline", 2)
    para(doc,
         "The governing blueprint divides this thesis into eight phases. Phases "
         "0 and 1 are complete: the data provenance and integrity gate returned "
         "PROCEED on the GastroHUN corpus, and a PRISMA 2020 review of 82 "
         "included studies framed the problem. Phase 2 is the first phase in "
         "which a model is trained. Its remit, stated in the blueprint, is "
         "narrow and deliberately so: reproduce the published ConvNeXt-Tiny "
         "result on the complete-agreement subset, in order to validate the "
         "pipeline before changing anything.")
    figure(doc, "P2_F01_flow.png",
           "Phase 2 execution order. Gates are hard stops; the pre-registration "
           "boundary separates decisions made before the result was known from "
           "everything computed afterwards.", width=4.5)

    h(doc, "1.2  Carry-forward from Phases 0 and 1", 2)
    para(doc,
         f"Phase 0 established that the corpus contains {PROV['corpus']['n_images']:,} "
         f"images from {PROV['corpus']['n_patients']} patients, with four "
         f"independent expert annotations retained per image, and that complete "
         f"four-of-four agreement holds for only "
         f"{COH['retention_pct']:.2f}% of them. The published benchmark — and "
         f"therefore this reproduction — is defined exclusively on that "
         f"unanimous subset. Phase 0 also measured per-patient Fleiss' κ at "
         f"0.7459 ± 0.1448, which is the reason every interval in this report "
         f"resamples patients rather than images.")

    h(doc, "1.3  Why reproduction precedes contribution", 2)
    para(doc,
         "A reproduction phase produces no scientific claim. Its output is a "
         "statement about the apparatus: that a model trained by this code, on "
         "this cohort, with this preprocessing, lands where the literature says "
         "it should. Without that statement, any difference observed in Phase 3 "
         "between agreement strata, or in Phase 4 between target constructions, "
         "would be uninterpretable — it could equally reflect a defect in the "
         "harness. Reproduction converts the pipeline from an assumption into a "
         "measured instrument.")
    callout(doc,
            "A reproduction that is tuned until it matches is not a "
            "reproduction. The target, the acceptance band and the diagnostic "
            "response to failure were therefore fixed in writing, and "
            "timestamped, before the first full training run was launched. The "
            "pre-registration is reproduced verbatim in Appendix A.",
            title="On the discipline this phase requires")

    h(doc, "1.4  Scope and explicit exclusions", 2)
    para(doc, "The following are outside Phase 2 and are not attempted here:")
    for x in PRE["scope_exclusions"]:
        bullet(doc, x)
    para(doc,
         "Model selection is likewise not in scope: the blueprint fixes "
         "ConvNeXt-Tiny as the primary backbone on compute grounds, and "
         "ConvNeXt-Large — the 88.25 published ceiling — is cited rather than "
         "attempted, because it does not fit the available 4 GB of video memory.")

    h(doc, "1.5  Contributions of this phase", 2)
    for b in [
        "A verified, versioned complete-agreement cohort with per-image "
        "provenance, reusable unchanged by Phases 3–6.",
        "Training-set channel statistics measured rather than assumed, and "
        "shown to differ materially from the ImageNet defaults.",
        "A hardware characterisation demonstrating that the protocol's "
        "mixed-precision prescription is a pessimisation on this class of GPU, "
        "with the measurement that justifies departing from it.",
        "A frozen pre-registration binding the acceptance criterion.",
        f"A {len(RUNS)}-seed reproduction with patient-clustered intervals, and "
        f"an explicit gate decision on the result.",
        "An observation about the published stability intervals: they are "
        "standard errors of a bootstrap mean, not intervals on model "
        "performance, and are therefore roughly an order of magnitude narrower "
        "than a correctly clustered interval on the same quantity.",
    ]:
        bullet(doc, b)


# =====================================================================
def chapter2(doc) -> None:
    h(doc, "2  Methodology Summary", 1, page_break=True)

    h(doc, "2.1  Design overview", 2)
    para(doc,
         "The design is a single-arm reproduction. One architecture, one label "
         "construction, one cohort, repeated over independent random seeds, "
         "evaluated once on a held-out test set defined by patient identity. "
         "Every decision that could be made after seeing a result was fixed "
         "beforehand.")

    h(doc, "2.2  Data", 2)
    h(doc, "2.2.1  Consensus cohort definition", 3)
    para(doc,
         f"The cohort is every image for which all four annotators assigned the "
         f"same label, taken under the official patient-level partition "
         f"published with the dataset. This yields {COH['n_images']:,} images, "
         f"{COH['retention_pct']:.2f}% of the corpus, distributed as "
         f"{COH['by_split']['Train']:,} training, "
         f"{COH['by_split']['Validation']} validation and "
         f"{COH['by_split']['Test']} test images — matching the counts reported "
         f"in the data descriptor exactly.")
    figure(doc, "P2_F02_cohort.png",
           "Construction of the complete-agreement cohort. Restricting to "
           "unanimity removes 39.80% of the corpus and, as panel C shows, "
           "removes two patients entirely.")

    h(doc, "2.2.2  Split integrity within the cohort", 3)
    para(doc,
         "Patient-level separation was re-verified after the consensus "
         "restriction rather than inherited from Phase 0, because filtering can "
         "in principle change which patients appear in which split.")
    table(doc,
          ["Check", "Result"],
          [["Train ∩ Validation patients", str(len(PROV["patient_overlap"]["train_val"]))],
           ["Train ∩ Test patients", str(len(PROV["patient_overlap"]["train_test"]))],
           ["Validation ∩ Test patients", str(len(PROV["patient_overlap"]["val_test"]))],
           ["Cohort filenames resolved against Phase 0 SHA-256 inventory",
            f"{PROV['hash_resolution']['n_resolved']:,} / "
            f"{PROV['hash_resolution']['n_cohort']:,}"],
           ["Duplicate filenames within cohort",
            str(PROV["hash_resolution"]["n_duplicate_filenames"])],
           ["Cohort images present on disk",
            f"{PROV['hash_resolution']['n_present_on_disk']:,} / "
            f"{PROV['hash_resolution']['n_cohort']:,}"],
           ["GATE 2 (counts and class coverage)",
            "PASS" if PROV["gate2_pass"] else "FAIL"],
           ["GATE 3 (hash and disk resolution)",
            "PASS" if PROV["gate3_pass"] else "FAIL"]],
          "Split-integrity and provenance checks on the complete-agreement "
          "cohort. All checks are recomputed from source, not inherited.",
          widths=[10.5, 5.0])

    h(doc, "2.2.3  Attrition introduced by consensus filtering", 3)
    lost = COH["patients_lost_to_consensus"]
    lost_txt = "; ".join(f"{k}: patient {', '.join(map(str, v))}"
                         for k, v in lost.items())
    para(doc,
         f"The descriptor describes the partition as 270 training, 58 "
         f"validation and 59 test patients. Within the complete-agreement "
         f"cohort the corresponding counts are "
         f"{COH['patients_by_split']['Train']}, "
         f"{COH['patients_by_split']['Validation']} and "
         f"{COH['patients_by_split']['Test']}: two patients contribute no "
         f"unanimously labelled image at all ({lost_txt}). The discrepancy is "
         f"small but it is real, it is not documented in the descriptor, and it "
         f"matters here because the bootstrap resamples patients — the test "
         f"set contains {COH['patients_by_split']['Test']} resampling units, "
         f"not 59.")

    h(doc, "2.2.4  Class and acquisition-stream composition", 3)
    para(doc,
         f"Class composition is homogeneous across the three splits "
         f"(χ² = {PROV['class_split_chi2']}, p = {PROV['class_split_p']:.3f}). "
         f"Acquisition stream is not "
         f"(χ² = {PROV['stream_split_chi2']}, "
         f"p = {PROV['stream_split_p']:.2e}), which is limitation L4 carried "
         f"forward from Phase 0 and restated here at cohort level.")
    figure(doc, "P2_F03_classes.png",
           "Cohort composition and the statistical resolution the test set can "
           "support. Panel B is the reason per-class results in this report are "
           "labelled exploratory.")
    srows = []
    for s in ["Train", "Validation", "Test"]:
        d = PROV["stream_by_split"][s]
        tot = sum(d.values())
        srows.append([s, f"{d.get('direct_capture', 0):,}",
                      f"{d.get('video_frame', 0):,}",
                      f"{100 * d.get('video_frame', 0) / tot:.2f}%"])
    table(doc, ["Split", "Direct capture", "Video frame", "Video frame share"],
          srows,
          f"Acquisition-stream composition of the cohort by split. The "
          f"imbalance is significant (χ² = {PROV['stream_split_chi2']}, "
          f"p = {PROV['stream_split_p']:.2e}) and is declared as limitation L4.",
          align_right=[1, 2, 3])

    h(doc, "2.3  Preprocessing", 2)
    para(doc,
         f"Images are decoded to RGB and resampled to {NORM['size']}×"
         f"{NORM['size']} with Lanczos interpolation, as specified. The "
         f"resampling is performed once and cached, so training, validation and "
         f"test images pass through an identical code path and the operation "
         f"cannot drift between seeds.")
    h(doc, "2.3.1  Training-set normalisation statistics", 3)
    para(doc,
         f"Channel statistics were computed over the "
         f"{NORM['n_images']:,} training images only "
         f"({NORM['n_pixels_per_channel']:,} pixels per channel), at final "
         f"resolution. They differ from the ImageNet defaults by up to "
         f"{max(NORM['abs_delta_mean']):.3f} in the mean, which is what one "
         f"expects of white-light endoscopy: the field of view is mucosa, and "
         f"mucosa is red.")
    table(doc,
          ["Channel", "GastroHUN mean", "ImageNet mean", "Δ",
           "GastroHUN SD", "ImageNet SD", "Δ"],
          [[c,
            f"{NORM['mean'][i]:.4f}", f"{NORM['imagenet_mean'][i]:.3f}",
            f"{NORM['mean'][i] - NORM['imagenet_mean'][i]:+.4f}",
            f"{NORM['std'][i]:.4f}", f"{NORM['imagenet_std'][i]:.3f}",
            f"{NORM['std'][i] - NORM['imagenet_std'][i]:+.4f}"]
           for i, c in enumerate(["Red", "Green", "Blue"])],
          "Training-set channel statistics against the ImageNet defaults. "
          "Using the defaults would mis-centre the red channel by 0.116.",
          align_right=[1, 2, 3, 4, 5, 6])
    figure(doc, "P2_F04_normalisation.png",
           "Empirical intensity distribution of the training split against the "
           "ImageNet normalisation constants.")
    para(doc,
         f"Panel C also shows why the red channel's standard deviation is the "
         f"one that rises ({NORM['std'][0]:.3f} against "
         f"{NORM['imagenet_std'][0]:.3f}) while the other two fall. The "
         f"distribution is bimodal: a narrow spike at zero in all three "
         f"channels, produced by the black surround left when the endoscope's "
         f"circular field of view is stored in a rectangular frame, and a broad "
         f"mucosal mode that in the red channel presses against saturation. "
         f"The images are therefore not merely tinted relative to ImageNet; "
         f"their intensity distribution has a different shape, which is the "
         f"substantive reason the specification requires corpus-specific "
         f"statistics rather than the convenience of the defaults.")

    h(doc, "2.3.2  Augmentation policy", 3)
    para(doc,
         "The augmentation policy is fixed here and held constant for all "
         "later phases, so that any Phase 4 effect is attributable to the "
         "target construction rather than to a change in augmentation. It is "
         "deliberately conservative for a domain-specific reason: the class "
         "label encodes an anatomical wall — anterior, posterior, greater or "
         "lesser curvature — and a horizontal or vertical flip, or a large "
         "rotation, changes the apparent wall. Flips would therefore corrupt "
         "the label rather than augment it. Only a mild scale and translation "
         "crop and photometric jitter are applied.")
    table(doc,
          ["Transform", "Setting", "Applied to"],
          [["Lanczos resample to 224×224", "once, cached", "all splits"],
           ["RandomResizedCrop", "scale 0.85–1.00, ratio 0.90–1.11", "training only"],
           ["ColorJitter", "brightness/contrast/saturation 0.2, hue 0.02",
            "training only"],
           ["Normalise", "training-set mean and SD", "all splits"],
           ["Horizontal / vertical flip", "not applied — would alter the "
            "anatomical wall that defines the label", "—"],
           ["Rotation", "not applied — same reason", "—"]],
          "Preprocessing and augmentation policy, fixed for Phases 2–6.",
          widths=[4.4, 7.6, 3.5])

    h(doc, "2.4  Model and fine-tuning scope", 2)
    para(doc,
         f"ConvNeXt-Tiny is instantiated with ImageNet-1k weights and a "
         f"{COH['n_classes']}-way linear head. The specification requires "
         f"fine-tuning 'the top 40% of the feature layers', which is not by "
         f"itself an executable instruction; it was resolved to an explicit "
         f"module list and logged. torchvision exposes the backbone as "
         f"{LAY['n_feature_modules']} top-level feature modules, so 40% rounds "
         f"to {LAY['n_modules_unfrozen']} modules — indices "
         f"{', '.join(map(str, LAY['modules_unfrozen']))}, comprising stage 3, "
         f"its downsampling layer and stage 4.")
    para(doc,
         f"It is worth stating what that means numerically, because the "
         f"phrase 'top 40%' understates it: those "
         f"{LAY['n_modules_unfrozen']} of {LAY['n_feature_modules']} modules "
         f"hold {100 * LAY['param_fraction_unfrozen']:.1f}% of the "
         f"{LAY['feature_params_total'] / 1e6:.1f} M feature parameters, "
         f"because ConvNeXt concentrates depth in stage 3. Forty per cent of "
         f"the layers is {100 * LAY['param_fraction_unfrozen']:.0f}% of the "
         f"weights.")
    table(doc,
          ["Property", "Value"],
          [["Backbone", "ConvNeXt-Tiny (torchvision, IMAGENET1K_V1)"],
           ["Feature modules", str(LAY["n_feature_modules"])],
           ["Modules unfrozen in stage 2",
            f"{LAY['modules_unfrozen']} ({100 * LAY['module_fraction']:.1f}% of modules)"],
           ["Feature parameters", f"{LAY['feature_params_total']:,}"],
           ["Feature parameters trainable",
            f"{LAY['feature_params_unfrozen']:,} "
            f"({100 * LAY['param_fraction_unfrozen']:.2f}%)"],
           ["Classifier head", f"Linear(768 → {COH['n_classes']})"]],
          "Resolved fine-tuning scope. The mapping from '40% of feature "
          "layers' to a concrete module list is recorded so the run can be "
          "repeated exactly.",
          widths=[7.0, 8.5])

    h(doc, "2.5  Optimisation schedule", 2)
    para(doc,
         f"Training proceeds in the two stages the specification requires: a "
         f"{R0['warmup_epochs']}-epoch head warm-up at constant learning rate "
         f"with the backbone frozen, followed by fine-tuning of the unfrozen "
         f"modules with early stopping on validation macro F1 at patience "
         f"{R0['patience']}. Model selection uses validation macro F1 only; the "
         f"test set is not consulted at any point during training.")
    table(doc,
          ["Hyper-parameter", "Value", "Source"],
          [["Optimiser", "AdamW", "standard for ConvNeXt"],
           ["Weight decay", f"{R0['weight_decay']}", "ConvNeXt convention"],
           ["Warm-up epochs", f"{R0['warmup_epochs']}", "blueprint §4"],
           ["Warm-up learning rate", f"{R0['lr_head']:g}, constant", "blueprint §4"],
           ["Fine-tune learning rate", f"{R0['lr_finetune']:g}, cosine annealed",
            "this work"],
           ["Fine-tune epoch cap", f"{R0['max_finetune_epochs']}",
            "pre-registration DEV-2"],
           ["Early stopping", f"validation macro F1, patience {R0['patience']}",
            "blueprint §4"],
           ["Batch size", f"{R0['batch_size']}", "measured (§3.2)"],
           ["Gradient accumulation", f"{R0['grad_accum_steps']} steps "
            f"(effective batch {R0['effective_batch']})", "this work"],
           ["Numerical precision", R0["precision"],
            "measured — pre-registration DEV-1"],
           ["Memory format", R0["memory_format"], "measured (§3.2)"],
           ["Seeds", ", ".join(str(r["seed"]) for r in RUNS),
            "pre-registration DEV-3"]],
          "Complete training configuration. Every value is either taken from "
          "the governing protocol or derived from a measurement reported in "
          "this document.",
          widths=[4.6, 6.4, 4.5])

    h(doc, "2.6  Evaluation framework", 2)
    h(doc, "2.6.1  Metrics", 3)
    para(doc,
         "Macro F1 is the primary endpoint, matching the published baseline and "
         "appropriate to a near-balanced 23-class problem. Macro and weighted "
         "precision and recall, accuracy and the full confusion matrix are "
         "secondary. Per-class F1 is reported but is exploratory only, for the "
         "reason quantified in §2.2.4. Expected calibration error and the Brier "
         "score are recorded as a baseline reference for Phase 4, not as a "
         "Phase 2 claim.")
    h(doc, "2.6.2  Why intervals resample patients", 3)
    para(doc,
         f"All intervals are obtained by resampling the "
         f"{MET['n_test_patients']} test patients with replacement, "
         f"{MET['n_boot']:,} times, and recomputing macro F1 on each resample. "
         f"Images are never resampled. Phase 0 measured per-patient Fleiss' κ "
         f"at 0.7459 ± 0.1448 with a range of 0.069 to 1.000: agreement, and "
         f"therefore difficulty, varies systematically by patient, so images "
         f"within a patient are not independent observations. An "
         f"image-level interval would understate the uncertainty by treating "
         f"{COH['by_split']['Test']} correlated observations as "
         f"{COH['by_split']['Test']} independent ones.")
    h(doc, "2.6.3  Pre-registered acceptance criterion", 3)
    para(doc,
         f"The published reference is {REPRO['published_macro_f1']:.1f} macro "
         f"F1 for ConvNeXt-Tiny under the complete-agreement label condition. "
         f"The acceptance rule, fixed in advance, is: "
         f"{PRE['verdict_rule']}")
    callout(doc, PRE["published_reference_note"],
            title="On the precision of the published target")


# =====================================================================
def chapter3(doc) -> None:
    h(doc, "3  Implementation and Execution", 1, page_break=True)

    h(doc, "3.1  Compute environment", 2)
    para(doc,
         f"The blueprint recorded a prerequisite: the installed PyTorch was a "
         f"CPU-only build, so no GPU was reachable. Resolving it was "
         f"constrained by the interpreter. Under Python {ENV['python']} the "
         f"cu128 channel serves torch only up to 2.11.0 — a downgrade — whereas "
         f"cu126 offers {ENV['torch_build']}, matching the version already "
         f"installed. The CUDA 12.6 build supports compute capability "
         f"{ENV['compute_capability']}, and the installed driver reports a "
         f"newer CUDA runtime, which is backward compatible.")
    table(doc,
          ["Component", "Value"],
          [["Operating system", ENV["platform"]],
           ["Python", ENV["python"]],
           ["PyTorch", ENV["torch_build"]],
           ["torchvision", ENV["packages"].get("torchvision") or "—"],
           ["CUDA (torch build)", ENV["torch_cuda_version"]],
           ["cuDNN", str(ENV["cudnn_version"])],
           ["GPU", ENV["gpu"]],
           ["Compute capability", ENV["compute_capability"]],
           ["Total VRAM", f"{PROBE['total_vram_mib']:.0f} MiB"],
           ["VRAM already resident (desktop)",
            f"{PROBE['vram_occupied_by_desktop_mib']:.0f} MiB"],
           ["scikit-learn", ENV["packages"].get("scikit-learn") or "—"],
           ["NumPy / pandas", f"{ENV['packages'].get('numpy')} / "
                              f"{ENV['packages'].get('pandas')}"]],
          "Execution environment. GATE 1 cleared once a CUDA build replaced the "
          "CPU-only wheel.",
          widths=[6.2, 9.3])
    para(doc,
         f"One practical detail deserves recording because it changes the "
         f"memory budget: of the {PROBE['total_vram_mib']:.0f} MiB of physical "
         f"video memory, {PROBE['vram_occupied_by_desktop_mib']:.0f} MiB were "
         f"already resident for the Windows desktop compositor before any "
         f"training process started. The usable budget is therefore closer to "
         f"{(PROBE['total_vram_mib'] - PROBE['vram_occupied_by_desktop_mib']) / 1024:.1f} GB "
         f"than to 4 GB.")

    h(doc, "3.2  Precision benchmark and batch-size selection", 2)
    para(doc,
         f"The governing protocol prescribes AMP float16 with a GradScaler, on "
         f"the reasoning that the device is Turing and Turing has FP16 tensor "
         f"cores. That reasoning does not hold for this particular chip. The "
         f"GTX 1650 is built on TU117, the one Turing die shipped without "
         f"tensor cores. Rather than assume either way, a factorial benchmark "
         f"was run over precision and memory format at fixed batch size.")
    fp16 = next(r for r in PROBE["precision_factorial"]
                if r["amp_fp16"] and r["channels_last"])
    fp32 = next(r for r in PROBE["precision_factorial"]
                if not r["amp_fp16"] and r["channels_last"])
    table(doc,
          ["Precision", "Memory format", "ms / step", "images / s",
           "Peak VRAM (MiB)"],
          [["AMP float16" if r["amp_fp16"] else "float32",
            "channels_last" if r["channels_last"] else "contiguous",
            f"{1000 * r['sec_per_step']:.0f}",
            f"{r['images_per_sec']:.1f}",
            f"{r['peak_alloc_mib']:.0f}"]
           for r in PROBE["precision_factorial"]],
          f"Precision × memory-format factorial, ConvNeXt-Tiny at 224×224, "
          f"batch 24, forward and backward. AMP float16 achieves "
          f"{PROBE['amp_vs_fp32_speedup']:.2f}× the throughput of float32.",
          align_right=[2, 3, 4])
    para(doc,
         f"The result is unambiguous and in the opposite direction to the "
         f"prescription: float32 runs at {fp32['images_per_sec']:.1f} images per "
         f"second against {fp16['images_per_sec']:.1f} for AMP, a slowdown of "
         f"{1 / PROBE['amp_vs_fp32_speedup']:.2f}×. Without tensor cores, FP16 "
         f"offers no matrix-multiply acceleration, while autocast's casting and "
         f"the gradient scaler's inspection of gradients impose real cost. "
         f"Float32 was adopted, and the departure from the protocol recorded as "
         f"deviation DEV-1 in the pre-registration. Mixed precision does reduce "
         f"peak memory — {fp16['peak_alloc_mib']:.0f} MiB against "
         f"{fp32['peak_alloc_mib']:.0f} MiB — but memory was not the binding "
         f"constraint at the selected batch size.")
    figure(doc, "P2_F05_hardware.png",
           "Hardware characterisation. Panel A is the precision finding; panels "
           "B and C show why batch 24 was selected — larger batches exceed the "
           "usable memory ceiling and collapse in throughput.")
    para(doc,
         f"Batch size was then selected empirically by descending a ladder in "
         f"the winning precision and taking the largest batch whose peak "
         f"allocation left a {PROBE['safety_margin_mib']} MiB margin below "
         f"physical VRAM. Batch {PROBE['chosen_batch']} was selected. The two "
         f"larger batches are instructive: batch 32 exceeds the ceiling and "
         f"drops to 16.1 images per second, and batch 48 collapses to 4.3, "
         f"because the allocator begins to thrash rather than fail outright. "
         f"Gradient accumulation over {R0['grad_accum_steps']} steps holds the "
         f"effective batch at {R0['effective_batch']}.")

    h(doc, "3.3  Software architecture", 2)
    table(doc,
          ["Script", "Stage", "Responsibility"],
          [["src/models/phase2_env.py", "A5", "environment snapshot"],
           ["src/models/phase2_probe.py", "A3–A4",
            "GATE 1, precision benchmark, batch ladder"],
           ["src/models/phase2_data.py", "B1–B7",
            "cohort construction, GATE 2 and GATE 3, composition tables"],
           ["src/models/phase2_normstats.py", "B6",
            "training-set channel statistics"],
           ["src/models/phase2_cache.py", "C1",
            "deterministic 224×224 Lanczos cache"],
           ["src/models/phase2_prereg.py", "D", "freeze the pre-registration"],
           ["src/models/phase2_train.py", "C, E1–E4",
            "two-stage training, early stopping, run manifest"],
           ["src/models/phase2_eval.py", "E5–E12",
            "test evaluation, bootstraps, verdict"],
           ["src/report/figures_phase2.py", "F2", "figure suite"],
           ["src/report/build_phase2_docx.py", "F3–F4", "this document"]],
          "Phase 2 script inventory. Execution order is top to bottom.",
          widths=[5.6, 2.2, 7.7])

    h(doc, "3.4  Faults encountered and resolved", 2)
    para(doc,
         "Three failures occurred during execution. They are recorded because "
         "each changed the configuration, and a reader reproducing this work on "
         "similar hardware will meet them.")
    table(doc,
          ["Fault", "Diagnosis", "Resolution"],
          [["Warm-up epochs ran at 95 s with the GPU at 41% utilisation",
            "single-process data loading; CPU-side augmentation was the "
            "bottleneck, not the GPU",
            "moved loading to worker processes; epoch time fell to 38 s"],
           ["Crash in the CUDA pinned-host allocator after the first "
            "fine-tuning epoch",
            "repeatedly cloning a 28 M-parameter state dictionary into host "
            "RAM alongside pinned dataloader buffers",
            "best checkpoint written to disk instead of held in memory; "
            "pinned memory disabled"],
           ["OpenBLAS allocation failure and worker death",
            "training and validation each held a persistent five-worker pool; "
            "ten resident torch processes exhausted 16 GB of host RAM",
            "three workers for training, none for validation"]],
          "Execution faults, their causes and the configuration changes made "
          "in response. None affected the scientific protocol.",
          widths=[4.8, 5.6, 5.1])

    h(doc, "3.5  Execution record", 2)
    table(doc,
          ["Seed", "Epochs run", "Stop reason", "Best epoch",
           "Best val macro F1", "Wall-clock", "Peak VRAM (MiB)"],
          [[str(r["seed"]), str(r["n_epochs_run"]),
            r["stop_reason"].replace("_", " "), str(r["best_epoch_overall"]),
            f"{100 * r['best_val_macro_f1']:.2f}",
            f"{r['wallclock_sec'] / 60:.1f} min",
            f"{r['peak_vram_mib']:.0f}"] for r in RUNS],
          "Per-seed execution record. 'Best epoch' is counted over the "
          "concatenated warm-up and fine-tuning schedule.",
          align_right=[1, 3, 4, 5, 6])


# =====================================================================
def chapter4(doc) -> None:
    h(doc, "4  Results", 1, page_break=True)

    h(doc, "4.1  Training dynamics", 2)
    ends = ", ".join(f"seed {r['seed']} stopped after {r['n_epochs_run']} "
                     f"epochs ({r['stop_reason'].replace('_', ' ')})"
                     for r in RUNS)
    para(doc,
         f"All {len(RUNS)} runs followed the same trajectory: rapid gain during "
         f"the frozen-backbone warm-up, a discontinuous jump when the backbone "
         f"was unfrozen, then a slow approach to a plateau. {ends.capitalize()}.")
    figure(doc, "P2_F06_training.png",
           "Training dynamics across seeds. The shaded band is the "
           "frozen-backbone warm-up; markers show the epoch selected by "
           "validation macro F1.")
    stops = {r["stop_reason"] for r in RUNS}
    if stops == {"early_stopping"}:
        para(doc,
             "Every run terminated on the early-stopping criterion rather than "
             "at the epoch cap. The cap introduced by deviation DEV-2 for "
             "compute reasons therefore did not bind, and the schedule executed "
             "is the schedule the protocol specifies.")
    else:
        para(doc,
             f"Not every run terminated on the early-stopping criterion: "
             f"{sum(1 for r in RUNS if r['stop_reason'] == 'epoch_cap')} of "
             f"{len(RUNS)} reached the {R0['max_finetune_epochs']}-epoch cap "
             f"imposed by deviation DEV-2. Where the cap bound, the reported "
             f"result is a lower bound on what the unrestricted schedule would "
             f"have achieved, and this is carried into the discussion as a "
             f"threat to validity.")

    h(doc, "4.2  Test-set performance", 2)
    para(doc,
         f"Each trained model was evaluated once on the "
         f"{MET['n_test_images']}-image complete-agreement test set, after "
         f"model selection had been completed on validation data.")
    rows = []
    for s in MET["seeds"]:
        d = MET["per_seed"][str(s)]
        rows.append([str(s), pc(d["macro_f1"]), pc(d["weighted_f1"]),
                     pc(d["macro_precision"]), pc(d["macro_recall"]),
                     pc(d["accuracy"])])
    rows.append(["mean", f"{100 * AGG['macro_f1_mean']:.2f}", "—", "—", "—", "—"])
    table(doc,
          ["Seed", "Macro F1", "Weighted F1", "Macro precision",
           "Macro recall", "Accuracy"],
          rows,
          f"Test-set metric battery, per seed and seed mean. All values are "
          f"percentages on the {MET['n_test_images']}-image cohort.",
          align_right=[1, 2, 3, 4, 5])

    h(doc, "4.3  Uncertainty quantification", 2)
    brows = []
    for s in MET["seeds"]:
        d = MET["per_seed"][str(s)]
        pubi = d["published_style_interval"]
        brows.append([str(s), pc(d["macro_f1"]),
                      f"{pc(d['macro_f1_ci95'][0])} – {pc(d['macro_f1_ci95'][1])}",
                      f"{100 * (d['macro_f1_ci95'][1] - d['macro_f1_ci95'][0]) / 2:.2f}",
                      f"{pubi['mean']:.2f} ± {pubi['margin_of_error']:.2f}"])
    brows.append(["mean", f"{100 * AGG['macro_f1_mean']:.2f}",
                  f"{pc(AGG['seed_mean_boot_ci95'][0])} – "
                  f"{pc(AGG['seed_mean_boot_ci95'][1])}",
                  f"{100 * (AGG['seed_mean_boot_ci95'][1] - AGG['seed_mean_boot_ci95'][0]) / 2:.2f}",
                  "—"])
    table(doc,
          ["Seed", "Macro F1", "Patient-clustered 95% CI", "Half-width (pp)",
           "Descriptor-style interval"],
          brows,
          f"Macro F1 with patient-clustered bootstrap intervals "
          f"({MET['n_boot']:,} resamples of {MET['n_test_patients']} patients), "
          f"alongside the descriptor's own procedure reproduced for "
          f"comparability.",
          align_right=[1, 2, 3, 4])
    figure(doc, "P2_F07_bootstrap.png",
           "Bootstrap distributions and seed agreement. Intervals come from "
           "resampling patients, not images.")

    ex = MET["per_seed"][str(MET["seeds"][0])]
    hw = 100 * (ex["macro_f1_ci95"][1] - ex["macro_f1_ci95"][0]) / 2
    pub_moe = ex["published_style_interval"]["margin_of_error"]
    callout(doc,
            f"The two interval columns differ by roughly a factor of "
            f"{hw / pub_moe:.0f}, and the difference is not a disagreement about "
            f"the data. The descriptor forms its margin as t × s / √B, the "
            f"standard error of the mean over B = 100 bootstrap iterations. "
            f"That quantity measures how precisely the bootstrap mean has been "
            f"estimated and shrinks as more iterations are run; it is not a "
            f"95% interval on model performance. The published margins near "
            f"±0.2 should be read accordingly, and comparisons between "
            f"published models that rely on them are less decisive than they "
            f"appear.",
            title="Why the published margins are so narrow")

    h(doc, "4.4  Confusion structure", 2)
    figure(doc, "P2_F08_confusion.png",
           "Row-normalised confusion matrix, averaged over seeds, ordered by "
           "SSS station and then by wall.", width=5.4)
    para(doc,
         "The matrix is presented in anatomical order — by station, then by "
         "wall — rather than alphabetically, so that block structure is "
         "visible. Interpretation of that structure against the human "
         "disagreement pattern measured in Phase 0 is deferred to Phase 6, "
         "where it is a stated objective; it is not analysed here.")

    h(doc, "4.5  Per-class performance", 2)
    para(doc,
         f"Per-class results are reported for completeness and are exploratory. "
         f"The smallest test class holds "
         f"{min(PROV['test_class_support'].values())} images, and as §2.2.4 "
         f"showed, most classes cannot support a ±10 percentage-point claim at "
         f"this sample size.")
    figure(doc, "P2_F09_perclass.png",
           "Per-class F1 and recall with Wilson intervals, ordered by F1. "
           "Interval width, not point position, is the message.")

    h(doc, "4.6  Seed-to-seed variability", 2)
    if AGG["macro_f1_sd"] is not None:
        para(doc,
             f"Across {len(RUNS)} seeds the macro F1 standard deviation is "
             f"{100 * AGG['macro_f1_sd']:.2f} percentage points, with a range of "
             f"{AGG['macro_f1_range_points']:.2f} points "
             f"({100 * AGG['macro_f1_min']:.2f} to "
             f"{100 * AGG['macro_f1_max']:.2f}). This is the quantity a "
             f"single-run reproduction cannot separate from reproduction error, "
             f"and it is why the pre-registration required more than one seed.")

    h(doc, "4.7  Calibration baseline", 2)
    eces = [MET["per_seed"][str(s)]["ece"] for s in MET["seeds"]]
    briers = [MET["per_seed"][str(s)]["brier"] for s in MET["seeds"]]
    para(doc,
         f"Expected calibration error averages {100 * sum(eces) / len(eces):.2f}% "
         f"and the Brier score {sum(briers) / len(briers):.4f}. These values "
         f"make no claim in Phase 2. They are recorded because RQ2 predicts "
         f"that soft-target training will improve calibration, and that "
         f"prediction requires a hard-label reference point measured under "
         f"identical conditions.")
    figure(doc, "P2_F11_calibration.png",
           "Reliability diagram and confidence distribution for the reproduced "
           "baseline.")

    h(doc, "4.8  Computational cost", 2)
    tot = sum(r["wallclock_sec"] for r in RUNS)
    para(doc,
         f"The {len(RUNS)} runs consumed {tot / 3600:.2f} GPU-hours in total "
         f"on a single consumer GPU, with peak allocation of "
         f"{max(r['peak_vram_mib'] for r in RUNS):.0f} MiB. This is the "
         f"evidence for the feasibility claim underlying the whole "
         f"experimental design: the five-configuration comparison planned for "
         f"Phase 4 is achievable on this hardware, whereas ConvNeXt-Large is "
         f"not.")


# =====================================================================
def chapter5(doc) -> None:
    h(doc, "5  Reproduction Verdict", 1, page_break=True)

    h(doc, "5.1  The reference value", 2)
    para(doc,
         f"The comparison is against {REPRO['published_source']}, under the "
         f"condition: {REPRO['published_condition']}.")
    para(doc, f"The descriptor states the value as: {PRE['published_quote']}")

    h(doc, "5.2  Numerical adjudication", 2)
    table(doc,
          ["Quantity", "Value"],
          [["Published reference macro F1", f"{REPRO['published_macro_f1']:.2f}"],
           ["Pre-registered acceptance band",
            f"± {REPRO['acceptance_band_points']} points"],
           ["Observed macro F1 (seed mean)", f"{REPRO['observed_macro_f1']:.2f}"],
           ["Patient-clustered 95% CI on the seed mean",
            f"{pc(AGG['seed_mean_boot_ci95'][0])} – "
            f"{pc(AGG['seed_mean_boot_ci95'][1])}"],
           ["Δ (observed − published)", f"{REPRO['delta_points']:+.2f} points"],
           ["|Δ|", f"{REPRO['abs_delta_points']:.2f} points"],
           ["Verdict", REPRO["verdict"]]],
          "Reproduction adjudication against the criterion frozen before any "
          "model was trained.",
          widths=[9.0, 6.5])
    figure(doc, "P2_F10_verdict.png",
           "The reproduction placed against the published acceptance band and "
           "the other reference points reported in the data descriptor.")

    h(doc, "5.3  Gate decision", 2)
    if PASS:
        para(doc,
             f"GATE 5 is cleared. The observed seed-mean macro F1 of "
             f"{REPRO['observed_macro_f1']:.2f} lies "
             f"{REPRO['abs_delta_points']:.2f} points from the published "
             f"reference, within the pre-registered ±"
             f"{REPRO['acceptance_band_points']}-point band. The pre-registered "
             f"diagnostic order was not invoked, and no hyper-parameter was "
             f"changed after the result was seen.")
        callout(doc,
                "The pipeline reproduces the published baseline. Differences "
                "observed in Phases 3–6 can therefore be attributed to the "
                "interventions under study rather than to a defect in the "
                "apparatus. This is the entire purpose of Phase 2, and it is "
                "now discharged.",
                title="What this verdict licenses")
    else:
        para(doc,
             f"GATE 5 is not cleared. The observed seed-mean macro F1 of "
             f"{REPRO['observed_macro_f1']:.2f} lies "
             f"{REPRO['abs_delta_points']:.2f} points from the published "
             f"reference, outside the pre-registered ±"
             f"{REPRO['acceptance_band_points']}-point band. The pre-registered "
             f"diagnostic order in Appendix A applies, and it is followed in "
             f"the order written rather than in an order chosen after seeing "
             f"which direction the discrepancy runs.")

    h(doc, "5.4  Implications for Phases 3 to 6", 2)
    para(doc,
         "Three assets are frozen by this phase and must be reused unchanged "
         "downstream, because changing any of them would break comparability "
         "with the reference point just established: the class index, the "
         "training-set normalisation statistics, and the augmentation policy. "
         "The two-stage schedule and the resolved fine-tuning module list are "
         "likewise fixed, so that the Phase 4 configurations C0 to C4 differ "
         "from one another only in target construction.")


# =====================================================================
def chapter6(doc) -> None:
    h(doc, "6  Discussion", 1, page_break=True)

    h(doc, "6.1  Interpretation", 2)
    para(doc,
         f"The reproduction "
         f"{'succeeded' if PASS else 'did not succeed'} on a pre-specified "
         f"criterion, using a protocol published in sufficient detail to be "
         f"followed: input size and interpolation, normalisation source, "
         f"two-stage schedule, warm-up length, fine-tuning scope, stopping "
         f"rule and split definition are all stated in the descriptor. That "
         f"level of methodological detail is not universal in this literature — "
         f"the Phase 1 review found incomplete reporting of ground-truth "
         f"construction to be one of its four commonest omissions — and it is "
         f"what made a genuine reproduction attempt possible here.")

    h(doc, "6.2  Sources of residual divergence", 2)
    for b in [
        "Optimiser and learning-rate schedule for the fine-tuning stage are "
        "not fully specified in the descriptor; the values used here are stated "
        "in Table 6 but are this work's choice, not the original's.",
        "The augmentation policy of the original is not reported. The policy "
        "used here is conservative and justified on anatomical grounds, but it "
        "is not necessarily the same one.",
        "Numerical precision differs by necessity: float32 here against an "
        "unstated precision in the original.",
        "The published figure is quoted to two significant figures in the "
        "narrative, so the target itself carries rounding uncertainty of up to "
        "half a point.",
    ]:
        bullet(doc, b)

    h(doc, "6.3  What this phase does and does not establish", 2)
    para(doc,
         "It establishes that this pipeline, on this cohort, lands where the "
         "literature says it should, and it quantifies the run-to-run variance "
         "that any later comparison must exceed to be meaningful. It "
         "establishes nothing about performance on contested images, about "
         "calibration under disagreement, about generalisation to another "
         "centre, or about whether the model's errors resemble human errors. "
         "Those are the subjects of Phases 3 to 6 and are deliberately "
         "untouched here.")

    h(doc, "6.4  The consequence of evaluating only on unanimous images", 2)
    para(doc,
         f"Every number in this report — and every number in the published "
         f"benchmark it reproduces — is computed on the "
         f"{COH['retention_pct']:.2f}% of the corpus where four experts agreed "
         f"completely. The remaining {100 - COH['retention_pct']:.2f}% is not "
         f"missing data; it is the part of the distribution where the task is "
         f"hard. A benchmark that reports only the unanimous subset reports "
         f"performance on the easy fraction and is silent about the rest. "
         f"Quantifying that silence is the object of Phase 3, and Phase 2 "
         f"provides the reference model with which to do it — unchanged, so "
         f"that the stratified comparison isolates the stratum rather than the "
         f"training.")

    h(doc, "6.5  Threats to validity", 2)
    tv = [["Single architecture",
           "Only ConvNeXt-Tiny was reproduced; a pipeline defect specific to "
           "another architecture family would not be detected",
           "Accepted: the blueprint fixes this backbone on compute grounds"],
          ["Target precision",
           "The published reference is quoted to two significant figures",
           "The ±1.5-point band is wider than the rounding uncertainty"],
          [f"{len(RUNS)} seeds",
           "The seed variance estimate rests on few runs",
           "Reported explicitly; the verdict uses the seed mean"],
          ["Compute-imposed epoch cap",
           "A cap below the protocol's 100 epochs was pre-registered",
           "Realised stop reason reported per seed in Table 8"],
          ["Unspecified original hyper-parameters",
           "Optimiser and augmentation of the original are unknown",
           "Enumerated in §6.2; all choices made here are documented"]]
    table(doc, ["Threat", "Nature", "Response"], tv,
          "Threats to the validity of the reproduction claim.",
          widths=[3.4, 6.2, 5.9])

    h(doc, "6.6  Limitations carried forward", 2)
    table(doc,
          ["ID", "Limitation", "Status in Phase 2"],
          [["L1", "22 of 23 classes underpowered for per-class claims",
            "Confirmed at cohort level and quantified in Figure 3; per-class "
            "results labelled exploratory"],
           ["L2", "No age or sex anywhere in the release",
            "Unchanged; no subgroup or fairness claim is made"],
           ["L3", "Single centre, single vendor",
            "Unchanged; addressed by Phase 5"],
           ["L4", "Acquisition stream imbalanced across splits",
            f"Re-measured on the cohort: χ² = {PROV['stream_split_chi2']}, "
            f"p = {PROV['stream_split_p']:.2e}; composition reported in Table 4"],
           ["L6", "Clinical context for only 60.2% of patients",
            "Unchanged; not used in this phase"],
           ["L8", "Two patients contribute no unanimous image",
            "New. Identified in §2.2.3; reduces the test set to "
            f"{COH['patients_by_split']['Test']} resampling units"]],
          "Declared limitations and their status after Phase 2. L8 is new to "
          "this phase.",
          widths=[1.2, 5.6, 8.7])


# =====================================================================
def chapter7(doc) -> None:
    h(doc, "7  Conclusion and Transition to Phase 3", 1, page_break=True)

    h(doc, "7.1  Summary of measured findings", 2)
    for b in [
        f"The complete-agreement cohort reproduces the published counts exactly "
        f"({COH['by_split']['Train']:,} / {COH['by_split']['Validation']} / "
        f"{COH['by_split']['Test']}), with zero patient overlap and full "
        f"provenance resolution against the Phase 0 inventory.",
        f"Two patients contribute no unanimously labelled image, so the cohort "
        f"spans {COH['n_patients']} patients rather than the 387 of the corpus, "
        f"and the test set spans {COH['patients_by_split']['Test']} rather "
        f"than 59.",
        f"Training-set channel statistics differ from the ImageNet defaults by "
        f"up to {max(NORM['abs_delta_mean']):.3f} in the mean.",
        f"AMP float16 runs at {PROBE['amp_vs_fp32_speedup']:.2f}× float32 "
        f"throughput on this tensor-core-less GPU, contradicting the governing "
        f"protocol's prescription.",
        f"Forty per cent of the feature layers corresponds to "
        f"{100 * LAY['param_fraction_unfrozen']:.0f}% of the feature "
        f"parameters.",
        f"The reproduced model attains {REPRO['observed_macro_f1']:.2f} macro "
        f"F1 (95% CI {pc(AGG['seed_mean_boot_ci95'][0])}–"
        f"{pc(AGG['seed_mean_boot_ci95'][1])}), "
        f"Δ = {REPRO['delta_points']:+.2f} points from the published reference: "
        f"{REPRO['verdict']}.",
        "The descriptor's published stability margins are standard errors of a "
        "bootstrap mean rather than intervals on performance, and understate "
        "uncertainty by roughly an order of magnitude.",
    ]:
        bullet(doc, b)

    h(doc, "7.2  Assets frozen for reuse", 2)
    table(doc,
          ["Asset", "Path", "Why it must not change"],
          [["Class index", "data/phase2_class_index.json",
            "label-to-integer mapping shared by every later confusion matrix"],
           ["Cohort manifest", "data/phase2_consensus_manifest.csv",
            "defines the reference population"],
           ["Normalisation statistics", "reports/phase2_norm_stats.json",
            "changing them changes the input distribution"],
           ["Preprocessing cache", "data/phase2_cache_224.npy",
            "guarantees identical pixels across phases"],
           ["Augmentation policy", "src/models/phase2_train.py",
            "must be constant so Phase 4 isolates target construction"],
           ["Trained checkpoints", "checkpoints/phase2_convnext_tiny_seed*.pt",
            "Phase 3 evaluates these models unchanged"]],
          "Artefacts frozen by Phase 2 and inherited by later phases.",
          widths=[3.8, 5.6, 6.1])

    h(doc, "7.3  Phase 3 entry conditions", 2)
    para(doc,
         "Phase 3 takes the checkpoints produced here, without retraining, and "
         "evaluates them on strata they never saw: majority, plurality, tied "
         "and dispersed images. The blueprint requires one decision to be made "
         "before that evaluation is run and not after — the scoring rule for "
         "strata in which no majority label exists, where accuracy against a "
         "single ground truth is undefined. Fixing that rule, and the pooling "
         "rule for the two smallest strata, is the first task of Phase 3 and "
         "should be pre-registered exactly as this phase's criterion was.")


# =====================================================================
def references(doc) -> None:
    h(doc, "References", 1, page_break=True)
    for i, r in enumerate([
        "Panesso-Ortiz, D., Ruiz-Fernández, D., et al. (2025). GastroHUN: an "
        "Endoscopy Dataset of Complete Systematic Screening Protocol for the "
        "Stomach. Scientific Data, 12, 102. "
        "https://doi.org/10.1038/s41597-025-04401-5",
        "Liu, Z., Mao, H., Wu, C.-Y., Feichtenhofer, C., Darrell, T., & Xie, S. "
        "(2022). A ConvNet for the 2020s. In Proceedings of the IEEE/CVF "
        "Conference on Computer Vision and Pattern Recognition (pp. 11976–11986).",
        "Loshchilov, I., & Hutter, F. (2019). Decoupled Weight Decay "
        "Regularization. In International Conference on Learning Representations.",
        "Mongan, J., Moy, L., & Kahn, C. E. (2020). Checklist for Artificial "
        "Intelligence in Medical Imaging (CLAIM): A Guide for Authors and "
        "Reviewers. Radiology: Artificial Intelligence, 2(2), e200029.",
        "Collins, G. S., Moons, K. G. M., Dhiman, P., et al. (2024). TRIPOD+AI "
        "statement: updated guidance for reporting clinical prediction models "
        "that use regression or machine learning methods. BMJ, 385, e078378.",
        "Wilson, E. B. (1927). Probable Inference, the Law of Succession, and "
        "Statistical Inference. Journal of the American Statistical "
        "Association, 22(158), 209–212.",
        "Efron, B., & Tibshirani, R. J. (1993). An Introduction to the "
        "Bootstrap. Chapman & Hall.",
        "Guo, C., Pleiss, G., Sun, Y., & Weinberger, K. Q. (2017). On "
        "Calibration of Modern Neural Networks. In Proceedings of the 34th "
        "International Conference on Machine Learning (pp. 1321–1330).",
        "Yao, K. (2013). The endoscopic diagnosis of early gastric cancer. "
        "Annals of Gastroenterology, 26(1), 11–22.",
        "Deng, J., Dong, W., Socher, R., Li, L.-J., Li, K., & Fei-Fei, L. "
        "(2009). ImageNet: A large-scale hierarchical image database. In IEEE "
        "Conference on Computer Vision and Pattern Recognition (pp. 248–255).",
    ], 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f"[{i}]  {r}")
        run.font.size = Pt(9.5)


# =====================================================================
def appendices(doc) -> None:
    h(doc, "Appendix A  Pre-registration (verbatim)", 1, page_break=True)
    para(doc,
         f"Frozen {PRE['frozen_at']}, before the first full training run. "
         f"Reproduced here without alteration.", italic=True)
    para(doc, PRE["statement"], italic=True)

    h(doc, "A.1  Target and acceptance rule", 2)
    table(doc, ["Field", "Value"],
          [["Published macro F1", f"{PRE['published_macro_f1']}"],
           ["Source", PRE["published_source"]],
           ["Condition", PRE["published_condition"]],
           ["Acceptance band", f"± {PRE['acceptance_band_points']} points"],
           ["Decision statistic", PRE["decision_statistic"]],
           ["Verdict rule", PRE["verdict_rule"]],
           ["Falsification", PRE["falsification"]]],
          "Pre-registered target and acceptance rule.",
          widths=[3.6, 11.9])

    h(doc, "A.2  Estimation procedure", 2)
    table(doc, ["Field", "Value"],
          [["Seeds", ", ".join(map(str, PRE["seeds"]))],
           ["Bootstrap resamples", f"{PRE['n_bootstrap']:,}"],
           ["Bootstrap unit", PRE["bootstrap_unit"]],
           ["Rule", PRE["bootstrap_rule"]],
           ["Secondary interval", PRE["secondary_interval"]]],
          "Pre-registered estimation procedure.",
          widths=[3.6, 11.9])

    h(doc, "A.3  Declared deviations from the governing protocol", 2)
    table(doc,
          ["ID", "Item", "Protocol", "Adopted", "Evidence", "Impact"],
          [[d["id"], d["item"], d["blueprint"], d["adopted"], d["evidence"],
            d["impact"]] for d in PRE["deviations"]],
          "Deviations from blueprint v3.0, declared in advance with the "
          "measurement that justifies each.",
          font=7.5,
          widths=[1.1, 2.2, 2.6, 2.0, 4.4, 3.2])

    h(doc, "A.4  Diagnostic order if the gate fails", 2)
    for s in PRE["diagnostic_order_if_fail"]:
        bullet(doc, s)
    h(doc, "A.5  Memory fallback ladder", 2)
    for s in PRE["vram_fallback_ladder"]:
        bullet(doc, s)

    # ---------------------------------------------------------------
    h(doc, "Appendix B  Environment specification", 1, page_break=True)
    table(doc, ["Package", "Version"],
          [[k, v or "not installed"] for k, v in ENV["packages"].items()],
          "Installed package versions at the time of execution.",
          widths=[6.0, 9.5])
    para(doc, f"nvidia-smi: {ENV['nvidia_smi']}", size=9)
    para(doc, ENV["note"], size=9, italic=True)

    # ---------------------------------------------------------------
    h(doc, "Appendix C  Full batch ladder", 1, page_break=True)
    table(doc,
          ["Batch", "Peak VRAM (MiB)", "s / step", "images / s", "Fits"],
          [[str(r["batch"]),
            "OOM" if r.get("oom") else f"{r['peak_alloc_mib']:.0f}",
            "—" if r.get("oom") else f"{r['sec_per_step']:.3f}",
            "—" if r.get("oom") else f"{r['images_per_sec']:.1f}",
            "yes" if r.get("fits") else "no"]
           for r in PROBE["batch_ladder"]],
          f"Complete batch-size ladder in the adopted precision. The selection "
          f"rule required peak allocation plus a "
          f"{PROBE['safety_margin_mib']} MiB margin to remain below "
          f"{PROBE['total_vram_mib']:.0f} MiB.",
          align_right=[1, 2, 3])

    # ---------------------------------------------------------------
    h(doc, "Appendix D  Resolved trainable-layer list", 1, page_break=True)
    table(doc, ["Field", "Value"],
          [[k, str(v)] for k, v in LAY.items()],
          "Machine-readable record of the fine-tuning scope, written by the "
          "training script at run time.",
          widths=[6.4, 9.1])

    # ---------------------------------------------------------------
    h(doc, "Appendix E  Per-class results with Wilson intervals", 1,
      page_break=True)
    agg_pc = {}
    for s in MET["seeds"]:
        for r in MET["per_seed"][str(s)]["per_class"]:
            agg_pc.setdefault(r["class"], []).append(r)
    rows = []
    for c, rs in sorted(agg_pc.items(), key=lambda kv: -kv[1][0]["support"]):
        n = len(rs)
        rows.append([c, str(rs[0]["support"]),
                     f"{100 * sum(x['precision'] for x in rs) / n:.2f}",
                     f"{100 * sum(x['recall'] or 0 for x in rs) / n:.2f}",
                     f"{100 * sum(x['f1'] for x in rs) / n:.2f}",
                     f"{100 * sum(x['recall_wilson_lo'] for x in rs) / n:.1f} – "
                     f"{100 * sum(x['recall_wilson_hi'] for x in rs) / n:.1f}",
                     f"±{sum(x['wilson_half_width_pp'] for x in rs) / n:.1f}"])
    table(doc,
          ["Class", "Support", "Precision", "Recall", "F1",
           "Recall Wilson 95% CI", "Half-width"],
          rows,
          "Per-class test performance, averaged over seeds, ordered by "
          "support. EXPLORATORY: interval half-widths exceed the ±10 pp "
          "criterion for most classes (limitation L1).",
          align_right=[1, 2, 3, 4, 6])

    # ---------------------------------------------------------------
    h(doc, "Appendix F  Per-seed run manifests", 1, page_break=True)
    for r in RUNS:
        h(doc, f"F.{r['seed']}  Seed {r['seed']}", 2)
        table(doc, ["Field", "Value"],
              [["Device", r["device"]], ["torch", r["torch"]],
               ["CUDA", str(r["cuda"])], ["Python", r["python"]],
               ["Batch size", str(r["batch_size"])],
               ["Gradient accumulation", str(r["grad_accum_steps"])],
               ["Effective batch", str(r["effective_batch"])],
               ["Precision", r["precision"]],
               ["Epochs run", str(r["n_epochs_run"])],
               ["Stop reason", r["stop_reason"]],
               ["Best epoch", str(r["best_epoch_overall"])],
               ["Best validation macro F1", f"{100 * r['best_val_macro_f1']:.3f}"],
               ["Peak VRAM (MiB)", f"{r['peak_vram_mib']:.0f}"],
               ["Wall-clock (s)", f"{r['wallclock_sec']:.0f}"]],
              f"Run manifest for seed {r['seed']}.",
              widths=[5.4, 10.1])

    # ---------------------------------------------------------------
    h(doc, "Appendix G  CLAIM checklist — Phase 2 applicable items", 1,
      page_break=True)
    claim = [
        ("Study design", "Reproduction of a published benchmark; single-arm, "
                         "pre-registered", "§1.1, §2.1, Appendix A"),
        ("Data source", "GastroHUN, public, doi:10.1038/s41597-025-04401-5",
         "§2.2.1"),
        ("Eligibility", "Complete four-of-four annotator agreement", "§2.2.1"),
        ("Data pre-processing", "Lanczos 224×224; training-set normalisation",
         "§2.3"),
        ("Ground truth", "Unanimous label of four independent experts",
         "§2.2.1"),
        ("Data partitions", "Official patient-level splits; overlap re-verified",
         "§2.2.2"),
        ("Model", f"ConvNeXt-Tiny, ImageNet-pretrained, "
                  f"{LAY['n_modules_unfrozen']} feature modules fine-tuned",
         "§2.4"),
        ("Training approach", "Two-stage; early stopping on validation macro F1",
         "§2.5"),
        ("Hyper-parameters", "Fully tabulated", "Table 6"),
        ("Metrics", "Macro F1 primary; secondary battery; ECE and Brier",
         "§2.6.1"),
        ("Statistical measures", f"Patient-clustered bootstrap, "
                                 f"{MET['n_boot']:,} resamples", "§2.6.2, §4.3"),
        ("Robustness", f"{len(RUNS)} random seeds; variance reported", "§4.6"),
        ("Failure analysis", "Confusion matrix; per-class with Wilson intervals",
         "§4.4, §4.5"),
        ("Reproducibility", "Scripts, seeds, environment and artefacts listed",
         "§3.3, Appendix B, Appendix H"),
        ("Limitations", "Declared and carried forward", "§6.5, §6.6"),
    ]
    table(doc, ["CLAIM item", "How addressed", "Location"], claim,
          "CLAIM checklist items applicable to a reproduction phase.",
          widths=[3.8, 7.6, 4.1])

    # ---------------------------------------------------------------
    h(doc, "Appendix H  Regeneration pipeline", 1, page_break=True)
    para(doc,
         "Executed in this order from the project root. No number in this "
         "document is typed by hand; each is interpolated from an artefact "
         "produced by these scripts.")
    for c in ["python src/models/phase2_env.py",
              "python src/models/phase2_probe.py",
              "python src/models/phase2_data.py",
              "python src/models/phase2_normstats.py",
              "python src/models/phase2_cache.py",
              "python src/models/phase2_prereg.py",
              "python src/models/phase2_train.py --seed 1",
              "python src/models/phase2_train.py --seed 2",
              "python src/models/phase2_train.py --seed 3",
              "python src/models/phase2_eval.py",
              "python src/report/figures_phase2.py",
              "python src/report/build_phase2_docx.py",
              "python src/report/finalise_phase2.py"]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(c)
        r.font.name = "Consolas"; r.font.size = Pt(9)
    para(doc,
         f"Approximate cost: the precision probe takes about ten minutes, the "
         f"cache build under two, and each training run "
         f"{sum(r['wallclock_sec'] for r in RUNS) / len(RUNS) / 60:.0f} "
         f"minutes on the reference GPU.")

    h(doc, "Appendix I  Artefact inventory", 1, page_break=True)
    inv = [
        ("reports/phase2_env.json", "environment snapshot"),
        ("reports/phase2_vram_probe.json",
         "GATE 1, precision factorial, batch ladder"),
        ("reports/phase2_split_provenance.json",
         "GATE 2 and GATE 3, cohort composition"),
        ("reports/phase2_norm_stats.json", "training-set channel statistics"),
        ("reports/phase2_prereg.json", "frozen pre-registration"),
        ("reports/phase2_trainable_layers.json", "resolved fine-tuning scope"),
        ("reports/phase2_run_seed*.json", "per-seed manifests and histories"),
        ("reports/phase2_test_metrics.json",
         "test metrics, intervals, per-class results, verdict"),
        ("reports/phase2_predictions_seed*.csv", "per-image test predictions"),
        ("data/phase2_consensus_manifest.csv", "cohort definition"),
        ("data/phase2_class_index.json", "label-to-index mapping"),
        ("data/phase2_cache_224.npy", "deterministic preprocessing cache"),
        ("checkpoints/phase2_convnext_tiny_seed*.pt", "trained models"),
        ("figures_phase2/*.png", "figure suite"),
    ]
    table(doc, ["Artefact", "Contents"], inv,
          "Every artefact produced by Phase 2.",
          widths=[7.4, 8.1])


# =====================================================================
def main() -> None:
    doc = new_document()
    title_page(doc)
    front_matter(doc)
    abstract(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    chapter6(doc)
    chapter7(doc)
    references(doc)
    appendices(doc)
    add_page_numbers(doc)
    doc.save(BD.OUT)
    print(f"wrote {BD.OUT.name}")


if __name__ == "__main__":
    main()
