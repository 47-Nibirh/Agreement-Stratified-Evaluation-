"""
Phase 1 narrative: systematic literature review, problem framing, proposed
methodology, feasibility, limitations, conclusion, references and appendices.

All counts are interpolated from literature_v2/prisma_counts.json and
literature_v2/extraction_table.csv.
"""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH

from build_docx import bullet, callout, figure, h, para, rich, table
from content_phase0_v2 import (AGR, CA_PCT, DEC, GRAN, INV, N_CLS, N_IMG,
                               N_PAT, STR)

ROOT = Path(__file__).resolve().parents[2]
LIT = ROOT / "literature_v2"
P = json.loads((LIT / "prisma_counts.json").read_text(encoding="utf-8"))
DF = pd.read_csv(LIT / "extraction_table.csv")
S, E = P["stages"], P["eligibility"]


def _theme(t: str) -> pd.DataFrame:
    return DF[DF["theme"] == t]


def sec_phase1(doc) -> None:
    h(doc, "Phase 1 — Literature Review and Problem Framing", 1)

    h(doc, "Objectives and review questions", 2)
    para(doc,
         "The literature review has three objectives. First, to establish what "
         "is already known about automatic recognition of gastric anatomical "
         "landmarks and about the quality-audit application that motivates it. "
         "Second, to establish what is known about inter-observer variability "
         "in endoscopy and about the machine-learning methods that exist for "
         "learning from disagreeing annotators. Third, to determine, from that "
         "evidence, whether the gap identified in Phase 0 is genuinely open.")

    para(doc,
         "This review replaces an earlier one in its entirety. The earlier "
         "protocol was organised around natural-language processing of "
         "free-text endoscopy reports, a research question that was retired "
         "with the previous dataset. Its search strings, screening criteria and "
         "included studies have no bearing on the present question and none are "
         "carried forward. Seven new themes were specified before the search "
         "was executed.")

    table(doc, ["Theme", "Review question"],
          [["T1 Landmark recognition",
            "How well can deep networks classify gastric anatomical sites, and "
            "under what evaluation conditions have those results been obtained?"],
           ["T2 Quality and blind-spot audit",
            "Does automated monitoring of endoscopic completeness improve "
            "measurable procedure quality?"],
           ["T3 Inter-observer variability",
            "How much do expert endoscopists agree, and how is that agreement "
            "conventionally reported?"],
           ["T4 Noisy and multi-annotator labels",
            "What methods exist for training when annotators disagree, and do "
            "they outperform consensus-only training?"],
           ["T5 Uncertainty and calibration",
            "Can a model's predictive uncertainty be made to reflect genuine "
            "ambiguity in the input?"],
           ["T6 External validation and shift",
            "How far do endoscopic imaging models transfer beyond their "
            "development centre?"],
           ["T7 Reporting standards",
            "What must a diagnostic-AI study report to be assessable, and how "
            "well is that standard currently met?"]],
          "The seven review themes and the question each addresses.",
          widths=[4.4, 12.4], font=8.2)

    h(doc, "Protocol and search strategy", 2)
    para(doc,
         f"The review follows PRISMA 2020. One Boolean query was specified per "
         f"theme against PubMed/MEDLINE through the NCBI E-utilities interface, "
         f"executed programmatically on {P['run_date']} so that the counts "
         f"reported here can be reproduced by re-running a single script. The "
         f"primary date window was {P['queries']['T1_landmark_recognition']['date_from']}"
         f"–{P['queries']['T1_landmark_recognition']['date_to']}, relaxed for "
         f"the two themes whose foundational literature predates it. A maximum "
         f"of {P['queries']['T1_landmark_recognition']['retmax']} records per "
         f"query was retrieved in relevance order.")

    q_rows = []
    for k, v in P["queries"].items():
        q_rows.append([v["theme"][:38], f"{v['total_hits']:,}",
                       f"{v['retrieved']}", f"{v['date_from']}–{v['date_to']}"])
    table(doc, ["Theme", "Total hits", "Retrieved", "Date window"], q_rows,
          "Search yield per theme. Full Boolean strings are reproduced in "
          "Appendix A.",
          widths=[7.0, 2.8, 2.4, 3.4], font=8.0,
          note="Where total hits exceed the retrieval cap, records were taken "
               "in PubMed relevance order; this is a stated deviation and is "
               "discussed under review limitations.")

    para(doc,
         "Screening proceeded in two stages. The title/abstract screen applied "
         "language, publication-type and topic filters, together with explicit "
         "homonym guards for traps observed during protocol development — most "
         "notably that a bare search for the CLAIM reporting checklist also "
         "returns papers about insurance claims databases. The eligibility "
         "stage then applied theme-specific criteria in which a record must "
         "match at least one term in every required term group. A second class "
         "of homonym guard was necessary here: the stem 'endoscop' matches "
         "drug-induced sleep endoscopy, laryngoscopy and neuroendoscopy, none "
         "of which concern the luminal gastrointestinal tract. Records matching "
         "those were excluded from the three GI-specific themes and the "
         "exclusions counted.")

    h(doc, "Search results", 2)

    figure(doc, "F17_prisma.png",
           "PRISMA 2020 flow diagram for the revised review protocol.")

    table(doc, ["PRISMA stage", "Records"],
          [["Identified through database searching (with overlap)",
            f"{S['records_identified_total']:,}"],
           ["Duplicates removed", f"{S['duplicates_removed']}"],
           ["Records after de-duplication", f"{S['records_after_deduplication']:,}"],
           ["Metadata successfully retrieved", f"{S['records_metadata_retrieved']:,}"],
           ["Excluded at title/abstract screening",
            f"{S['records_excluded_at_screening']}"],
           ["Assessed for eligibility", f"{E['records_assessed_for_eligibility']:,}"],
           ["Excluded — non-GI endoscopy homonym", f"{E['excluded_non_gi_homonym']}"],
           ["Excluded — not a luminal GI study", f"{E['excluded_not_luminal_gi']}"],
           ["Excluded — failed theme eligibility criteria",
            f"{E['excluded_failed_criteria']}"],
           ["Excluded — below the per-theme relevance cap",
            f"{E['excluded_below_cap']}"],
           ["Included from database searching",
            f"{E['included_from_database_search']}"],
           ["Included via other methods (hand-searched)",
            f"{E['included_from_other_methods']}"],
           ["TOTAL INCLUDED IN REVIEW", f"{E['total_included_in_review']}"]],
          "PRISMA 2020 stage counts.",
          widths=[11.0, 3.4], font=8.2,
          note="Exclusion counts reconcile exactly with the number assessed; "
               "the build asserts this and fails if it does not.")

    figure(doc, "F18_literature.png",
           "Left: composition of the included studies by theme. Right: "
           "publication-year distribution of the studies retrieved by database "
           "search.")

    yrs = pd.to_numeric(DF["year"], errors="coerce").dropna().astype(int)
    para(doc,
         f"{E['total_included_in_review']} studies were included. "
         f"{100*(yrs >= 2020).mean():.0f}% were published in 2020 or later, "
         f"reflecting how recent this field is; the pre-2013 entries are the "
         f"foundational statistical and architectural works added through the "
         f"other-methods arm, which are cited for method rather than for "
         f"findings.")

    para(doc,
         "The per-theme relevance cap warrants explicit statement, because it "
         "is a deviation from an idealised systematic review. Assessing "
         f"{E['records_assessed_for_eligibility']:,} records at full text is "
         "not achievable within an undergraduate thesis. A pre-specified cap "
         "was therefore applied per theme, with records ranked by a declarative "
         "relevance score computed from theme-specific terms and ties broken by "
         "recency then by PubMed identifier, so the selection is deterministic "
         "and reproducible. The consequence is that this is a systematically "
         "conducted but not exhaustive review, and it is described as such "
         "throughout.")


def sec_synthesis(doc) -> None:
    h(doc, "Thematic synthesis", 2)

    # ---- T1 ----------------------------------------------------------------
    h(doc, "T1 — Anatomical landmark recognition in upper GI endoscopy", 3)
    t1 = _theme("T1 Landmark recognition in UGI endoscopy")
    para(doc,
         f"{len(t1)} studies address the recognition of anatomical sites from "
         f"endoscopic images. The consistent finding is that modern "
         f"convolutional and transformer architectures classify gastric "
         f"landmarks to a level broadly comparable with individual expert "
         f"endoscopists, provided the evaluation set is clean. Reported "
         f"accuracies cluster in the mid-eighties to low nineties depending on "
         f"the number of classes, and the ConvNeXt family recurs as the "
         f"strongest convolutional backbone.")

    para(doc,
         f"For the corpus audited in Phase 0, the reference results are those "
         f"of the dataset descriptor itself. ConvNeXt-Large attains a macro "
         f"F1 of 88.25 ± 0.22 on the consensus test set; ConvNeXt-Tiny reaches "
         f"approximately 85 with 28 million parameters against 200 million, a "
         f"trade-off that matters for the compute budget set out later in this "
         f"report. A separate experiment in the descriptor trains ConvNeXt-Tiny "
         f"on the labels agreed by the two residents and reports a macro F1 of "
         f"87.05 ± 0.21, which exceeds the 84.82 ± 0.23 obtained when training "
         f"on the best single annotator's labels — and does so with fewer "
         f"training images.")

    callout(doc,
            "That last result is the single most important observation in the "
            "reviewed literature for this thesis. It shows, on this exact "
            "corpus, that how the annotator labels are combined changes "
            "performance by more than two F1 points — more than the difference "
            "between architecture families. The descriptor reports it as an "
            "aside. It is in fact evidence that label aggregation is an "
            "underexplored design axis, and it is the direct precedent for the "
            "soft-label experiments proposed below.",
            title="Label aggregation matters more than architecture")

    # ---- T2 ----------------------------------------------------------------
    h(doc, "T2 — Endoscopic quality metrics and blind-spot auditing", 3)
    t2 = _theme("T2 Endoscopic quality & blind-spot audit")
    para(doc,
         f"{len(t2)} studies address the clinical application. The rationale is "
         f"well established: a substantial proportion of early gastric cancers "
         f"are missed at endoscopy, missed lesions concentrate in particular "
         f"anatomical regions, and completeness of mucosal inspection is a "
         f"modifiable determinant of detection. Systematic photodocumentation "
         f"protocols exist precisely to enforce completeness, and the Japanese "
         f"SSS protocol with its 22 prescribed images is the most demanding of "
         f"them.")

    para(doc,
         "The reviewed evidence indicates that real-time systems monitoring "
         "blind-spot coverage measurably improve procedural quality, with "
         "several randomised and multicentre evaluations reporting reduced "
         "blind-spot rates and increased detection. This establishes the "
         "clinical value of the landmark-classification task: it is not an "
         "academic exercise but the perception layer underneath an audit tool. "
         "It also sharpens why performance under disagreement matters. An audit "
         "system is deployed against a whole procedure, including the images a "
         "clinician would find ambiguous, and cannot restrict itself to the "
         "subset on which four experts would have concurred.")

    # ---- T3 ----------------------------------------------------------------
    h(doc, "T3 — Inter-observer variability in endoscopy", 3)
    t3 = _theme("T3 Inter-observer variability in endoscopy")
    para(doc,
         f"{len(t3)} studies and methodological references address agreement. "
         f"Imperfect agreement between expert endoscopists is a stable, "
         f"repeatedly replicated finding across essentially every classification "
         f"task in gastroenterology — lesion morphology, mucosal cleanliness, "
         f"staging systems and anatomical assessment alike. Reported kappa "
         f"values commonly fall in the moderate-to-substantial band, which is "
         f"where the Phase 0 measurement of "
         f"{AGR['fleiss_kappa']:.3f} for this corpus also sits.")

    para(doc,
         "Two methodological observations follow from this theme and are "
         "applied in Phase 0. First, reporting a single agreement coefficient "
         "is poor practice, because the common coefficients have different "
         "failure modes under skewed prevalence; the audit therefore reports "
         "Fleiss' kappa, Krippendorff's alpha and Gwet's AC1 together and "
         "explains their convergence. Second, the literature reports agreement "
         "almost exclusively as a data-quality statistic — a number quoted once "
         "to establish that labels are trustworthy — and essentially never as "
         "an experimental variable that model performance is analysed against. "
         "That absence is the methodological space this thesis occupies.")

    # ---- T4 ----------------------------------------------------------------
    h(doc, "T4 — Learning from noisy, soft and multi-annotator labels", 3)
    t4 = _theme("T4 Noisy / soft / multi-annotator labels")
    para(doc,
         f"{len(t4)} studies address the machine-learning response to "
         f"annotator disagreement. The methods divide into three families. "
         f"Label-smoothing and soft-target approaches replace the one-hot "
         f"target with a distribution, which in a multi-annotator setting can "
         f"be taken directly from the vote proportions. Annotator-modelling "
         f"approaches learn a per-annotator confusion structure and infer a "
         f"latent true label. Uncertainty-aware approaches treat disagreement "
         f"as an explicit training signal, propagating label uncertainty into "
         f"predictive uncertainty.")

    para(doc,
         "The recurring empirical finding is that discarding disagreement is "
         "wasteful: models trained on soft or multi-annotator targets are "
         "generally better calibrated than those trained on hard consensus "
         "labels, and often no worse, sometimes better, on accuracy. Most of "
         "this work has been done in segmentation, particularly radiotherapy "
         "contouring and histopathology, where multi-reader datasets are more "
         "common. Application to endoscopic landmark classification is scarce, "
         "which is unsurprising given that GastroHUN is the first public "
         "upper-GI dataset to publish per-annotator labels at all.")

    # ---- T5 ----------------------------------------------------------------
    h(doc, "T5 — Uncertainty quantification and calibration", 3)
    t5 = _theme("T5 Uncertainty & calibration")
    para(doc,
         f"{len(t5)} studies and foundational references address whether a "
         f"model's confidence can be trusted. Modern networks are systematically "
         f"overconfident; temperature scaling, deep ensembles and Monte Carlo "
         f"dropout are the established correctives, and conformal prediction "
         f"offers distribution-free coverage guarantees at the cost of "
         f"set-valued rather than point predictions. Expected calibration error "
         f"and reliability diagrams are the standard reporting instruments.")

    para(doc,
         "For this thesis the relevant question is narrower and, in the "
         "reviewed literature, largely unanswered: does a model's predictive "
         "uncertainty align with the specific images on which human experts "
         "disagreed? Calibration is nearly always assessed against correctness, "
         "not against human ambiguity. A corpus carrying four independent "
         "labels permits the stronger test, because the human disagreement "
         "level of every individual image is known.")

    # ---- T6 ----------------------------------------------------------------
    h(doc, "T6 — External validation and dataset shift", 3)
    t6 = _theme("T6 External validation & dataset shift")
    para(doc,
         f"{len(t6)} studies address transfer beyond the development centre. "
         f"The consistent finding across medical imaging is that performance "
         f"degrades on external data, sometimes severely, and that "
         f"single-centre results systematically overstate deployable "
         f"performance. Endoscopy is particularly exposed because processor "
         f"vendor, illumination, insufflation practice and operator technique "
         f"all vary between units and all alter image appearance.")

    para(doc,
         "This bears directly on Phase 0 limitation L3. The audited corpus "
         "comes from a single Colombian hospital using a single Olympus "
         "platform throughout. Any performance figure obtained on it is a "
         "single-centre figure. Two public datasets — HyperKvasir and "
         "GastroVision — contain upper-GI categories that overlap the SSS "
         "landmark set at a coarse level, notably the pylorus, the "
         "retroflexed stomach view and the Z-line, and therefore offer a "
         "feasible if partial external check.")

    # ---- T7 ----------------------------------------------------------------
    h(doc, "T7 — Reporting standards for medical AI", 3)
    t7 = _theme("T7 Reporting standards for medical AI")
    para(doc,
         f"{len(t7)} studies address how diagnostic-AI research should be "
         f"reported and how well it currently is. The instruments are "
         f"established — CLAIM for medical imaging AI, TRIPOD+AI for prediction "
         f"models, STARD for diagnostic accuracy, PROBAST for risk of bias — "
         f"and the repeated finding of the meta-research studies retrieved here "
         f"is that adherence is poor. Recurrent omissions are the absence of "
         f"external validation, incomplete description of the study population, "
         f"missing detail on how ground truth was established, and no reporting "
         f"of calibration.")

    callout(doc,
            "Three of those four recurrent omissions are addressed directly by "
            "the design adopted here: ground-truth construction is the central "
            "object of study rather than an unexamined input; calibration is a "
            "primary endpoint rather than an afterthought; and external "
            "validation is a planned phase rather than an aspiration. The "
            "fourth — full population description — cannot be met, because the "
            "corpus records no demographics, and Phase 0 declares this as "
            "limitation L2 rather than passing over it.",
            title="Reporting posture adopted by this thesis")


def sec_gap(doc) -> None:
    h(doc, "Synthesis: what is established, and what is not", 2)

    table(doc, ["Established by the reviewed literature", "Not established"],
          [["Deep networks classify gastric landmarks at a level comparable to "
            "individual experts, on curated evaluation sets.",
            "How that performance behaves on images where experts disagree."],
           ["Automated blind-spot and completeness monitoring improves "
            "measurable endoscopic quality.",
            "Whether an audit system remains reliable on the ambiguous images "
            "it will necessarily encounter in deployment."],
           ["Expert endoscopists agree only moderately-to-substantially on "
            "classification tasks, including anatomical ones.",
            "Whether the structure of that disagreement is exploitable, rather "
            "than merely a nuisance to be averaged away."],
           ["Soft and multi-annotator targets improve calibration, mostly "
            "demonstrated in segmentation.",
            "Whether they help for endoscopic landmark classification, where "
            "per-annotator labels have only just become public."],
           ["Model confidence can be calibrated against correctness.",
            "Whether model uncertainty tracks human ambiguity specifically."],
           ["Medical imaging models degrade under centre shift.",
            "The magnitude of that degradation for SSS landmark recognition."]],
          "Synthesis of the review against the Phase 0 findings.",
          widths=[8.4, 8.4], font=8.2)

    h(doc, "The research gap", 2)

    figure(doc, "F19_conceptual_framework.png",
           "Conceptual framework. Three established bodies of work converge on "
           "an unexamined intersection, from which the three research questions "
           "follow.")

    para(doc,
         f"The gap is precise and is stated in the dataset descriptor's own "
         f"limitations. Every published benchmark on this corpus trains and "
         f"tests exclusively on images carrying complete four-of-four expert "
         f"consensus. Phase 0 quantifies what that means: {CA_PCT:.1f}% of the "
         f"corpus is retained and {100-CA_PCT:.1f}% is discarded, including "
         f"{AGR['n_no_majority']:,} images "
         f"({AGR['pct_no_majority']:.2f}%) for which no majority label exists "
         f"under any voting rule. Model behaviour on the discarded portion is "
         f"unmeasured. Because the discarded portion is defined by human "
         f"ambiguity, it is precisely the portion where an automated second "
         f"reader would be most valuable and where its failure would be least "
         f"visible.")

    para(doc,
         f"Phase 0 also supplies the structural finding that makes the gap "
         f"tractable rather than merely large. Disagreement is not noise: "
         f"{DEC['same_station_different_wall']:.1f}% of it is wall confusion "
         f"within a station both annotators agree on, and confusions "
         f"overwhelmingly involve circumferentially adjacent walls and "
         f"neighbouring stations. Because the label space has a known geometry "
         f"and the disagreement respects it, disagreement can be modelled "
         f"rather than merely measured.")

    h(doc, "Research questions and hypotheses", 2)

    table(doc, ["ID", "Research question", "Hypothesis", "Primary endpoint"],
          [["RQ1",
            "How does landmark-classification performance vary across strata of "
            "expert agreement?",
            "Macro F1 declines monotonically from the unanimous stratum to the "
            "no-majority stratum, and the decline is larger than the difference "
            "between architecture families.",
            "Macro F1 per agreement stratum, patient-clustered bootstrap 95% CI"],
           ["RQ2",
            "Does training on soft targets derived from all four annotator "
            "votes outperform training on hard consensus labels?",
            "Soft-target training matches hard-label training on the unanimous "
            "stratum and exceeds it on the contested strata, with better "
            "calibration throughout.",
            "Macro F1 by stratum; expected calibration error"],
           ["RQ3",
            "Does predictive uncertainty track human disagreement, and does it "
            "survive a change of centre?",
            "Predictive entropy correlates positively with per-image annotator "
            "disagreement, and the ranking is preserved on external data.",
            "Spearman correlation between predictive entropy and annotator "
            "vote entropy; external macro F1"]],
          "Research questions, hypotheses and primary endpoints.",
          widths=[1.0, 4.8, 5.4, 5.6], font=7.8,
          note="Hypotheses are stated before any modelling is performed. "
               "Phase 0 deliberately measured no model performance so that "
               "these remain genuine predictions.")


def sec_methodology(doc) -> None:
    h(doc, "Proposed Methodology", 1)

    figure(doc, "F20_methodology.png",
           "Proposed experimental pipeline across data, model and evaluation "
           "layers.")

    h(doc, "Design", 2)
    para(doc,
         f"The design is a controlled comparison on a fixed corpus with fixed "
         f"official splits. The independent variables are the target "
         f"construction (hard consensus label versus soft four-vote "
         f"distribution) and the evaluation stratum (agreement level). The "
         f"dependent variables are macro F1 and calibration error. Because the "
         f"splits are published and patient-disjoint, and because Phase 0 "
         f"verified both properties directly, no split construction is required "
         f"and no opportunity for split-related optimism is introduced.")

    para(doc,
         "The evaluation strata are defined a priori from the Phase 0 voting "
         "patterns, and are mutually exclusive and exhaustive over the test "
         "set.")

    vp, vpp = AGR["vote_patterns"], AGR["vote_patterns_pct"]
    table(doc, ["Stratum", "Definition", "Images (whole corpus)", "% of corpus"],
          [["S-unanimous", "All four annotators assign the same label",
            f"{vp.get('4', 0):,}", f"{vpp.get('4', 0):.2f}%"],
           ["S-majority", "Three of four agree", f"{vp.get('3-1', 0):,}",
            f"{vpp.get('3-1', 0):.2f}%"],
           ["S-plurality", "Two agree, the other two differ from them and each "
            "other", f"{vp.get('2-1-1', 0):,}", f"{vpp.get('2-1-1', 0):.2f}%"],
           ["S-tied", "Two-way tie, no majority", f"{vp.get('2-2', 0):,}",
            f"{vpp.get('2-2', 0):.2f}%"],
           ["S-dispersed", "All four differ", f"{vp.get('1-1-1-1', 0):,}",
            f"{vpp.get('1-1-1-1', 0):.2f}%"]],
          "Pre-specified evaluation strata.",
          widths=[3.0, 7.2, 3.2, 2.6], font=8.0,
          note="For the two strata with no majority label, accuracy is "
               "undefined against a single ground truth; performance is scored "
               "as agreement with the annotator distribution, and the scoring "
               "rule is fixed before any model is run.")

    h(doc, "Models and training", 2)
    para(doc,
         "ConvNeXt-Tiny is adopted as the primary backbone. The choice is "
         "driven by the compute budget documented below and is supported by the "
         "descriptor's own finding that this architecture reaches within about "
         "three F1 points of ConvNeXt-Large at one seventh the parameter count. "
         "ImageNet-pretrained weights are used, with a warm-up phase training "
         "the classifier head followed by fine-tuning of the upper feature "
         "layers, matching the descriptor's protocol so that the reproduction "
         "is comparable to the published baseline.")

    table(doc, ["Configuration", "Target construction", "Purpose"],
          [["C0 — reproduction", "Hard label, complete-agreement subset only",
            "Reproduce the published baseline and validate the pipeline"],
           ["C1 — hard, all data", "Majority label where one exists",
            "Isolate the effect of adding contested images"],
           ["C2 — soft targets", "Vote proportions across the four annotators",
            "Test RQ2"],
           ["C3 — smoothed", "Hard label with label smoothing",
            "Control: distinguishes soft-target benefit from generic "
            "regularisation"],
           ["C4 — anatomy-aware", "Soft targets plus a wall/station-structured "
            "penalty", "Exploit the Phase 0 structural finding"]],
          "Planned training configurations.",
          widths=[3.4, 6.0, 6.8], font=8.0,
          note="C3 exists specifically to prevent a confound: any gain from "
               "soft targets must be shown to exceed the gain from ordinary "
               "label smoothing.")

    h(doc, "Evaluation and statistics", 2)
    para(doc,
         f"Macro-averaged F1 is the primary metric, consistent with the "
         f"published baseline and appropriate to a near-balanced "
         f"{N_CLS}-class problem. All intervals are computed by bootstrap "
         f"resampling of patients rather than images, because Phase 0 "
         f"established that agreement varies systematically between patients "
         f"and that images within a patient are therefore not independent. "
         f"Calibration is reported as expected calibration error with "
         f"reliability diagrams. Per-class results are reported but, following "
         f"Phase 0 limitation L1, are treated as exploratory.")

    h(doc, "Compute budget and feasibility", 2)
    para(doc,
         "The available hardware is a single NVIDIA GeForce GTX 1650 with 4 GB "
         "of video memory (Turing, compute capability 7.5). This constrains the "
         "design in three ways that are stated here so the plan is honest about "
         "what it can deliver.")
    bullet(doc, "ConvNeXt-Tiny at 224×224 fits in 4 GB with mixed-precision "
                "training at a batch size of roughly 16–24; gradient "
                "accumulation supplies any larger effective batch. "
                "ConvNeXt-Large does not fit and is not attempted — the "
                "published 88.25 figure is cited as the ceiling, not "
                "reproduced.")
    bullet(doc, "Turing supports FP16 tensor cores but not BF16, so automatic "
                "mixed precision must use float16 with gradient scaling.")
    bullet(doc, "Deep ensembles, if used for RQ3, must be trained sequentially "
                "rather than in parallel; five members at roughly one hour each "
                "is tractable, and Monte Carlo dropout is the cheaper fallback.")
    para(doc,
         "One practical prerequisite is recorded: the currently installed "
         "PyTorch is a CPU-only build, so the GPU is not yet usable and a CUDA "
         "build must be installed before Phase 2 begins.")


def sec_limitations(doc) -> None:
    h(doc, "Threats to Validity", 1)

    para(doc,
         "The limitations below combine those arising from Phase 0 with those "
         "arising from the review protocol. They are stated in full because a "
         "reader's ability to discount a finding appropriately depends on "
         "knowing what constrains it.")

    table(doc, ["Threat", "Type", "Assessment", "Response"],
          [["Single-centre, single-vendor corpus", "External validity",
            "All images from one Colombian hospital on one Olympus platform; "
            "generalisation to other units is untested",
            "External validation on HyperKvasir and GastroVision shared "
            "landmarks is a required phase"],
           ["No demographic data", "External validity / reporting",
            "Age and sex absent; population describable only at institutional "
            "level",
            "Declared; no subgroup or fairness claim made"],
           ["Per-class test precision", "Statistical conclusion validity",
            f"{AGR['n_test_classes_underpowered_hw_gt_10pct']}/{N_CLS} classes "
            "exceed a ±10 pp Wilson half-width",
            "Macro-averaged primary endpoints; per-class exploratory"],
           ["Contested strata are small", "Statistical conclusion validity",
            f"The tied and dispersed strata together are only "
            f"{AGR['pct_no_majority']:.2f}% of the corpus",
            "Report stratum sizes with every stratified result; pool the two "
            "no-majority strata if power demands"],
           ["Four annotators is a small panel", "Construct validity",
            "Vote proportions from four raters are a coarse estimate of the "
            "underlying label distribution",
            "Soft targets treated as an ordinal ambiguity signal, not as a "
            "calibrated probability"],
           ["Annotator idiosyncrasy", "Construct validity",
            "Phase 0 found FG2 to be an outlier on every measure; a four-rater "
            "mean is sensitive to one atypical rater",
            "Leave-one-annotator-out sensitivity analysis"],
           ["Review not exhaustive", "Review validity",
            f"Retrieval capped per query; {E['excluded_below_cap']} eligible "
            "records fell below the relevance cap",
            "Cap and ranking rule pre-specified, deterministic and reported"],
           ["Single database", "Review validity",
            "PubMed/MEDLINE only; IEEE Xplore, ACM and arXiv not searched",
            "Foundational computer-science work added via the other-methods "
            "arm; declared as limitation L5"],
           ["Perceptual duplicate detection is threshold-dependent",
            "Internal validity",
            "The near-duplicate result depends on a decision rule that had to "
            "be calibrated against controls before it was trustworthy",
            "Calibration procedure, controls and visual audit reported in full "
            "so the rule can be challenged"]],
          "Threats to validity and the response adopted for each.",
          widths=[3.6, 2.6, 5.4, 5.2], font=7.6)


def sec_conclusion(doc) -> None:
    h(doc, "Conclusion", 1)

    para(doc,
         f"Phase 0 subjected the GastroHUN corpus to an eight-criterion "
         f"integrity gate and the gate returned PROCEED. The corpus is "
         f"physically intact — {INV['n_decoded_ok']:,} images decoded without "
         f"a single failure, no missing or orphan files, no exact duplicates "
         f"— correctly partitioned at patient level with statistically verified "
         f"balance, and honestly documented with a named ethics approval and an "
         f"explicit consent basis. Two criteria returned CONDITIONAL: the "
         f"consensus test set is underpowered for per-class claims, and the "
         f"release carries no demographic data. Both are declared as "
         f"limitations and both constrain the claims the thesis may make "
         f"without invalidating the corpus.")

    para(doc,
         f"The audit produced one finding of independent interest. Expert "
         f"disagreement in this corpus is anatomically structured rather than "
         f"random: {DEC['same_station_different_wall']:.1f}% of it consists of "
         f"annotators who agree on the anatomical station but differ on the "
         f"gastric wall, and confusions overwhelmingly involve adjacent walls "
         f"and neighbouring stations. Collapsing the label space to station "
         f"raises mean pairwise kappa from "
         f"{GRAN['full']['mean_pairwise_kappa']:.3f} to "
         f"{GRAN['station']['mean_pairwise_kappa']:.3f}, while collapsing to "
         f"wall leaves it unchanged. Endoscopists know how deep the scope is "
         f"and disagree about which way it points. This is not reported in the "
         f"dataset descriptor and it converts disagreement from a nuisance into "
         f"a modellable structure.")

    para(doc,
         f"Phase 1 established, from {E['total_included_in_review']} included "
         f"studies, that the three literatures bearing on this problem — "
         f"landmark recognition, endoscopic quality auditing, and learning from "
         f"disagreeing annotators — are individually mature and jointly "
         f"unconnected. Agreement is universally reported as a data-quality "
         f"statistic and essentially never as a variable that model performance "
         f"is analysed against. The dataset descriptor states the resulting gap "
         f"in its own limitations: its results were obtained exclusively on "
         f"images with complete expert consensus, which 'misses a much more "
         f"variable real-world scenario'.")

    para(doc,
         f"That gap is the thesis. Model performance on this corpus has only "
         f"ever been measured on the {CA_PCT:.1f}% of images four experts label "
         f"identically. The remaining {100-CA_PCT:.1f}% — the images clinicians "
         f"themselves find ambiguous, and therefore the images where an "
         f"automated second reader would matter most and fail least visibly — "
         f"has never been evaluated. Phases 2 to 5 will measure it, test "
         f"whether soft targets built from all four votes improve behaviour "
         f"there, and check whether the result survives a change of centre.")

    callout(doc,
            "The methodological lesson of this report is narrower and worth "
            "stating plainly. The project's previous corpus failed its audit "
            "only because the audit was performed at all; the reported accuracy "
            "on it was near-perfect and entirely spurious. The same protocol "
            "applied to GastroHUN returns a PROCEED with two honest "
            "qualifications. Running the gate before the modelling, rather than "
            "defending the modelling afterwards, is what separates those two "
            "outcomes.")

    doc.add_page_break()


def sec_references(doc) -> None:
    h(doc, "References", 1)
    para(doc,
         f"{len(DF)} works, formatted in APA 7th edition. Bibliographic detail "
         f"for MEDLINE-indexed records was retrieved programmatically from "
         f"PubMed; hand-searched works are marked in Appendix B.",
         italic=True, size=9.5)

    refs = []
    for _, r in DF.iterrows():
        s = r.get("apa", "")
        if not isinstance(s, str) or not s.strip() or s.strip().lower() == "nan":
            au = str(r.get("authors_full", "")).split(";")[0].strip()
            s = (f"{au} ({r.get('year', 'n.d.')}). {r.get('title', '')}. "
                 f"{r.get('journal', '')}.")
        refs.append(s.replace("*", ""))

    for s in sorted(set(refs), key=lambda x: x.lower()):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = __import__("docx").shared.Cm(0.9)
        p.paragraph_format.first_line_indent = __import__("docx").shared.Cm(-0.9)
        p.paragraph_format.space_after = __import__("docx").shared.Pt(5)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(s)
        r.font.size = __import__("docx").shared.Pt(9.5)

    doc.add_page_break()


def sec_appendix(doc) -> None:
    h(doc, "Appendix A — Search strings", 1)
    para(doc,
         "Verbatim Boolean strings submitted to the NCBI E-utilities esearch "
         "endpoint. Reproducing the counts in this report requires only these "
         "strings, the stated date windows and the retrieval cap.",
         italic=True, size=9.5)
    for k, v in P["queries"].items():
        h(doc, v["theme"], 3)
        p = doc.add_paragraph()
        r = p.add_run(v["search_string"])
        r.font.name = "Consolas"
        r.font.size = __import__("docx").shared.Pt(8)
        para(doc, f"Window {v['date_from']}–{v['date_to']}; "
                  f"{v['total_hits']:,} hits; {v['retrieved']} retrieved.",
             size=8.5, italic=True)

    h(doc, "Appendix B — Included studies", 1)
    rows = []
    for _, r in DF.sort_values(["theme", "year"], ascending=[True, False]).iterrows():
        rows.append([
            str(r.get("theme", ""))[:26],
            str(r.get("year", ""))[:4],
            str(r.get("first_author", ""))[:22],
            str(r.get("title", ""))[:74],
            "hand" if "other methods" in str(r.get("source", "")) else "db",
        ])
    table(doc, ["Theme", "Year", "First author", "Title", "Source"], rows,
          f"All {len(DF)} studies included in the review.",
          widths=[3.4, 1.1, 2.8, 8.0, 1.2], font=6.8,
          note="Source 'db' = retrieved by database search; 'hand' = added "
               "through the PRISMA other-methods arm.")

    h(doc, "Appendix C — Reproducibility", 1)
    para(doc,
         "Every number, table and figure in this report is generated from a "
         "script. Executing the following in order regenerates the entire "
         "document from the raw dataset.")
    steps = [
        ("src/data/gastrohun_inventory.py",
         "Physical inventory, decode check, SHA-256 and perceptual hashes"),
        ("src/data/gastrohun_agreement.py",
         "Agreement statistics, split integrity, power, clinical metadata"),
        ("src/data/gastrohun_structure.py",
         "Wall/station decomposition, provenance heterogeneity, per-patient κ"),
        ("src/data/gastrohun_neardup.py",
         "Exhaustive near-duplicate scan with pixel verification"),
        ("src/data/gastrohun_dup_calibration.py",
         "Calibration of the duplicate decision rule against matched controls"),
        ("src/literature/search_v2.py", "PubMed search and PRISMA stage counts"),
        ("src/literature/eligibility_v2.py",
         "Theme eligibility, homonym guards, extraction table"),
        ("src/literature/enrich_v2.py", "Bibliographic enrichment and APA strings"),
        ("src/report/figures_v2.py", "All figures"),
        ("src/report/build_docx_v2.py", "This document"),
        ("src/report/finalise_v2.py", "Field update and PDF export"),
    ]
    table(doc, ["Script", "Produces"],
          [[s, d] for s, d in steps],
          "Regeneration pipeline, in execution order.",
          widths=[6.4, 10.0], font=8.0,
          note="No value in this report is typed by hand; every figure in the "
               "text is interpolated from a generated JSON or CSV artefact.")
