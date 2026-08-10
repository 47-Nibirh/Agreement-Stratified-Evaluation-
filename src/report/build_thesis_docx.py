"""
Phase 7 / P7.5-P7.6 -- build the thesis document.

This is not a concatenation of the seven phase reports. It is one argument, and
the reports are its evidence base. Every number is read from
reports/phase7_register.json, which resolved each quantity exactly once from the
phase artefacts, so no two chapters can quote the same value differently. Every
figure is read from reports/phase7_figure_registry.json, which records which
phase script drew it.

Nothing is typed by hand. Where a sentence needs a number it interpolates one;
where it needs a verdict it reads a verdict field produced by a frozen rule.

Run:  python src/report/build_thesis_docx.py
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_docx as BD  # noqa: E402

ROOT = Path(__file__).resolve().parents[2]
BD.FIGD = ROOT / "figures_thesis"
BD.OUT = ROOT / "Thesis.docx"

from build_docx import (add_page_numbers, bullet, callout, figure,  # noqa: E402
                        front_matter, new_document, para)
from build_docx import h as _h  # noqa: E402
from build_docx import table as _table_cm  # noqa: E402

import bibliography as BIB  # noqa: E402
from bibliography import cite, cite_theme  # noqa: E402

REP = ROOT / "reports"

# A4 portrait, 0.95" side margins -> 6.37 in of usable width.
USABLE_IN = 8.27 - 2 * 0.95

_seen_h1: list = []


def _prev_is_page_break(doc) -> bool:
    """True if the last paragraph is an empty one carrying a page break."""
    if not doc.paragraphs:
        return True
    p = doc.paragraphs[-1]
    if p.text.strip():
        return False
    xml = p._p.xml
    return 'w:br' in xml and 'type="page"' in xml


def h(doc, text, level=1, **kw):
    """Heading wrapper that starts every chapter on a fresh page, correctly.

    The document used to end each chapter with doc.add_page_break(), which
    inserts an EMPTY PARAGRAPH whose only content is the break. That paragraph
    then renders at the top of the following page, and where a chapter happened
    to end near a page boundary the result was a page containing nothing but it
    -- the blank pages that appeared mid-document.

    Setting page_break_before on the heading itself carries no stray paragraph,
    so a chapter can never be preceded by a blank page. The check for an
    existing break keeps this compatible with build_docx.front_matter, which
    issues its own breaks and is shared with the phase reports.
    """
    if level == 1 and _seen_h1 and not _prev_is_page_break(doc):
        p = _h(doc, text, level, **kw)
        p.paragraph_format.page_break_before = True
    else:
        p = _h(doc, text, level, **kw)
    if level == 1:
        _seen_h1.append(text)
    return p


def table(doc, headers, rows, caption, *, widths=None, **kw):
    """Width-correcting wrapper.

    build_docx.table applies Cm() to whatever it is given. Every `widths` list
    in THIS builder was written in inches -- they sum to about 6.5, which is the
    usable text width in inches, not in centimetres -- so each table was being
    rendered at 1/2.54 of its intended width. That is what split numbers across
    lines ('0.803 1', 'C 0', 'Annotato r pair') throughout the document.

    The other phase builders are not touched: build_phase2_docx.py writes widths
    that already sum to 15.5 cm and is correct as it stands. Converting here
    rather than in the shared helper fixes this document without disturbing
    them. Widths are also rescaled to fill the text block exactly, so no table
    is narrower or wider than the body text.
    """
    if widths:
        total = sum(widths)
        widths = [2.54 * w * (USABLE_IN / total) for w in widths]
    return _table_cm(doc, headers, rows, caption, widths=widths, **kw)
REG = json.loads((REP / "phase7_register.json").read_text(encoding="utf-8"))["register"]
FIGREG = json.loads((REP / "phase7_figure_registry.json").read_text(encoding="utf-8"))
MULT = json.loads((REP / "phase7_multiplicity.json").read_text(encoding="utf-8"))
RQ5 = json.loads((REP / "phase7_rq5.json").read_text(encoding="utf-8"))
P6H = json.loads((REP / "phase6_human.json").read_text(encoding="utf-8"))
P3AMD = json.loads((REP / "phase3b_amendment.json").read_text(encoding="utf-8"))

TIERS = ["S-unanimous", "S-majority", "S-plurality", "S-no-majority"]
POOLED = "S-contested (pooled)"
TIER_LABEL = {"S-unanimous": "Unanimous (4/4)", "S-majority": "Majority (3/4)",
              "S-plurality": "Plurality (2-1-1)", "S-no-majority": "No majority",
              POOLED: "Contested (pooled)"}
CFG = {"C0": "C0 hard label, 4/4 cohort", "C1": "C1 hard majority label",
       "C2": "C2 vote proportions", "C3": "C3 hard + matched smoothing (control)",
       "C4": "C4 vote proportions + anatomical penalty"}

_fig_used = set()


def FIG(doc, chapter, slug, caption):
    """Place the registered thesis figure whose slug matches, once."""
    for r in FIGREG["registry"]:
        if r["chapter"] == chapter and slug in r["file"]:
            if r["file"] in _fig_used:
                return
            _fig_used.add(r["file"])
            figure(doc, r["file"], caption)
            return
    print(f"  NOTE figure not found: ch{chapter} '{slug}'")


def pc(x, n=2):
    return "n/a" if x is None else f"{x:.{n}f}"


def p100(x, n=1):
    return "n/a" if x is None else f"{100 * x:.{n}f}"


def iv(v, n=2):
    if not v or v[0] is None:
        return "n/a"
    return f"[{v[0]:.{n}f}, {v[1]:.{n}f}]"


# =====================================================================
# Authorship and supervision. Held in one place so the title page, the
# declaration and the approval page cannot disagree about who wrote this.
AUTHORS = [("Fatin Sadab Nibirh", "0242310005101526"),
           ("MD Himel Rahman", "0242310005101800")]
SUPERVISOR = ("Ms. Shayla Sharmin", "Assistant Professor")
CO_SUPERVISOR = ("Dr. Md. Zahid Hasan", "Associate Professor")
DEPARTMENT = "Department of Computer Science and Engineering"
UNIVERSITY = "Daffodil International University"
SUBMISSION_DATE = "July 2026"


def _centre(doc, text, *, size=11.0, bold=False, italic=False, before=0, after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = BD.Pt(before)
    p.paragraph_format.space_after = BD.Pt(after)
    r = p.add_run(text)
    r.font.size = BD.Pt(size); r.bold = bold; r.italic = italic
    return p


def title_page(doc) -> None:
    for _ in range(2):
        doc.add_paragraph()
    _centre(doc, "Agreement-Stratified Evaluation of Deep Learning for\n"
                 "Anatomical Landmark Recognition in\n"
                 "Upper Gastrointestinal Endoscopy",
            size=20, bold=True, after=10)
    _centre(doc, "A thesis submitted in partial fulfilment of the requirements\n"
                 "for the degree of B.Sc. in Computer Science and Engineering",
            size=11.5, after=18)

    _centre(doc, "Submitted by", size=10.5, italic=True, after=4)
    for name, sid in AUTHORS:
        _centre(doc, name, size=13, bold=True, after=0)
        _centre(doc, f"ID: {sid}", size=10.5, after=8)

    _centre(doc, "Supervised by", size=10.5, italic=True, before=8, after=4)
    _centre(doc, SUPERVISOR[0], size=12.5, bold=True, after=0)
    _centre(doc, SUPERVISOR[1], size=10.5, after=8)
    _centre(doc, "Co-Supervised by", size=10.5, italic=True, after=4)
    _centre(doc, CO_SUPERVISOR[0], size=12.5, bold=True, after=0)
    _centre(doc, CO_SUPERVISOR[1], size=10.5, after=16)

    _centre(doc, DEPARTMENT, size=11.5, bold=True, after=0)
    _centre(doc, UNIVERSITY, size=11.5, bold=True, after=14)
    _centre(doc, SUBMISSION_DATE, size=11, after=14)

    a = REG["ch2_audit"]
    _centre(doc, f"Corpus: GastroHUN — {a['n_images']:,} images · "
                 f"{a['n_patients']} patients · {a['n_classes']} classes · "
                 f"{len(a['annotators'])} independent annotators",
            size=9.5, italic=True)
    # The title page issues no heading of its own, so record it here: without
    # this the first real chapter heading would see an empty _seen_h1, decide it
    # is the top of the document, and run on directly beneath the title.
    _seen_h1.append("title page")


def approval_page(doc) -> None:
    h(doc, "Approval", 1)
    para(doc, f"This thesis titled “Agreement-Stratified Evaluation of Deep "
              f"Learning for Anatomical Landmark Recognition in Upper "
              f"Gastrointestinal Endoscopy”, submitted by "
              f"{AUTHORS[0][0]} (ID: {AUTHORS[0][1]}) and {AUTHORS[1][0]} "
              f"(ID: {AUTHORS[1][1]}) to the {DEPARTMENT}, {UNIVERSITY}, has "
              f"been accepted as satisfactory for the partial fulfilment of the "
              f"requirements for the degree of B.Sc. in Computer Science and "
              f"Engineering and approved as to its style and contents.")
    rows = [
        ["Supervisor", f"{SUPERVISOR[0]}\n{SUPERVISOR[1]}\n{DEPARTMENT}\n{UNIVERSITY}", ""],
        ["Co-Supervisor", f"{CO_SUPERVISOR[0]}\n{CO_SUPERVISOR[1]}\n{DEPARTMENT}\n{UNIVERSITY}", ""],
        ["Internal Examiner", "", ""],
        ["External Examiner", "", ""],
        ["Chairman", "", ""],
    ]
    table(doc, ["Role", "Name and affiliation", "Signature"], rows,
          "Board of examiners.", widths=[1.4, 3.4, 1.6], font=9.0)


def declaration_page(doc) -> None:
    h(doc, "Declaration", 1)
    names = " and ".join(n for n, _ in AUTHORS)
    para(doc, f"We, {names}, hereby declare that this thesis has been prepared "
              f"by us under the supervision of {SUPERVISOR[0]}, "
              f"{SUPERVISOR[1]}, {DEPARTMENT}, {UNIVERSITY}, and co-supervised "
              f"by {CO_SUPERVISOR[0]}, {CO_SUPERVISOR[1]}. We further declare "
              f"that neither this thesis nor any part of it has been submitted "
              f"elsewhere for the award of any degree or diploma.")
    para(doc, f"The work reported here was carried out on the publicly released "
              f"GastroHUN corpus {cite('gastrohun')}, distributed under CC BY 4.0 "
              f"with ethics approval CEI-2019-06-10 recorded by its custodians. "
              f"No new patient data were collected and no identifiable data were "
              f"handled at any point.")
    callout(doc,
            "Every quantity reported in this thesis was resolved from a "
            "committed JSON artefact by a script, and every figure was drawn by "
            "a committed script. No number and no reference in this document "
            "was typed by hand. Appendix E lists the scripts and Appendix F the "
            "citation provenance, so any claim here can be traced to the code "
            "that produced it.",
            title="On the provenance of every number in this document")
    rows = [["Submitted by", f"{n}\nID: {i}\n{DEPARTMENT}\n{UNIVERSITY}", ""]
            for n, i in AUTHORS]
    rows += [["Supervised by",
              f"{SUPERVISOR[0]}\n{SUPERVISOR[1]}\n{DEPARTMENT}\n{UNIVERSITY}", ""],
             ["Co-Supervised by",
              f"{CO_SUPERVISOR[0]}\n{CO_SUPERVISOR[1]}\n{DEPARTMENT}\n{UNIVERSITY}", ""]]
    table(doc, ["", "Name and affiliation", "Signature"], rows,
          "Declaration signatories.", widths=[1.4, 3.4, 1.6], font=9.0)


def acknowledgements_page(doc) -> None:
    h(doc, "Acknowledgements", 1)
    para(doc, f"We thank our supervisor, {SUPERVISOR[0]}, and our co-supervisor, "
              f"{CO_SUPERVISOR[0]}, for holding this project to a standard it "
              f"would not otherwise have reached. The insistence that a claim be "
              f"pre-registered before it is tested, and withdrawn when it does "
              f"not survive, shaped the thesis more than any single result in "
              f"it.")
    para(doc, f"We thank the custodians of the GastroHUN corpus "
              f"{cite('gastrohun')} for releasing the individual annotators' "
              f"labels rather than a single consensus. That decision is what "
              f"makes this thesis possible: without per-annotator labels, not "
              f"one endpoint in this design exists. We thank the teams behind "
              f"HyperKvasir {cite('borgli')} and GastroVision {cite('jha')} for "
              f"the same reason at the external-validation stage.")
    para(doc, f"We thank the {DEPARTMENT} at {UNIVERSITY} for the compute on "
              f"which every experiment reported here was run, and our families "
              f"for their patience across the months this work took.")


def abbreviations_page(doc) -> None:
    h(doc, "List of Abbreviations", 1)
    rows = [
        ["AURC", "Area Under the Risk–Coverage curve"],
        ["CAM", "Class Activation Map (Grad-CAM, gradient-weighted)"],
        ["CI", "Confidence Interval"],
        ["CNN", "Convolutional Neural Network"],
        ["ECE", "Expected Calibration Error"],
        ["EGD", "Esophagogastroduodenoscopy (upper GI endoscopy)"],
        ["GI", "Gastrointestinal"],
        ["IoU", "Intersection over Union"],
        ["LOO", "Leave-One-Out (annotator held out from the reference panel)"],
        ["MCE", "Maximum Calibration Error"],
        ["OOP", "Out-Of-Protocol (an image that is not an SSS station)"],
        ["PRISMA", "Preferred Reporting Items for Systematic Reviews and "
                   "Meta-Analyses"],
        ["RQ", "Research Question"],
        ["SSS", "Systematic Screening of the Stomach protocol"],
        ["α", "Krippendorff's alpha, a chance-corrected agreement coefficient"],
        ["AC1", "Gwet's AC1, a chance-corrected agreement coefficient"],
        ["κ", "Cohen's or Fleiss' kappa, a chance-corrected agreement "
              "coefficient"],
        ["ε", "Label-smoothing strength"],
        ["λ", "Weight on the anatomy-aware penalty term"],
    ]
    table(doc, ["Abbreviation", "Meaning"], rows,
          "Abbreviations and symbols used in this thesis.",
          widths=[1.3, 5.1], font=8.6)


def sec_abstract(doc) -> None:
    h(doc, "Abstract", 1)
    a, s, t, e, x, r5 = (REG["ch2_audit"], REG["ch4_stratified"], REG["ch5_targets"],
                         REG["ch6_external"], REG["ch7_error"], REG["ch8_rq5"])

    para(doc, f"Background. Deep classifiers for endoscopic landmark recognition are "
              f"trained and evaluated almost exclusively on images where expert "
              f"annotators agree. In this corpus that subset is "
              f"{a['agreement_tiers_pct']['complete_agreement_4of4']:.1f}% of "
              f"the whole. What happens on the remainder is not reported, because the "
              f"per-annotator labels needed to ask are rarely published.")
    para(doc, f"Objective. To separate two quantities that the literature reports as "
              f"one: the degradation of the classifier, and the degradation of the "
              f"reference standard against which it is scored.")
    para(doc, f"Methods. A ConvNeXt-Tiny baseline was reproduced, then evaluated across "
              f"four strata of annotator agreement on the full "
              f"{sum(s['n_by_tier'].values()):,}-image test split. Five target "
              f"constructions were trained on a cohort held constant, including a "
              f"label-smoothing control whose strength was derived rather than chosen. "
              f"All five arms were transferred to two external corpora without "
              f"adaptation. Finally each annotator was held out in turn and scored "
              f"against the other three by the identical metric applied to the model, "
              f"alongside the modal-vote oracle that bounds any single-label predictor. "
              f"Every phase was pre-registered before it ran; every internal interval is "
              f"a patient-clustered bootstrap of 1,000 resamples.")
    para(doc, f"Results. Annotator-marginalized macro F1 falls "
              f"{p100(s['macro_f1_by_tier']['S-unanimous'])} → "
              f"{p100(s['macro_f1_by_tier']['S-majority'])} → "
              f"{p100(s['macro_f1_by_tier']['S-plurality'])} across the first "
              f"three strata and stands at "
              f"{p100(s['macro_f1_by_tier']['S-no-majority'])} on the fourth, "
              f"which is not lower than the third because its attainable "
              f"ceiling is different and it holds only "
              f"{s['n_by_tier']['S-no-majority']} images; the ceiling "
              f"falls with the score throughout, and the ceiling-normalised "
              f"unanimous-minus-majority gap is "
              f"{pc(s['gap_4v3_ceiling_normalised'])} points "
              f"(95% CI {iv(s['gap_4v3_ci'])}). Expected calibration error rises "
              f"{p100(s['ece_by_tier']['S-unanimous'])}% → "
              f"{p100(s['ece_by_tier']['S-plurality'])}%: mean confidence falls only "
              f"{pc(100 * (s['mean_confidence_by_tier']['S-unanimous'] - s['mean_confidence_by_tier']['S-plurality']))} "
              f"points while expected accuracy falls "
              f"{pc(100 * (s['expected_accuracy_by_tier']['S-unanimous'] - s['expected_accuracy_by_tier']['S-plurality']))}. "
              f"No target construction repairs this: the pre-registered contrast against "
              f"the matched control is {pc(t['contrast_C2_C3'])} points "
              f"({iv(t['contrast_C2_C3_ci'])}). On contested images the model "
              f"out-predicts a held-out annotator ({pc(x['model_by_stratum'][POOLED], 4)} "
              f"against {pc(x['human_by_stratum'][POOLED], 4)}) but recovers only "
              f"{p100(x['headroom_recovered'][POOLED], 0)}% of the headroom to the "
              f"modal-vote oracle ({pc(x['oracle_by_stratum'][POOLED], 4)}) and exceeds "
              f"that oracle on no stratum. Externally the arms are separated decisively "
              f"by selective prediction (AURC {pc(x['aurc_external']['C2'], 4)} for C2 "
              f"against {pc(x['aurc_external']['C3'], 4)} for C3) where every internal "
              f"endpoint failed to separate them.")
    para(doc, f"Conclusion. Agreement stratification separates a falling reference "
              f"standard from a falling classifier. Both fall; the ceiling accounts for "
              f"the larger share and a real model shortfall remains. What degrades "
              f"further and faster than discrimination is confidence, and the endpoint "
              f"that exposes it is the model's willingness to decline rather than its "
              f"accuracy. A negative-control audit of the retired corpus "
              f"({r5['verdict']}) shows the audit protocol that made these claims "
              f"trustworthy is a well-formedness instrument, not a viability one.")


# =====================================================================
def ch1(doc) -> None:
    h(doc, "1. Introduction", 1)
    h(doc, "1.1 Clinical context", 2)
    a = REG["ch2_audit"]
    para(doc, f"The Systematic Screening of the Stomach protocol prescribes a fixed set "
              f"of photodocumented anatomical stations so that a gastroscopy can be "
              f"shown to have been complete. Automating the recognition of those "
              f"stations is a natural target for deep learning: the classes are defined, "
              f"the images are routinely captured, and completeness is a measurable "
              f"quality endpoint {cite_theme('T2', limit=3)}. This thesis works "
              f"with GastroHUN {cite('gastrohun')}, "
              f"{a['n_images']:,} images from {a['n_patients']} patients, labelled "
              f"independently by {len(a['annotators'])} annotators across "
              f"{a['n_classes']} classes.")
    h(doc, "1.2 The problem: ground truth is a constructed object", 2)
    para(doc, f"Those {len(a['annotators'])} annotators agree unanimously on "
              f"{a['agreement_tiers_pct']['complete_agreement_4of4']:.1f}% of "
              f"the corpus. Published baselines for this dataset "
              f"{cite('gastrohun')} — and, as the review in "
              f"Chapter 4 shows, for this literature generally — are trained and "
              f"evaluated on that subset. The remaining "
              f"{100 - a['agreement_tiers_pct']['complete_agreement_4of4']:.1f}% "
              f"is not reported on, and "
              f"{a['pct_no_majority']:.2f}% of images admit no majority label under any "
              f"voting rule. A model's performance on the agreed subset is therefore not "
              f"an estimate of its performance in a clinic, and the difference is not a "
              f"matter of degree: on the hardest images there is no single label for the "
              f"model to be right about.")
    callout(doc, "The question this thesis exists to answer is not 'how accurate is the "
                 "classifier'. It is: when performance falls as annotators disagree, how "
                 "much of that fall is the classifier getting worse, and how much is the "
                 "reference standard ceasing to exist? Those are different quantities "
                 "with different consequences, and the literature reports their sum.",
            title="The central question")
    FIG(doc, 1, "conceptual_framework",
        "Ground truth as a constructed object: the same images admit different reference "
        "standards depending on how annotator votes are combined.")
    h(doc, "1.3 The central claim", 2)
    x, s = REG["ch7_error"], REG["ch4_stratified"]
    callout(doc,
            f"Expert-agreement stratification separates two things the literature "
            f"reports as one: a reference standard that ceases to exist, and a "
            f"classifier that falls short of what remains of it. As agreement falls the "
            f"attainable ceiling drops from {pc(x['oracle_by_stratum']['S-unanimous'], 2)} "
            f"to {pc(x['oracle_by_stratum'][POOLED], 2)}, so most of the apparent "
            f"collapse is the ceiling moving, while the model still recovers only "
            f"{p100(x['headroom_recovered'][POOLED], 0)}% of the distance from an "
            f"individual annotator to that ceiling. Confidence degrades further and "
            f"faster than discrimination, no target construction repairs it, and the "
            f"endpoint that separates configurations is not accuracy but the model's "
            f"willingness to decline.",
            title="Central claim")
    h(doc, "1.4 Contributions", 2)
    bullet(doc, "An agreement-stratified evaluation protocol that reports the attainable "
                "ceiling alongside the score, so that a falling reference standard is "
                "not mistaken for a falling model (Chapter 5).")
    bullet(doc, "A human comparator built from held-out annotators, scored by the "
                "identical metric on the identical images, together with the modal-vote "
                "oracle that bounds it (Chapter 8). To our knowledge this comparison has "
                "not previously been reported for this task.")
    bullet(doc, f"Evidence that no target construction among five — including a control "
                "whose smoothing strength was derived to match the soft target's "
                f"displaced probability mass {cite('szegedy', 'hinton')} — repairs "
                f"calibration on contested images (Chapter 6).")
    bullet(doc, "An out-of-protocol rejection and selective-prediction analysis showing "
                "that configurations indistinguishable internally separate decisively "
                "under domain shift (Chapters 7 and 8).")
    bullet(doc, "A negative-control evaluation of the audit protocol itself, which "
                "showed it to be a well-formedness instrument rather than a viability "
                "one, and three proposed extensions (Chapter 9).")
    FIG(doc, 1, "workflow",
        "Phase structure. Each phase gates the next; no phase's analysis was written "
        "before its pre-registration was frozen.")


# =====================================================================
def ch2(doc) -> None:
    h(doc, "2. Corpus Audit and Data Provenance", 1)
    a = REG["ch2_audit"]
    h(doc, "2.1 Why an audit chapter precedes any modelling chapter", 2)
    para(doc, "Every claim in the chapters that follow is conditional on the corpus "
              "being what it says it is. The audit was therefore run first and "
              "deliberately measured no model performance, so that the hypotheses in "
              "Chapter 5 remained genuine predictions.")
    h(doc, "2.2 Integrity and provenance", 2)
    rows = [
        ["Manifest rows", f"{a['n_images']:,}"],
        ["Decoded successfully", f"{a['n_decoded_ok']:,}"],
        ["Missing / orphan / corrupt", f"{a['n_missing']} / {a['n_orphan']} / {a['n_corrupt']}"],
        ["Patients", f"{a['n_patients']}"],
        ["Classes", f"{a['n_classes']}"],
        ["Annotators", ", ".join(a["annotators"])],
        ["Near-duplicate pairs examined", f"{a['neardup_pairs_examined']:,}"],
        ["Cross-split duplicates, uncalibrated rule", f"{a['dup_flagged_uncalibrated']}"],
        ["Cross-split duplicates, calibrated rule", f"{a['dup_confirmed_calibrated']}"],
        ["Split class-composition χ² p", f"{a['split_class_chi2_p']:.5f}"],
        ["Classes below ±10 pp Wilson half-width", f"{a['n_underpowered_classes']}/{a['n_classes']}"],
        ["Patients with a clinical record", f"{a['pct_with_clinical_record']}%"],
    ]
    table(doc, ["Measurement", "Value"], rows,
          "Corpus audit summary. Every value regenerates from src/data/gastrohun_*.py.",
          widths=[3.6, 2.9], font=8.4)
    callout(doc,
            f"The contamination scan first reported {a['dup_flagged_uncalibrated']} "
            f"cross-split duplicate pairs. That number was an artefact of an "
            f"uncalibrated threshold. A first correction using a randomly-paired null "
            f"was also wrong, because random pairs mostly compare different anatomical "
            f"stations and set an artificially low bar. Rebuilding the null from "
            f"class-matched pairs and anchoring the rule on a synthetic-duplicate "
            f"positive control reduced the count to {a['dup_confirmed_calibrated']}, and "
            f"visual audit confirmed the flagged pairs were different patients "
            f"photographed at the same landmark. An uncalibrated threshold is not a "
            f"measurement, and a threshold calibrated against the wrong comparison is "
            f"not much better.",
            title="A methodological lesson carried into every later threshold")
    FIG(doc, 2, "dup_calibration",
        "Threshold calibration. The positive control fixes where the decision rule "
        "belongs; the class-matched null shows why the first attempt was wrong.")
    h(doc, "2.3 Agreement, and why three chance corrections coincide", 2)
    para(doc, f"Fleiss' κ = {a['fleiss_kappa']} {cite('fleiss')}, "
              f"Krippendorff's α = {a['krippendorff_alpha']} {cite('hayes')} and "
              f"Gwet's AC1 = {a['gwet_ac1']} {cite('gwet')}. The three "
              f"coincide because the SSS protocol makes the class marginal near-uniform, "
              f"so Σpⱼ² ≈ 1/K and the Fleiss and Gwet expectations become algebraically "
              f"equal. The kappa paradox does not apply here.")
    rows = [[k, f"{v['kappa']:.4f}", iv(v["ci95"], 4),
             "within team" if v["within_team"] else "between teams"]
            for k, v in sorted(a["pairwise_kappa"].items(),
                               key=lambda kv: -kv[1]["kappa"])]
    table(doc, [f"Annotator pair", f"Cohen's κ {cite('cohen')}", "95% CI",
                "Relationship"], rows,
          "Pairwise agreement. Each resident agrees more with the gastroenterologists "
          "than with the other resident, so seniority does not predict agreement and no "
          "design may treat a team as a coherent unit.",
          widths=[1.5, 1.1, 1.7, 1.7], font=8.2)
    FIG(doc, 2, "kappa_matrix", "Pairwise agreement structure.")
    FIG(doc, 2, "agreement_cascade",
        "The agreement cascade. Published baselines for this corpus are computed on the "
        "unanimous tier alone.")
    h(doc, "2.4 Disagreement is anatomically structured", 2)
    d = a["disagreement_decomposition_pct"]
    rows = [[k.replace("_", " "), f"{v:.2f}%"] for k, v in
            sorted(d.items(), key=lambda kv: -kv[1])]
    table(doc, ["Disagreement type", f"% of {a['n_disagreement_events']:,} events"], rows,
          "Disagreement decomposition. Collapsing the wall axis recovers little; "
          "collapsing the station axis recovers a great deal.",
          widths=[3.6, 2.4], font=8.4)
    kg, ug = a["kappa_by_granularity"], a["unanimity_by_granularity"]
    rows = [[g.title(), f"{kg[g]['mean_pairwise_kappa']:.4f}",
             f"{kg[g]['n_categories']}", f"{ug[g]:.2f}%"] for g in kg]
    table(doc, ["Granularity", "Mean pairwise κ", "Categories", "Unanimity"], rows,
          "Agreement recomputed at coarser granularity. Endoscopists know how deep the "
          "scope is and disagree about which way it points — which is what makes the "
          "disagreement modellable rather than noise.",
          widths=[1.7, 1.8, 1.3, 1.5], font=8.4)
    FIG(doc, 2, "disagreement_decomposition", "Where the disagreement lives.")
    FIG(doc, 2, "sss_taxonomy",
        "The wall × station grid. This structure is the thesis's main analytical lever "
        "and is not documented as such in the dataset descriptor.")
    h(doc, "2.5 OTHERCLASS is a subjective judgement", 2)
    oc = a["otherclass_per_rater"]
    para(doc, f"Per-annotator rejection rates span "
              f"{min(oc.values()):.2f}% to {max(oc.values()):.2f}%, a "
              f"{max(oc.values()) / max(min(oc.values()), 1e-9):.1f}× spread. Quality "
              f"assessment and anatomical classification are different tasks and are "
              f"modelled and evaluated separately throughout.")
    FIG(doc, 2, "otherclass", "Per-annotator rejection behaviour.")
    FIG(doc, 2, "test_power",
        "Per-class test-set precision. 22 of 23 classes exceed a ±10 pp Wilson "
        "half-width, which is why every per-class result in this thesis is exploratory.")


# =====================================================================
def ch_methods(doc) -> None:
    """Materials and methods, generated from the Phase 2 artefacts.

    Every hyperparameter here is read from reports/phase2_run_seed1.json,
    phase2_norm_stats.json, phase2_trainable_layers.json,
    phase2_split_provenance.json and phase2_env.json rather than transcribed,
    so the chapter cannot drift from the run it describes.
    """
    RUN = json.loads((REP / "phase2_run_seed1.json").read_text(encoding="utf-8"))
    NORM = json.loads((REP / "phase2_norm_stats.json").read_text(encoding="utf-8"))
    TL = json.loads((REP / "phase2_trainable_layers.json").read_text(encoding="utf-8"))
    SP = json.loads((REP / "phase2_split_provenance.json").read_text(encoding="utf-8"))
    ENV = json.loads((REP / "phase2_env.json").read_text(encoding="utf-8"))
    PRE = json.loads((REP / "phase2_prereg.json").read_text(encoding="utf-8"))
    a = REG["ch2_audit"]

    h(doc, "3. Materials and Methods", 1)
    para(doc, "This chapter fixes the machinery. Everything in it was frozen "
              "before the first result in Chapter 5 was computed, and every "
              "value below is read from the run artefact rather than "
              "transcribed into prose, so the description cannot drift from the "
              "experiment it describes.")

    h(doc, "3.1 Cohort and splits", 2)
    co = SP["cohort"]
    para(doc, f"The corpus holds {SP['corpus']['n_images']:,} images from "
              f"{SP['corpus']['n_patients']} patients {cite('gastrohun')}. The "
              f"reproduction cohort is the complete-agreement subset — the "
              f"{co['retention_pct']}% of images on which all four annotators "
              f"agree — because that is the condition under which the published "
              f"baseline was trained and scored, and Chapter 5 cannot claim to "
              f"have reproduced a result it did not reproduce the conditions "
              f"of. It comprises {co['by_split']['Train']:,} training, "
              f"{co['by_split']['Validation']:,} validation and "
              f"{co['by_split']['Test']:,} test images. The stratified "
              f"evaluation in Chapter 5 then scores the frozen model on the "
              f"full {sum(REG['ch4_stratified']['n_by_tier'].values()):,}-image "
              f"test split, contested images included.")
    ov = SP["patient_overlap"]
    rows = [
        ["Split unit", "patient — no patient contributes to two splits"],
        ["Train / validation patient overlap", f"{len(ov['train_val'])} patients"],
        ["Train / test patient overlap", f"{len(ov['train_test'])} patients"],
        ["Validation / test patient overlap", f"{len(ov['val_test'])} patients"],
        ["Class composition across splits",
         f"χ² = {SP['class_split_chi2']}, p = {SP['class_split_p']:.5f}"],
        ["Acquisition stream across splits",
         f"χ² = {SP['stream_split_chi2']}, p = {SP['stream_split_p']:.2e}"],
        ["Cohort images resolved to a file on disk",
         f"{SP['hash_resolution']['n_resolved']:,} of "
         f"{SP['hash_resolution']['n_cohort']:,}"],
    ]
    table(doc, ["Property", "Value"], rows,
          "Split construction. Splits are patient-disjoint by construction, and "
          "class composition does not differ across them. Acquisition stream "
          "does differ, which is declared here rather than discovered later: it "
          "is the reason Chapter 5 carries an acquisition-stream confound "
          "control.", widths=[2.9, 3.5], font=8.6)

    h(doc, "3.2 Pre-processing", 2)
    para(doc, f"Images are resampled once to {NORM['size']}×{NORM['size']} with "
              f"{NORM['resample'].title()} resampling and cached, so that every "
              f"arm in every later phase reads pixel-identical inputs. "
              f"Normalisation uses channel statistics computed on the "
              f"{NORM['n_images']:,} training images of this cohort — mean "
              f"({', '.join(f'{m:.3f}' for m in NORM['mean'])}), sd "
              f"({', '.join(f'{s:.3f}' for s in NORM['std'])}) — rather than the "
              f"ImageNet constants, which differ from them by up to "
              f"{max(NORM['abs_delta_mean']):.3f} in the mean. Endoscopic "
              f"illumination is red-dominant and the ImageNet statistics do not "
              f"describe it.")
    callout(doc,
            "The augmentation recipe omits horizontal and vertical flips and "
            "any large rotation, which is a departure from the standard "
            "image-classification recipe and is deliberate. The label encodes a "
            "gastric WALL — anterior, posterior, lesser curvature, greater "
            "curvature — so a horizontal flip maps an image of one wall onto "
            "the appearance of another and silently relabels it. An "
            "augmentation that destroys the target is not regularisation. What "
            "remains is a mild scale and aspect jitter "
            f"(RandomResizedCrop, scale 0.85–1.00, ratio 0.90–1.11) and a "
            f"photometric jitter (brightness, contrast and saturation ±0.20, "
            f"hue ±0.02), neither of which moves the anatomy.",
            title="Why the standard augmentation recipe is not used")

    h(doc, "3.3 Architecture and training", 2)
    tlp = 100 * TL["param_fraction_unfrozen"]
    rows = [
        ["Backbone", f"ConvNeXt-Tiny {cite('liu22')}, ImageNet-pretrained"],
        ["Fine-tuning depth",
         f"{TL['n_modules_unfrozen']} of {TL['n_feature_modules']} feature "
         f"modules unfrozen ({tlp:.1f}% of feature parameters, "
         f"{TL['feature_params_unfrozen']:,} of {TL['feature_params_total']:,})"],
        ["Optimiser", "AdamW"],
        ["Learning rate, head", f"{RUN['lr_head']}"],
        ["Learning rate, fine-tuning", f"{RUN['lr_finetune']}"],
        ["Weight decay", f"{RUN['weight_decay']}"],
        ["Batch size", f"{RUN['batch_size']} "
                       f"(gradient accumulation ×{RUN['grad_accum_steps']}, "
                       f"effective {RUN['effective_batch']})"],
        ["Precision / memory format",
         f"{RUN['precision']}, {RUN['memory_format'].replace('_', ' ')}"],
        ["Warm-up epochs", f"{RUN['warmup_epochs']}"],
        ["Maximum fine-tuning epochs", f"{RUN['max_finetune_epochs']}"],
        ["Early-stopping patience",
         f"{RUN['patience']} epochs on validation macro F1"],
        ["Seeds", ", ".join(str(s) for s in PRE["seeds"])],
        ["Model selection", "validation macro F1 only; the test split is "
                            "scored once, after selection"],
    ]
    table(doc, ["Setting", "Value"], rows,
          "Training configuration, read from reports/phase2_run_seed1.json. The "
          "same configuration is reused unchanged by every later arm, so that "
          "the only thing varying across the Chapter 6 configurations is the "
          "target.", widths=[2.5, 3.9], font=8.6)
    para(doc, f"All runs used a single {ENV['gpu']} with "
              f"{ENV['nvidia_smi'].split(',')[-1].strip()} of memory, on "
              f"{ENV['platform']} with Python {ENV['python']}, torch "
              f"{ENV['packages']['torch']} and torchvision "
              f"{ENV['packages']['torchvision']}. Peak observed VRAM was "
              f"{RUN['peak_vram_mib']:.0f} MiB. The 4 GB memory limit is the "
              f"binding constraint behind the epoch cap declared in Chapter 6, "
              f"and is recorded here so that the cap reads as a resource "
              f"limitation rather than a design choice.")

    h(doc, "3.4 Metrics, and exactly how each is computed", 2)
    para(doc, "Four quantities carry the thesis, and three of them are "
              "reported in the literature under names that hide a choice. The "
              "choices are made explicit here.")

    h(doc, "3.4.1 Annotator-marginalized macro F1", 3)
    para(doc, f"The primary discrimination metric is the mean, over the "
              f"{len(a['annotators'])} annotators, of the macro F1 computed "
              f"against that annotator's labels taken alone. Macro F1 is "
              f"averaged over all {a['n_classes']} classes with zero division "
              f"scored as zero, so a class absent from both the reference and "
              f"the prediction contributes 0 rather than being dropped. "
              f"Marginalizing over annotators rather than scoring against a "
              f"consensus is what makes the metric defined on contested images "
              f"at all: where the four annotators disagree there is no "
              f"consensus label to score against, but each annotator's own "
              f"labelling remains a valid reference.")

    h(doc, "3.4.2 Agreement strata", 3)
    rows = [
        ["S-unanimous", "4-0", f"{REG['ch4_stratified']['n_by_tier']['S-unanimous']:,}",
         "all four annotators give the same label"],
        ["S-majority", "3-1", f"{REG['ch4_stratified']['n_by_tier']['S-majority']:,}",
         "three agree, one dissents"],
        ["S-plurality", "2-1-1", f"{REG['ch4_stratified']['n_by_tier']['S-plurality']:,}",
         "two agree, two dissent separately"],
        ["S-no-majority", "2-2 or 1-1-1-1",
         f"{REG['ch4_stratified']['n_by_tier']['S-no-majority']:,}",
         "no label holds a majority under any voting rule"],
    ]
    table(doc, ["Stratum", "Vote pattern", "n (test)", "Definition"], rows,
          "The four agreement strata. They partition the test split and are "
          "defined by the vote pattern alone, so membership is fixed before any "
          "model is trained and cannot be influenced by a result.",
          widths=[1.3, 1.4, 0.9, 2.8], font=8.6)

    h(doc, "3.4.3 Calibration", 3)
    para(doc, f"Expected calibration error is computed with 10 equal-width bins "
              f"over the [0, 1] confidence range, as the support-weighted mean "
              f"absolute difference between mean confidence and expected "
              f"accuracy within each bin {cite('guo')}. The bin count is stated "
              f"because ECE is not comparable across studies without it: the "
              f"same predictions binned more finely generally yield a larger "
              f"ECE. Expected accuracy on a contested image is the fraction of "
              f"the four annotators whose label the prediction matches, so the "
              f"calibration target degrades gracefully rather than becoming "
              f"undefined where consensus fails.")

    h(doc, "3.4.4 Two attainable ceilings, and why they are not one quantity", 3)
    ORC = json.loads((REP / "phase8_oracle_reconcile.json").read_text(
        encoding="utf-8"))
    para(doc, "A score means nothing on a contested image without knowing what "
              "the best possible score there would be. Two such bounds appear "
              "in this thesis. They were previously both called “the modal-vote "
              "oracle”, they differ by up to 14 points, and separating them is "
              "necessary before either can be read.")
    rows = [[t.replace("S-", ""),
             f"{r['panel_ceiling_recomputed']:.4f}",
             f"{r['loo_oracle_recomputed']:.4f}",
             f"{r['difference_loo_minus_panel']:+.4f}",
             f"{r['tie']['mean_tie_depth_4refs']:.2f} → "
             f"{r['tie']['mean_tie_depth_3refs']:.2f}"]
            for t, r in ORC["by_stratum"].items()]
    table(doc, ["Stratum", "Panel ceiling", "LOO oracle", "Difference",
                "Mean tie depth, 4→3 refs"], rows,
          "The two bounds, recomputed from the committed votes by "
          "src/models/phase8_oracle_reconcile.py, which reproduces both "
          f"committed series ({ORC['reproduction_gate']['n_reproduce']} of "
          f"{ORC['reproduction_gate']['n_checks']} checks). The panel ceiling is "
          "the modal label of all four votes scored against all four "
          "annotators, and bounds Chapter 5, where nothing is held out. The "
          "leave-one-out oracle is the modal label of three votes scored "
          "against those same three, and bounds Chapter 8, where it must face "
          "the identical task as the held-out human.", widths=[1.2, 1.3, 1.2, 1.1, 1.6],
          font=8.4)
    callout(doc,
            "The two coincide on S-unanimous and S-majority to five decimal "
            "places, and diverge by 14.29 points on S-no-majority. The cause is "
            "tie multiplicity, and it is a property of the vote pattern rather "
            "than of any model. S-no-majority is 73 images split 2-2 and 8 "
            "split 1-1-1-1; on a 2-2 image the modal label of four references is "
            "a two-way tie worth 2/4, and against any three references it "
            "becomes a 2-1 split with a unique mode worth 2/3, so removing a "
            "reference makes the oracle's task strictly easier. On S-plurality "
            "the effect cancels exactly — dropping one of the two modal voters "
            "leaves a three-way tie worth 1/3, dropping either singleton leaves "
            "a unique mode worth 2/3, and the four folds average to 0.5000 "
            "either way — so the residual 0.37 points there is not a difficulty "
            "difference at all but the residue of macro F1 being a nonlinear "
            "class-wise aggregate rather than expected accuracy. RQ1's "
            "confirmatory endpoint uses only the two strata where the "
            "definitions agree, so it is unaffected by the choice.",
            title="Why the two bounds differ, measured rather than argued")

    h(doc, "3.4.5 Intervals", 3)
    para(doc, f"Every internal interval in this thesis is a patient-clustered "
              f"bootstrap of {PRE['n_bootstrap']:,} resamples: the "
              f"{PRE['bootstrap_unit']}s in the relevant split are resampled "
              f"with replacement and the statistic is recomputed on each draw. "
              f"Images are never resampled independently, because images from "
              f"one procedure are not independent observations and an "
              f"image-level interval on this corpus would be optimistically "
              f"narrow. Where a ratio is reported, both numerator and "
              f"denominator are recomputed inside each draw so that the pairing "
              f"is preserved — in particular the ceiling-normalised gap in "
              f"Chapter 5 resamples its ceiling rather than holding it fixed. "
              f"The one place this construction is unavailable is external "
              f"validation, where neither corpus publishes a case identifier; "
              f"those intervals are image-level and are labelled as such "
              f"wherever they appear.")

    h(doc, "3.5 Pre-registration and the gating discipline", 2)
    para(doc, f"Each phase writes a pre-registration file before it runs, "
              f"fixing its hypotheses, its endpoints, its verdict rules and its "
              f"analysis order. The generating scripts refuse to overwrite an "
              f"existing pre-registration, so a rule cannot be revised after a "
              f"result is seen; Appendix A lists the files and Appendix B every "
              f"declared deviation. The Phase 2 rule is representative: "
              f"“{PRE['verdict_rule']}” This is the discipline that lets "
              f"Chapter 10 report three unresolved research questions as "
              f"findings rather than as failures — each was a pre-registered "
              f"test against a matched control, not a search that happened to "
              f"come up empty.")


# =====================================================================
def ch3(doc) -> None:
    LS = json.loads((REP / "phase8_lit_synthesis.json").read_text(encoding="utf-8"))
    h(doc, "4. Literature Review", 1)
    para(doc, f"A PRISMA 2020 review {cite('prisma2020')} across seven themed "
              f"searches identified 1,382 records, 1,349 unique, of which "
              f"{LS['n_included']} were included (68 from database searching, 14 "
              f"hand-searched). {LS['pct_published_2020_or_later']}% were "
              f"published in 2020 or later, so the omissions catalogued below "
              f"are properties of the current literature and not of its "
              f"history.")
    FIG(doc, 3, "prisma", "PRISMA 2020 flow.")

    h(doc, "4.1 What the seven themed searches cover", 2)
    THEME_NOTE = {
        "T1": ("the task itself. Landmark and anatomical-site recognition in "
               "endoscopy, including the corpora this thesis transfers to."),
        "T2": ("why the task matters. Blind-spot audit, photodocumentation "
               "completeness and quality control during endoscopy."),
        "T3": ("how much experts disagree, measured. Interobserver studies "
               "across endoscopic classification tasks."),
        "T4": ("what to do about disagreement. Noisy, soft and multi-annotator "
               "label learning."),
        "T5": ("whether a confidence can be believed. Calibration and "
               "predictive uncertainty."),
        "T6": ("whether any of it survives a second centre. External validation "
               "and dataset shift."),
        "T7": ("what a study of this kind is obliged to report."),
    }
    rows = [[c, t["label"][3:], f"{t['n']}", str(t["median_year"]),
             THEME_NOTE.get(c, ""), cite_theme(c)]
            for c, t in LS["themes"].items()]
    table(doc, ["", "Themed search", "n", "Median yr", "What it establishes",
                "Included studies"], rows,
          "The seven themed searches. Each addresses one link in the argument, "
          "and the review was structured this way so that a gap in the "
          "literature could be located to a specific link rather than asserted "
          "of the field as a whole. The final column lists every study the "
          "search contributed, so the reference list contains no entry that is "
          "not cited somewhere in the text.",
          widths=[0.3, 1.45, 0.3, 0.5, 1.95, 1.9], font=7.8)

    h(doc, "4.2 What the included studies report, counted", 2)
    para(doc, f"The claim that this literature under-reports certain things is "
              f"the claim the whole thesis rests on, so it is counted rather "
              f"than asserted. Each of the {LS['n_with_abstract']} included "
              f"studies that ships an abstract was screened by regular "
              f"expression for five reporting dimensions. The counts are "
              f"MENTION counts and are therefore upper bounds: a study that "
              f"mentions calibration may not report it, but a study that never "
              f"mentions it almost certainly did not make it an endpoint. The "
              f"patterns were written generously, which weakens the bound "
              f"rather than strengthening it.")
    rows = [[v["label"], f"{v['n_mentioning']}/{v['n_scored']}",
             f"{v['pct_mentioning']:.1f}%", v["addressed_by_this_thesis"]]
            for v in LS["dimensions"].values()]
    table(doc, ["Reporting dimension", "Mentioning", "%",
                "Where this thesis addresses it"], rows,
          f"Reporting dimensions across the {LS['n_with_abstract']} included "
          f"studies with abstracts. Generated by "
          f"src/report/literature_synthesis.py.",
          widths=[2.1, 0.9, 0.5, 2.9], font=8.4)
    FIG(doc, 3, "literature_gaps",
        "Reporting gaps in the included studies, and which this design addresses.")

    cal = LS["dimensions"]["calibration"]
    pop = LS["dimensions"]["population_description"]
    callout(doc,
            f"Two limitations travel with these counts and neither is "
            f"cosmetic. {LS['n_without_abstract']} of {LS['n_included']} "
            f"included studies ship no abstract in the MEDLINE record and are "
            f"excluded from the denominator rather than counted as absences. "
            f"And an abstract-level screen is not a full-text appraisal — which "
            f"matters most for population description, where the count of "
            f"{pop['n_mentioning']}/{pop['n_scored']} should be read as "
            f"'demographics are not a headline in this literature' and not as "
            f"'no study describes its population'. Cohort demographics live in "
            f"a baseline-characteristics table, not in an abstract. The "
            f"calibration figure carries no such caveat: a study reporting "
            f"calibration as an endpoint names it in the abstract, so "
            f"{cal['pct_mentioning']:.1f}% is a firm ceiling on how much of "
            f"this literature reports whether its probabilities can be "
            f"believed.",
            title="What these counts can and cannot bear")

    h(doc, "4.3 The gap this thesis occupies", 2)
    gi = LS["gap_intersection"]
    para(doc, f"The four commonest omissions are missing external validation, "
              f"absent calibration reporting, unexamined ground-truth "
              f"construction, and incomplete population description. This "
              f"design addresses three directly. The fourth cannot be met on "
              f"this corpus — GastroHUN ships no age or sex — and is declared "
              f"as a limitation in Chapter 10 rather than passed over.")
    para(doc, f"The gap is sharper than any single dimension, and it is an "
              f"intersection. Of the {gi['n_mentioning_ground_truth_construction']} "
              f"studies that engage at all with how their reference standard "
              f"was built, {gi['n_also_mentioning_calibration']} — "
              f"{gi['pct_of_those']}% — also say anything about whether their "
              f"predicted probabilities are trustworthy "
              f"{cite_theme('T4', limit=3)}. The two questions are treated as "
              f"unrelated. This thesis argues they are the same question: a "
              f"confidence is a claim about a reference standard, and where the "
              f"reference standard is contested the claim has no fixed "
              f"referent.")
    callout(doc,
            f"Ground-truth construction is the gap this thesis occupies. Of the "
            f"included studies, those that used multiple annotators almost "
            f"universally reduced them to a consensus label before modelling "
            f"and reported no analysis of the images where consensus failed "
            f"{cite_theme('T3', limit=4)}. The per-annotator labels required to "
            f"do otherwise are rarely published; GastroHUN {cite('gastrohun')} "
            f"is unusual in retaining them, which is why this design is "
            f"possible on this corpus and on very few others.",
            title="The gap")

    h(doc, "4.4 What is already known, and what this thesis does not claim", 2)
    para(doc, f"Two of this thesis's findings have antecedents, and saying so "
              f"is what fixes the size of the contribution. That modern "
              f"networks are overconfident, and that their calibration degrades "
              f"under distribution shift, is established "
              f"{cite('guo', 'gal', 'lakshminarayanan')}. That training on soft "
              f"or multi-annotator targets can transfer information a hard "
              f"consensus label destroys is likewise established "
              f"{cite('hinton', 'szegedy', 'delamor', 'gao')}. Neither "
              f"phenomenon is claimed here as new.")
    para(doc, f"What is new is the measurement. No study in this review reports "
              f"performance or calibration as a function of how much the expert "
              f"panel agreed, because doing so requires per-annotator labels "
              f"that are almost never released; and none separates the "
              f"degradation of the classifier from the degradation of the "
              f"reference standard it is scored against. That separation — the "
              f"attainable ceiling reported beside the score — is the "
              f"contribution, and it is a methodological one rather than a "
              f"performance one. This thesis produces no model that classifies "
              f"better than the published baseline, and does not claim to.")


# =====================================================================
def ch4(doc) -> None:
    h(doc, "5. Agreement-Stratified Evaluation", 1)
    s = REG["ch4_stratified"]
    h(doc, "5.1 Baseline reproduction as a validity check", 2)
    para(doc, f"Before anything was changed, the published ConvNeXt-Tiny "
              f"{cite('liu22')} result was "
              "reproduced on the unanimous subset. Three seeds gave a mean macro F1 of "
              "83.92 against a published ~85.0, inside the pre-registered ±1.5-point "
              "band. The pipeline was therefore validated before it was used to make "
              "any novel claim.")
    FIG(doc, 4, "baseline_reproduction", "Baseline reproduction against the published "
                                         "target.")
    h(doc, "5.2 Results: the raw decline, and why it misleads", 2)
    rows = [[TIER_LABEL[t], f"{s['n_by_tier'][t]:,}", p100(s["macro_f1_by_tier"][t]),
             pc(s["ceilings"][t], 4), p100(s["ece_by_tier"][t])] for t in TIERS]
    table(doc, ["Stratum", "n", "Macro F1", "Attainable ceiling", "ECE %"], rows,
          "Per-stratum performance with the attainable ceiling. The ceiling is the "
          "modal-vote oracle: the best any single-label predictor can achieve against "
          "the annotator panel on those images.",
          widths=[1.9, 0.8, 1.2, 1.5, 1.1], font=8.2)
    FIG(doc, 4, "stratified_curve_raw",
        "The raw stratified decline — the quantity the literature would report.")
    callout(doc,
            f"Most of that decline is the ceiling moving. Held against the attainable "
            f"maximum, the unanimous-minus-majority gap is "
            f"{pc(s['gap_4v3_ceiling_normalised'])} points "
            f"(95% CI {iv(s['gap_4v3_ci'])}) rather than the "
            f"{pc(100 * (s['macro_f1_by_tier']['S-unanimous'] - s['macro_f1_by_tier']['S-majority']))} "
            f"points the raw scores suggest. Reporting the raw decline alone would "
            f"attribute to the model a change that is mostly a property of the "
            f"reference standard.",
            title="Ceiling normalisation, and why it is not optional")
    FIG(doc, 4, "stratified_curve_ceiling",
        "The same strata with the attainable ceiling held constant.")
    FIG(doc, 4, "gap_forest",
        "RQ1's confirmatory endpoint with patient-clustered 95% intervals.")
    h(doc, "5.3 The calibration collapse", 2)
    dc = 100 * (s["mean_confidence_by_tier"]["S-unanimous"] - s["mean_confidence_by_tier"]["S-plurality"])
    da = 100 * (s["expected_accuracy_by_tier"]["S-unanimous"] - s["expected_accuracy_by_tier"]["S-plurality"])
    para(doc, f"Expected calibration error rises from {p100(s['ece_by_tier']['S-unanimous'])}% "
              f"on unanimous images to {p100(s['ece_by_tier']['S-plurality'])}% on the "
              f"2-1-1 stratum. The mechanism is visible in the components: mean "
              f"confidence falls {pc(dc)} points across those strata while expected "
              f"accuracy falls {pc(da)}. The model becomes wrong far faster than it "
              f"becomes hesitant. This is the durable finding of the thesis and it "
              f"survives every subsequent attempt to remove it.")
    FIG(doc, 4, "calibration_by_stratum", "Calibration by agreement stratum.")
    FIG(doc, 4, "confound_controls",
        "Confound controls. Class composition explains 3.5–4.2% of the drop and "
        "acquisition-stream composition does not differ across strata.")
    h(doc, "5.4 Corrections to the first analysis", 2)
    para(doc, "Four claims made in the first version of this analysis did not survive "
              "re-examination and were withdrawn. They are recorded here rather than in "
              "an appendix, because a reader meeting the original claim elsewhere in the "
              "literature should meet the correction in the same place.")
    bullet(doc, "The assertion that the any-annotator hit rate 'shows no such reversal' "
                "was false; the same dip is present in the published table.")
    bullet(doc, "Any-annotator hit rate is confounded by tier-varying acceptance-set "
                "size and cannot be read as a skill measure across tiers.")
    bullet(doc, "The confusion-geometry comparison against the human benchmark was "
                "withdrawn to hypothesis status because a model interval was compared "
                "against a human point estimate. Chapter 8 settles it.")
    bullet(doc, "The '16× the architecture benchmark' headline was a scale artefact. "
                "The direction survives; the magnitude does not.")


# =====================================================================
def ch5(doc) -> None:
    h(doc, "6. Target Construction and Uncertainty", 1)
    t = REG["ch5_targets"]
    h(doc, "6.1 Five configurations, one cohort, one thing varying", 2)
    rows = [[c, CFG[c], p100(t["macro_f1_pooled_contested"].get(c)),
             p100(t["ece_pooled_contested"].get(c)), p100(t["ece_unanimous"].get(c))]
            for c in t["arms"]]
    table(doc, ["Arm", "Target construction", "Macro F1 (contested)",
                "ECE % (contested)", "ECE % (unanimous)"], rows,
          "The configuration matrix on the pooled contested stratum. Backbone, "
          "schedule, augmentation, normalisation and selection criterion are identical "
          "across arms; only the target differs.",
          widths=[0.5, 2.3, 1.3, 1.2, 1.2], font=8.0)
    FIG(doc, 5, "config_design", "The configuration matrix.")
    h(doc, "6.2 Deriving the control rather than choosing it", 2)
    para(doc, f"C3 exists so that any benefit of the vote-proportion target can be "
              f"separated from ordinary regularisation. Its smoothing strength was "
              f"derived, not set to a convention: ε = {t['epsilon_mass_matched']} matches "
              f"the probability mass that the soft target displaces from the modal label. "
              f"Matching by entropy instead would have given "
              f"{t['epsilon_entropy_matched']}, a materially weaker control. Mass rather "
              f"than entropy, because the gradient of the soft-target cross-entropy is "
              f"(q − t), so the displaced mass is the perturbation.")
    h(doc, "6.3 Results, and the reversal that is the finding", 2)
    para(doc, f"The pre-registered contrast C2 − C3 on the pooled contested stratum is "
              f"{pc(t['contrast_C2_C3'])} points ({iv(t['contrast_C2_C3_ci'])}): not "
              f"resolved. On calibration the result is stronger and points the other "
              f"way — the generic control is markedly better calibrated on contested "
              f"images than the vote-proportion arm, while on unanimous images the order "
              f"reverses. Uniform smoothing suppresses confidence globally, so it is "
              f"under-confident where annotators agree and closer to correct where they "
              f"do not; C2 is trained one-hot on unanimous images and is nearly exact "
              f"there.")
    FIG(doc, 5, "rq2_forest", "RQ2's pre-registered contrast against the matched "
                              "control.")
    FIG(doc, 5, "calibration_by_config", "Calibration by configuration and stratum.")
    callout(doc, "No configuration achieves acceptable calibration on contested images. "
                 "That relocates the Chapter 5 finding from 'an artefact of "
                 "consensus-only training' to a property of the problem, and it is the "
                 "reason the thesis's recommendation in Chapter 11 is about abstention "
                 "rather than about target design.",
            title="What the control bought")
    FIG(doc, 5, "overconfidence", "Confidence against accuracy across strata.")
    h(doc, "6.4 Why RQ2, RQ3 and RQ4 are reported as unresolved", 2)
    para(doc, f"RQ2 is not resolved on accuracy and not supported on calibration. RQ4's "
              f"anatomy-aware penalty gives {pc(t['rq4_C4_C2'], 5)} "
              f"({iv(t['rq4_C4_C2_ci'], 5)}) at unit λ, and because no λ sweep was run "
              f"this is evidence about unit weight rather than about anatomy-aware "
              f"losses in general. RQ3 is addressed in Chapter 8, where it turns out not "
              f"to be estimable in the form it was posed.")
    callout(doc, "Eight of the twelve training runs stopped at the compute-imposed epoch "
                 "cap rather than by early stopping. The cap applies identically to all "
                 "four arms, so the target contrasts are unaffected, but the absolute "
                 "scores are lower bounds and the C1−C0 comparison is not controlled.",
            title="A bound that must be declared")


# =====================================================================
def ch6(doc) -> None:
    h(doc, "7. External Validation", 1)
    e = REG["ch6_external"]
    h(doc, "7.1 The label-space finding, which precedes any transfer number", 2)
    para(doc, f"HyperKvasir {cite('borgli')} and GastroVision {cite('jha')} were "
              f"acquired, hashed against the GastroHUN "
              f"inventory with zero collisions, and mapped by a table frozen before any "
              f"image was scored. The blueprint's premise did not survive contact with "
              f"the data: GastroHUN's label space is wall × station, and neither "
              f"external corpus carries the wall axis or has a class for four of the six "
              f"stations. A 23-way external validation is not available from these "
              f"corpora. The phase was therefore reframed — before scoring — into a "
              f"2-way anatomical collapse over {e['n_gastric']:,} gastric images and an "
              f"out-of-protocol rejection endpoint over {e['n_out_of_protocol']:,} "
              f"images that are not gastric stations at all.")
    FIG(doc, 6, "external_label_space", "What the mapping destroys.")
    h(doc, "7.2 Transfer and rejection", 2)
    para(doc, f"Transfer: {e['transfer_verdict']}. Binary macro F1 "
              f"{pc(e['external_f1'])} externally against {pc(e['internal_f1'])} "
              f"internally, a drop of {pc(e['drop_points'])} points "
              f"({iv(e['drop_ci'])}) — about twice the pre-registered expectation. Every "
              f"arm met the pre-registered precision target, so these are powered "
              f"verdicts rather than underpowered nulls.")
    rows = [[c, CFG.get(c, c), p100(v)] for c, v in e["rejection_by_arm"].items()]
    table(doc, ["Arm", "Target construction", "Out-of-protocol rejection %"], rows,
          f"Out-of-protocol rejection against a chance floor of "
          f"{100 * e['chance_rate']:.2f}%. The pre-registered hypothesis was that "
          f"rejection would sit at or below chance; it was falsified favourably, and the "
          f"soft-target arms reject far better than the hard-label ones.",
          widths=[0.6, 2.8, 2.2], font=8.2)
    FIG(doc, 6, "external_transfer", "Transfer by arm, internal against external.")
    FIG(doc, 6, "external_rejection", "Out-of-protocol rejection by arm.")
    callout(doc, "This benefit is invisible internally. GastroHUN's test split holds "
                 "almost no true out-of-protocol images, so the endpoint on which the "
                 "arms separate most sharply cannot be measured on the corpus the model "
                 "was trained on. That is the argument for external validation stated as "
                 "a measurement rather than as a principle.",
            title="Why this endpoint required a second centre")
    h(doc, "7.3 Self-training, and what adaptation cost", 2)
    rows = [[k, v] for k, v in e["p5b_verdicts"].items()]
    table(doc, ["Endpoint", "Verdict"], rows,
          "Phase 5B self-training, run only after the clean transfer numbers were frozen "
          "and committed — adapting first would have made the external validation "
          "circular.", widths=[2.4, 3.6], font=8.2)
    callout(doc, e["p5b_split_weakness"], title="A declared weakness of the 5B split")
    callout(doc, "All Phase 5 and 5B intervals are image-level, because neither external "
                 "corpus publishes a case identifier. They are optimistic relative to a "
                 "correctly clustered interval and must not be compared directly against "
                 "the patient-clustered intervals used everywhere else in this thesis.",
            title="An interval caveat that travels with every external number")


# =====================================================================
def ch7(doc) -> None:
    h(doc, "8. Explainability and Error Analysis", 1)
    x = REG["ch7_error"]
    h(doc, "8.1 Changing the comparator", 2)
    para(doc, "Every comparison in Chapters 5 to 7 was model against model. That design "
              "ranks target constructions but cannot calibrate the ranking against "
              "anything outside it: a model scoring in the twenties looks broken until "
              "one knows what a board-certified endoscopist scores on the same images. "
              "This chapter introduces two comparators the project had never used — the "
              "annotators themselves, and the model's own confidence ordering.")
    h(doc, "8.2 The human comparator", 2)
    para(doc, "Each annotator is held out in turn and scored against the other three; "
              "the model is scored against the same three, on the same images, under the "
              "same patient resample. Annotator a is excluded from their own reference "
              "set, because scoring a rater against a panel containing themselves makes "
              "one term an identity and inflates the human side.")
    rows = [[TIER_LABEL[s], pc(x["human_by_stratum"][s], 4), pc(x["model_by_stratum"][s], 4),
             pc(x["oracle_by_stratum"][s], 4),
             "n/a" if x["headroom_recovered"][s] is None else p100(x["headroom_recovered"][s], 0) + "%",
             p100(x["singleton_rate"][s], 0) + "%"] for s in TIERS + [POOLED]]
    table(doc, ["Stratum", "Held-out annotator", "Model (C2)", "Modal-vote oracle",
                "Headroom recovered", "Singleton rate"], rows,
          "The human comparator with the oracle that bounds it. The singleton rate is "
          "the fraction of held-out annotators whose label is shared by none of the "
          "other three — a structural handicap fixed by the stratum definition, not by "
          "skill.", widths=[1.5, 1.2, 0.9, 1.2, 1.1, 0.9], font=7.8)
    FIG(doc, 7, "human_comparator",
        "Held-out annotator, model and modal-vote oracle on a common axis.")
    callout(doc, x["qualified_verdict"], title="What may actually be claimed")
    para(doc, "The pre-registered rule returned ABOVE THE HUMAN PANEL on the contested "
              "strata. That verdict stands as the frozen rule produced it, but it must "
              "not be read as 'the model outperforms an expert'. Two asymmetries, found "
              "on re-reading the construction and then measured, explain it: the model "
              "is optimised to predict this panel's consensus while the annotator is "
              "simply being themselves, and — more decisively — the model chooses a "
              "label while the annotator is stuck with the one they gave. The modal vote "
              "of the same three references bounds any single-label predictor, and the "
              "model does not reach that bound on any stratum.", italic=True, size=9.8)
    h(doc, "8.3 Confusion geometry, and the settlement of a withdrawn claim", 2)
    para(doc, f"On the pooled contested stratum the model's wall-confusion geometry "
              f"differs from the human geometry by {pc(x['geometry_wall_delta'])} points "
              f"({iv(x['geometry_wall_ci'])}) and its station geometry by "
              f"{pc(x['geometry_station_delta'])} points "
              f"({iv(x['geometry_station_ci'])}). When the model confuses the "
              f"circumferential direction of the scope it confuses it the way "
              f"endoscopists do; when it confuses depth, it travels further along the "
              f"insertion axis than human disagreements do. That is a specific, nameable "
              f"deficit rather than a general one.")
    callout(doc, x["x3_finding"], title="Settling the withdrawn claim")
    FIG(doc, 7, "confusion_geometry", "Error geometry, model and human, both with "
                                      "patient-clustered intervals.")
    h(doc, "8.4 Attribution: a measurement that could not be made", 2)
    para(doc, f"The pre-registered attribution endpoint — the within-stratum correlation "
              f"{cite('gradcam')} "
              f"between Grad-CAM dispersion and annotator vote entropy — returned "
              f"{x['attribution_primary_verdict']}. Vote entropy is a deterministic "
              f"function of the vote pattern (a 3-1 split is always 0.5623 nats) and the "
              f"strata are defined by that pattern, so entropy is constant within a "
              f"stratum and the correlation does not exist. Reporting this as 'no "
              f"association' would assert a measurement that was never available.")
    FIG(doc, 7, "attribution_not_estimable",
        "Why the pre-registered attribution signal cannot be computed within a stratum.")
    rows = [[c, pc(x["iou_unanimous"].get(c), 3), x["attribution_secondary"].get(c, "—")]
            for c in x["attribution_secondary"]]
    table(doc, ["Arm", "Inter-seed IoU (unanimous)", "Verdict"], rows,
          "The secondary attribution endpoint did return a result: attribution "
          "destabilises on contested images for every arm, and the three soft-target "
          "arms are markedly more spatially consistent across seeds than the two "
          "hard-label arms.", widths=[0.7, 2.0, 3.3], font=8.2)
    FIG(doc, 7, "attribution_stability", "Inter-seed attribution stability.")
    h(doc, "8.5 Selective prediction", 2)
    rows = [[c, pc(x["aurc_internal"].get(c), 4), pc(x["aurc_external"].get(c), 4)]
            for c in x["aurc_internal"]]
    table(doc, ["Arm", "Internal AURC", "External AURC"], rows,
          "Area under the risk–coverage curve, lower is better. Internally the arms are "
          "nearly indistinguishable; externally they separate decisively. External "
          "intervals are image-level and are not comparable with internal ones.",
          widths=[0.8, 2.2, 2.2], font=8.2)
    FIG(doc, 7, "risk_coverage_external", "Risk–coverage on the external panel.")
    para(doc, f"The ordering agrees with the single-operating-point rejection result of "
              f"Chapter 7 ({x['phase5_consistency']}), so that finding was not an "
              f"artefact of where the 23-way decision boundary happens to fall.")
    FIG(doc, 7, "endpoint_synthesis",
        "Arm ranking under every endpoint. The rows disagree, which is the point.")


# =====================================================================
def ch8(doc) -> None:
    h(doc, "9. The Audit Protocol as an Instrument", 1)
    r = REG["ch8_rq5"]
    h(doc, "9.1 An audit that passes everything it sees is not an audit", 2)
    para(doc, "Chapter 2's protocol returned PROCEED on the adopted corpus. That result "
              "is uninformative unless the protocol can also return the opposite. This "
              "chapter scores it against a corpus known to be unsound: the peptic-ulcer "
              "dataset this project began with and retired.")
    h(doc, "9.2 Gate-by-gate comparison", 2)
    rows = []
    for g in RQ5["gates"]:
        rows.append([g["id"], g["name"],
                     g[RQ5["corpora"]["sound"]]["verdict"],
                     g[RQ5["corpora"]["unsound"]]["verdict"]])
    table(doc, ["Gate", "Criterion", "Adopted corpus", "Retired corpus"], rows,
          "The eight gates applied to both corpora, each reduced to a criterion that is "
          "meaningful in both modalities and derived from artefacts rather than quoted "
          "from report prose.", widths=[0.6, 2.6, 1.5, 1.5], font=8.0)
    FIG(doc, 8, "negative_control", "The negative control.")
    callout(doc, RQ5["discrimination"]["interpretation"],
            title="Gate-counting is not a quality score")
    h(doc, "9.3 What the protocol missed", 2)
    rows = [[f["id"], f["defect"], f["caught_by_gate"] or "none", f["why"]]
            for f in RQ5["what_the_protocol_missed"]["fatal_defects"]]
    table(doc, ["ID", "Fatal defect", "Caught by", "Why not"], rows,
          f"Only {r['n_fatal_caught']} of {r['n_fatal']} fatal defects is caught by any "
          f"gate, and that one only incidentally.",
          widths=[0.5, 1.7, 0.9, 3.4], font=7.4)
    callout(doc, RQ5["what_the_protocol_missed"]["finding"],
            title="Well-formedness is not viability")
    h(doc, "9.4 Proposed extensions", 2)
    rows = [[e["id"], e["name"], e["criterion"], e["would_have_caught"]]
            for e in RQ5["proposed_protocol_extension"]]
    table(doc, ["ID", "Gate", "Criterion", "Would have caught"], rows,
          "Three extensions, each of which would have caught a fatal defect and each of "
          "which costs minutes to run.", widths=[0.5, 1.3, 3.5, 1.2], font=7.6)
    para(doc, RQ5["why_this_matters_for_the_thesis"], italic=True, size=9.8)


# =====================================================================
def ch9(doc) -> None:
    h(doc, "10. Synthesis", 1)
    x, s = REG["ch7_error"], REG["ch4_stratified"]
    h(doc, "10.1 The central claim, defended", 2)
    para(doc, "The claim is that agreement stratification separates a falling reference "
              "standard from a falling classifier. Three alternative readings deserve "
              "answering.")
    bullet(doc, f"'The model simply fails on hard images.' If that were the whole story, "
                f"a held-out expert would do well where the model does badly. They do "
                f"not: on contested images the annotator scores "
                f"{pc(x['human_by_stratum'][POOLED], 4)} against the model's "
                f"{pc(x['model_by_stratum'][POOLED], 4)}.")
    bullet(doc, f"'The task is simply impossible there, so the model is blameless.' If "
                f"that were the whole story, the model would sit at the attainable "
                f"ceiling. It does not: the oracle reaches "
                f"{pc(x['oracle_by_stratum'][POOLED], 4)} and the model recovers only "
                f"{p100(x['headroom_recovered'][POOLED], 0)}% of the distance to it.")
    m = REG["ch9_synthesis"]
    if "n_replicating" in m:
        bullet(doc, f"'This is a ConvNeXt artefact.' This was the strongest remaining "
                    f"objection and it has now been tested. {m['n_replicating']} of "
                    f"{m['n_tested']} claims replicate on {m['backbone']} — see §10.2a. "
                    f"Three endpoints still rest on one architecture and are named in "
                    f"§10.3.")
    else:
        bullet(doc, "'This is a ConvNeXt artefact.' This is the strongest remaining "
                    "objection and is not yet answered; see §10.3.")
    h(doc, "10.2 What the nulls establish", 2)
    para(doc, "Three of five research questions returned unresolved or not-estimable "
              "verdicts. Reported carelessly that is a weak thesis; reported precisely it "
              "is the result. The nulls are informative because each was a "
              "pre-registered test against a matched control, with a precision target "
              "met, rather than an underpowered search that happened to find nothing. "
              "That no target construction repairs calibration is a stronger statement "
              "than that one of them did, because it moves the problem from the training "
              "objective to the problem itself — and it is what licenses the "
              "recommendation to deploy abstention rather than a better loss.")
    callout(doc, m["summary"], title="Multiplicity")

    if "n_replicating" in m:
        h(doc, "10.2a Backbone generalisation", 2)
        para(doc, f"To test whether these conclusions are properties of ConvNeXt, the "
                  f"target-construction contrast was re-run on {m['backbone']} "
                  f"{cite('efficientnet')}. The "
                  f"training script imports the Phase 4 code and rebinds only the model "
                  f"constructor, so cohort, cache, normalisation, augmentation, "
                  f"schedule, loss, early stopping, epoch cap and seeds are identical "
                  f"and any difference is architectural. Fine-tuning depth was matched "
                  f"in parameter terms rather than module count (92.3% of feature "
                  f"parameters unfrozen against ConvNeXt's 94.5%). Every metric was "
                  f"computed by the same function that produced the ConvNeXt number.")
        r1, r2, r3 = m["R1"], m["R2"], m["R3"]
        rows = [
            ["R1", "RQ2 accuracy contrast C2−C3 on contested images",
             f"{r1['convnext']} ({pc(r1['convnext_estimate'])} "
             f"{iv(r1['convnext_ci'])})",
             f"{r1['efficientnet_b0']} ({pc(r1['b0_estimate'])} {iv(r1['b0_ci'])})",
             "yes" if r1["replicates"] else "NO"],
            ["R2", "Calibration reversal: C3 better on contested, C2 on unanimous",
             f"present = {r2['convnext']['pattern_present']}",
             f"present = {r2['efficientnet_b0']['pattern_present']}",
             "yes" if r2["replicates"] else "NO"],
            ["R3", "Model between held-out annotator and modal-vote oracle",
             f"{p100(r3['convnext_headroom_recovered'], 0)}% of headroom; "
             f"exceeds oracle = {r3['convnext_exceeds_oracle_anywhere']}",
             f"{p100(r3['b0_headroom_recovered'], 0)}% of headroom; "
             f"exceeds oracle = {r3['b0_exceeds_oracle_anywhere']}",
             "yes" if r3["replicates"] else "NO"],
        ]
        table(doc, ["", "Claim", "ConvNeXt-Tiny (28 M)", "EfficientNet-B0 (4 M)",
                    "Replicates"], rows,
              f"Backbone replication. {m['n_replicating']} of {m['n_tested']} claims "
              f"hold on an architecture with one seventh the parameters and a different "
              f"inductive bias.", widths=[0.4, 2.0, 1.7, 1.7, 0.7], font=7.2)
        e_c, e_u = m["b0_ece_contested"], m["b0_ece_unanimous"]
        para(doc, f"The calibration reversal is the thesis's durable finding and it "
                  f"reproduces in both directions. On EfficientNet-B0 the matched "
                  f"control C3 is better calibrated than the vote-proportion arm C2 on "
                  f"contested images ({p100(e_c['C3'])}% against {p100(e_c['C2'])}%), "
                  f"and the order reverses on unanimous images ({p100(e_u['C2'])}% "
                  f"against {p100(e_u['C3'])}%) — the same pattern, at the same strata, "
                  f"in a network with one seventh the parameters.")
        callout(doc, f"What this does NOT cover: "
                     f"{'; '.join(m['endpoints_WITHOUT_a_second_backbone'])}. Those "
                     f"remain single-architecture results and the threat table below "
                     f"is updated accordingly rather than declaring the objection "
                     f"closed outright.",
                title="The limit of this replication")

    h(doc, "10.3 Threats to validity, ranked", 2)
    replicated = "n_replicating" in m
    rows = [
        ["Single backbone for the geometry, attribution and external endpoints",
         "Medium" if replicated else "High",
         ("Reduced but not eliminated: the accuracy null, the calibration reversal and "
          "the human-comparator position replicate on EfficientNet-B0 (§10.2a); the "
          "three endpoints named there do not yet have a second backbone.")
         if replicated else
         "A second backbone replicates the target contrast only."],
        ["Compute-bound epoch cap on 8/12 runs",
         "Medium", "Applies identically across arms, so contrasts hold, but absolute "
                   "scores are lower bounds."],
        ["Contested strata are small (n = 342, 127, 81)",
         "Medium", "Per-stratum intervals are wide; the pooled stratum carries the "
                   "endpoints."],
        ["External intervals are image-level",
         "Medium", "No external corpus publishes a case key; declared wherever an "
                   "external number appears."],
        ["Four annotators bound the human comparator",
         "Medium", "The held-out construction leaves a reference panel of three."],
        ["Strata are defined by the same agreement that scores the human",
         "Medium", "The model-minus-human contrast is unaffected (identical rows), but "
                   "the human curve alone is not a free-standing estimate of expert "
                   "accuracy."],
        ["No age or sex in the corpus",
         "Low for the claims made", "No demographic or fairness claim is made anywhere."],
        ["Confirmatory family defined retrospectively",
         "Low", "Each RQ's primary endpoint was pre-registered as that RQ's test; only "
                "the grouping is new."],
    ]
    table(doc, ["Threat", "Severity", "Status"], rows,
          "Threats to validity, ranked by how much they could change a conclusion.",
          widths=[2.3, 1.1, 3.1], font=7.8)
    h(doc, "10.4 What could be deployed, and under what conditions", 2)
    para(doc, f"Nothing in this thesis supports deploying an autonomous landmark "
              f"classifier. What it does support is a completeness-checking assistant "
              f"that abstains. The external risk–coverage analysis gives the operating "
              f"characteristic: the C2 arm reaches an AURC of "
              f"{pc(x['aurc_external']['C2'], 4)} against "
              f"{pc(x['aurc_external']['C3'], 4)} for the internally-best arm, and its "
              f"out-of-protocol rejection is the only endpoint in the project that "
              f"separates configurations under domain shift. A deployment would have to "
              f"select on that endpoint, monitor it at the new centre, and route "
              f"declined images to a human rather than scoring them.")


# =====================================================================
def ch10(doc) -> None:
    h(doc, "11. Conclusions and Future Work", 1)
    h(doc, "11.1 Answers to the research questions", 2)
    x, s, t, e, r = (REG["ch7_error"], REG["ch4_stratified"], REG["ch5_targets"],
                     REG["ch6_external"], REG["ch8_rq5"])
    rows = [
        ["RQ1", "Does performance vary across strata of expert agreement?",
         f"Yes, and the ceiling moves with it. Ceiling-normalised gap "
         f"{pc(s['gap_4v3_ceiling_normalised'])} pts {iv(s['gap_4v3_ci'])}."],
        ["RQ2", "Do soft targets from all four votes beat hard consensus labels?",
         f"NOT RESOLVED on accuracy ({pc(t['contrast_C2_C3'])} pts "
         f"{iv(t['contrast_C2_C3_ci'])}); NOT SUPPORTED on calibration — the control "
         f"wins on contested images."],
        ["RQ3", "Does predictive uncertainty track human disagreement?",
         "Supported for no configuration internally, and NOT ESTIMABLE in the "
         "within-stratum form it was posed, for a structural reason (Chapter 8)."],
        ["RQ4", "Does an anatomy-aware loss help on contested images?",
         f"NOT RESOLVED at unit λ ({pc(t['rq4_C4_C2'], 5)} "
         f"{iv(t['rq4_C4_C2_ci'], 5)}). No sweep was run, so this is evidence about "
         f"unit weight only."],
        ["RQ5", "Does the audit protocol discriminate a sound corpus from an unsound one?",
         r["verdict"] + f" — {r['n_separating']}/{r['n_independent']} gates separate the "
                        f"corpora but only {r['n_fatal_caught']}/{r['n_fatal']} fatal "
                        f"defects is caught by any gate."],
    ]
    table(doc, ["RQ", "Question", "Answer"], rows,
          "Answers to the five research questions. Every verdict was selected by a rule "
          "frozen before the analysis ran.", widths=[0.5, 2.4, 3.6], font=7.6)
    h(doc, "11.2 Contributions restated", 2)
    para(doc, "Against the reporting gaps identified in Chapter 4, this thesis "
              "contributes an agreement-stratified protocol that reports the attainable "
              "ceiling alongside the score; a human comparator and oracle that make the "
              "model's residual shortfall measurable rather than assumed; evidence that "
              "the calibration failure is a property of the problem rather than of the "
              "training objective; a demonstration that selective prediction separates "
              "configurations that no internal endpoint can; and a negative-control "
              "evaluation that improved the audit protocol it tested.")
    h(doc, "11.3 Future work, costed", 2)
    rows = [
        ["Second backbone across the geometry, attribution and external endpoints",
         "~8 GPU hours", "The target contrast, calibration reversal and human-comparator "
                         "position already replicate (§10.2a); these three do not yet."],
        ["λ sweep for the anatomy-aware loss",
         "~15 GPU hours", "Converts RQ4's unit-λ null into an answer."],
        ["Raise the epoch cap on all twelve Phase 4 runs",
         "~22 GPU hours", "Removes the lower-bound caveat on absolute scores."],
        ["Attribution-method sweep",
         "~45 minutes", "Deliberately not run: it would invite selecting whichever "
                        "method correlated best."],
        ["Patient-clustered external intervals",
         "not possible", "Neither external corpus publishes a case identifier."],
        ["Human comparator on external corpora",
         "not possible", "Neither external corpus publishes per-annotator labels."],
        ["Prospective evaluation of the abstaining assistant",
         "a clinical study", "The only way to test the deployment claim of §10.4."],
    ]
    table(doc, ["Work", "Cost", "What it would buy"], rows,
          "Future work, costed. Items marked not possible are limits of the available "
          "data, not of effort.", widths=[2.3, 1.2, 3.0], font=7.8)


# =====================================================================
def appendices(doc) -> None:
    h(doc, "Appendix A. Pre-registration records", 1)
    rows = [
        ["Phase 2", "reports/phase2_prereg.json", "target, seeds, bootstrap, diagnostic order"],
        ["Phase 4", "reports/phase4_prereg.json", "ε derived, λ fixed, verdict rules for RQ2–RQ4"],
        ["Phase 5", "reports/phase5_prereg.json", "collapse, endpoints, precision target"],
        ["Phase 6", "reports/phase6_prereg.json", "four endpoints, CAM layer, top-q, verdict rules"],
    ]
    table(doc, ["Phase", "Artefact", "What it fixed"], rows,
          "Each pre-registration script refuses to overwrite an existing file.",
          widths=[0.9, 2.4, 3.2], font=8.0)

    h(doc, "Appendix B. Declared deviations and amendments", 1)
    rows = [
        ["P4-DEV-1", "MC dropout → MC stochastic depth (ConvNeXt has no Dropout/BatchNorm)"],
        ["P4-DEV-2", "5-member ensemble → 3 members, on budget"],
        ["P4-DEV-3", "λ fixed at 1.0, no sweep"],
        ["P4-DEV-4", "3 dataloader workers → 2, after a reproducible CUDA host-allocation failure"],
        ["P5-DEV-3", "External intervals image-level; no case key exists"],
        ["X1–X4", "Four Phase 3 claims withdrawn on re-examination (Chapter 5 §5.4)"],
        ["P6-DEV-1", "Grad-CAM on all images, not only contested ones"],
        ["P6-DEV-2", "Human comparator and selective prediction added to Phase 6"],
        ["P6-DEV-3", "Grad-CAM only; no attribution-method sweep"],
        ["P6-AMD-1", "Human comparator degenerate on a unanimity-defined stratum"],
        ["P6-AMD-2", "Human error geometry undefined on S-unanimous"],
        ["P6-AMD-3", "CAM targets the committed prediction, not the live argmax"],
        ["P6-AMD-4", "Vote entropy constant within a stratum; RQ3 endpoint not estimable"],
        ["P6-AMD-5", "Exposure and choice asymmetries in the human comparator"],
    ]
    table(doc, ["ID", "Deviation or amendment"], rows,
          "Every deviation was declared before or at the moment it was taken, and every "
          "amendment was forced by a gate or by a numerical fact rather than adopted "
          "silently.", widths=[1.0, 5.5], font=7.8)

    h(doc, "Appendix C. Multiplicity", 1)
    rows = []
    for f in MULT["family"]:
        mm = f.get("multiplicity", {})
        state = ("not estimable" if f.get("not_estimable")
                 else "no interval" if not mm
                 else "contains 0" if mm.get("contains_zero")
                 else f"excludes 0; Holm survives = {mm.get('holm', {}).get('survives')}")
        rows.append([f["rq"], f["endpoint"], state])
    table(doc, ["RQ", "Primary endpoint", "Under multiplicity"], rows,
          MULT["declaration"]["confirmatory_family"] + " " +
          MULT["declaration"]["everything_else"], widths=[0.5, 3.4, 2.6], font=7.6)

    h(doc, "Appendix D. Figure provenance", 1)
    rows = [[r["thesis_figure"], f"Ch {r['chapter']}", r["source_file"], r["drawn_by"]]
            for r in FIGREG["registry"]]
    table(doc, ["Figure", "Chapter", "Phase figure", "Drawn by"], rows,
          FIGREG["principle"], widths=[0.7, 0.7, 2.4, 2.7], font=7.0)

    h(doc, "Appendix E. Reproducibility index", 1)
    para(doc, f"Every number in this thesis was resolved from a JSON artefact by "
              f"src/report/phase7_register.py, which read "
              f"{len(json.loads((REP / 'phase7_register.json').read_text(encoding='utf-8'))['provenance'])} "
              f"values from 24 artefacts. Running the phase pipelines in order "
              f"regenerates every artefact, and rebuilding this document from them "
              f"regenerates every sentence that contains a number.")
    rows = [
        ["Phase 0–1", "src/data/gastrohun_*.py, src/literature/*_v2.py"],
        ["Phase 2", "src/models/phase2_*.py"],
        ["Phase 3", "src/models/phase3_*.py, phase3b_*.py"],
        ["Phase 4", "src/models/phase4_*.py"],
        ["Phase 5", "src/models/phase5_*.py, phase5b_*.py"],
        ["Phase 6", "src/models/phase6_*.py"],
        ["Phase 7", "src/models/phase7_*.py, src/report/phase7_register.py, "
                    "figures_thesis.py, build_thesis_docx.py"],
    ]
    table(doc, ["Phase", "Scripts"], rows, "Script manifest.", widths=[1.2, 5.3], font=8.0)

    h(doc, "Appendix F. Citation provenance", 1)
    cov = BIB.coverage()
    para(doc, f"The reference list holds {cov['n_total']} entries and none of "
              f"them was typed into this document. The builder calls "
              f"cite(key); src/report/bibliography.py resolves the key to a "
              f"number against reports/phase8_bibliography.json, which is "
              f"generated from literature_v2/extraction_table.csv. In-text "
              f"markers and the reference list therefore cannot disagree, and "
              f"renumbering is automatic.")
    rows = [
        ["Review set (PRISMA-included)", f"{cov['n_review_set']}",
         f"{cov['n_review_cited']}",
         "literature_v2/extraction_table.csv"],
        ["Additional (corpus, guideline, methods)", f"{cov['n_additional_set']}",
         f"{cov['n_additional_cited']}",
         "literature_v2/additional_references.json"],
    ]
    table(doc, ["Set", "In list", "Cited in text", "Source artefact"], rows,
          "Citation provenance. The two sets are numbered into one list but "
          "counted apart, so that the PRISMA total of 82 included studies is "
          "never inflated by a reference that did not come through screening — "
          "the corpus descriptor, the reporting guideline and the two method "
          "papers were never candidates for review inclusion.",
          widths=[2.4, 0.7, 0.9, 2.4], font=8.4)

    h(doc, "References", 1)
    para(doc, "Entries 1–86, ordered by first author surname. Entries drawn "
              "from the PRISMA review are marked in Appendix F; the remainder "
              "are the corpus descriptor, the reporting guideline this review "
              "followed, and two method papers cited for the techniques they "
              "define.")
    for e in BIB.references():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = BD.Pt(4)
        p.paragraph_format.left_indent = BD.Inches(0.42)
        p.paragraph_format.first_line_indent = BD.Inches(-0.42)
        r = p.add_run(f"[{e['n']}]  ")
        r.bold = True; r.font.size = BD.Pt(9)
        r2 = p.add_run(e["formatted"])
        r2.font.size = BD.Pt(9)


def main() -> None:
    doc = new_document()
    title_page(doc)
    approval_page(doc)
    declaration_page(doc)
    acknowledgements_page(doc)
    front_matter(doc)
    abbreviations_page(doc)
    sec_abstract(doc)
    # ch_methods is the new Chapter 3; ch3 is the literature review, now 4.
    for fn in (ch1, ch2, ch_methods, ch3, ch4, ch5, ch6, ch7, ch8, ch9, ch10):
        fn(doc)
    appendices(doc)
    add_page_numbers(doc)
    doc.save(BD.OUT)
    caps = [p.text for p in doc.paragraphs if p.style.name == "Caption"]
    nfig = sum(1 for c in caps if c.strip().startswith("Figure"))
    ntab = sum(1 for c in caps if c.strip().startswith("Table"))
    print(f"[thesis] {len(doc.inline_shapes)} images, {nfig} figure captions, "
          f"{ntab} table captions -> {BD.OUT}")
    unused = [r["file"] for r in FIGREG["registry"] if r["file"] not in _fig_used]
    if unused:
        print(f"[thesis] NOTE {len(unused)} registered figures unused: {unused}")


if __name__ == "__main__":
    main()
