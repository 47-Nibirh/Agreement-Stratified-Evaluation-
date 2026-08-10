"""
Build the Phase-I Progress Report by filling the university template in place.

The template is not re-created. It is opened, its placeholder paragraphs are
replaced with real content, its tables are populated, and everything else --
page setup, the identification and thematic-area tables, the certificate block,
the back-matter instruction page and the departmental notices -- is left exactly
as the department issued it.

Numbering is assigned as the document is written, so figures and tables are
numbered in the order they appear and the cross-references in the prose match.

Run:  python src/report/build_phase1_docx.py
Out:  Phase-I_Progress_Report.docx
"""
from __future__ import annotations

import copy
import re
import sys
from pathlib import Path

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(Path(__file__).resolve().parent))

import content_phase1_report as C      # noqa: E402
from phase1_facts import ADMIN, facts  # noqa: E402

TEMPLATE = ROOT / "2. Phase-I Progress Report Tamplate..docx"
OUTPUT = ROOT / "Phase-I_Progress_Report.docx"

FONT = "Times New Roman"
BODY_PT = Pt(12)
CAP_PT = Pt(10)
TBL_PT = Pt(9.5)
SHADE = "D9D9D9"       # the template's own grey; no new colour is introduced


# --------------------------------------------------------------------------
# low-level helpers
# --------------------------------------------------------------------------
def _style_run(run, size=BODY_PT, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = size
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rf)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(attr), FONT)
    return run


def _blank_paragraph_like(model: Paragraph) -> Paragraph:
    """A new empty paragraph cloned from `model`, inserted directly after it."""
    new_p = copy.deepcopy(model._p)
    for child in list(new_p):
        if child.tag in (qn("w:r"), qn("w:hyperlink"), qn("w:bookmarkStart"),
                         qn("w:bookmarkEnd")):
            new_p.remove(child)
    model._p.addnext(new_p)
    return Paragraph(new_p, model._parent)


def _clear(par: Paragraph) -> Paragraph:
    for r in list(par.runs):
        r._element.getparent().remove(r._element)
    return par


def _body_format(par: Paragraph, justify=True):
    pf = par.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    pf.space_after = Pt(8)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15
    pf.left_indent = None
    pf.first_line_indent = None
    # a numbered-list template paragraph must lose its numbering when reused
    ppr = par._p.get_or_add_pPr()
    for tag in ("w:numPr", "w:ind"):
        el = ppr.find(qn(tag))
        if el is not None:
            ppr.remove(el)
    return par


def _shade(cell, fill=SHADE):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def _cell_text(cell, text, bold=False, size=TBL_PT, align=None):
    cell.text = ""
    par = cell.paragraphs[0]
    pf = par.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    if align is not None:
        pf.alignment = align
    _style_run(par.add_run(str(text)), size=size, bold=bold)


def _repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    el = trpr.makeelement(qn("w:tblHeader"), {})
    trpr.append(el)


# --------------------------------------------------------------------------
# cross-reference registry
# --------------------------------------------------------------------------
class Refs:
    """Numbers figures and tables once, in document order, and resolves the
    `{{fig:id}}` / `{{tbl:id}}` tokens the prose uses.

    The caption number and the in-text reference therefore come from the same
    place. Inserting a figure renumbers both or neither; there is no way for
    them to drift apart, which is how the first draft of this document ended up
    citing Figure 6 for the architecture diagram that had become Figure 7.
    """

    TOKEN = re.compile(r"\{\{(fig|tbl):([a-z0-9_]+)\}\}")

    def __init__(self):
        self.maps = {"fig": {}, "tbl": {}}
        self.next = {"fig": 1, "tbl": 1}

    def assign(self, kind, key):
        m = self.maps[kind]
        if key not in m:
            m[key] = self.next[kind]
            self.next[kind] += 1
        return m[key]

    def scan(self, blocks):
        """Register every numbered block in a section, in order."""
        for b in blocks:
            if b["t"] in ("fig", "tbl"):
                self.assign(b["t"], b["id"])

    def sub(self, text):
        def repl(m):
            kind, key = m.group(1), m.group(2)
            if key not in self.maps[kind]:
                raise KeyError(f"unresolved {kind} reference: {key}")
            return str(self.maps[kind][key])
        return self.TOKEN.sub(repl, text)


# --------------------------------------------------------------------------
# block writer
# --------------------------------------------------------------------------
class Writer:
    """Writes content blocks after a given anchor paragraph, in order."""

    def __init__(self, doc, anchor: Paragraph, refs: Refs):
        self.doc = doc
        self.cursor = anchor
        self.refs = refs

    def _new(self) -> Paragraph:
        p = _blank_paragraph_like(self.cursor)
        self.cursor = p
        return p

    def para(self, text, justify=True, bold=False, size=BODY_PT):
        p = _body_format(self._new(), justify=justify)
        _style_run(p.add_run(self.refs.sub(text)), size=size, bold=bold)
        return p

    def numbered(self, items):
        for i, it in enumerate(items, 1):
            p = _body_format(self._new())
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.first_line_indent = Inches(-0.35)
            _style_run(p.add_run(f"{i}."), bold=True)
            _style_run(p.add_run(f"\t{self.refs.sub(it)}"))

    def figure(self, fid, path, cap, width):
        p = self._new()
        _body_format(p, justify=False)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        p.add_run().add_picture(str(ROOT / path), width=Inches(width))
        n = self.refs.assign("fig", fid)
        c = self._new()
        _body_format(c, justify=False)
        c.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(12)
        _style_run(c.add_run(f"Figure {n}. "), size=CAP_PT, bold=True)
        _style_run(c.add_run(self.refs.sub(cap)), size=CAP_PT, italic=True)

    def table(self, tid, cap, head, rows, widths):
        n = self.refs.assign("tbl", tid)
        c = self._new()
        _body_format(c, justify=False)
        c.paragraph_format.space_after = Pt(4)
        _style_run(c.add_run(f"Table {n}. "), size=CAP_PT, bold=True)
        _style_run(c.add_run(self.refs.sub(cap)), size=CAP_PT, italic=True)

        t = self.doc.add_table(rows=1, cols=len(head))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        for j, (h, w) in enumerate(zip(head, widths)):
            _cell_text(t.rows[0].cells[j], h, bold=True)
            _shade(t.rows[0].cells[j])
            t.rows[0].cells[j].width = Inches(w)
        _repeat_header(t.rows[0])
        for r in rows:
            cells = t.add_row().cells
            for j, (v, w) in enumerate(zip(r, widths)):
                _cell_text(cells[j], v)
                cells[j].width = Inches(w)
        # move the table from the document end to the cursor
        self.cursor._p.addnext(t._tbl)
        spacer = _blank_paragraph_like(Paragraph(t._tbl.getnext(), self.cursor._parent)) \
            if False else None
        # a trailing empty paragraph keeps the next block off the table border
        after = copy.deepcopy(self.cursor._p)
        for child in list(after):
            if child.tag == qn("w:r"):
                after.remove(child)
        t._tbl.addnext(after)
        self.cursor = Paragraph(after, self.cursor._parent)
        _body_format(self.cursor)
        self.cursor.paragraph_format.space_after = Pt(10)
        return t

    def blocks(self, items):
        for b in items:
            if b["t"] == "p":
                self.para(b["text"])
            elif b["t"] == "num":
                self.numbered(b["items"])
            elif b["t"] == "fig":
                self.figure(b["id"], b["path"], b["cap"], b.get("w", 6.2))
            elif b["t"] == "tbl":
                self.table(b["id"], b["cap"], b["head"], b["rows"], b["w"])
            else:
                raise ValueError(f"unknown block type {b['t']}")


# --------------------------------------------------------------------------
# template-specific fills
# --------------------------------------------------------------------------
def fill_identification(TT, F):
    t = TT[0]
    _cell_text(t.rows[0].cells[1], ADMIN["title"], size=Pt(11), bold=True)

    members = ADMIN["members"]
    cell = t.rows[1].cells[1]
    cell.text = ""
    for i, (name, sid) in enumerate(members):
        par = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        par.paragraph_format.space_after = Pt(2)
        _style_run(par.add_run(f"Name: {name}"), size=Pt(11))
        _style_run(par.add_run(f"     Student ID: {sid}"), size=Pt(11))

    for row, (name, desig) in ((t.rows[2], ADMIN["supervisor"]),
                               (t.rows[3], ADMIN["co_supervisor"])):
        cell = row.cells[1]
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        _style_run(p1.add_run(f"Name: {name}"), size=Pt(11))
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        _style_run(p2.add_run(f"Designation: {desig}"), size=Pt(11))

    _cell_text(t.rows[4].cells[1], ADMIN["submission_date"], size=Pt(11))


UNCHECKED, CHECKED = "☐", "☒"


def fill_thematic(TT, F):
    """Tick the applicable thematic areas and list the software stack.

    Each checkbox cell is wrapped in a `w:sdt` content control, so it is not a
    direct `w:tc` child of its row and python-docx's `row.cells` never returns
    it. The rows are therefore walked at the XML level: for each `w:sdt` the
    label is the text of the next `w:tc` sibling, and ticking means setting the
    control's state *and* swapping the visible glyph, since Word renders the
    glyph rather than deriving it from the state.
    """
    t = TT[1]
    wanted = {a.strip().lower() for a in ADMIN["thematic_areas"]}
    ticked, seen = 0, []
    for tr in t._tbl.iter(qn("w:tr")):
        sdt = tr.find(qn("w:sdt"))
        if sdt is None:
            continue
        # the checkbox control is the row's last child; the label is the
        # nearest preceding w:tc
        label_tc = sdt.getprevious()
        while label_tc is not None and label_tc.tag != qn("w:tc"):
            label_tc = label_tc.getprevious()
        if label_tc is None:
            continue
        label = "".join(n.text or "" for n in label_tc.iter(qn("w:t"))).strip()
        seen.append(label)
        if label.lower() not in wanted:
            continue
        for chk in sdt.iter(qn("w14:checked")):
            chk.set(qn("w14:val"), "1")
        for tnode in sdt.iter(qn("w:t")):
            if tnode.text and UNCHECKED in tnode.text:
                tnode.text = tnode.text.replace(UNCHECKED, CHECKED)
                ticked += 1
    missing = wanted - {s.lower() for s in seen}
    if missing:
        raise SystemExit(f"thematic areas not found in the template: {missing}\n"
                         f"template offers: {seen}")
    # software / tools row
    last = t.rows[len(t.rows) - 1]
    _cell_text(last.cells[1], ADMIN["software"], size=Pt(10))
    return ticked


EXAMPLE_PROJECT = "AI-Based Crop Disease Detection System"


def fill_co_table(TT, F):
    """The template ships the CO/PO rows carrying an example project name.

    That name sits in red as a fill-me-in marker and is split across runs, so
    the cell is rewritten wholesale with the real title and normal body colour
    rather than patched run by run.
    """
    t = TT[2]
    title = ADMIN["title"]
    replaced = 0
    for row in t.rows[1:]:
        cell = row.cells[1]
        txt = " ".join(cell.text.split())
        if EXAMPLE_PROJECT in txt:
            txt = txt.replace(EXAMPLE_PROJECT, title)
            replaced += 1
        _cell_text(cell, txt, size=Pt(10), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0, 0, 0)
    if replaced == 0:
        raise SystemExit("the CO table's example project name was not found; "
                         "check EXAMPLE_PROJECT against the template")
    return replaced


def fill_simple_table(TT, idx, rows, widths=None):
    t = TT[idx]
    for i, data in enumerate(rows, start=1):
        if i >= len(t.rows):
            t.add_row()
        for j, v in enumerate(data):
            if j < len(t.columns):
                _cell_text(t.rows[i].cells[j], v)
                if widths:
                    t.rows[i].cells[j].width = Inches(widths[j])
    for c in t.rows[0].cells:
        for p in c.paragraphs:
            for r in p.runs:
                _style_run(r, size=TBL_PT, bold=True)
        _shade(c)
    return t


EST_FILL = "2E74B5"    # the template's own "estimated period" blue
ACT_FILL = "70AD47"    # the template's own "actual period" green


def _clear_shading(cell):
    tcpr = cell._tc.get_or_add_tcPr()
    for shd in tcpr.findall(qn("w:shd")):
        tcpr.remove(shd)


def fill_gantt(TT, F):
    """Populate the Gantt grid the template provides.

    The grid is not one row per task. It gives four task slots, each a pair of
    rows sharing a vertically merged label cell: an upper row shaded in the
    template's blue for the *estimated* period and a lower row in its green for
    the *actual* period, matching the legend beneath the table. Writing one task
    per row -- which the first pass did -- writes the second task's label into
    the first task's merged cell and loses half of them.
    """
    t = TT[5]
    weeks = [int(c.text.strip()) for c in t.rows[1].cells[1:]
             if c.text.strip().isdigit()]
    tasks = C.TIMELINE_TASKS
    n_slots = (len(t.rows) - 2) // 2
    if len(tasks) > n_slots:
        raise SystemExit(f"the template's Gantt has {n_slots} task slots but "
                         f"{len(tasks)} tasks were supplied")

    for i, (label, est, act) in enumerate(tasks):
        est_row, act_row = t.rows[2 + 2 * i], t.rows[3 + 2 * i]
        _cell_text(est_row.cells[0], label, size=Pt(8))
        for row, (w0, w1), fill in ((est_row, est, EST_FILL),
                                    (act_row, act, ACT_FILL)):
            for j, wk in enumerate(weeks):
                cell = row.cells[1 + j]
                _cell_text(cell, "")
                _clear_shading(cell)          # drop the template's example bars
                if w0 <= wk <= w1:
                    _shade(cell, fill)

    # unused slots are removed rather than left as blank task lines
    for row in list(t.rows[2 + 2 * len(tasks):]):
        row._tr.getparent().remove(row._tr)
    for c in list(t.rows[0].cells) + list(t.rows[1].cells):
        for p in c.paragraphs:
            for r in p.runs:
                _style_run(r, size=Pt(8), bold=True)

    # The legend table beneath is left exactly as issued: its two coloured
    # cells are the key to the bars above, not fields to fill in.


def caption_before(table, model: Paragraph, refs, tid, cap):
    """Give a template-owned table the same numbered caption as a built one.

    The template ships the challenges, next-steps and Gantt tables already
    formatted, so they are filled rather than rebuilt -- but they still occupy
    numbers in the table sequence and a reader has to be able to find them.
    """
    n = refs.assign("tbl", tid)
    p = copy.deepcopy(model._p)
    for child in list(p):
        if child.tag == qn("w:r"):
            p.remove(child)
    table._tbl.addprevious(p)
    par = Paragraph(p, model._parent)
    _body_format(par, justify=False)
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after = Pt(4)
    _style_run(par.add_run(f"Table {n}. "), size=CAP_PT, bold=True)
    _style_run(par.add_run(refs.sub(cap)), size=CAP_PT, italic=True)


def set_list_level(par: Paragraph, ilvl: int):
    """Force a template heading's outline level.

    The issued template puts "Objectives" at level 0 while its three siblings
    under Project Overview -- Introduction, Background Study and Gap Analysis --
    are at level 1. Left alone it numbers as a top-level section, which pushes
    Methodology to 3, Progress Achieved to 4 and every later section along with
    them. Correcting the level restores the template's own intended structure.
    """
    ppr = par._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        return False
    lvl = numpr.find(qn("w:ilvl"))
    if lvl is None:
        return False
    lvl.set(qn("w:val"), str(ilvl))
    return True


def set_heading(par: Paragraph, text=None):
    """Restyle an existing template heading, keeping its list numbering."""
    if text is not None:
        _clear(par)
        _style_run(par.add_run(text), size=Pt(13), bold=True)
    else:
        for r in par.runs:
            _style_run(r, size=r.font.size or Pt(13), bold=True)
    par.paragraph_format.space_before = Pt(10)
    par.paragraph_format.space_after = Pt(6)


# --------------------------------------------------------------------------
# main
# --------------------------------------------------------------------------
def main() -> None:
    F = facts()
    doc = docx.Document(str(TEMPLATE))
    paras = doc.paragraphs
    TT = list(doc.tables)      # template tables, captured before any insertion

    # --- header block ----------------------------------------------------
    _clear(paras[5])
    _style_run(paras[5].add_run(f"reporting period – {ADMIN['reporting_period']}"),
               size=Pt(13), bold=True)
    paras[5].runs[0].font.small_caps = True

    fill_identification(TT, F)
    ticked = fill_thematic(TT, F)
    fill_co_table(TT, F)

    # "Objectives" belongs under Project Overview like its three siblings
    if not set_list_level(paras[22], 1):
        raise SystemExit("expected a numbered 'Objectives' heading at P22")

    # --- 1. Project Overview ---------------------------------------------
    # Placeholder paragraph -> section content. Indices are the template's.
    plan = [
        (15, C.introduction(F)),
        (18, C.background(F)),
        (21, C.gap_analysis(F)),
        (23, C.objectives(F)),
        (28, C.research_design(F)),
        (31, C.data_collection(F)),
        (34, C.analysis_techniques(F)),
        (39, C.completed_tasks(F)),
        (42, C.results(F)),
    ]
    challenges_blocks = C.challenges(F)
    timeline_blocks = C.timeline(F)
    resources_blocks = C.resources(F)
    management_blocks = C.management(F)
    future_blocks = C.future(F)
    conclusion_blocks = C.conclusion(F)
    appendix_blocks = C.appendix(F)

    # Register every numbered object in document order BEFORE writing any of
    # it, so that a reference appearing earlier in the prose than its target
    # still resolves. The three template-owned tables are registered at the
    # position they physically occupy.
    k = Refs()
    for _, blocks in plan:
        k.scan(blocks)
    k.scan(challenges_blocks)
    k.assign("tbl", "challenges")
    k.assign("tbl", "nextsteps")
    k.assign("tbl", "gantt")
    for blocks in (timeline_blocks, resources_blocks, management_blocks,
                   future_blocks, conclusion_blocks, appendix_blocks):
        k.scan(blocks)
    # Content must be written in document order for the figure/table counters
    # to come out right, and each anchor must be resolved before its own
    # section inserts paragraphs after it -- so resolve all anchors first.
    anchors = {i: paras[i] for i, _ in plan}
    for idx, blocks in plan:
        anchor = anchors[idx]
        _clear(anchor)
        _body_format(anchor)
        # the anchor becomes the first block rather than being left empty
        w = Writer(doc, anchor, k)
        first = blocks[0]
        if first["t"] == "p":
            _style_run(anchor.add_run(k.sub(first["text"])))
            w.blocks(blocks[1:])
        else:
            w.blocks(blocks)
            anchor._p.getparent().remove(anchor._p)

    # --- 4. Challenges ----------------------------------------------------
    anchor = paras[45]
    _clear(anchor)
    _body_format(anchor)
    ch = challenges_blocks
    _style_run(anchor.add_run(k.sub(ch[0]["text"])))
    Writer(doc, anchor, k).blocks(ch[1:])
    caption_before(TT[3], anchor, k, "challenges",
                   "Issues and challenges encountered during the reporting period, "
                   "with the strategy adopted for each.")
    fill_simple_table(TT, 3, C.challenges_table(F), widths=[0.45, 2.6, 3.15])
    _cell_text(TT[3].rows[0].cells[0], "S. No.", bold=True)
    _cell_text(TT[3].rows[0].cells[1], "Issues and challenges", bold=True)
    _cell_text(TT[3].rows[0].cells[2], "Strategies or plans", bold=True)
    for c in TT[3].rows[0].cells:
        _shade(c)

    # --- 5. Next steps ----------------------------------------------------
    anchor = paras[48]
    _clear(anchor)
    _body_format(anchor)
    _style_run(anchor.add_run(k.sub(
        "The tasks below carry the project from a measurement of the problem to an "
        "intervention on it. Each is scheduled against the phase structure in "
        "Figure {{fig:workflow}} and each ends in a pre-registered verdict rather "
        "than an open-ended investigation.")))
    caption_before(TT[4], anchor, k, "nextsteps",
                   "Tasks and milestones planned for the next phase, with estimated "
                   "completion dates.")
    fill_simple_table(TT, 4, C.next_steps_table(F), widths=[0.45, 4.4, 1.35])
    _cell_text(TT[4].rows[0].cells[0], "S. No.", bold=True)
    _cell_text(TT[4].rows[0].cells[1], "Next task", bold=True)
    _cell_text(TT[4].rows[0].cells[2],
               "Estimated completion (MM-YY)", bold=True)
    for c in TT[4].rows[0].cells:
        _shade(c)

    # --- 6. Updated timeline ---------------------------------------------
    anchor = paras[51]
    _clear(anchor)
    _body_format(anchor)
    _style_run(anchor.add_run(k.sub(timeline_blocks[0]["text"])))
    caption_before(TT[5], anchor, k, "gantt",
                   "Updated project timeline by calendar week. For each task the "
                   "upper bar is the estimated period and the lower bar the actual "
                   "period, keyed to the legend beneath the table.")
    fill_gantt(TT, F)

    # --- 7-10 narrative sections -----------------------------------------
    for idx, blocks in ((55, resources_blocks), (58, management_blocks),
                        (61, future_blocks), (64, conclusion_blocks)):
        anchor = paras[idx]
        _clear(anchor)
        _body_format(anchor)
        _style_run(anchor.add_run(k.sub(blocks[0]["text"])))
        Writer(doc, anchor, k).blocks(blocks[1:])

    # --- References -------------------------------------------------------
    anchor = paras[67]
    _clear(anchor)
    _body_format(anchor, justify=False)
    anchor.paragraph_format.space_after = Pt(4)
    _style_run(anchor.add_run(
        "References are given in IEEE style and cited by number in the text."),
        size=Pt(11), italic=True)
    w = Writer(doc, anchor, k)
    for i, ref in enumerate(C.REFERENCES, 1):
        p = _body_format(w._new(), justify=True)
        p.paragraph_format.left_indent = Inches(0.42)
        p.paragraph_format.first_line_indent = Inches(-0.42)
        p.paragraph_format.space_after = Pt(6)
        _style_run(p.add_run(f"[{i}]"), size=Pt(11))
        _style_run(p.add_run(f"\t{ref}"), size=Pt(11))
    # drop the template's two empty reference stubs
    for stub in (paras[68], paras[69]):
        stub._p.getparent().remove(stub._p)

    # --- Appendix ---------------------------------------------------------
    anchor = paras[72]
    _clear(anchor)
    _body_format(anchor)
    ap = appendix_blocks
    _style_run(anchor.add_run(k.sub(ap[0]["text"])))
    Writer(doc, anchor, k).blocks(ap[1:])

    # --- headings and body typography ------------------------------------
    for i in (6, 8, 10, 66, 71):
        set_heading(paras[i])
    for i in (12, 25, 36, 44, 47, 50, 54, 57, 60, 63, 14, 17, 20, 22, 27, 30,
              33, 38, 41):
        set_heading(paras[i])

    doc.save(str(OUTPUT))
    print(f"wrote {OUTPUT.relative_to(ROOT)}")
    print(f"  figures numbered: {k.next['fig'] - 1}")
    print(f"  tables numbered: {k.next['tbl'] - 1}")
    print(f"  thematic areas ticked: {ticked}")


if __name__ == "__main__":
    main()
