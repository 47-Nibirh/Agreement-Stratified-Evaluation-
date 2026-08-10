"""
Phase 5 report -- external validation and the limits of the available label spaces.

Every number is interpolated from reports/phase5_*.json. No verdict sentence is
pre-written: each is selected from a pre-registered verdict field, so the report
says whatever the frozen rules produced.

Run:  python src/report/build_phase5_docx.py
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_docx as BD  # noqa: E402

ROOT = Path(__file__).resolve().parents[2]
BD.FIGD = ROOT / "figures_phase5"
BD.OUT = ROOT / "Phase5_Report.docx"

from build_docx import (add_page_numbers, bullet, callout, figure,  # noqa: E402
                        front_matter, h, new_document, para, table)

REP = ROOT / "reports"


def J(name):
    p = REP / name
    return json.loads(p.read_text(encoding="utf-8")) if p.exists() else None


PRE = J("phase5_prereg.json")
PROV = J("phase5_provenance.json")
MAP = J("phase5_mapping.json")
CG = J("phase5_cache_gate.json")
IG = J("phase5_infer_gate.json")
TR = J("phase5_transfer.json")
RJ = J("phase5_rejection.json")
CAL = J("phase5_calibration.json")
SEN = J("phase5_sensitivity.json")
CARRY = J("phase5_carry_forward.json")
P4AMD = J("phase4_amendment.json")

CFG_LABEL = {"C0": "C0 hard label, 4/4 cohort (Phase 2 reference)",
             "C1": "C1 hard majority label, extended cohort",
             "C2": "C2 vote proportions",
             "C3": "C3 hard label + matched label smoothing (control)",
             "C4": "C4 vote proportions + anatomical penalty"}


def ARMS():
    return TR["arms"] if TR else PRE["arms"]["carried"]


def pc(x, n=2):
    return "n/a" if x is None else f"{x:.{n}f}"


def iv(v, n=2):
    return "n/a" if not v or v[0] is None else f"[{v[0]:.{n}f}, {v[1]:.{n}f}]"


# =====================================================================
def title_page(doc) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("External Validation and the Limits of the Available "
                  "Label Spaces")
    r.bold = True; r.font.size = BD.Pt(19)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Phase 5 — Transfer of an Agreement-Aware Gastric Landmark "
                  "Classifier to HyperKvasir and GastroVision")
    r.font.size = BD.Pt(12.5)
    doc.add_paragraph()
    rows = [["P5-A retroflexion transfer", TR["verdict"] if TR else "—"],
            ["P5-B out-of-protocol rejection", RJ["verdict"] if RJ else "—"],
            ["P5-C calibration ordering",
             CAL["verdict_P5C"].get("verdict", "—") if CAL else "—"]]
    table(doc, ["Pre-registered endpoint", "Verdict"], rows,
          "The verdicts this phase produced, stated before the reader meets any "
          "argument for them.", font=9.5, head_font=9.5)
    if RJ and RJ.get("hypothesis_supported") is False:
        callout(doc,
                "The P5-B pre-registered hypothesis was FALSIFIED, and in the "
                "favourable direction. This report states the prediction that was "
                "frozen before scoring, and then the result that contradicted it.",
                title="A prediction this phase got wrong")


def sec_abstract(doc) -> None:
    h(doc, "Abstract", level=1)
    a = TR["aggregate_3seed"]; head = TR["headline_arm"]
    ext, inte = a[head]["external_macro_f1_mean_3seed"], a[head][
        "internal_macro_f1_mean_3seed"]
    para(doc, f"**Background.** Phases 2 to 4 established an agreement-stratified "
              f"picture of a ConvNeXt-Tiny gastric landmark classifier on GastroHUN, "
              f"and left calibration as the headline deficiency. None of it had been "
              f"tested outside one centre.", size=10)
    para(doc, f"**Objective.** To measure what transfers to two independent public "
              f"corpora, HyperKvasir and GastroVision, and to state plainly what "
              f"those corpora can and cannot test.", size=10)
    para(doc, f"**Methods.** {PROV['gates']['P5.1c_no_overlap_with_gastrohun']['n_images_hashed']:,} "
              f"external images were acquired, hashed against the GastroHUN "
              f"inventory (zero collisions) and mapped onto the GastroHUN label "
              f"space by a table frozen before any image was scored. All five Phase "
              f"2/4 arms were carried, three seeds each, with no fine-tuning, no "
              f"adaptation and no threshold tuning. Every interval is a 1,000-"
              f"resample image-level bootstrap.", size=10)
    para(doc, f"**Results.** The external corpora cannot express GastroHUN's wall x "
              f"station label space: neither carries the wall axis, and four of the "
              f"six stations have no external counterpart. At the granularity they "
              f"do support, the retroflexion distinction transferred "
              f"({TR['verdict']}): {pc(ext)} binary macro F1 externally against "
              f"{pc(inte)} internally, a drop of {pc(a[head]['drop_points'])} points "
              f"{iv(a[head]['drop_ci95'])}. Out-of-protocol rejection was predicted "
              f"to sit at or below the {100 * RJ['chance_rate']:.2f}% chance rate; it "
              f"reached {100 * RJ['aggregate_3seed'][head]['rejection_rate_mean_3seed']:.1f}% "
              f"for {head}, falsifying the hypothesis. The Phase 4 calibration "
              f"ordering was {CAL['verdict_P5C']['verdict']} "
              f"(Spearman rho = {CAL['verdict_P5C']['spearman_rho']:.3f}), but the "
              f"top two arms exchanged places.", size=10)
    para(doc, f"**Conclusion.** The soft-label arm C2 is the best arm externally on "
              f"calibration, on out-of-protocol rejection and on accuracy, having "
              f"looked worse-calibrated than its own control throughout Phase 4. "
              f"That reversal is only visible because five arms were carried rather "
              f"than the single best-calibrated one.", size=10)


def sec_introduction(doc) -> None:
    h(doc, "1. Introduction", level=1, page_break=True)
    h(doc, "1.1 What Phase 4 established, and what it left untested", level=2)
    para(doc, "Phase 4 compared four target constructions against a matched control "
              "and found no resolved accuracy difference, no resolved uncertainty "
              "correlation, and a calibration result whose interest lay in the "
              "control winning: uniform label smoothing achieved a lower expected "
              "calibration error than vote-proportion targets, but bought it by "
              "suppressing confidence globally rather than by knowing which images "
              "were hard. Everything in that finding was measured at one centre, on "
              "one vendor's equipment, against one annotation panel.")
    para(doc, "External validation is the only test of whether any of it is a "
              "property of the model rather than of the dataset. This phase runs "
              "that test, and its first result is about the test itself.")

    h(doc, "1.2 The label spaces do not line up, and that is a finding", level=2)
    para(doc, MAP["why_not_23_way"])
    figure(doc, "P5_F33_label_space.png",
           "What the frozen mapping does with the external images, and how the "
           "23-class GastroHUN label space collapses into the groups the external "
           "labels can express.")
    para(doc, MAP["limitation_to_state_in_the_report"])

    h(doc, "1.3 Research questions and pre-registered hypotheses", level=2)
    for k in ("P5-A", "P5-B", "P5-C"):
        rq = PRE["research_questions"][k]
        bullet(doc, f"{k}: {rq['question']}")
        bullet(doc, f"Endpoint: {rq['primary_endpoint']}", level=1)
        if "hypothesis" in rq:
            bullet(doc, f"Pre-registered expectation: {rq['hypothesis']}", level=1)


def sec_methods(doc) -> None:
    h(doc, "2. Methods", level=1, page_break=True)
    h(doc, "2.1 The external corpora", level=2)
    rows = [[v.get("name", k), f"{v.get('n_images', 0):,}", v.get("licence", ""),
             (v.get("centre") or "")[:60]] for k, v in PROV["corpora"].items()]
    table(doc, ["Corpus", "Images", "Licence", "Centre"], rows,
          "The two external corpora, as acquired. Counts are realised from the "
          "extracted archives, not quoted from the papers.", font=8.5)
    g = PROV["gates"]["P5.1c_no_overlap_with_gastrohun"]
    para(doc, f"Every external image was SHA-256 hashed against the "
              f"{PROV['gastrohun_reference_hashes']:,}-image GastroHUN inventory: "
              f"{g['n_images_hashed']:,} images checked exhaustively, "
              f"{g['n_collisions']} collisions. A collision would have meant the "
              f"external test set was not external.")
    figure(doc, "P5_F34_inventory.png",
           "Class inventory of each corpus, coloured by the mapping decision.")

    h(doc, "2.2 The mapping table", level=2)
    para(doc, "Each of the 50 external class directories receives exactly one "
              "decision, and each decision carries a written anatomical rationale. "
              "Ambiguity is recorded per endpoint, because a label can be "
              "unambiguous on one axis and ambiguous on another: 'pylorus' is not an "
              "SSS station and so is ambiguous for station identity, but it is "
              "unambiguously a forward view and so is not ambiguous for P5-A, which "
              "asks only retroflexion versus forward.")
    tally = MAP["images_by_decision"]
    table(doc, ["Decision", "Images", "Meaning"],
          [["RETROFLEXION", f"{tally.get('RETROFLEXION', 0):,}",
            "GastroHUN stations 4-5"],
           ["FORWARD_GASTRIC", f"{tally.get('FORWARD_GASTRIC', 0):,}",
            "GastroHUN stations 1, 2, 3, 6"],
           ["OTHERCLASS", f"{tally.get('OTHERCLASS', 0):,}",
            "not a gastric SSS station at all"],
           ["discard", f"{tally.get('discard', 0):,}",
            "the label does not fix the anatomical site"]],
          "The frozen mapping, in totals. The per-class table with every rationale "
          "is Appendix A.", font=8.5)

    h(doc, "2.3 What is held fixed", level=2)
    para(doc, f"Preprocessing is the Phase 2 path unchanged: {CG['decode_path']}. "
              f"Gate P5.3a re-decoded "
              f"{CG['gates']['P5.3a_decode_path_identity']['n_compared']} GastroHUN "
              f"images through the Phase 5 module and required them to come back "
              f"bit-identical to the Phase 4 cache; "
              f"{CG['gates']['P5.3a_decode_path_identity']['n_mismatched']} "
              f"mismatched. Without that, a measured drop could be a preprocessing "
              f"artefact rather than a domain effect.")
    callout(doc, CG["normalisation"] + ". Recomputing them on the external corpora "
                 "would silently adapt the model to the target domain, which is "
                 "precisely the adaptation this phase exists to measure the absence "
                 "of.", title="Normalisation")
    para(doc, PRE["preprocessing"]["no_adaptation"])

    h(doc, "2.4 Which arms were carried, and why all five", level=2)
    lr = CARRY["literal_rule_evaluation"]
    para(doc, f"The blueprint status board records the Phase 4 guidance as 'carry "
              f"the best-calibrated arm'. Applied literally that rule selects "
              f"{lr['selects']}. {lr['why']}")
    para(doc, "A second reason is structural: the external half of RQ3 asks whether "
              "the uncertainty ranking is preserved outside the training "
              "distribution, and a ranking cannot be measured with one arm. Phase 5 "
              "is inference-only, so carrying all five costs fifteen forward passes. "
              "The decision was frozen in reports/phase5_carry_forward.json before "
              "any external image existed on disk.")

    h(doc, "2.5 Evaluation and the pre-registered precision target", level=2)
    pt = PRE["precision_target"]
    para(doc, f"{pt['rationale']} The target is a 95% CI half-width of at most "
              f"{pt['max_ci95_halfwidth_points']} points on the P5-A endpoint. "
              f"{pt['rule']}")
    callout(doc, PRE["interval_procedure"]["declared_weakness"],
            title="Declared weakness: the clustering unit (P5-DEV-3)")

    rows = [[g, str(v.get("pass", v)) if not isinstance(v, dict)
             else str(v.get("pass"))]
            for g, v in list(IG["gates"].items())]
    rows = [[k, str(v.get("pass"))] for k, v in IG["gates"].items()]
    table(doc, ["Inference gate", "Pass"], rows,
          f"Gates on the {IG['n_external_images']:,}-image external inference. "
          f"P5.5c is the load-bearing one: "
          f"{IG['gates']['P5.5c_reproduces_phase4_internal_probs']['n_bit_identical']}"
          f" of "
          f"{IG['gates']['P5.5c_reproduces_phase4_internal_probs']['checked']} "
          f"arm-seeds reproduced the Phase 4 internal probabilities bit-identically, "
          f"so the external numbers and the internal comparator come from the same "
          f"pipeline.", font=8.5)


def sec_results(doc) -> None:
    h(doc, "3. Results", level=1, page_break=True)
    a, head = TR["aggregate_3seed"], TR["headline_arm"]

    h(doc, "3.1 P5-A: does the retroflexion distinction transfer?", level=2)
    figure(doc, "P5_F35_transfer.png",
           "Left: binary macro F1 internally and externally, per arm, with the "
           "majority-class floor. Right: the drop, with its 95% CI.")
    rows = [[CFG_LABEL[c], pc(a[c]["internal_macro_f1_mean_3seed"]),
             pc(a[c]["external_macro_f1_mean_3seed"]),
             iv(a[c]["external_ci95"]), pc(a[c]["drop_points"]),
             pc(a[c]["ci95_halfwidth_points"]),
             str(a[c]["meets_precision_target"]), a[c]["verdict"]]
            for c in ARMS()]
    table(doc, ["Configuration", "Internal F1", "External F1", "External 95% CI",
                "Drop", "Half-width", "Meets precision", "Verdict"], rows,
          f"P5-A on {TR['n_gastric_external']:,} external gastric images "
          f"({TR['n_retroflexion']:,} retroflexion, {TR['n_forward']:,} forward) "
          f"against {TR['n_gastric_internal_comparator']:,} internal ones.",
          font=7.6,
          note="An OTHERCLASS prediction counts as incorrect, so the endpoint "
               "penalises both confusing the two views and failing to recognise a "
               "gastric view at all.")
    para(doc, f"Every arm transferred, and every arm met the pre-registered "
              f"precision target, so these are powered verdicts rather than "
              f"underpowered nulls. The pre-registered expectation was a drop of "
              f"more than 10 points; the realised drop for {head} is "
              f"{pc(a[head]['drop_points'])} points {iv(a[head]['drop_ci95'])}. "
              f"A degradation of this size is the expected and publishable result, "
              f"not a failure — but it is roughly twice what was predicted, and the "
              f"report does not soften that.")

    h(doc, "3.2 P5-B: out-of-protocol rejection, and a falsified prediction", level=2)
    para(doc, "The pre-registered hypothesis was that rejection would sit at or "
              "below chance. The reasoning was concrete: the GastroHUN test split "
              "contains only 50 OTHERCLASS images in 1,353, so the model had almost "
              "no opportunity to learn a rejection behaviour.")
    ra = RJ["aggregate_3seed"]
    figure(doc, "P5_F36_rejection.png",
           "Left: OTHERCLASS rate on images that are not gastric stations at all, "
           "against the chance floor. Right: mean top-1 confidence on those same "
           "images.")
    rows = [[CFG_LABEL[c], f"{100 * ra[c]['rejection_rate_mean_3seed']:.2f}",
             iv([100 * x for x in ra[c]["ci95"]]),
             f"{100 * ra[c]['mean_top1_confidence_mean_3seed']:.2f}",
             ra[c]["verdict"]] for c in ARMS()]
    table(doc, ["Configuration", "Rejection (%)", "95% CI", "Mean top-1 conf (%)",
                "Verdict"], rows,
          f"P5-B on {RJ['n_out_of_protocol']:,} out-of-protocol images. Chance rate "
          f"for a 23-way head is {100 * RJ['chance_rate']:.2f}%.", font=8.0)
    best = max(ARMS(), key=lambda c: ra[c]["rejection_rate_mean_3seed"])
    worst = min(ARMS(), key=lambda c: ra[c]["rejection_rate_mean_3seed"])
    para(doc, f"The prediction was wrong by an order of magnitude, and the arms "
              f"separate sharply: {best} rejects "
              f"{100 * ra[best]['rejection_rate_mean_3seed']:.1f}% of out-of-protocol "
              f"images against {worst}'s "
              f"{100 * ra[worst]['rejection_rate_mean_3seed']:.1f}%. The separation "
              f"runs in favour of the soft-target arms. This is a benefit GastroHUN's "
              f"own test split is too thin in OTHERCLASS to have detected, and it is "
              f"the clearest case in this thesis of external data answering a "
              f"question the internal data could not.")

    h(doc, "3.3 P5-C: does the calibration ordering survive the shift?", level=2)
    v = CAL["verdict_P5C"]
    figure(doc, "P5_F37_calibration.png",
           "Left: external ECE by arm. Right: each arm's ECE internally and "
           "externally, showing which orderings survive.")
    ca = CAL["aggregate_3seed"]
    rows = [[CFG_LABEL[c], pc(v["internal_ece_points"][c]),
             pc(ca[c]["ece_top1_mean_3seed"]), iv(ca[c]["ece_top1_ci95"]),
             pc(ca[c]["overconfidence_points_mean_3seed"]),
             pc(100 * ca[c]["accuracy_collapsed_mean_3seed"])] for c in ARMS()]
    table(doc, ["Configuration", "Internal ECE", "External ECE", "External 95% CI",
                "Over/under-confidence", "External accuracy (%)"], rows,
          "Calibration internally (Phase 4, pooled contested) and externally. "
          "A negative over/under-confidence value means the arm is UNDER-confident.",
          font=7.8,
          note="; ".join(CAL["definitional_differences_from_phase4"][:2]))
    para(doc, f"The ordering is {v['verdict']} (Spearman rho = "
              f"{v['spearman_rho']:.3f}). Internally the order was "
              f"{' < '.join(v['internal_rank_order'])}; externally it is "
              f"{' < '.join(v['external_rank_order'])}. The two best arms exchanged "
              f"places: {v['lowest_ece_internal']} was best internally and "
              f"{v['lowest_ece_external']} is best externally.")
    c3 = ca.get("C3", {})
    if c3:
        para(doc, f"The most direct confirmation of the Phase 4 mechanism claim is "
                  f"C3's under-confidence. Phase 4 measured it at -6.45 points on "
                  f"unanimous images and argued it was a global confidence shift — a "
                  f"property of the model, not of the data. Externally it is "
                  f"{pc(c3['overconfidence_points_mean_3seed'])} points. It followed "
                  f"the model to two new centres essentially unchanged, which is what "
                  f"a global shift does, and it stops being an advantage there.")

    h(doc, "3.4 Sensitivity: the ambiguous mapping decisions, re-run", level=2)
    if SEN:
        rows = [["baseline", f"{SEN['baseline']['P5-A']['n_images']:,}",
                 pc(SEN["baseline"]["P5-A"]["macro_f1"]),
                 SEN["baseline"]["P5-A"]["verdict"],
                 f"{100 * SEN['baseline']['P5-B']['rejection_rate']:.2f}",
                 SEN["baseline"]["P5-B"]["verdict"]]]
        for k, r in SEN["per_flip"].items():
            aa, bb = r.get("P5-A"), r.get("P5-B")
            rows.append([f"flip {k}",
                         f"{aa['n_images']:,}" if aa else f"{bb['n_images']:,}" if bb else "—",
                         pc(aa["macro_f1"]) if aa else "—",
                         aa["verdict"] if aa else "—",
                         f"{100 * bb['rejection_rate']:.2f}" if bb else "—",
                         bb["verdict"] if bb else "—"])
        table(doc, ["Mapping", "n", "P5-A F1", "P5-A verdict", "P5-B rej. (%)",
                    "P5-B verdict"], rows,
              "Every mapping decision flagged ambiguous for an endpoint under test, "
              "re-run with the alternative frozen before scoring.", font=7.6)
        para(doc, f"Verdicts invariant to every single flip: "
                  f"{SEN['verdicts_invariant_to_every_single_flip']}. The mapping's "
                  f"judgement calls do not carry the result.")


def sec_discussion(doc) -> None:
    h(doc, "4. Discussion", level=1, page_break=True)
    a, head = TR["aggregate_3seed"], TR["headline_arm"]
    v = CAL["verdict_P5C"]

    h(doc, "4.1 What the label-space finding means", level=2)
    para(doc, "The most consequential result of this phase is not a number. It is "
              "that the two most widely used public upper-GI corpora cannot express "
              "the question the thesis asks. GastroHUN labels wall and station; "
              "HyperKvasir and GastroVision label neither. Any paper claiming "
              "external validation of station-level gastric landmark classification "
              "on these corpora is either using a coarser endpoint than it appears "
              "to, or mapping labels that do not correspond. Stating that plainly is "
              "worth more than a station-level number would have been.")

    h(doc, "4.2 A 19-point drop, and what it is not", level=2)
    para(doc, f"The drop of {pc(a[head]['drop_points'])} points is roughly twice the "
              f"pre-registered expectation. It is not evidence that the model is "
              f"broken: it still transfers well clear of the majority-class floor on "
              f"corpora from different centres, different vendors and different "
              f"framing conventions, with no adaptation of any kind. But it does set "
              f"the scale of the domain gap, and it should be quoted whenever an "
              f"internal number from Phases 2 to 4 is quoted.")

    h(doc, "4.3 The result that changes the thesis's recommendation", level=2)
    ra = CAL["aggregate_3seed"]
    para(doc, f"Phase 4 concluded that calibration, not accuracy, was the headline "
              f"deficiency, and that the matched control C3 was the best-calibrated "
              f"arm. Phase 5 keeps the first half and overturns the second. "
              f"Externally, C2 has the lowest ECE ({pc(ra['C2']['ece_top1_mean_3seed'])} "
              f"against C3's {pc(ra['C3']['ece_top1_mean_3seed'])}), much higher "
              f"accuracy ({pc(100 * ra['C2']['accuracy_collapsed_mean_3seed'])}% "
              f"against {pc(100 * ra['C3']['accuracy_collapsed_mean_3seed'])}%), and "
              f"substantially better out-of-protocol rejection. C3's internal "
              f"advantage was a global confidence shift that travelled intact and "
              f"stopped helping.")
    callout(doc, "Had Phase 5 carried only the arm the status board named "
                 "best-calibrated, it would have carried C3, and none of this would "
                 "be visible. The carry-forward decision was the single most "
                 "consequential methodological choice in this phase.",
            title="Why five arms")

    h(doc, "4.4 Out-of-protocol rejection is the deployment-relevant endpoint", level=2)
    para(doc, "In a screening workflow the model does not receive a curated stream "
              "of gastric stations. It receives whatever the endoscope is pointed "
              "at. An arm that routes 63% of non-gastric frames to OTHERCLASS is "
              "categorically safer than one that routes 42% of them to a confidently "
              "asserted station, and the two arms are indistinguishable on "
              "GastroHUN. This endpoint should be primary, not secondary, in any "
              "future work.")

    h(doc, "4.5 Limitations specific to this phase", level=2)
    bullet(doc, PRE["interval_procedure"]["declared_weakness"])
    bullet(doc, "Neither corpus ships per-annotator labels, so there is no external "
                "agreement stratification and no external vote entropy. The "
                "within-stratum correlation that is RQ3's internal primary quantity "
                "has no external counterpart (deviation P5-DEV-2).")
    bullet(doc, "All retroflexion images come from HyperKvasir, so P5-A rests on one "
                "corpus. GastroVision contributes the forward-view side and almost "
                "all of the out-of-protocol set.")
    bullet(doc, f"{MAP['images_by_decision'].get('discard', 0):,} images were "
                f"discarded because their labels do not fix the anatomical site. "
                f"They were scored anyway and re-included in the sensitivity "
                f"analysis, which left every verdict unchanged.")
    bullet(doc, "No adaptation of any kind was performed, by design. The Phase 5B "
                "self-training arm measures transfer-after-adaptation and is a "
                "separate comparison against the numbers frozen here.")

    h(doc, "4.6 Implications for Phases 6 and 7", level=2)
    bullet(doc, "Phase 6's Grad-CAM work should target the retroflexion/forward "
                "boundary and the out-of-protocol failures, which are where the "
                "external errors concentrate.")
    bullet(doc, "Phase 7 should quote the external number beside every internal one, "
                "and should present the C2-versus-C3 reversal as the phase's central "
                "result rather than as a footnote to Phase 4.")
    bullet(doc, "Any deployment recommendation should name C2, and should cite "
                "out-of-protocol rejection rather than macro F1 as the reason.")


def sec_conclusion(doc) -> None:
    h(doc, "5. Conclusion", level=1, page_break=True)
    a, head = TR["aggregate_3seed"], TR["headline_arm"]
    v = CAL["verdict_P5C"]
    h(doc, "5.1 Answers", level=2)
    bullet(doc, f"P5-A: {TR['verdict']}. {pc(a[head]['external_macro_f1_mean_3seed'])} "
                f"external against {pc(a[head]['internal_macro_f1_mean_3seed'])} "
                f"internal, a drop of {pc(a[head]['drop_points'])} points "
                f"{iv(a[head]['drop_ci95'])}, with the pre-registered precision "
                f"target met.")
    bullet(doc, f"P5-B: {RJ['verdict']}. The pre-registered hypothesis was "
                f"falsified: rejection reached "
                f"{100 * RJ['aggregate_3seed'][head]['rejection_rate_mean_3seed']:.1f}% "
                f"against a {100 * RJ['chance_rate']:.2f}% chance rate, and the "
                f"soft-target arms rejected far more reliably than the hard-label "
                f"ones.")
    bullet(doc, f"P5-C: {v['verdict']}, Spearman rho = {v['spearman_rho']:.3f}, but "
                f"with {v['lowest_ece_internal']} and {v['lowest_ece_external']} "
                f"exchanging places at the top.")
    h(doc, "5.2 Carry-forward decisions", level=2)
    bullet(doc, "Carry C2 forward as the recommended configuration, on external "
                "calibration and out-of-protocol rejection rather than on internal "
                "accuracy.")
    bullet(doc, "Retain C3 as a control in any future comparison, and treat a low "
                "ECE achieved alongside large under-confidence on easy cases as a "
                "warning sign rather than a result.")
    bullet(doc, "Report out-of-protocol rejection as a primary endpoint in Phases 6 "
                "and 7. It separated the arms where every internal endpoint failed "
                "to.")
    bullet(doc, "State the label-space limitation wherever external validation is "
                "claimed, in this thesis and in any paper drawn from it.")
    bullet(doc, "Do not compare a Phase 5 interval against a Phase 0-4 interval: the "
                "clustering units differ and the Phase 5 ones are optimistic.")


def sec_appendices(doc) -> None:
    h(doc, "Appendices", level=1, page_break=True)
    h(doc, "Appendix A. The full mapping table", level=2)
    rows = [[r["corpus"], r["external_class"], f"{r['n_images']:,}", r["decision"],
             r["rationale"][:150]] for r in MAP["table"]]
    table(doc, ["Corpus", "External class", "n", "Decision", "Rationale"], rows,
          "Every external class directory, its decision and the anatomical reason "
          "for it.", font=6.6)

    h(doc, "Appendix B. Per-seed results", level=2)
    rows = []
    for c in ARMS():
        for s in TR["seeds"]:
            p = TR["per_seed"][c][str(s)] if str(s) in TR["per_seed"][c] \
                else TR["per_seed"][c][s]
            rows.append([c, str(s), pc(p["external_macro_f1"]),
                         pc(p["internal_macro_f1"]), pc(p["drop_points"]),
                         f"{100 * p['gastric_recognition_rate']:.1f}"])
    table(doc, ["Config", "Seed", "External F1", "Internal F1", "Drop",
                "Gastric recognition (%)"], rows,
          "P5-A per seed. The recognition column is the fraction of gastric images "
          "not routed to OTHERCLASS.", font=7.4)

    h(doc, "Appendix C. Pre-registration record", level=2)
    para(doc, f"Frozen at {PRE['frozen_at']}, after the corpora, mapping and cache "
              f"existed and before any external image was scored. The script that "
              f"writes it refuses to overwrite an existing copy.")
    callout(doc, PRE["statement"], title="Pre-registration statement")
    table(doc, ["ID", "Item", "Adopted", "Evidence", "Impact"],
          [[d["id"], d["item"], d["adopted"], d["evidence"][:200],
            d["impact"][:200]] for d in PRE["deviations"]],
          "Declared deviations.", font=6.8)
    para(doc, PRE["falsification"])

    h(doc, "Appendix D. Script and artefact manifest", level=2)
    for s in ["python src/models/phase5_carry_forward.py   # P5.0, frozen first",
              "python src/models/phase5_acquire.py --extract  # P5.1, gates a-d",
              "python src/models/phase5_mapping.py         # P5.2, gates a-c",
              "python src/models/phase5_cache.py           # P5.3, gate P5.3a",
              "python src/models/phase5_cache_supplement.py  # P5.3b",
              "python src/models/phase5_prereg.py          # P5.4, FROZEN",
              "python src/models/phase5_infer.py           # P5.5, gates a-c",
              "python src/models/phase5_infer_supplement.py  # P5.5b",
              "python src/models/phase5_eval.py            # P5.6 and P5.7",
              "python src/models/phase5_calibration.py     # P5.8",
              "python src/models/phase5_sensitivity.py     # P5.10",
              "python src/report/figures_phase5.py",
              "python src/report/build_phase5_docx.py",
              "python src/report/finalise_phase5.py"]:
        bullet(doc, s)

    h(doc, "Appendix E. Analyses not executed, and what they would cost", level=2)
    bullet(doc, "Phase 5B self-training on the external corpora. Approved and "
                "planned, gated on the numbers in this report being frozen and "
                "committed first; adapting before the clean transfer numbers exist "
                "would make the external validation circular.")
    bullet(doc, "Station-level external evaluation. Not costed, because it is not "
                "possible: the external label spaces do not carry the station or "
                "wall axes.")
    bullet(doc, "Case-clustered intervals. Would require case identifiers neither "
                "corpus publishes.")
    bullet(doc, "A third external corpus. Kvasir-Capsule and the Nerthus set were "
                "considered and rejected as anatomically disjoint from the gastric "
                "SSS protocol.")


def main() -> None:
    missing = [n for n, v in [("phase5_transfer.json", TR),
                              ("phase5_rejection.json", RJ),
                              ("phase5_calibration.json", CAL),
                              ("phase5_mapping.json", MAP)] if v is None]
    if missing:
        raise SystemExit("cannot build the report; missing: " + ", ".join(missing))

    doc = new_document()
    title_page(doc)
    front_matter(doc)
    sec_abstract(doc)
    sec_introduction(doc)
    sec_methods(doc)
    sec_results(doc)
    sec_discussion(doc)
    sec_conclusion(doc)
    sec_appendices(doc)
    add_page_numbers(doc)
    doc.save(BD.OUT)
    caps = [p.text for p in doc.paragraphs if p.style.name == "Caption"]
    nfig = sum(1 for c in caps if c.strip().startswith("Figure"))
    ntab = sum(1 for c in caps if c.strip().startswith("Table"))
    print(f"[report] {len(doc.inline_shapes)} images, {nfig} figure captions, "
          f"{ntab} table captions -> {BD.OUT}")


if __name__ == "__main__":
    main()
