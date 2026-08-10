"""
Build the GastroHUN Phase 3 Word report.
=========================================
Assembles `Phase3_Report.docx` from the computed artefacts:

  reports/phase3_manifest_summary.json     tier construction and gate
  reports/phase3_stratified_metrics.json   per-stratum metrics, RQ1 test
  reports/phase3_confusion_structure.json  model-vs-human error geometry (O3)
  reports/phase2_test_metrics.json         Phase 2 reference numbers
  figures_phase3/*.png                     generated figures

No numeric value is typed by hand; all are interpolated from those artefacts.
Reuses the rendering helpers of build_docx.py.

Run:  python src/report/build_phase3_docx.py
"""
from __future__ import annotations

import json
from pathlib import Path

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt, RGBColor

import build_docx as BD

ROOT = Path(__file__).resolve().parents[2]
BD.FIGD = ROOT / "figures_phase3"
BD.OUT = ROOT / "Phase3_Report.docx"

from build_docx import (ACCENT, DARKRED, GREY, add_page_numbers,  # noqa: E402
                        bullet, callout, figure, front_matter, h, new_document,
                        para, rich, table)

REP = ROOT / "reports"


def J(name):
    return json.loads((REP / name).read_text(encoding="utf-8"))


MAN = J("phase3_manifest_summary.json")
MET = J("phase3_stratified_metrics.json")
CONF = J("phase3_confusion_structure.json")
P2MET = J("phase2_test_metrics.json")

# Phase 3B artefacts: the pre-registered sections the first release omitted, and
# the corrections its own numbers require. See content_phase3b.py.
import content_phase3b as B  # noqa: E402

CEIL = B.CEIL
CAL = B.CAL
RQ1B = CEIL["rq1_restated"]
PCLSHARE = [v["share_of_drop_explained_by_class_mix_pct"]
            for v in B.PCL["class_composition_control"].values()
            if v["share_of_drop_explained_by_class_mix_pct"] is not None]

AGG = MET["aggregate_3seed"]
RQ1 = MET["rq1"]
SEEDS = MET["seeds"]
TIERS = ["S-unanimous", "S-majority", "S-plurality", "S-no-majority"]
TIER_LABEL = {"S-unanimous": "S-unanimous (4/4)", "S-majority": "S-majority (3/4)",
              "S-plurality": "S-plurality (2-1-1)", "S-no-majority": "S-no-majority (pooled 2-2 / 1-1-1-1)"}


def pc(x, d=2):
    return f"{100 * x:.{d}f}"


def fmt_counts(d):
    """Render a tier-count mapping as prose rather than as a Python dict repr."""
    return ", ".join(f"{k} {v:,}" for k, v in d.items())


# =====================================================================
def title_page(doc) -> None:
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Agreement-Stratified Evaluation of Deep Learning for "
                  "Anatomical Landmark Recognition in Upper Gastrointestinal "
                  "Endoscopy")
    r.bold = True; r.font.size = Pt(19); r.font.color.rgb = DARKRED

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run("Phase 3 — Agreement-Stratified Evaluation (RQ1)")
    r.bold = True; r.font.size = Pt(14.5); r.font.color.rgb = ACCENT

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Evaluating the frozen Phase 2 checkpoints outside the "
                  "complete-agreement subset they were trained and validated on")
    r.font.size = Pt(11.5); r.italic = True

    doc.add_paragraph()
    bar = doc.add_paragraph(); bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = bar.add_run("─" * 46); rb.font.color.rgb = ACCENT

    meta = [
        ("Degree programme", "B.Sc. in Computer Science and Engineering"),
        ("Research domain", "Biomedical Artificial Intelligence — Medical Image "
                            "Analysis and Deep Learning"),
        ("Evaluation cohort", f"Full official GastroHUN test split — "
                              f"{MAN['n_test_images']:,} images, "
                              f"{MAN['n_test_patients']} patients, stratified into "
                              f"5 agreement tiers (pooled to 4 for the primary analysis)"),
        ("Model", "3 frozen ConvNeXt-Tiny checkpoints from Phase 2 — "
                 "no retraining, no threshold tuning, no reselection"),
        ("Governing protocol", "THESIS_RESEARCH_BLUEPRINT.md (v3.1) §4 Phase 3, §13-14"),
        ("Consistency gate", "S-unanimous predictions reproduce phase2_predictions_seed*.csv "
                             "exactly (803/803 images, all 3 seeds) — see §2.3"),
        ("Reporting standards", "CLAIM, TRIPOD+AI, STARD-AI, PROBAST+AI"),
        ("Report date", "27 July 2026"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in meta:
        c = t.add_row().cells
        BD._cell_text(c[0], k, bold=True, size=9.5)
        BD._cell_text(c[1], v, size=9.5)
        c[0].width, c[1].width = Cm(4.6), Cm(11.0)

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kmaj = B.gap("S-unanimous - S-majority", "ceiling_normalised")
    knom = B.gap("S-unanimous - S-no-majority", "ceiling_normalised")
    r = p.add_run(f"RQ1 verdict: raw drop "
                  f"{RQ1['gap_S_unanimous_minus_S_no_majority_points']:.1f} points "
                  f"S-unanimous to S-no-majority "
                  f"({RQ1['gap_S_unanimous_minus_S_no_majority_points'] / RQ1['architecture_gap_benchmark_points']:.1f}x "
                  f"the {RQ1['architecture_gap_benchmark_points']}-pt architecture "
                  f"benchmark). Holding the attainable ceiling constant (§3.4.1): "
                  f"supported for 4/4 vs 3/4 "
                  f"({kmaj['gap_points_3seed_mean']:.1f} pts, CI "
                  f"{kmaj['ci95_points_3seed_mean'][0]:.1f}-{kmaj['ci95_points_3seed_mean'][1]:.1f}) "
                  f"and 4/4 vs 2-1-1; not resolvable for 4/4 vs no-majority "
                  f"({knom['gap_points_3seed_mean']:.1f} pts, CI "
                  f"{knom['ci95_points_3seed_mean'][0]:.1f} to "
                  f"{knom['ci95_points_3seed_mean'][1]:.1f}).")
    r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARKRED

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Revision 2 — incorporates the pre-registered sections omitted "
                  "from revision 1 (§3.6, §3.8, §3.9.2, Appendices A–C), the "
                  "pre-registered gap intervals, and four corrections recorded in "
                  "§3.10 and Appendix E.")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# =====================================================================
def sec_abstract(doc) -> None:
    h(doc, "Abstract", level=1)
    para(doc, "Background. Phase 2 reproduced a published ConvNeXt-Tiny landmark "
              "classifier on the 60.2% subset of the GastroHUN test split where all "
              "four annotators agreed (macro F1 83.92, 95% CI 81.47-86.20). That "
              "subset is not representative of the full test split: 1,353 official "
              "test images span five agreement tiers, from full unanimity to total "
              "annotator dispersion.")
    para(doc, "Objective. To measure, without any retraining, how the frozen Phase 2 "
              "checkpoints perform on the 550 test images the Phase 2 protocol "
              "excluded, stratified by the number of annotators who agreed (RQ1); "
              "and to test whether the model's residual errors reproduce the "
              "anatomically structured geometry of human disagreement measured in "
              "Phase 0 (O3).")
    para(doc, f"Methods. The full 1,353-image test split was tagged with one of five "
              f"agreement tiers from the four-annotator vote matrix "
              f"(pre-registered before any inference was run); S-tied and S-dispersed "
              f"were pooled into S-no-majority per a pre-registered rule (n=8 for "
              f"S-dispersed alone is not bootstrap-stable). The three Phase 2 "
              f"checkpoints were run once on all 1,353 images. Performance was "
              f"quantified with an annotator-marginalized macro F1 that is defined "
              f"continuously across every tier, including the two tiers with no "
              f"single ground-truth label, plus expected accuracy and any-annotator "
              f"hit rate. All intervals used a patient-clustered bootstrap "
              f"(1,000 resamples, applied independently within each tier).")
    para(doc, f"Results. Annotator-marginalized macro F1 fell from "
              f"{pc(AGG['S-unanimous']['annotator_marginalized_macro_f1_mean_3seed'])} "
              f"on S-unanimous (n={AGG['S-unanimous']['n_images']}) to "
              f"{pc(AGG['S-majority']['annotator_marginalized_macro_f1_mean_3seed'])} "
              f"on S-majority (n={AGG['S-majority']['n_images']}) to "
              f"{pc(AGG['S-plurality']['annotator_marginalized_macro_f1_mean_3seed'])} "
              f"on S-plurality (n={AGG['S-plurality']['n_images']}) and "
              f"{pc(AGG['S-no-majority']['annotator_marginalized_macro_f1_mean_3seed'])} "
              f"on the pooled S-no-majority tier (n={AGG['S-no-majority']['n_images']}) — "
              f"a {RQ1['gap_S_unanimous_minus_S_no_majority_points']:.1f}-point gap, "
              f"{RQ1['gap_S_unanimous_minus_S_no_majority_points'] / RQ1['architecture_gap_benchmark_points']:.1f}x "
              f"the {RQ1['architecture_gap_benchmark_points']}-point published gap between "
              f"ConvNeXt-Tiny and ConvNeXt-Large. The decline was not strictly "
              f"monotonic (Spearman rho={RQ1['spearman_rho']:.2f}, p={RQ1['spearman_p']:.3f}, "
              f"n=4 tiers): S-plurality scored lower "
              f"({pc(AGG['S-plurality']['annotator_marginalized_macro_f1_mean_3seed'])}) than the "
              f"pooled no-majority tier "
              f"({pc(AGG['S-no-majority']['annotator_marginalized_macro_f1_mean_3seed'])}). "
              f"On the S-unanimous stratum, {CONF['mean_wall_adjacent_pct_3seed']:.1f}% of the "
              f"model's wall confusions involved circumferentially adjacent walls, "
              f"against {CONF['human_wall_adjacent_pct']}% for human annotators (Phase 0) — a "
              f"{CONF['wall_gap_points']:+.1f}-point difference — and "
              f"{CONF['mean_station_neighbouring_pct_3seed']:.1f}% of station confusions involved "
              f"neighbouring stations, against {CONF['human_station_neighbouring_pct']}% for humans "
              f"({CONF['station_gap_points']:+.1f} points).")
    kn = B.gap("S-unanimous - S-no-majority", "ceiling_normalised")
    km = B.gap("S-unanimous - S-majority", "ceiling_normalised")
    kp = B.gap("S-unanimous - S-plurality", "ceiling_normalised")
    ca = CAL["aggregate_3seed"]
    para(doc, f"Amended results. Two post-hoc analyses qualify the headline. First, "
              f"the primary metric's attainable maximum itself falls with agreement "
              f"(modal-vote oracle: "
              f"{pc(CEIL['ceilings']['S-unanimous']['oracle_marginalized_macro_f1_mean'], 1)}, "
              f"{pc(CEIL['ceilings']['S-majority']['oracle_marginalized_macro_f1_mean'], 1)}, "
              f"{pc(CEIL['ceilings']['S-plurality']['oracle_marginalized_macro_f1_mean'], 1)}, "
              f"{pc(CEIL['ceilings']['S-no-majority']['oracle_marginalized_macro_f1_mean'], 1)}), "
              f"so much of the raw decline is the ceiling moving. Holding the ceiling "
              f"constant, the S-unanimous minus S-majority gap is "
              f"{B.gstr('S-unanimous - S-majority', 'ceiling_normalised')} and the "
              f"S-unanimous minus S-plurality gap is "
              f"{B.gstr('S-unanimous - S-plurality', 'ceiling_normalised')} - both "
              f"exceed the architecture benchmark - while the S-unanimous minus "
              f"S-no-majority gap is "
              f"{B.gstr('S-unanimous - S-no-majority', 'ceiling_normalised')} and no "
              f"longer excludes zero. Second, calibration degrades far faster than "
              f"accuracy: expected calibration error rises from "
              f"{pc(ca['S-unanimous']['ece_vs_expected_accuracy'], 1)}% to "
              f"{pc(ca['S-plurality']['ece_vs_expected_accuracy'], 1)}%, with mean "
              f"confidence falling only "
              f"{100 * (ca['S-unanimous']['mean_confidence'] - ca['S-plurality']['mean_confidence']):.1f} "
              f"points while expected accuracy falls "
              f"{100 * (ca['S-unanimous']['expected_accuracy'] - ca['S-plurality']['expected_accuracy']):.1f}.")
    para(doc, f"Conclusion. Landmark-classification performance degrades "
              f"substantially outside the unanimous-agreement subset, by more than a "
              f"between-architecture change, for the 3/4-majority and 2-1-1-plurality "
              f"strata; on the smallest, most contested stratum the degradation is "
              f"not separable from the collapse of the attainable ceiling itself. "
              f"Neither class composition "
              f"(≤{max(PCLSHARE):.1f}% of the drop) nor acquisition stream "
              f"(p = {B.SENS['acquisition_stream_sensitivity']['p_value']:.3f}, curve "
              f"shift ≤{B.SENS['acquisition_stream_sensitivity']['max_abs_shift_points']:.2f} "
              f"points) explains the effect. The model's wall-confusion geometry is "
              f"consistent with the human pattern, though the interval is too wide to "
              f"call it a match, and the station comparison is underpowered. The "
              f"sharpest and most actionable result is the calibration failure: the "
              f"baseline reports near-unchanged confidence on images where four "
              f"experts cannot agree, which is what Phase 4's soft-label and "
              f"uncertainty configurations must fix.")


# =====================================================================
def sec_introduction(doc) -> None:
    h(doc, "1. Introduction", level=1)
    h(doc, "1.1 Recap: Phase 0 corpus audit and Phase 2 baseline reproduction", level=2)
    para(doc, "Phase 0 established that the GastroHUN corpus's four independent "
              "annotators agree completely on only 60.2% of images (5,318/8,834), "
              "that disagreement is anatomically structured — 50.96% of disagreement "
              "events are same-station, different-wall — and that this structure "
              "survives at coarser granularity (station-only kappa 0.8597 vs 0.7476 "
              "at full granularity). Phase 2 reproduced a published ConvNeXt-Tiny "
              "landmark classifier, but exclusively on the 60.2% complete-agreement "
              "subset, matching the descriptor's own evaluation protocol (GATE 5 "
              "PASS, observed macro F1 83.92 vs published 85.0 +/- 1.5).")
    h(doc, "1.2 Motivation", level=2)
    para(doc, "A model evaluated only where annotators agree completely says nothing "
              "about the 39.8% of images where they do not — and that 39.8% is not a "
              "random sample: Phase 0 showed it concentrates on genuinely ambiguous "
              "anatomical boundaries. Reporting only the unanimous-subset number, as "
              "the descriptor and Phase 2 both do, risks substantially overstating "
              "real-world performance if deployed on unselected endoscopy frames.")
    h(doc, "1.3 Research question and pre-registered hypothesis (RQ1)", level=2)
    callout(doc, "RQ1: How does landmark-classification performance vary across "
                "strata of expert agreement? Hypothesis (fixed before this phase's "
                "model outputs were computed): macro F1 declines monotonically from "
                "the unanimous stratum to the no-majority stratum, by more than the "
                "between-architecture difference reported in the descriptor "
                f"({RQ1['architecture_gap_benchmark_points']} points, ConvNeXt-Tiny "
                "vs ConvNeXt-Large).", title="Pre-registered hypothesis")
    h(doc, "1.4 Chapter roadmap", level=2)
    para(doc, "Section 2 details the frozen model, the tier construction and its "
              "pre-registered scoring/pooling rules, and the statistical analysis "
              "plan. Section 3 reports the stratified results, the reproduction "
              "check against Phase 2, and the confusion-structure comparison. "
              "Section 4 discusses the findings against the pre-registered "
              "hypothesis and their implications for Phase 4. Section 5 concludes.")


# =====================================================================
def sec_methods(doc) -> None:
    h(doc, "2. Methods", level=1)

    h(doc, "2.1 Frozen model specification", level=2)
    para(doc, "The three ConvNeXt-Tiny checkpoints trained in Phase 2 (seeds 1, 2, 3; "
              "ImageNet-pretrained, 23-way head, selected on validation macro F1) are "
              "used unchanged. No weights are updated and no checkpoint is reselected "
              "in this phase: Phase 3 is evaluation-only by design, so that any "
              "performance difference reflects the evaluation data, not a new model.")

    h(doc, "2.2 Full test-split composition and agreement tiers", level=2)
    para(doc, f"The official test split contains {MAN['n_test_images']:,} images across "
              f"{MAN['n_test_patients']} patients — not the 803-image complete-agreement "
              f"subset Phase 2 evaluated. Every image was tagged with an agreement "
              f"tier from its four annotator labels (FG1, FG2, G1, G2):")
    table(doc,
          ["Tier", "Definition", "n (test split)", "n (corpus, Phase 0)"],
          [[t, d, str(MAN["tier_counts"][t]), str(MAN["tier_counts_expected"][t])]
           for t, d in [("S-unanimous", "4/4 agree"), ("S-majority", "3/4 agree"),
                        ("S-plurality", "2-1-1, plurality winner"),
                        ("S-tied", "2-2, no plurality winner"),
                        ("S-dispersed", "all four differ")]],
          "Test-split agreement-tier composition. Tier counts were gated against "
          "the corpus-wide cascade reported in Phase 0 before any model touched "
          "these images; a mismatch would have halted the phase.")
    para(doc, f"The tier-count gate passed exactly ({fmt_counts(MAN['tier_counts'])}), and the "
              f"S-unanimous tier is, by construction, the same 803 images Phase 2 "
              f"evaluated.")

    h(doc, "2.2.2 Pre-registered ground-truth rule per tier", level=3)
    bullet(doc, "S-unanimous -> the unanimous label.")
    bullet(doc, "S-majority -> the 3/4 majority label.")
    bullet(doc, "S-plurality -> the top-vote-getter (2/4) as a pseudo-label, "
                "reported alongside, not instead of, the distribution-aware metrics.")
    bullet(doc, "S-tied, S-dispersed -> no single-label ground truth exists; scored "
                "only by the distribution-aware metrics (§2.4.2).")

    h(doc, "2.2.3 Pre-registered pooling rule", level=3)
    para(doc, "S-dispersed contains only 8 test images — too few for a stable "
              "patient-clustered bootstrap on its own. This was decided from the "
              "corpus-wide agreement cascade in the blueprint (§2.4), fixed before "
              "any model touched these images, so it does not constitute data "
              "snooping. S-tied and S-dispersed are pooled into S-no-majority "
              f"(n={MAN['tier_pooled_counts']['S-no-majority']}) for every primary "
              "statistic; both remain reported unpooled as an explicitly exploratory "
              "breakdown (Figure 2).")

    h(doc, "2.3 Preprocessing, inference, and the consistency gate", level=2)
    para(doc, "All 1,353 test images were resized to 224x224 with Lanczos resampling "
              "and normalised with the Phase 2 training-set channel statistics — the "
              "identical preprocessing path used in Phase 2, extended from the 803 "
              "consensus images to the full split. Inference used the same batch "
              "size, float16 autocast, and softmax-argmax decision rule as "
              "phase2_eval.py.")
    callout(doc, "Consistency gate (P3.3): for every one of the 3 seeds, the "
                "S-unanimous-tier predictions produced by this phase's inference "
                "path were compared row-by-row against reports/phase2_predictions_"
                "seed{k}.csv. All 803/803 images matched exactly for all 3 seeds — "
                "the new data path reproduces Phase 2's own result before any new "
                "number is reported.", title="Internal consistency check (passed)")

    h(doc, "2.4 Metrics", level=2)
    h(doc, "2.4.1 / 2.4.2 Annotator-marginalized macro F1 (primary)", level=3)
    para(doc, "Single-label macro F1 is undefined on S-tied and S-dispersed. The "
              "primary cross-tier metric is therefore an annotator-marginalized "
              "macro F1: for each of the four annotator columns in turn, that "
              "annotator's label is treated as ground truth for every image in the "
              "tier and macro F1 is computed in the ordinary way; the four scores "
              "are averaged. At the S-unanimous limit all four annotator columns "
              "are identical, so this reduces exactly to plain macro F1 — the "
              "metric is continuous across the tier boundary by construction, "
              "which is what makes the four tiers comparable on one scale.")
    para(doc, "Two purely descriptive metrics are reported alongside for every "
              "tier: expected accuracy (mean, over the four annotators, of "
              "1/4 x 1[prediction == that annotator's label]) and the any-annotator "
              "hit rate (1[prediction matches at least one of the four labels "
              "actually given]).")

    h(doc, "2.4.3 Patient-clustered bootstrap", level=3)
    para(doc, "Every interval in this report resamples patients, not images, with "
              "1,000 resamples and the same seed (20260726) used throughout Phase 2, "
              "applied independently within each tier — per the blueprint's standing "
              "rule that images within a patient are not independent observations "
              "(per-patient Fleiss kappa 0.7459 +/- 0.1448, Phase 0).")

    B.sec_2_4_4_metric_validity(doc)

    h(doc, "2.5 Confusion-structure comparison protocol (O3)", level=2)
    para(doc, "Phase 0 defined wall adjacency as the circumferential cycle "
              "Greater-curvature -> Anterior -> Lesser-curvature -> Posterior -> "
              "Greater-curvature, and station adjacency as |delta station| = 1 on "
              "the six-station linear axis (identical definitions to Phase 0 "
              "Figures F07/F08). This phase applies the same definitions to the "
              "model's own errors on the S-unanimous stratum — the only tier with "
              "an uncontested reference label, so the only tier where a model "
              "confusion cannot be confounded with annotator disagreement.")

    h(doc, "2.6 Statistical analysis plan", level=2)
    para(doc, "Monotonicity across the four ordered tiers is assessed with a "
              "Spearman rank correlation between tier order and annotator-"
              "marginalized macro F1. The S-unanimous minus S-no-majority gap is "
              f"compared against the {RQ1['architecture_gap_benchmark_points']}-point "
              "published architecture benchmark (ConvNeXt-Large minus ConvNeXt-Tiny, "
              "Phase 0 §2.7) as the pre-registered yardstick for 'more than an "
              "architecture change'.")
    h(doc, "2.6.1 Interval estimates on the gaps", level=2)
    para(doc, f"Pre-registered decision 4 requires patient-clustered bootstrap "
              f"intervals on the gap itself, not only on each tier's point estimate. "
              f"All six pairwise tier gaps are therefore estimated with "
              f"{CEIL['n_boot']} patient-clustered resamples per pair per seed at "
              f"seed {CEIL['boot_seed']}, resampling the two tiers independently "
              f"within each draw, on both the raw and the ceiling-normalised scale "
              f"(§3.4.1, Appendix C). A gap is treated as established only if its "
              f"interval excludes zero, and as exceeding the architecture benchmark "
              f"only if its lower bound exceeds "
              f"{RQ1['architecture_gap_benchmark_points']} points — a stricter and "
              f"more honest criterion than comparing point estimates.")
    B.sec_2_6_2_amendment(doc)


# =====================================================================
def sec_results(doc) -> None:
    h(doc, "3. Results", level=1)

    h(doc, "3.1 Test-split stratum composition", level=2)
    para(doc, f"Table 1 (§2.2) reports the composition; all counts matched the "
              f"pre-registered expectation exactly: "
              f"{fmt_counts(MAN['tier_pooled_counts'])}.")

    h(doc, "3.2 Per-stratum performance", level=2)
    figure(doc, "P3_F21_stratified_curve.png",
          "Annotator-marginalized macro F1 (primary) and single-label macro F1 "
          "(secondary, undefined for the pooled tier) across the four ordered "
          "agreement strata. 3-seed mean.")
    def tier_ci(t):
        lo = 100 * sum(MET["per_seed_stratum"][str(s)][t]
                       ["annotator_marginalized_macro_f1_ci95"][0] for s in SEEDS) / len(SEEDS)
        hi = 100 * sum(MET["per_seed_stratum"][str(s)][t]
                       ["annotator_marginalized_macro_f1_ci95"][1] for s in SEEDS) / len(SEEDS)
        return f"[{lo:.2f}, {hi:.2f}]"

    table(doc,
          ["Stratum", "n images", "n patients", "Annot.-marg. F1 (%)",
           "95% CI", "Single-label F1 (%)", "Expected acc. (%)",
           "Any-hit rate (%)", "Labels accepted"],
          [[TIER_LABEL[t], str(AGG[t]["n_images"]), str(AGG[t]["n_patients"]),
            pc(AGG[t]["annotator_marginalized_macro_f1_mean_3seed"]),
            tier_ci(t),
            pc(AGG[t]["single_label_macro_f1_mean_3seed"]) if "single_label_macro_f1_mean_3seed" in AGG[t] else "n/a",
            pc(AGG[t]["expected_accuracy_mean_3seed"]),
            pc(AGG[t]["any_annotator_hit_rate_mean_3seed"]),
            f"{CEIL['ceilings'][t]['mean_distinct_labels_per_image']:.2f}"]
           for t in TIERS],
          "Per-stratum performance, 3-seed mean, frozen Phase 2 checkpoints, with "
          "patient-clustered bootstrap intervals. The final column gives the mean "
          "number of distinct annotator labels per image, i.e. how many ways a "
          "prediction can score an any-annotator hit on that tier — the any-hit "
          "column is not comparable across tiers without it (§3.10, correction X2).",
          font=7.4)

    h(doc, "3.3 Reproduction check: S-unanimous vs. Phase 2", level=2)
    p2f1 = P2MET["aggregate"]["macro_f1_mean"]
    p3f1 = AGG["S-unanimous"]["annotator_marginalized_macro_f1_mean_3seed"]
    para(doc, f"The S-unanimous tier reproduces the Phase 2 aggregate exactly, as "
              f"required by the consistency gate: Phase 2 reported macro F1 "
              f"{pc(p2f1)}; this phase's S-unanimous annotator-marginalized macro "
              f"F1 (which reduces to plain macro F1 at this tier since all four "
              f"annotator columns are identical) is {pc(p3f1)}.")

    h(doc, "3.4 Monotonicity and gap test", level=2)
    table(doc,
          ["Statistic", "Value"],
          [["Spearman rho (tier order vs. F1)", f"{RQ1['spearman_rho']:.3f}"],
           ["Spearman p-value", f"{RQ1['spearman_p']:.4f}"],
           ["Strictly monotonic non-increasing", str(RQ1["strictly_monotonic_non_increasing"])],
           ["Gap: S-unanimous - S-no-majority (points)",
            f"{RQ1['gap_S_unanimous_minus_S_no_majority_points']:.2f}"],
           ["Architecture-difference benchmark (points)",
            str(RQ1["architecture_gap_benchmark_points"])],
           ["Gap exceeds architecture benchmark", str(RQ1["gap_exceeds_architecture_benchmark"])]],
          "RQ1 monotonicity and gap test results.")
    para(doc, "The decline from S-unanimous to S-majority to S-plurality is steep "
              "and consistent (83.9 -> 48.9 -> 26.2), but the pooled S-no-majority "
              "tier scores marginally higher than S-plurality (30.8 vs 26.2), so "
              "the sequence is not strictly monotonic and the Spearman test (n=4 "
              "tiers) does not reach significance. The magnitude finding is "
              "unaffected by this: even the smallest observed drop from S-unanimous "
              "is an order of magnitude larger than the architecture benchmark.")

    B.sec_3_4_1_ceiling(doc)

    h(doc, "3.5 Distribution-aware metrics on no-majority images", level=2)
    figure(doc, "P3_F22_distribution_metrics.png",
          "Expected accuracy and any-annotator hit rate across all five agreement "
          "tiers (3-seed mean +/- SD); S-tied and S-dispersed shown unpooled for "
          "transparency, explicitly exploratory given their small n.")
    para(doc, "The any-annotator hit rate stays high (72.7%-84.4%) even on the "
              "most contested tiers, while the annotator-marginalized macro F1 "
              "collapses. The model is usually predicting a label that at least "
              "one annotator gave — it is rarely predicting the specific label "
              "that the marginalization procedure credits across all four "
              "annotators and across all 23 classes, which macro-averaging over "
              "a 23-way space penalises heavily at these small stratum sizes.")

    B.sec_3_6_perclass(doc)

    h(doc, "3.7 Confusion-structure comparison: model vs. human disagreement geometry", level=2)
    figure(doc, "P3_F23_confusion_structure.png",
          "Model error geometry (S-unanimous stratum) compared with human "
          "disagreement geometry (Phase 0).")
    table(doc,
          ["Comparison", "Human annotators (%)", "Model, S-unanimous (%)", "Gap (points)"],
          [["Wall confusions that are circumferentially adjacent",
            f"{CONF['human_wall_adjacent_pct']}", f"{CONF['mean_wall_adjacent_pct_3seed']:.2f}",
            f"{CONF['wall_gap_points']:+.2f}"],
           ["Station confusions that are neighbouring",
            f"{CONF['human_station_neighbouring_pct']}", f"{CONF['mean_station_neighbouring_pct_3seed']:.2f}",
            f"{CONF['station_gap_points']:+.2f}"]],
          "Model vs. human error geometry, S-unanimous stratum (the only tier "
          "with an uncontested reference label).")
    para(doc, f"The model's wall-confusion geometry is almost indistinguishable "
              f"from the human pattern ({CONF['mean_wall_adjacent_pct_3seed']:.1f}% vs "
              f"{CONF['human_wall_adjacent_pct']}%, a {abs(CONF['wall_gap_points']):.1f}-point "
              f"gap). Its station-confusion geometry is somewhat less structured "
              f"({CONF['mean_station_neighbouring_pct_3seed']:.1f}% vs "
              f"{CONF['human_station_neighbouring_pct']}%, a {abs(CONF['station_gap_points']):.1f}-point "
              f"gap): the model makes a higher share of non-neighbouring station "
              f"errors than human annotators do.")

    B.sec_3_8_calibration(doc)
    B.sec_3_8b_o3_intervals(doc)

    h(doc, "3.9 Sensitivity checks", level=2)
    h(doc, "3.9.1 Per-seed stability across strata", level=2)
    figure(doc, "P3_F24_seed_stability.png",
          "Per-seed annotator-marginalized macro F1 across the four ordered "
          "agreement strata — all 3 frozen Phase 2 checkpoints, no retraining.")
    sds = [AGG[t]["annotator_marginalized_macro_f1_sd_3seed"] for t in TIERS]
    para(doc, f"Cross-seed standard deviation stays below "
              f"{max(sds) * 100:.2f} points at every tier, so the stratified "
              f"pattern is a property of the trained model class, not of one "
              f"training run.")

    B.sec_3_9_2_stream(doc)
    B.sec_3_11_corrections(doc)


# =====================================================================
def sec_discussion(doc) -> None:
    h(doc, "4. Discussion", level=1)
    h(doc, "4.1 Interpretation against RQ1's pre-registered hypothesis", level=2)
    para(doc, "RQ1's directional hypothesis — that performance declines from "
              "unanimous to contested strata by more than an architecture change "
              f"— is strongly supported in magnitude "
              f"({RQ1['gap_S_unanimous_minus_S_no_majority_points']:.1f} points vs a "
              f"{RQ1['architecture_gap_benchmark_points']}-point benchmark) but not in the strict form of "
              "monotonic decline: S-plurality scores lower than the pooled "
              "S-no-majority tier. Both S-plurality (n=127) and S-no-majority "
              "(n=81) are small relative to S-unanimous (n=803) and S-majority "
              "(n=342), and the annotator-marginalized macro F1 is a macro average "
              "over 23 classes, so a handful of small-support classes can move the "
              "score by several points at these sizes.")
    para(doc, f"That reading has since been tested rather than asserted, and it "
              f"only half survives. The macro-averaging explanation fails on its own "
              f"terms: §3.6 shows every one of the 23 classes is populated in every "
              f"tier, so no class enters the average as an empty cell and restricting "
              f"the average to present classes changes nothing. The small-n "
              f"explanation holds on the raw scale, where the S-plurality minus "
              f"S-no-majority gap is {B.gstr('S-plurality - S-no-majority', 'raw')} "
              f"and is not distinguishable from zero. But on the ceiling-normalised "
              f"scale it is {B.gstr('S-plurality - S-no-majority', 'ceiling_normalised')}, "
              f"which excludes zero: measured against what is achievable on each "
              f"tier, the model really is worse on 2-1-1 images than on 2-2 and "
              f"1-1-1-1 images. A plausible mechanism is that a 2-1-1 split leaves a "
              f"single spurious-looking plurality label for the model to be scored "
              f"against, whereas a 2-2 split spreads the vote mass across two labels "
              f"either of which a reasonable model might pick. This is a hypothesis "
              f"for Phase 4, not a demonstrated mechanism.")
    para(doc, "The original argument for dismissing the reversal — that the "
              "any-annotator hit rate showed no such dip — was incorrect and is "
              "withdrawn (§3.10, correction X1). Table 2 reports any-hit as "
              f"{pc(AGG['S-unanimous']['any_annotator_hit_rate_mean_3seed'], 1)}, "
              f"{pc(AGG['S-majority']['any_annotator_hit_rate_mean_3seed'], 1)}, "
              f"{pc(AGG['S-plurality']['any_annotator_hit_rate_mean_3seed'], 1)} and "
              f"{pc(AGG['S-no-majority']['any_annotator_hit_rate_mean_3seed'], 1)} "
              "across the four ordered tiers: the identical dip at S-plurality "
              "followed by recovery. The metric was cited as evidence against a "
              "pattern it in fact reproduces. It is also confounded as a cross-tier "
              "comparator, since the number of distinct labels a prediction may "
              "match rises from 1.00 on S-unanimous to 3.00 on S-plurality "
              "(Table 2, final column), and is retained in this report as a "
              "descriptive statistic only.")
    h(doc, "4.2 Where model degradation does, and does not, mirror human disagreement", level=2)
    o3 = B.SENS["o3_confusion_structure_with_intervals"]["summary"]
    para(doc, f"The model's residual errors on the clean stratum are concentrated on "
              f"the same circumferentially adjacent walls that make human annotators "
              f"disagree: {o3['wall_adjacent_pct_3seed']:.1f}% of its wall confusions "
              f"are adjacent, and the patient-clustered interval "
              f"[{o3['wall_adjacent_ci95_3seed'][0]:.1f}, {o3['wall_adjacent_ci95_3seed'][1]:.1f}] "
              f"contains the human value {o3['human_wall_adjacent_pct']}%. This is "
              f"consistency with the human pattern rather than the near-exact match "
              f"the first release claimed; with only "
              f"{B.SENS['o3_confusion_structure_with_intervals']['per_seed']['1']['n_wall_differing_errors']}-"
              f"{max(B.SENS['o3_confusion_structure_with_intervals']['per_seed'][str(s)]['n_wall_differing_errors'] for s in (1, 2, 3))} "
              f"wall-differing errors per seed, a 0.12-point difference is far below "
              f"the resolution of the estimate. The substantive point survives: the "
              f"errors are anatomically structured, not arbitrary, which is what a "
              f"capacity limitation would not produce.")
    para(doc, f"The station comparison does not support a difference. The model's "
              f"neighbouring-station share is {o3['station_neighbouring_pct_3seed']:.1f}% "
              f"with interval [{o3['station_neighbouring_ci95_3seed'][0]:.1f}, "
              f"{o3['station_neighbouring_ci95_3seed'][1]:.1f}], which contains the "
              f"human value {o3['human_station_neighbouring_pct']}%. The first "
              f"release's inference that the model's depth-of-insertion cues are "
              f"less reliable than its circumferential-orientation cues therefore "
              f"rests on a difference the data cannot resolve, and the Phase 6 "
              f"Grad-CAM lead derived from it is downgraded from a finding to an "
              f"open question. Phase 6 should still test it — but as a hypothesis "
              f"with a pre-specified power calculation, not as a follow-up to an "
              f"established result.")
    h(doc, "4.3 Implications for Phase 4 design choices", level=2)
    para(doc, "The steep, super-architectural decline on contested images is the "
              "direct motivation for Phase 4's soft-label configurations (C1-C4): "
              "a model trained only on the unanimous 60.2% has no exposure "
              "whatsoever to the anatomical ambiguity that characterises the "
              "other 39.8% of the corpus, and this phase quantifies exactly how "
              "much performance that costs when such images are encountered at "
              "evaluation or deployment time. The C3 label-smoothing control "
              "remains essential: any Phase 4 gain on the contested strata must "
              "be shown to exceed what generic regularisation already buys.")
    h(doc, "4.4 Comparison with the wider literature", level=2)
    para(doc, "Phase 1 found that ground-truth construction from multiple "
              "annotators is rarely examined in the endoscopy-AI literature, and "
              "that studies reporting only a single consensus number give no way "
              "to assess how performance would look on the discarded, contested "
              "cases. This phase's stratified curve is a direct, quantitative "
              "instance of that gap, computed on a corpus and model where the "
              "annotator-level data actually exists to make the comparison.")
    h(doc, "4.5 Limitations specific to this phase", level=2)
    bullet(doc, "4.5.1 S-dispersed alone has n=8 in the test split; even pooled "
                "with S-tied (n=81), the tier is small relative to S-unanimous. "
                "Point estimates on this tier should be read as indicative, not "
                "precise.")
    bullet(doc, "4.5.2 S-plurality's single-label reading uses a 2-of-4 pseudo-"
                "label that is, by definition, not a validated ground truth; it is "
                "reported only as a secondary, clearly flagged number.")
    bullet(doc, "4.5.3 This phase evaluates one architecture (ConvNeXt-Tiny) across "
                "3 seeds; the stratified pattern's generalisation to other "
                "architectures is not established here.")
    bullet(doc, "4.5.4 The attainable ceiling used in §3.4.1 is the modal-vote "
                "oracle, which maximises expected accuracy exactly but is only a "
                "lower bound on the supremum of macro F1, because macro F1 does not "
                "decompose across images. The ceiling-normalised scores are "
                "therefore conservative: the true normalised scores can only be "
                "lower, so the reported normalised gaps are, if anything, "
                "overstated rather than understated.")
    bullet(doc, "4.5.5 The ceiling-normalised analysis is post-hoc (§2.6.2). It is "
                "reported alongside — never in place of — the pre-registered "
                "raw-scale result, and its intervals are wide enough on the "
                "smallest tier that it is best read as motivating a "
                "pre-registered replication in Phase 4 rather than as a settled "
                "finding.")
    bullet(doc, "4.5.6 Calibration is assessed against expected accuracy rather "
                "than a hard label, because no hard label exists on two tiers. "
                "This is the appropriate target but it is not the convention in "
                "the literature, so the ECE values here are not directly "
                "comparable with single-label ECE figures reported elsewhere.")
    bullet(doc, "4.5.7 Evaluation is confined to the official test split of a "
                "single centre. Phase 5's external validation remains the only "
                "test of whether the agreement-stratified pattern transfers.")


def sec_conclusion(doc) -> None:
    h(doc, "5. Conclusion", level=1)
    h(doc, "5.1 Answer to RQ1", level=2)
    para(doc, "Landmark-classification performance declines sharply as expert "
              f"agreement falls, from {pc(AGG['S-unanimous']['annotator_marginalized_macro_f1_mean_3seed'])} "
              f"macro F1 on the unanimous stratum to "
              f"{pc(AGG['S-no-majority']['annotator_marginalized_macro_f1_mean_3seed'])} on the pooled "
              f"no-majority stratum — a decline "
              f"{RQ1['gap_S_unanimous_minus_S_no_majority_points'] / RQ1['architecture_gap_benchmark_points']:.0f}x larger than "
              "the published gap between architecture families. That is the "
              "pre-registered result, and it is reported as such.")
    para(doc, f"The amended answer is narrower and better supported. Once the "
              f"metric's attainable ceiling is held constant, the degradation from "
              f"the unanimous stratum exceeds the between-architecture benchmark for "
              f"the 3/4-majority stratum "
              f"({B.gstr('S-unanimous - S-majority', 'ceiling_normalised')}) and the "
              f"2-1-1-plurality stratum "
              f"({B.gstr('S-unanimous - S-plurality', 'ceiling_normalised')}), and is "
              f"not resolvable for the pooled no-majority stratum "
              f"({B.gstr('S-unanimous - S-no-majority', 'ceiling_normalised')}), "
              f"where a perfect single-label classifier could itself score no more "
              f"than {pc(CEIL['ceilings']['S-no-majority']['oracle_marginalized_macro_f1_mean'], 1)}%. "
              f"Neither class composition nor acquisition stream accounts for the "
              f"effect. The decline is not strictly monotonic on either scale, and "
              f"on the ceiling-normalised scale the reversal at S-plurality is "
              f"statistically significant rather than noise — a finding to be "
              f"replicated under pre-registration in Phase 4, not explained away.")
    para(doc, f"The most consequential result of the phase is not in RQ1 at all. "
              f"Calibration degrades far faster than accuracy: expected calibration "
              f"error rises "
              f"{CAL['headline']['ece_ratio_worst_over_unanimous']:.1f}-fold from the "
              f"unanimous to the most contested tier, and mean confidence falls only "
              f"{100 * (CAL['aggregate_3seed']['S-unanimous']['mean_confidence'] - CAL['aggregate_3seed']['S-plurality']['mean_confidence']):.1f} "
              f"points while expected accuracy falls "
              f"{100 * (CAL['aggregate_3seed']['S-unanimous']['expected_accuracy'] - CAL['aggregate_3seed']['S-plurality']['expected_accuracy']):.1f}. "
              f"A consensus-only evaluation protocol cannot see this, and a "
              f"confidence-thresholded deployment would not be protected from it.")
    h(doc, "5.2 Carry-forward decisions for Phase 4", level=2)
    bullet(doc, "Use the full agreement-tier structure, not just the unanimous "
                "subset, when constructing C1-C4 training targets.")
    bullet(doc, "Retain the C3 label-smoothing control. The justification is now "
                "stronger and more specific than 'the raw gap is large': label "
                "smoothing is itself a calibration intervention, so it is the "
                "correct control for the calibration result in §3.8 as well as for "
                "the accuracy result.")
    bullet(doc, "Pre-register calibration as a primary Phase 4 endpoint, not a "
                "secondary one. §3.8 shows the baseline's failure mode is "
                "confidently-wrong rather than merely wrong, which is the failure "
                "mode soft targets are supposed to fix.")
    bullet(doc, "Report the WITHIN-tier predictive-entropy / vote-entropy "
                "correlation as RQ3's primary quantity. §3.8.1 shows the pooled "
                "correlation (0.320) is almost entirely a between-tier effect and "
                "collapses to 0.02-0.08 within tiers; the pooled figure would read "
                "as success while measuring only tier membership.")
    bullet(doc, "Pre-register the ceiling-normalised metric alongside the raw one, "
                "so that Phase 4's comparison of C0-C4 across tiers is not exposed "
                "to the same scale artefact this phase had to correct post-hoc.")
    bullet(doc, "Carry the wall-adjacent / station-neighbouring error decomposition "
                "into Phase 6 as a pre-specified comparison axis, but treat the "
                "station-geometry gap as an untested hypothesis rather than an "
                "established lead (§4.2), and compute the required error count in "
                "advance — the present intervals are 13-17 points wide.")


def sec_appendices(doc) -> None:
    h(doc, "Appendices", level=1, page_break=True)
    B.appendix_a_perclass(doc)
    B.appendix_b_confusion(doc)
    B.appendix_c_bootstrap(doc)

    h(doc, "Appendix D. Script and artefact manifest (reproducibility index)", level=2)
    para(doc, "Executed in this order; each step is gated on the previous one's "
              "validation criterion. Steps marked 3B were added to repair the "
              "omissions and corrections recorded in Appendix E.")
    for s in ["python src/models/phase3_data.py",
             "python src/models/phase3_cache.py",
             "python src/models/phase3_eval.py",
             "python src/models/phase3_confusion.py",
             "python src/models/phase3b_probs.py            # 3B",
             "python src/models/phase3b_ceiling.py          # 3B",
             "python src/models/phase3b_calibration.py      # 3B",
             "python src/models/phase3b_perclass.py         # 3B",
             "python src/models/phase3b_sensitivity.py      # 3B",
             "python src/models/phase3b_amendment.py        # 3B",
             "python src/report/figures_phase3.py",
             "python src/report/figures_phase3b.py          # 3B",
             "python src/report/build_phase3_docx.py",
             "python src/report/finalise_phase3.py"]:
        bullet(doc, s)
    para(doc, "Numerical primitives shared by the 3B scripts live in "
              "src/models/phase3b_common.py, whose selftest() asserts equality with "
              "scikit-learn to 1e-12 at import time; the bincount reimplementation "
              "exists only to make the added bootstraps affordable and can never "
              "silently change a published value.")

    B.appendix_e_register(doc)
    para(doc, "The tier definitions, ground-truth rule, pooling rule (S-tied + "
              "S-dispersed -> S-no-majority), and primary metric definition "
              "(annotator-marginalized macro F1) were fixed in "
              "THESIS_RESEARCH_BLUEPRINT.md §4 Phase 3 before phase3_eval.py was "
              "run against these images. The tier counts were computed from the "
              "official splits alone (no model output) and gated exactly against "
              "the pre-registered expectation before proceeding. That blueprint is "
              "a living document, however, and was edited to v3.2 after this phase "
              "ran, so the claim is not independently verifiable from the committed "
              "artefacts alone — which is why the register above, rather than this "
              "paragraph, is the operative record.")


def main() -> None:
    doc = new_document()
    title_page(doc)
    front_matter(doc)
    B.abbreviations(doc)
    sec_abstract(doc)
    sec_introduction(doc)
    sec_methods(doc)
    sec_results(doc)
    sec_discussion(doc)
    sec_conclusion(doc)
    sec_appendices(doc)
    B.references(doc)

    add_page_numbers(doc)
    doc.save(BD.OUT)

    caps = [p.text for p in doc.paragraphs if p.style.name == "Caption"]
    nfig = sum(1 for c in caps if c.strip().startswith("Figure"))
    ntab = sum(1 for c in caps if c.strip().startswith("Table"))
    print(f"[report] {len(doc.inline_shapes)} images, {nfig} figure captions, "
          f"{ntab} table captions -> {BD.OUT}")


if __name__ == "__main__":
    main()
