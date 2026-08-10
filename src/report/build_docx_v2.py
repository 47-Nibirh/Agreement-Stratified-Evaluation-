"""
Build the GastroHUN Phase 0 / Phase 1 Word report.
==================================================
Assembles `Phase0_Phase1_Report.docx` from the computed artefacts:

  reports/gastrohun_inventory.json        physical inventory
  reports/gastrohun_agreement.json        agreement, splits, power
  reports/gastrohun_structure.json        wall/station decomposition
  reports/gastrohun_neardup.json          contamination scan
  reports/gastrohun_dup_calibration.json  calibrated duplicate rule
  reports/phase0_results.json             retired corpus (negative control)
  literature_v2/extraction_table.csv      included studies with APA strings
  literature_v2/prisma_counts.json        PRISMA stage counts
  figures_v2/*.png                        generated figures

Reuses the rendering helpers of build_docx.py, repointed at the new figure
directory and output filename. No numeric value is typed by hand.

Run:  python src/report/build_docx_v2.py
"""

from __future__ import annotations

import json
from pathlib import Path

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt

import build_docx as BD

ROOT = Path(__file__).resolve().parents[2]

# Repoint the shared renderer at this report's assets before anything imports
# the helpers that close over them.
BD.FIGD = ROOT / "figures_v2"
BD.OUT = ROOT / "Phase0_Phase1_Report.docx"

from build_docx import (ACCENT, DARKRED, GREY, _cell_text, add_page_numbers,  # noqa: E402
                        front_matter, new_document)

INV = json.loads((ROOT / "reports" / "gastrohun_inventory.json").read_text(encoding="utf-8"))
AGR = json.loads((ROOT / "reports" / "gastrohun_agreement.json").read_text(encoding="utf-8"))


def title_page(doc) -> None:
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Agreement-Stratified Evaluation of Deep Learning for "
                  "Anatomical Landmark Recognition in Upper Gastrointestinal "
                  "Endoscopy")
    r.bold = True; r.font.size = Pt(19); r.font.color.rgb = DARKRED

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run("Phase 0 — Data Provenance and Integrity Gate\n"
                  "Phase 1 — Literature Review and Problem Framing")
    r.bold = True; r.font.size = Pt(13.5); r.font.color.rgb = ACCENT

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Interim Technical Report")
    r.font.size = Pt(11.5); r.italic = True

    doc.add_paragraph()
    bar = doc.add_paragraph(); bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = bar.add_run("─" * 46); rb.font.color.rgb = ACCENT

    meta = [
        ("Degree programme", "B.Sc. in Computer Science and Engineering"),
        ("Research domain", "Biomedical Artificial Intelligence — Medical Image "
                            "Analysis and Machine Learning"),
        ("Corpus under audit",
         f"GastroHUN 'Labeled Images' — {AGR['n_images']:,} images, "
         f"{AGR['n_patients']} patients, {AGR['n_classes']} classes"),
        ("Corpus provenance",
         "Hospital Universitario Nacional de Colombia; Sci Data 12:102 (2025); "
         "doi:10.1038/s41597-025-04401-5"),
        ("Ethics approval", "CEI-2019-06-10 (Ethics Committee, HUN); informed "
                            "consent obtained"),
        ("Audited payload", f"{INV['bytes']['total_gb']:.2f} GB, "
                            f"{INV['n_decoded_ok']:,} images verified"),
        ("Reporting standards", "PRISMA 2020, CLAIM, TRIPOD+AI, STARD-AI"),
        ("Report date", "26 July 2026"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in meta:
        c = t.add_row().cells
        _cell_text(c[0], k, bold=True, size=9.5)
        _cell_text(c[1], v, size=9.5)
        c[0].width, c[1].width = Cm(4.6), Cm(11.0)

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("STATUS: PROCEED — the integrity gate cleared, with two "
                  "conditional criteria carried forward as declared limitations.")
    r.bold = True; r.font.size = Pt(10.5)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("This report supersedes the previous Phase 0 / Phase 1 report "
                  "in full. No content from that report is carried forward.")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def main() -> None:
    doc = new_document()
    title_page(doc)
    front_matter(doc)

    import content_phase0_v2 as C0
    import content_phase0b_v2 as C0B
    import content_phase1_v2 as C1

    C0.sec_executive_summary(doc)
    C0.sec_phase0(doc)
    C0.sec_contamination(doc)
    C0B.sec_agreement(doc)
    C0B.sec_structure(doc)
    C0B.sec_splits_power(doc)
    C0B.sec_gate_verdict(doc)

    C1.sec_phase1(doc)
    C1.sec_synthesis(doc)
    C1.sec_gap(doc)
    C1.sec_methodology(doc)
    C1.sec_limitations(doc)
    C1.sec_conclusion(doc)
    C1.sec_references(doc)
    C1.sec_appendix(doc)

    add_page_numbers(doc)
    doc.save(BD.OUT)

    caps = [p.text for p in doc.paragraphs if p.style.name == "Caption"]
    nfig = sum(1 for c in caps if c.strip().startswith("Figure"))
    ntab = sum(1 for c in caps if c.strip().startswith("Table"))
    words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"[report] {len(doc.inline_shapes)} images, {nfig} figure captions, "
          f"{ntab} table captions, ~{words:,} words -> {BD.OUT}")


if __name__ == "__main__":
    main()
